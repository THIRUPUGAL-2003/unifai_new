# -*- coding: utf-8 -*-
"""
Full Document Generator for UnifAI Enterprise Architecture, Headers & Features Guide
Outputs:
  - document/UnifAI_Architecture_and_Features_Guide.md
  - document/UnifAI_Architecture_and_Features_Guide.docx
  - document/UnifAI_Architecture_and_Features_Guide.pdf
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

# Ensure UTF-8 output encoding for Windows console
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

DOC_DIR = os.path.join(os.getcwd(), "document")
os.makedirs(DOC_DIR, exist_ok=True)

MD_PATH = os.path.join(DOC_DIR, "UnifAI_Architecture_and_Features_Guide.md")
DOCX_PATH = os.path.join(DOC_DIR, "UnifAI_Architecture_and_Features_Guide.docx")
PDF_PATH = os.path.join(DOC_DIR, "UnifAI_Architecture_and_Features_Guide.pdf")

# ==============================================================================
# 1. MARKDOWN CONTENT GENERATION
# ==============================================================================
print("Generating Markdown documentation...")

md_content = """# UnifAI Enterprise Architecture, Headers & Features Master Guide
## விரிவான கட்டமைப்பு, ஹெடர்ஸ் மற்றும் அம்சங்களின் கையேடு (Tamil & English Technical Manual)

**Document Version:** 2.0 (Enterprise Edition)  
**Classification:** Technical Architecture & System Design Manual  
**Generated At:** 2026-09-05  
**Platform:** UnifAI Unified AI Gateway & Governance Control Plane  

---

## Table of Contents (பொருளடக்கம்)
1. [Executive Summary & Platform Overview (செயல் சுருக்கம்)](#1-executive-summary--platform-overview)
2. [End-to-End Request Lifecycle & System Architecture (கோரிக்கை வாழ்க்கைச் சுழற்சி)](#2-end-to-end-request-lifecycle--system-architecture)
3. [Exhaustive HTTP Headers Reference Guide (ஹெடர்ஸ் முழு விவரக் கையேடு)](#3-exhaustive-http-headers-reference-guide)
   - 3.1 Authentication & Identity Headers
   - 3.2 Model Routing & Selection Headers
   - 3.3 Semantic Caching & Performance Headers
   - 3.4 MCP (Model Context Protocol) & Tool Headers
   - 3.5 Observability, Tracing & Debugging Headers
   - 3.6 Protocol Compatibility & Transformation Headers
   - 3.7 Upstream & Gateway Response Headers (`x-unifai-*`)
4. [In-Depth Feature Catalog (7 Core Pillars - 38 Features)](#4-in-depth-feature-catalog)
   - 4.1 Observability (கண்காணிப்பு & செயல்திறன்)
   - 4.2 Models (மாடல் மேலாண்மை & ரூட்டிங்)
   - 4.3 MCP Gateway (டூல்ஸ் & ஏஜென்ட்கள்)
   - 4.4 Governance (அணுகல் கட்டுப்பாடு & பாதுகாப்பு)
   - 4.5 Guardrails (உள்ளடக்கப் பாதுகாப்பு & கொள்கை அமலாக்கம்)
   - 4.6 Adaptive Routing (ஸ்மார்ட் தேர்வு & அறிவு களஞ்சியம்)
   - 4.7 Global Settings (கட்டமைப்பு & செயல்திறன் ட்யூனிங்)
5. [Cross-Feature Interconnection & Data Flow Matrix (இணைப்பு வரைபடம்)](#5-cross-feature-interconnection--data-flow-matrix)
6. [Technology Stack & Programming Languages Deep Dive (தொழில்நுட்ப கட்டமைப்பு)](#6-technology-stack--programming-languages-deep-dive)
7. [Enterprise Production Scenarios & Playbooks (நடைமுறை பயன்பாடுகள்)](#7-enterprise-production-scenarios--playbooks)

---

# 1. Executive Summary & Platform Overview
### செயல் சுருக்கம் & தளம் கண்ணோட்டம்

UnifAI is a high-performance, enterprise-grade **Unified AI Gateway, Router, Governance, Guardrails, and Observability Control Plane**. In modern organizations, AI consumption is fragmented across multiple providers (OpenAI, Anthropic Claude, Google Gemini, AWS Bedrock, Mistral, Groq, and self-hosted Ollama/vLLM models). 

UnifAI unifies all LLM interactions behind a single secure endpoint, providing:
1. **Cost Optimization (செலவு குறைப்பு):** Up to 80% cost reduction through **Complexity Routing** (sending easy questions to smaller models) and **Semantic Caching** (zero-cost answers for repeated queries).
2. **High Availability & Zero Downtime (நம்பகத்தன்மை):** Real-time **Circuit Breaker** and automatic fallback routing when primary LLM providers experience outages or latency spikes.
3. **Enterprise Governance & Security (அணுகல் கட்டுப்பாடு):** Virtual API Keys, Role-Based Access Control (RBAC), multi-tenant customer attribution, and strict rate limits/budgets per team.
4. **Data Loss Prevention (DLP) & Guardrails (உள்ளடக்கப் பாதுகாப்பு):** Input and output content inspection via Google CEL rules and providers (Presidio DLP, Llama Guard) to block PII leaks and prompt injections.
5. **Agentic Extensibility via MCP Gateway (டூல்ஸ் பயன்பாடு):** Standardized connection to external tools, databases, and APIs using the Model Context Protocol (MCP).
6. **Full-Stack Observability (கண்காணிப்பு):** Real-time logging of prompts, completions, tokens, latency, TTFT, and export connectors to Datadog, New Relic, Kafka, and BigQuery.

---

# 2. End-to-End Request Lifecycle & System Architecture
### கோரிக்கை வாழ்க்கைச் சுழற்சி & கணினி கட்டமைப்பு

When an application, SDK, or developer makes an API request to UnifAI, the request traverses 7 distinct architectural layers:

```
[ CLIENT APPLICATION / SDK / BROWSER AI ]
                   │
                   ▼  (1) HTTP POST with `x-uf-vk: sk-uf-...` & payload
┌────────────────────────────────────────────────────────────────────────┐
│ 1. TRANSPORT & GOVERNANCE LAYER                                        │
│ • FastHTTP router parses headers and TLS connection                    │
│ • Virtual Key (`sk-uf-*`) validated against PostgreSQL / KV Store      │
│ • Check User, Team, Business Unit, and Customer membership             │
│ • Enforce Rate Limits (RPM, TPM) & Financial Monthly Budgets           │
│ • Verify Access Profile (allowed models & allowed MCP tool groups)     │
└──────────────────────────────────┬─────────────────────────────────────┘
                                   ▼  (2) Authorized Context
┌────────────────────────────────────────────────────────────────────────┐
│ 2. PRE-LLM GUARDRAILS & SAFETY LAYER                                   │
│ • Evaluate Google CEL (Common Expression Language) Rules               │
│ • Provider scan: Presidio DLP (PII redaction), Llama Guard, Regex      │
│ • Prompt Injection & Jailbreak detection                               │
│ • Short-circuit with 400 Bad Request if policy violated                │
└──────────────────────────────────┬─────────────────────────────────────┘
                                   ▼  (3) Safe, Sanitized Prompt
┌────────────────────────────────────────────────────────────────────────┐
│ 3. SEMANTIC CACHING LAYER                                              │
│ • Generate vector embedding of prompt or exact SHA256 key              │
│ • Search Qdrant / PgVector / Redis cache partition (`x-uf-cache-key`)   │
│ • If Cosine Similarity >= Threshold (`x-uf-cache-threshold`):          │
│   ───► [CACHE HIT] Return cached response directly (Latency < 20ms)    │
└──────────────────────────────────┬─────────────────────────────────────┘
                                   ▼  (4) Cache Miss -> Needs LLM
┌────────────────────────────────────────────────────────────────────────┐
│ 4. INTELLIGENT ROUTING & LOAD BALANCING LAYER                          │
│ • Complexity Router analyzes prompt (Simple, Medium, Complex, Reason)   │
│ • Routing Rules match CEL conditions (tier, department, tags)          │
│ • Adaptive Router checks real-time latency & error rate weights        │
│ • Circuit Breaker checks provider health; routes to fallback if down   │
│ • Resolved target model & provider determined                          │
└──────────────────┬────────────────────────────────┬────────────────────┘
                   ▼                                ▼
┌────────────────────────────────────┐ ┌─────────────────────────────────┐
│ 5A. UPSTREAM LLM PROVIDER          │ │ 5B. MCP GATEWAY EXECUTION       │
│ • OpenAI, Anthropic, Bedrock, etc. │ │ • Model Context Protocol Tools  │
│ • API key from rotation pool       │ │ • OAuth token session injected  │
│ • Streaming SSE / Full response    │ │ • Sandboxed Starlark execution  │
└──────────────────┬─────────────────┘ └────────────────┬────────────────┘
                   └────────────────┬───────────────────┘
                                    ▼  (5) Raw Generated AI Output
┌────────────────────────────────────────────────────────────────────────┐
│ 6. POST-LLM GUARDRAILS & OUTPUT FILTERING                              │
│ • Scan response for hallucinated credentials or toxic outputs          │
│ • PII masking on generated text before sending to client               │
└──────────────────────────────────┬─────────────────────────────────────┘
                                   ▼  (6) Final Validated Response
┌────────────────────────────────────────────────────────────────────────┐
│ 7. OBSERVABILITY, METRICS & TELEMETRY                                  │
│ • Append complete record to LLM Logs & MCP Logs                        │
│ • Calculate exact token cost via Pricing Overrides                     │
│ • Update real-time Metrics Dashboard via WebSockets                    │
│ • Stream telemetry event to Connectors (Datadog, Kafka, BigQuery, OTel)│
│ • Attach `x-unifai-*` response headers and return HTTP 200 OK to Client│
└────────────────────────────────────────────────────────────────────────┘
```

---

# 3. Exhaustive HTTP Headers Reference Guide
### ஹெடர்ஸ் முழு விவரக் கையேடு

Headers play a central role in UnifAI. They allow developers to configure request routing, caching, security, debugging, and logging per request without modifying application payload structures.

## 3.1 Authentication & Identity Headers (அடையாளம் & பாதுகாப்பு)

| Header Name | Type / Format | என்ன செய்கிறது & எதற்கு பயன்படும் (Why & Use Case) |
| :--- | :--- | :--- |
| **`x-uf-vk`** | `String`<br>`sk-uf-abc123xyz` | **Virtual Key Header**: Authenticates the caller. Binds the request to a specific Team, User, Budget quota, and Model permissions profile. Alternately passed as `Authorization: Bearer sk-uf-...`. |
| **`x-uf-api-key`** | `String` | **Platform Admin API Key**: Used to access UnifAI control-plane endpoints (creating virtual keys, configuring routing rules, modifying cluster configs). |
| **`x-uf-customer-id`** | `String` | **Customer Identifier**: In multi-tenant B2B applications, routes costs and log attribution to a specific customer account under a shared team virtual key. |
| **`x-uf-customer-name`** | `String` | **Customer Human-Readable Name**: Used alongside customer-id for display in audit trails and cost analytics. |
| **`x-uf-direct-key`** | `Boolean`<br>`true / false` | **Direct Key Pass-Through**: When set to `true`, tells UnifAI to ignore registered provider keys and use the raw API key supplied in standard provider headers (e.g., `x-api-key` or `x-goog-api-key`). |
| **`X-UnifAI-Temp-Token`** | `String`<br>`JWT / UUID` | **Temporary Session Token**: Used for short-lived interactive authentications, such as MCP user consent callbacks and UI workspace sessions. |

## 3.2 Model Routing & Selection Headers (ரூட்டிங் & மாடல் தேர்வு)

| Header Name | Type / Format | என்ன செய்கிறது & எதற்கு பயன்படும் (Why & Use Case) |
| :--- | :--- | :--- |
| **`x-uf-provider`** | `String`<br>`openai, anthropic, bedrock...` | **Explicit Provider Override**: Forces UnifAI to route the request to a specific upstream provider, bypassing automated routing rules. |
| **`x-uf-model`** | `String`<br>`gpt-4o, claude-3-7-sonnet...` | **Explicit Model Override**: Specifies the target model name when not defined or overridden in the JSON body. |
| **`x-uf-api-key-id`** | `String`<br>`key_12345` | **Key Pool Pinning**: When a provider has multiple configured credentials (e.g., 5 Azure OpenAI deployments in different regions), pins the request to a specific key ID. |
| **`x-uf-circuit-breaker-bypass`** | `Boolean`<br>`true` | **Circuit Breaker Bypass**: For administrative diagnostic probes, forces the request through a provider even if its circuit breaker is currently TRIPPED. |

## 3.3 Semantic Caching & Performance Headers (கேச்சிங் & செயல்திறன்)

| Header Name | Type / Format | என்ன செய்கிறது & எதற்கு பயன்படும் (Why & Use Case) |
| :--- | :--- | :--- |
| **`x-uf-cache-key`** | `String`<br>`tenant-42:feature-chat` | **Cache Partition Key**: Isolates cached responses into distinct logical namespaces so that different customers or features never collide. |
| **`x-uf-cache-ttl`** | `Integer`<br>`Seconds (e.g. 3600)` | **Cache Time-To-Live Override**: Sets how long this specific response should remain cached before automatic expiration. |
| **`x-uf-cache-threshold`** | `Float`<br>`0.0 to 1.0 (e.g. 0.88)` | **Semantic Similarity Threshold**: Sets the minimum cosine similarity required for a vector match. Higher = stricter match; Lower = more frequent hits. |
| **`x-uf-cache-type`** | `String`<br>`direct / semantic` | **Cache Mode Selection**: `direct` checks exact SHA-256 string hash; `semantic` performs embedding vector similarity search. |
| **`x-uf-cache-no-store`** | `Boolean`<br>`true` | **No-Store Directive**: Allows reading from cache if a hit exists, but prevents writing new LLM outputs into the cache. |

## 3.4 MCP (Model Context Protocol) & Tool Headers (டூல்ஸ் பயன்பாடு)

| Header Name | Type / Format | என்ன செய்கிறது & எதற்கு பயன்படும் (Why & Use Case) |
| :--- | :--- | :--- |
| **`x-uf-mcp-include-clients`** | `String`<br>`*` or `github,postgres` | **Allowed MCP Servers**: Comma-separated list of MCP server names to inject into the LLM request. `*` enables all permitted servers. |
| **`x-uf-mcp-include-tools`** | `String`<br>`exec_query,read_file` | **Allowed Tool Whitelist**: Specific function names within allowed MCP servers that the model is permitted to execute. |
| **`x-uf-mcp-session-id`** | `String`<br>`sess_usr_9981` | **Per-User MCP Session**: Associates stateful OAuth credentials (e.g., individual user's GitHub or Google Drive auth) with the tool execution. |
| **`x-uf-eh-*`** | `String`<br>`x-uf-eh-custom-id: val` | **Extra Headers to MCP Server**: Prefixed headers forwarded verbatim to upstream MCP servers if present in their `allowed_extra_headers` config. |

## 3.5 Observability, Tracing & Debugging Headers (கண்காணிப்பு & பிழைத்திருத்தம்)

| Header Name | Type / Format | என்ன செய்கிறது & எதற்கு பயன்படும் (Why & Use Case) |
| :--- | :--- | :--- |
| **`x-uf-session-id`** | `String`<br>`session_abc123` | **Conversation Session ID**: Correlates multiple sequential LLM calls into a single conversational session for tracing and analytics. |
| **`x-uf-dim-<name>`** | `String`<br>`x-uf-dim-env: prod` | **Custom Log Dimensions**: Arbitrary key-value metadata attached to internal log records for filtering (e.g., `x-uf-dim-department: hr`). |
| **`x-uf-lh-<name>`** | `String`<br>`x-uf-lh-correlation-id` | **Captured Request Header**: Explicitly instructs UnifAI to store the specified request header into LLM log records. |
| **`x-uf-disable-content-logging`** | `Boolean`<br>`true` | **Content Privacy Mode**: Drops prompt and completion text from logs, capturing only token metrics, cost, latency, and model metadata. |
| **`x-uf-store-raw-request-response`**| `Boolean`<br>`true` | **Raw Byte Capture**: Stores exact upstream provider wire payloads in internal storage for low-level protocol inspection. |
| **`x-uf-send-back-raw-request`** | `Boolean`<br>`true` | **Debug Echo Request**: Returns the exact JSON payload sent to the upstream provider in a debug response header. |
| **`x-uf-send-back-raw-response`** | `Boolean`<br>`true` | **Debug Echo Response**: Returns the unparsed upstream provider response in a debug header. |
| **`traceparent`** | `W3C Trace Format` | **Distributed Tracing Header**: Connects the UnifAI request span to an existing distributed OpenTelemetry trace across microservices. |
| **`x-uf-log-repo-id`** | `String` | **Target Evaluation Repo**: Directs telemetry logs to a specific repository in evaluation platforms like Maxim AI. |

## 3.6 Protocol Compatibility Headers (பரிமாற்ற வடிவமைப்பு)

| Header Name | Type / Format | என்ன செய்கிறது & எதற்கு பயன்படும் (Why & Use Case) |
| :--- | :--- | :--- |
| **`x-uf-compat`** | `String` | **API Format Bridge**: Configures real-time translation between OpenAI and non-OpenAI formats (e.g., Anthropic native messages, Bedrock Converse). |

## 3.7 Upstream & Gateway Response Headers (`x-unifai-*`)

When returning a response, UnifAI injects the following response headers:

| Response Header | Sample Value | விளக்கம் (Description & Utility) |
| :--- | :--- | :--- |
| **`x-unifai-provider`** | `anthropic` | Displays which model provider handled and completed the request. |
| **`x-unifai-original-model`** | `gpt-4o` | The model name requested by the client prior to any router transformation. |
| **`x-unifai-resolved-model`** | `claude-3-5-sonnet` | The actual model that executed the prompt after routing rules or complexity analysis. |
| **`x-unifai-fallback-index`** | `1` | `0` if primary provider succeeded; `1, 2...` indicating which backup fallback provider was triggered. |
| **`x-unifai-request-type`** | `chat` | Identifies the interaction type: `chat`, `completion`, `embedding`, `realtime`, `audio`. |

---

# 4. In-Depth Feature Catalog (7 Core Pillars - 38 Features)
### விரிவான அம்சங்களின் பட்டியல்

---

## 4.1 Observability (கண்காணிப்பு & செயல்திறன்)
Observability delivers real-time visibility into cost, performance, and behavior across all AI workloads.

1. **Dashboard (`/workspace/dashboard`)**:
   * *What it does:* Interactive analytics interface showing total API calls, token counts, aggregate dollar spend, latency percentiles (p50, p95, p99), error distributions, top active models, and virtual key usage.
   * *Why & Where used:* CTOs, engineering managers, and FinOps teams use it to track departmental spending, detect unexpected cost spikes, and monitor SLA performance.

2. **LLM Logs (`/workspace/logs`)**:
   * *What it does:* Full-fidelity request-response ledger capturing prompt text, generated completions, token counts, time-to-first-token (TTFT), duration, latency, cost, and guardrail verdicts.
   * *Why & Where used:* Engineers use it for root-cause debugging of failed requests, verifying prompt template behavior, and monitoring model response quality.

3. **MCP Logs (`/workspace/mcp-logs`)**:
   * *What it does:* Dedicated execution log for Model Context Protocol interactions. Records tool function names, input arguments, execution time, return payloads, and exceptions.
   * *Why & Where used:* AI agent developers use it to audit agent tool-calling loops, identify flaky API tools, and diagnose unexpected tool behavior.

4. **Browser AI (`/workspace/browser-ai`)**:
   * *What it does:* An enterprise browser extension and local DLP proxy that intercepts employee interactions on web AI portals (ChatGPT, Claude, Perplexity). Enforces DLP rules and audits prompts.
   * *Why & Where used:* Security and compliance teams use it to prevent employees from pasting proprietary intellectual property, customer PII, or API secrets into public AI interfaces.

5. **Connectors (`/workspace/observability`)**:
   * *What it does:* Streaming export pipelines that forward telemetry, logs, and trace events to external enterprise platforms (Datadog, New Relic, Apache Kafka, Google Cloud BigQuery, Google PubSub, OpenTelemetry OTLP).
   * *Why & Where used:* Enterprises use it to centralize AI metrics within their existing corporate SIEM, enterprise data lake, or security operations center (SOC).

6. **Logs Settings (`/workspace/config/logging`)**:
   * *What it does:* Centralized log management console to configure retention schedules (e.g., 30 days), sampling rates (e.g., 10% of high-volume calls), content redaction rules, and raw byte storage toggles.
   * *Why & Where used:* Compliance officers use it to satisfy GDPR, HIPAA, and SOC 2 data storage regulations.

---

## 4.2 Models (மாடல் மேலாண்மை & ரூட்டிங்)
The Models pillar manages upstream LLM integrations, intelligent traffic routing, and cost control.

7. **Model Catalog (`/workspace/model-catalog`)**:
   * *What it does:* Master directory of all configured models across all connected providers. Displays modalities (text, vision, audio), context limits, supported features (function calling, streaming), and live availability.
   * *Why & Where used:* Developers use it to discover available models in the workspace and verify their parameter capabilities.

8. **Model Providers (`/workspace/providers`)**:
   * *What it does:* Connection manager for upstream AI providers (OpenAI, Anthropic, AWS Bedrock, Google Vertex, Azure OpenAI, Mistral, Groq, Cohere, Replicate, Ollama). Stores API keys, endpoints, and timeouts.
   * *Why & Where used:* Platform administrators use it to securely add, rotate, and manage provider credentials.

9. **Budgets & Limits (`/workspace/model-limits`)**:
   * *What it does:* Financial and throughput quota engine. Enforces rate limits (Requests Per Minute - RPM, Tokens Per Minute - TPM) and hard/soft spending ceilings per Virtual Key, User, or Model.
   * *Why & Where used:* FinOps teams use it to prevent runaway bills from infinite loops or unexpected traffic surges.

10. **Routing Rules (`/workspace/routing-rules`)**:
    * *What it does:* Condition-based dynamic routing engine evaluated using Google Common Expression Language (CEL). Routes requests based on model aliases, user department, prompt keywords, or time of day.
    * *Why & Where used:* Enables model aliasing (e.g., routing `prod-chat` to `gpt-4o-mini` during peak hours or `claude-3-5-sonnet` for code queries).

11. **Complexity Router (`/workspace/complexity-router`)**:
    * *What it does:* Real-time prompt complexity classifier that buckets prompts into Simple, Medium, Complex, or Reasoning tiers based on keyword rules, token count, or lightweight classifiers.
    * *Why & Where used:* Slashes AI operational costs by 60–80% by routing simple greetings and lookups to cheap micro-models while reserving expensive frontier models for deep reasoning.

12. **Circuit Breaker (`/workspace/circuit-breaker`)**:
    * *What it does:* Automated fault detection and failover engine. Continuously tracks upstream error rates and latency. If a provider fails, it trips the circuit and reroutes traffic to healthy backup providers.
    * *Why & Where used:* Guarantees 99.99% application uptime even during major public cloud or OpenAI outages.

13. **Pricing Overrides (`/workspace/custom-pricing/overrides`)**:
    * *What it does:* Granular cost definition table supporting custom per-token pricing ($/1M prompt tokens, $/1M completion tokens) per model, provider, or enterprise negotiated discount rate.
    * *Why & Where used:* Allows enterprises with custom provider discounts or proprietary fine-tuned models to calculate accurate financial spend.

14. **Model Settings (`/workspace/custom-pricing`)**:
    * *What it does:* Global default pricing rules and fallback cost structures applied when specific provider overrides are not set.

---

## 4.3 MCP Gateway (டூல்ஸ் & ஏஜென்ட்கள்)
The Model Context Protocol (MCP) Gateway standardizes agentic tool integration and execution.

15. **MCP Catalog (`/workspace/mcp-registry`)**:
    * *What it does:* Comprehensive registry of all registered MCP servers, active server instances, transport types (stdio, SSE), and available tool definitions.
    * *Why & Where used:* Agent developers use it to inspect what tools and external APIs are currently available for autonomous agents to use.

16. **MCP Library (`/workspace/mcp-registry/library`)**:
    * *What it does:* One-click curated repository of pre-built, production-ready MCP servers (GitHub, PostgreSQL, Slack, Google Drive, Jira, Filesystem, Brave Search).
    * *Why & Where used:* Dramatically accelerates agent development by eliminating the need to write custom integration scrapers.

17. **Tool Groups (`/workspace/mcp-tool-groups`)**:
    * *What it does:* Logical grouping and security boundary mechanism that bundles specific tools together (e.g., `Finance-Tools`, `DevOps-Tools`, `Customer-Support-Tools`).
    * *Why & Where used:* Attached to Virtual Keys via Access Profiles to ensure agents only access the specific tools required for their job.

18. **Auth Sessions (`/workspace/mcp-sessions`)**:
    * *What it does:* Stateful credential manager that tracks per-user authenticated sessions for external MCP tools requiring individual user accounts.
    * *Why & Where used:* Enables an AI agent to act on behalf of a specific logged-in user without sharing global administrative credentials.

19. **OAuth Grants (`/workspace/oauth-grants`)**:
    * *What it does:* Downstream OAuth 2.0 authorization server that manages consent screens, token exchange, and refresh token cycles for third-party tools (Google, Microsoft, Slack).
    * *Why & Where used:* Allows end users to safely grant AI agents permission to read their calendar, email, or documents.

20. **MCP Settings (`/workspace/mcp-settings`)**:
    * *What it does:* Global execution settings for MCP servers, including connection timeouts, maximum tool execution concurrency, and process restart policies.

21. **Plugins (`/workspace/plugins`)**:
    * *What it does:* Modular extensibility system supporting custom Go and sandboxed Starlark plugins executed inside the UnifAI request lifecycle pipeline.
    * *Why & Where used:* Allows enterprises to inject proprietary business logic, custom header encryption, or custom rate-limiting algorithms.

---

## 4.4 Governance (அணுகல் கட்டுப்பாடு & பாதுகாப்பு)
The Governance pillar enforces enterprise identity, multi-tenancy, authorization, and regulatory compliance.

22. **Virtual Keys (`/workspace/governance/virtual-keys`)**:
    * *What it does:* Synthetic proxy API keys (`sk-uf-...`) issued to developers and internal applications. Bypasses exposure of raw provider keys and enforces individual quotas and permissions.
    * *Why & Where used:* Core mechanism for application onboarding. Allows revoking or updating keys without redeploying microservices.

23. **Users (`/workspace/governance/users`)**:
    * *What it does:* User account directory for developers, administrators, and stakeholders with workspace access.
    * *Why & Where used:* Manages individual platform access, passwords, and multi-factor authentication (MFA).

24. **Teams (`/workspace/governance/teams`)**:
    * *What it does:* Team-level grouping (e.g., Frontend Team, Search Engineering, Data Science). Virtual keys and budgets are assigned at the team level.
    * *Why & Where used:* Enables departmental chargebacks and shared quota management.

25. **Business Units (`/workspace/governance/business-units`)**:
    * *What it does:* Top-level enterprise grouping that encompasses multiple teams (e.g., Retail Division, Cloud Infrastructure, Corporate IT).
    * *Why & Where used:* Used by enterprise finance teams for macroscopic cost allocation and executive reporting.

26. **Customers (`/workspace/governance/customers`)**:
    * *What it does:* Multi-tenant client registry for B2B SaaS applications. Attaches usage to external customer accounts via `x-uf-customer-id`.
    * *Why & Where used:* Allows B2B software vendors to calculate the exact AI infrastructure cost generated by each paying customer.

27. **User Provisioning - SCIM (`/workspace/scim`)**:
    * *What it does:* System for Cross-domain Identity Management (SCIM) v2.0 endpoint. Automatically syncs users and team memberships from enterprise Identity Providers (Okta, Microsoft Entra ID / Azure AD).
    * *Why & Where used:* Eliminates manual user management; automatically deprovisions former employees when removed from corporate identity directories.

28. **Roles & Permissions - RBAC (`/workspace/governance/rbac`)**:
    * *What it does:* Fine-grained Role-Based Access Control matrix. Governs which users can view logs, create virtual keys, edit routing rules, or modify security settings.
    * *Why & Where used:* Ensures least-privilege security compliance required for SOC 2 and ISO 27001 certifications.

29. **Access Profiles (`/workspace/governance/access-profiles`)**:
    * *What it does:* Reusable security presets defining allowed models, allowed providers, and allowed MCP tool groups. Bound directly to Virtual Keys.
    * *Why & Where used:* Prevents unauthorized usage (e.g., guaranteeing a staging virtual key can only use cheap open-source models).

30. **Audit Logs (`/workspace/audit-logs`)**:
    * *What it does:* Immutable, cryptographically timestamped audit trail of all administrative actions (who created a key, who changed a route, who altered budget limits).
    * *Why & Where used:* Mandatory for enterprise compliance reviews, security forensics, and regulatory audits.

---

## 4.5 Guardrails (உள்ளடக்கப் பாதுகாப்பு & கொள்கை அமலாக்கம்)
Guardrails safeguard corporate data, enforce brand safety, and prevent adversarial exploits.

31. **Rules (`/workspace/guardrails/configuration`)**:
    * *What it does:* CEL-based safety policy definitions executed pre-LLM (input validation) and post-LLM (output validation). Includes PII redaction, prompt injection detection, and content moderation rules.
    * *Why & Where used:* Automatically intercepts and masks credit card numbers, Social Security Numbers, or Aadhaar numbers before they reach external LLMs.

32. **Providers (`/workspace/guardrails/providers`)**:
    * *What it does:* Engine configuration for underlying guardrail detection engines, including Microsoft Presidio DLP, Meta Llama Guard, AWS Bedrock Guardrails, Lakera AI, and regex engines.
    * *Why & Where used:* Connects UnifAI to state-of-the-art specialized safety classifiers.

33. **Cluster Config (`/workspace/cluster`)**:
    * *What it does:* Distributed cluster coordination manager. Synchronizes routing rules, rate limit counters, and guardrail policies across multiple UnifAI gateway nodes.
    * *Why & Where used:* Ensures consistent security and quota enforcement in multi-region, high-traffic distributed deployments.

---

## 4.6 Adaptive Routing (ஸ்மார்ட் தேர்வு & அறிவு களஞ்சியம்)
Dynamically optimizes model performance while managing prompt and skill assets.

34. **Adaptive Routing Dashboard & Settings (`/workspace/adaptive-routing`)**:
    * *What it does:* Intelligent routing dashboard powered by multi-armed bandit algorithms. Continuously balances traffic across equivalent models based on real-time latency, cost, and provider uptime.
    * *Why & Where used:* Automatically steers traffic to the fastest and cheapest provider without manual intervention.

35. **Prompt Repository (`/workspace/prompt-repo`)**:
    * *What it does:* Version-controlled enterprise prompt repository. Stores prompt templates, parameterized variables, model bindings, and test suites.
    * *Why & Where used:* Centralizes prompt engineering, preventing hardcoded prompt strings in application source code.

36. **Skills Repository (`/workspace/skills-repo`)**:
    * *What it does:* Standardized library of autonomous AI Agent instructions, persona definitions, and domain skill packages.
    * *Why & Where used:* Enables reuse of battle-tested agent behaviors across multiple development teams.

---

## 4.7 Global Settings (கட்டமைப்பு & செயல்திறன் ட்யூனிங்)
Platform-level infrastructure tuning, caching, and network security.

37. **Client Settings, Compatibility, Caching, Security & Performance Tuning (`/workspace/config/*`)**:
    * *Client Settings:* Global connection timeouts, keep-alive settings, and header forward allowlists (`x-uf-eh-*`).
    * *Compatibility:* Automatic JSON request-response format converter between OpenAI, Anthropic, Bedrock, and Gemini.
    * *Caching:* Semantic vector cache configuration using Redis or Qdrant for embedding-based similarity lookup.
    * *Security:* TLS certificates, IP address allowlists, CORS policies, and master encryption keys.
    * *API Keys:* UnifAI administrative API credentials.
    * *Performance Tuning:* Worker pool concurrency, memory pool sizes (`sync.Pool`), and buffer sizes.
    * *Feature Flags:* Dynamic runtime toggles to enable or test beta capabilities without restarting the gateway.

---

# 5. Cross-Feature Interconnection & Data Flow Matrix
### இணைப்பு வரைபடம் (எப்படி ஒன்றுடன் ஒன்று இணைகிறது?)

The following matrix illustrates the exact technical relationships between UnifAI modules:

| Source Module | Interconnected Modules | Technical Relationship & Data Flow |
| :--- | :--- | :--- |
| **Virtual Keys** | Users, Teams, Customers, Budgets, Access Profiles | Authenticates request; resolves quota limits, team ownership, and access profile. |
| **Access Profiles** | Model Catalog, MCP Tool Groups | Restricts the virtual key to specific approved models and MCP tool groups. |
| **Complexity Router** | Routing Rules, Model Catalog | Classifies prompt tier (Simple/Reasoning) and sets context for routing rules. |
| **Routing Rules** | Model Providers, Circuit Breaker | Evaluates CEL expressions; passes target provider to Circuit Breaker for health check. |
| **Circuit Breaker** | Routing Rules, Fallback Providers | If primary provider is TRIPPED, automatically switches to next configured fallback. |
| **Guardrails Rules** | Guardrail Providers, LLM Pipeline | Intercepts prompt pre-LLM and response post-LLM; runs PII/injection validation. |
| **Semantic Cache** | Vector Store, Models Pipeline | Intercepts request before LLM; returns cached hit or stores new completion on miss. |
| **MCP Tool Groups** | MCP Catalog, Virtual Keys | Injects approved tool JSON schemas into LLM prompt; audits calls in MCP Logs. |
| **Browser AI** | Guardrails, Observability | Intercepts browser web AI queries, runs DLP guardrail rules, logs to Browser AI logs. |
| **Pricing Overrides** | Dashboard, Budgets & Limits, LLM Logs | Computes exact dollar cost per token; updates budget balances and dashboard graphs. |
| **Connectors** | LLM Logs, MCP Logs, Audit Logs | Streams structured JSON telemetry to external platforms (Datadog, Kafka, BigQuery). |
| **SCIM Provisioning** | Users, Teams, RBAC | Automatically provisions enterprise IdP accounts and maps them into UnifAI roles. |

---

# 6. Technology Stack & Programming Languages Deep Dive
### தொழில்நுட்ப கட்டமைப்பு

The UnifAI codebase is engineered for ultra-low latency, high concurrency, and enterprise reliability.

```
┌────────────────────────────────────────────────────────────────────────┐
│                        UNIFAI FULL TECH STACK                          │
├───────────────────┬────────────────────────────────────────────────────┤
│ Gateway Core      │ Go (Golang 1.23+), Goroutines, sync.Pool, fasthttp │
│ JSON Serialization│ Sonic JSON (ByteDance ultra-fast assembly parser)  │
│ Rule Evaluation   │ Google CEL (Common Expression Language)            │
│ Agent Scripting   │ Starlark (Google Bazel deterministic language)     │
│ Frontend Web App  │ TypeScript 5.x, React 18, Vite, TanStack Router    │
│ UI & Styling      │ Tailwind CSS, Radix UI (Shadcn), Lucide Icons      │
│ State Management  │ Redux Toolkit (RTK Query), WebSockets              │
│ Local Proxy/Agent │ Python 3.11+, mitmproxy / asyncio (Browser Guard)  │
│ Relational DB     │ PostgreSQL (pgx v5 driver), SQLite, GORM           │
│ Caching & Vectors │ Redis (go-redis), Qdrant, pgvector                 │
│ Telemetry         │ OpenTelemetry (OTel), Prometheus, Datadog, Kafka   │
│ DevOps & Build    │ Docker, Docker Compose, Makefile, Nix / Flake      │
└───────────────────┴────────────────────────────────────────────────────┘
```

### Why These Specific Technologies Were Chosen:
1. **Go (Golang):** Selected for the core gateway engine due to its lightweight goroutine model, memory safety, and zero-allocation buffer pooling (`sync.Pool`), allowing UnifAI to handle 100,000+ concurrent requests with sub-millisecond proxy overhead.
2. **FastHTTP & Sonic JSON:** Standard Go `net/http` and `encoding/json` introduce substantial memory allocations under heavy load. FastHTTP avoids per-request heap allocations, while ByteDance's Sonic provides JIT-compiled JSON parsing.
3. **Google CEL (Common Expression Language):** Safe, non-Turing complete expression evaluation. Compiles routing rules and guardrail policies into bytecode, executing in microseconds without risk of infinite loops or security sandbox escapes.
4. **TypeScript & TanStack Router:** Provides strict compile-time type safety across all 38 UI views. TanStack Router provides preload-on-hover routing for instantaneous user navigation.
5. **Python for Browser Guard:** Python's mature networking and proxy ecosystem (`browser_ai_proxy.py`) allows rapid interception and regex-based DLP filtering of browser traffic.

---

# 7. Enterprise Production Scenarios & Playbooks
### நடைமுறை பயன்பாடுகள் & தயாரிப்பு காட்சிகள்

### Scenario A: Multi-Tenant B2B SaaS Cost Attribution
* **Challenge:** A software company provides an AI copilot to 500 enterprise customers using a single backend service. They need to track how much each customer costs.
* **Solution:** The backend makes calls using a single Team Virtual Key but attaches `x-uf-customer-id: cust_123` on every request. UnifAI's Governance and Pricing engine attributes token spend per customer, visible on the Customers dashboard and exported to BigQuery for monthly invoicing.

### Scenario B: Zero-Downtime High Availability Failover
* **Challenge:** Public cloud outages at OpenAI cause an e-commerce customer support chatbot to go down, hurting customer satisfaction.
* **Solution:** A Circuit Breaker is configured on `gpt-4o` with a 50% error threshold over a 30-second window, paired with a fallback route to `claude-3-5-sonnet`. During an outage, UnifAI automatically trips the breaker and redirects requests within 100 milliseconds. The client receives `x-unifai-fallback-index: 1`.

### Scenario C: Slashing Costs with Complexity Routing & Semantic Cache
* **Challenge:** A company receives 1,000,000 user queries daily. 60% are routine questions (e.g., "What are your hours?", "How do I reset my password?"), but all queries currently go to expensive frontier models ($15/1M tokens).
* **Solution:** 
  1. Semantic Caching is enabled (`x-uf-cache-threshold: 0.90`), answering 25% of queries directly from Redis cache at $0.00 cost and 15ms latency.
  2. The Complexity Router routes remaining simple queries to a compact model (`gpt-4o-mini` at $0.15/1M tokens), reserving frontier models for reasoning. Total AI bill decreases by 78%.

### Scenario D: Enterprise DLP with Browser AI
* **Challenge:** Employees accidentally paste proprietary software code, API secrets, or customer PII into web ChatGPT.
* **Solution:** UnifAI Browser AI agent is deployed to company laptops. When an employee attempts to paste text containing credit card numbers or API keys into ChatGPT, Browser AI intercepts the request, runs Guardrail DLP rules, blocks the submission, and logs the violation in the Audit Logs.

---

**End of Document**  
*UnifAI Enterprise Architecture & Feature Documentation — Confidential & Proprietary.*
"""

with open(MD_PATH, "w", encoding="utf-8") as f:
    f.write(md_content)

print("Markdown document created at:", MD_PATH)

# ==============================================================================
# 2. DOCX GENERATION
# ==============================================================================
print("Generating Word Document (.docx)...")

doc = Document()

# Set standard page margins (0.75 in)
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Setup Header & Footer
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "UnifAI Enterprise Architecture & Features Master Guide"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if hp.runs:
        hp.runs[0].font.name = "Segoe UI"
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(100, 116, 139)
        
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = "Confidential & Proprietary — UnifAI Control Plane Documentation"
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if fp.runs:
        fp.runs[0].font.name = "Segoe UI"
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor(100, 116, 139)

def style_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Dark Navy
    return p

def style_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235) # Blue
    return p

def style_heading_3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42) # Slate
    return p

def add_body_p(text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Segoe UI"
        r_bold.font.size = Pt(10)
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(15, 23, 42)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(10)
    run.italic = italic
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet_item(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Segoe UI"
        r_bold.font.size = Pt(10)
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(15, 23, 42)
    run = p.add_run(text)
    run.font.name = "Segoe UI"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_callout(text, title="NOTE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    # Light blue background
    shd = parse_xml(r'<w:shd {} w:fill="F0F9FF"/>'.format(nsdecls('w')))
    tcPr.append(shd)
    # Left border navy
    borders = parse_xml(r'<w:tcBorders {}><w:left w:val="single" w:sz="24" w:space="0" w:color="0284C7"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>'.format(nsdecls('w')))
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r_title = p.add_run(f"[{title}] ")
    r_title.bold = True
    r_title.font.name = "Segoe UI"
    r_title.font.size = Pt(9.5)
    r_title.font.color.rgb = RGBColor(2, 132, 199)
    
    r_text = p.add_run(text)
    r_text.font.name = "Segoe UI"
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(12, 74, 110)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_table(table, col_widths, col_names):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Format header row
    hdr_cells = table.rows[0].cells
    for i, name in enumerate(col_names):
        hdr_cells[i].text = name
        tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
        tcPr.append(shd)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.name = "Segoe UI"
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    # Format data rows
    for r_idx, row in enumerate(table.rows[1:]):
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            tcPr.append(shd)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = "Segoe UI"
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(15, 23, 42)
                
    # Apply widths
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = Inches(width)

# --- Title Section ---
title_p = doc.add_paragraph()
title_p.paragraph_format.space_before = Pt(20)
title_p.paragraph_format.space_after = Pt(4)
r_title = title_p.add_run("UnifAI Enterprise Architecture & Feature Guide")
r_title.font.name = "Segoe UI"
r_title.font.size = Pt(22)
r_title.bold = True
r_title.font.color.rgb = RGBColor(30, 58, 138)

sub_p = doc.add_paragraph()
sub_p.paragraph_format.space_after = Pt(14)
r_sub = sub_p.add_run("A Complete Technical Manual for Architecture, HTTP Headers, Core Pillars, and Tech Stack (Tamil & English Guide)")
r_sub.font.name = "Segoe UI"
r_sub.font.size = Pt(11)
r_sub.font.color.rgb = RGBColor(71, 85, 105)

add_callout(
    "UnifAI is an enterprise-grade AI Gateway, Router, Observability, Governance, and MCP Platform. "
    "It sits between your applications and AI model providers to slash costs up to 80%, guarantee 99.99% uptime, "
    "enforce enterprise security & DLP, and manage intelligent routing.",
    title="ARCHITECTURE SUMMARY"
)

# --- Section 1 ---
style_heading_1("1. Executive Summary & Core Value Proposition")
add_body_p(
    "Modern enterprises face fragmented AI integrations with multiple vendors (OpenAI, Anthropic Claude, "
    "Google Gemini, AWS Bedrock, Mistral, and local open-source models). UnifAI centralizes all LLM traffic through "
    "a high-performance Go-based gateway that provides unified governance, observability, and cost control."
)
add_bullet_item("Slashing AI spending by 60–80% via Complexity Routing and Semantic Caching.", "Cost Optimization: ")
add_bullet_item("Automatic failover via Circuit Breaker to backup providers during cloud outages.", "High Reliability: ")
add_bullet_item("Virtual API keys, multi-tenant customer scoping, and granular team budgets.", "Enterprise Governance: ")
add_bullet_item("Pre-LLM and Post-LLM content inspection (PII masking, prompt injection defense).", "Content Safety (Guardrails): ")
add_bullet_item("Standardized agentic execution via Model Context Protocol (MCP) tool catalog.", "Extensible Agents (MCP): ")
add_bullet_item("Real-time logging and streaming connectors to Datadog, New Relic, Kafka, and BigQuery.", "Full Observability: ")

# --- Section 2 ---
style_heading_1("2. End-to-End Request Lifecycle")
add_body_p(
    "Every AI request sent to UnifAI undergoes a 7-stage pipeline designed for sub-millisecond overhead:"
)
add_bullet_item("Incoming request authenticated via Virtual Key (sk-uf-*); validates rate limits and budget balance.", "Stage 1 - Transport & Governance: ")
add_bullet_item("CEL rules and Presidio DLP scan prompt for sensitive data or prompt injection attempts.", "Stage 2 - Pre-LLM Guardrails: ")
add_bullet_item("Checks vector store for cosine similarity. If cached hit exists, returns response immediately.", "Stage 3 - Semantic Caching: ")
add_bullet_item("Classifies query complexity (Simple/Reasoning), applies routing rules, and checks Circuit Breaker.", "Stage 4 - Intelligent Routing: ")
add_bullet_item("Dispatches request to upstream LLM (OpenAI, Claude) or executes MCP tools with OAuth sessions.", "Stage 5 - Upstream Execution: ")
add_bullet_item("Scans generated output for toxic content or accidental credential leakage before return.", "Stage 6 - Post-LLM Guardrails: ")
add_bullet_item("Stores trace in LLM/MCP logs, streams to Datadog/Kafka connectors, and returns response headers.", "Stage 7 - Telemetry & Response: ")

# --- Section 3 ---
style_heading_1("3. Exhaustive HTTP Headers Reference Guide")
add_body_p(
    "HTTP Headers allow developers to control authentication, routing, semantic caching, debugging, and logging "
    "per-request without changing the JSON payload."
)

style_heading_2("3.1 Authentication & Identity Headers")
headers_auth_data = [
    ("x-uf-vk", "sk-uf-...", "Virtual Key: Authenticates application, team, budget quota, and model permissions."),
    ("x-uf-api-key", "key_admin...", "Platform Admin Key: For control plane APIs (configuring routes, creating keys)."),
    ("x-uf-customer-id", "cust_42", "Customer Scoping: Attaches cost and log attribution to a specific end-customer."),
    ("x-uf-direct-key", "true / false", "Direct Key: Bypasses registered key pool to use raw API key provided in request."),
    ("X-UnifAI-Temp-Token", "UUID / JWT", "Temporary Session Token: For interactive MCP OAuth consent and UI actions.")
]
tbl_auth = doc.add_table(rows=len(headers_auth_data)+1, cols=3)
format_table(tbl_auth, [1.8, 1.4, 3.8], ["Header Name", "Sample Format", "Description & Technical Purpose"])
for idx, (h, fmt, desc) in enumerate(headers_auth_data):
    tbl_auth.rows[idx+1].cells[0].text = h
    tbl_auth.rows[idx+1].cells[1].text = fmt
    tbl_auth.rows[idx+1].cells[2].text = desc

style_heading_2("3.2 Routing & Model Selection Headers")
headers_route_data = [
    ("x-uf-provider", "anthropic, openai", "Force Provider: Bypasses dynamic routing rules to send request to specified provider."),
    ("x-uf-model", "claude-3-7-sonnet", "Force Model: Explicitly targets a specific model architecture."),
    ("x-uf-api-key-id", "key_prod_01", "Key Pinning: Selects a specific credential from a multi-key provider rotation pool."),
    ("x-uf-circuit-breaker-bypass", "true", "Breaker Bypass: Administrative probe override even if provider breaker is TRIPPED.")
]
tbl_route = doc.add_table(rows=len(headers_route_data)+1, cols=3)
format_table(tbl_route, [1.8, 1.4, 3.8], ["Header Name", "Sample Format", "Description & Technical Purpose"])
for idx, (h, fmt, desc) in enumerate(headers_route_data):
    tbl_route.rows[idx+1].cells[0].text = h
    tbl_route.rows[idx+1].cells[1].text = fmt
    tbl_route.rows[idx+1].cells[2].text = desc

style_heading_2("3.3 Semantic Caching Headers")
headers_cache_data = [
    ("x-uf-cache-key", "tenant-01:chat", "Partition Key: Scopes cached entries to a tenant, preventing data cross-contamination."),
    ("x-uf-cache-ttl", "3600 (seconds)", "TTL Override: How long this response remains cached before automatic invalidation."),
    ("x-uf-cache-threshold", "0.90", "Cosine Similarity Threshold: Minimum vector similarity score required for a cache hit."),
    ("x-uf-cache-type", "direct / semantic", "Cache Mode: 'direct' for SHA256 string hash; 'semantic' for vector embeddings."),
    ("x-uf-cache-no-store", "true", "No-Store: Allows reading existing cache hit, but skips writing new LLM response.")
]
tbl_cache = doc.add_table(rows=len(headers_cache_data)+1, cols=3)
format_table(tbl_cache, [1.8, 1.4, 3.8], ["Header Name", "Sample Format", "Description & Technical Purpose"])
for idx, (h, fmt, desc) in enumerate(headers_cache_data):
    tbl_cache.rows[idx+1].cells[0].text = h
    tbl_cache.rows[idx+1].cells[1].text = fmt
    tbl_cache.rows[idx+1].cells[2].text = desc

style_heading_2("3.4 MCP & Tool Execution Headers")
headers_mcp_data = [
    ("x-uf-mcp-include-clients", "github,postgres", "Allowed MCP Clients: Whitelist of MCP servers to inject into prompt (* for all)."),
    ("x-uf-mcp-include-tools", "exec_query,read_file", "Allowed Tools: Specific tool function names exposed to the LLM."),
    ("x-uf-mcp-session-id", "sess_usr_42", "OAuth Session ID: Associates individual user OAuth credentials with tool call."),
    ("x-uf-eh-*", "x-uf-eh-org-id: 12", "Extra Tool Headers: Forwarded directly to upstream MCP server if allowlisted.")
]
tbl_mcp = doc.add_table(rows=len(headers_mcp_data)+1, cols=3)
format_table(tbl_mcp, [1.8, 1.4, 3.8], ["Header Name", "Sample Format", "Description & Technical Purpose"])
for idx, (h, fmt, desc) in enumerate(headers_mcp_data):
    tbl_mcp.rows[idx+1].cells[0].text = h
    tbl_mcp.rows[idx+1].cells[1].text = fmt
    tbl_mcp.rows[idx+1].cells[2].text = desc

style_heading_2("3.5 Observability, Logging & Response Headers")
headers_obs_data = [
    ("x-uf-session-id", "chat_sess_99", "Tracing Session: Groups multiple sequential LLM calls into a single user session trace."),
    ("x-uf-dim-<key>", "x-uf-dim-env: prod", "Custom Dimension: Arbitrary metadata key/value pair attached to internal log records."),
    ("x-uf-disable-content-logging", "true", "Content Privacy: Drops prompt/response text from logs while retaining cost/latency."),
    ("traceparent", "00-4bf92...-01", "W3C Tracing: Propagates distributed OpenTelemetry trace context across services."),
    ("x-unifai-provider", "anthropic", "[Response Header] Identifies the actual model provider that served the request."),
    ("x-unifai-resolved-model", "claude-3-5-sonnet", "[Response Header] Identifies the model used after routing and complexity resolution."),
    ("x-unifai-fallback-index", "1", "[Response Header] Indicates if and which fallback provider was used (0 = primary).")
]
tbl_obs = doc.add_table(rows=len(headers_obs_data)+1, cols=3)
format_table(tbl_obs, [1.8, 1.4, 3.8], ["Header Name", "Sample Format", "Description & Technical Purpose"])
for idx, (h, fmt, desc) in enumerate(headers_obs_data):
    tbl_obs.rows[idx+1].cells[0].text = h
    tbl_obs.rows[idx+1].cells[1].text = fmt
    tbl_obs.rows[idx+1].cells[2].text = desc

# --- Section 4 ---
style_heading_1("4. In-Depth Feature Catalog (7 Core Pillars - 38 Features)")
add_body_p("A detailed breakdown of all 38 features across the UnifAI platform:")

pillars = [
    ("Pillar 1: Observability", [
        ("Dashboard", "Real-time metrics console displaying request volume, spend, p50/p95/p99 latency, and top models."),
        ("LLM Logs", "Detailed ledger of prompts, completions, tokens, TTFT, cost, and guardrail verdicts."),
        ("MCP Logs", "Granular audit log of tool execution calls, parameters, response data, and exceptions."),
        ("Browser AI", "Desktop agent and browser DLP proxy intercepting employee web AI usage to block IP leaks."),
        ("Connectors", "Streaming export pipelines forwarding telemetry to Datadog, New Relic, Kafka, and BigQuery."),
        ("Logs Settings", "Retention policies, sampling controls, and sensitive data redaction configurations.")
    ]),
    ("Pillar 2: Models & Intelligent Routing", [
        ("Model Catalog", "Inventory of all connected LLMs, context windows, pricing, and capability flags."),
        ("Model Providers", "Configuration console for API credentials and connection endpoints (OpenAI, Bedrock, etc.)."),
        ("Budgets & Limits", "Financial quota and rate-limiting engine enforcing RPM, TPM, and monthly cost ceilings."),
        ("Routing Rules", "Condition-based routing engine evaluated using Google CEL expressions and model aliases."),
        ("Complexity Router", "Multi-tier prompt classifier routing simple queries to cheap models and complex queries to frontier LLMs."),
        ("Circuit Breaker", "Automated failover monitor tripping unhealthy providers and rerouting traffic to backups."),
        ("Pricing Overrides", "Granular per-token price definitions for custom enterprise discount agreements."),
        ("Model Settings", "Global routing parameters, timeout defaults, and global cost configurations.")
    ]),
    ("Pillar 3: MCP Gateway (Tools & Agents)", [
        ("MCP Catalog", "Registry of all active Model Context Protocol servers and tool schemas."),
        ("MCP Library", "One-click curated installer for popular servers (GitHub, Postgres, Slack, Drive)."),
        ("Tool Groups", "Logical tool bundles assigned to Virtual Keys via Access Profiles to restrict agent permissions."),
        ("Auth Sessions", "Per-user credential manager tracking user-specific tool authorizations."),
        ("OAuth Grants", "OAuth 2.0 authorization server managing consent screens and token refresh cycles."),
        ("MCP Settings", "Gateway tool execution timeouts, process concurrency limits, and restart policies."),
        ("Plugins", "Modular extension framework for custom Go and Starlark plugins in the request pipeline.")
    ]),
    ("Pillar 4: Governance & Enterprise Compliance", [
        ("Virtual Keys", "Synthetic proxy keys (sk-uf-*) issued to applications to hide upstream provider credentials."),
        ("Users & Teams", "Organizational hierarchy managing developer accounts and departmental team budgets."),
        ("Business Units", "Top-level enterprise divisions grouping multiple teams for macroscopic financial chargeback."),
        ("Customers", "Multi-tenant B2B client registry enabling per-customer cost allocation and auditing."),
        ("User Provisioning (SCIM)", "Automated user lifecycle sync from Okta and Azure AD via SCIM v2.0 protocol."),
        ("Roles & Permissions (RBAC)", "Fine-grained permission matrix governing who can create keys, view logs, or edit routes."),
        ("Access Profiles", "Reusable security profiles defining allowed models and tool groups attached to Virtual Keys."),
        ("Audit Logs", "Immutable audit trail recording administrative configuration changes for SOC 2 compliance.")
    ]),
    ("Pillar 5: Guardrails & Content Security", [
        ("Rules", "CEL safety policies executing pre-LLM and post-LLM to mask PII and block prompt injections."),
        ("Providers", "Detection engines integrating Microsoft Presidio, Meta Llama Guard, and regex filters."),
        ("Cluster Config", "Distributed synchronization engine maintaining rule consistency across gateway nodes.")
    ]),
    ("Pillar 6: Adaptive Routing & Assets", [
        ("Adaptive Routing Dashboard", "Reinforcement learning traffic distributor optimizing latency and error rates."),
        ("Prompt Repository", "Version-controlled prompt template library supporting variables and test suites."),
        ("Skills Repository", "Centralized repository for reusable autonomous agent instructions and personas.")
    ]),
    ("Pillar 7: Global Settings & Engine Tuning", [
        ("Client Settings", "Default keep-alive parameters, timeouts, and header allowlists."),
        ("Compatibility", "Real-time protocol converter translating between OpenAI and non-OpenAI formats."),
        ("Caching", "Semantic vector caching engine configured with Redis or Qdrant for similarity search."),
        ("Security", "TLS certificates, IP address allowlists, CORS policies, and master secret encryption."),
        ("API Keys", "Platform administrative API keys for CI/CD and automated orchestration."),
        ("Performance Tuning", "Worker pool concurrency, zero-allocation memory pools, and connection tuning."),
        ("Feature Flags", "Runtime capability toggles to enable or test beta features dynamically.")
    ])
]

for p_title, p_items in pillars:
    style_heading_2(p_title)
    tbl_p = doc.add_table(rows=len(p_items)+1, cols=2)
    format_table(tbl_p, [2.2, 4.8], ["Feature Name", "Core Capabilities & Practical Utility"])
    for idx, (name, desc) in enumerate(p_items):
        tbl_p.rows[idx+1].cells[0].text = name
        tbl_p.rows[idx+1].cells[1].text = desc

# --- Section 5 ---
style_heading_1("5. Technology Stack & Programming Languages")
add_body_p(
    "The UnifAI platform is engineered with modern, cloud-native technologies optimized for zero-allocation memory overhead:"
)
tech_data = [
    ("Go (Golang 1.23+)", "Core Gateway Engine", "Goroutines, worker pools, zero-allocation sync.Pool for high concurrency."),
    ("FastHTTP & Sonic JSON", "HTTP & Serialization", "Replaces net/http with FastHTTP and ByteDance Sonic JIT-compiled JSON."),
    ("Google CEL", "Rule Evaluation", "Fast, sandboxed, memory-safe execution of routing and guardrail policies."),
    ("TypeScript & React 18", "Frontend Web UI", "Type-safe interface powered by Vite, TanStack Router, and Shadcn UI."),
    ("Python 3.11+", "Browser AI Proxy", "Desktop DLP proxy and mitmproxy daemon for browser prompt inspection."),
    ("Starlark", "Agent Code Sandbox", "Deterministic, secure scripting language for MCP code execution."),
    ("PostgreSQL & pgx", "Relational Database", "Stores governance metadata, audit logs, virtual keys, and budgets."),
    ("Redis & Qdrant", "Caching & Vector Store", "Powers sub-millisecond semantic similarity search and session tracking."),
    ("OpenTelemetry (OTel)", "Distributed Tracing", "Standardized telemetry exported to Datadog, New Relic, and Kafka.")
]
tbl_tech = doc.add_table(rows=len(tech_data)+1, cols=3)
format_table(tbl_tech, [1.8, 1.8, 3.4], ["Technology / Language", "System Component", "Architectural Role & Benefit"])
for idx, (tech, comp, role) in enumerate(tech_data):
    tbl_tech.rows[idx+1].cells[0].text = tech
    tbl_tech.rows[idx+1].cells[1].text = comp
    tbl_tech.rows[idx+1].cells[2].text = role

# --- Section 6 ---
style_heading_1("6. Production Architectural Scenarios")
add_body_p("Real-world architectural playbooks illustrating how UnifAI solves production challenges:")

add_body_p("A B2B SaaS application powers an AI assistant for 500 enterprise clients. Using a shared team virtual key, the backend appends x-uf-customer-id: cust_99 on each API call. UnifAI calculates precise token consumption per client, generating itemized monthly billing reports in BigQuery.", bold_prefix="Scenario 1 - B2B Multi-Tenant Cost Attribution: ")

add_body_p("An e-commerce customer support chatbot experiences a sudden OpenAI outage. The UnifAI Circuit Breaker detects that 5xx error rates exceed 50% over a 30-second window, automatically redirecting incoming prompts to Anthropic Claude-3.5-Sonnet within 100ms. Zero user sessions drop.", bold_prefix="Scenario 2 - High Availability Failover: ")

add_body_p("An enterprise handling 1,000,000 daily queries enables Semantic Caching and Complexity Routing. 25% of queries match cached answers (15ms latency, $0 cost). Of the remainder, 70% are simple questions routed to lightweight models, reducing monthly AI costs from $45,000 to $9,800.", bold_prefix="Scenario 3 - Cost Reduction via Complexity Routing: ")

add_body_p("Employees use web AI portals for drafting emails. The UnifAI Browser AI agent runs locally on work machines, scanning clipboard pastes for credit card numbers, passwords, or customer records, immediately blocking data leakage and alerting the security SOC.", bold_prefix="Scenario 4 - Enterprise Data Loss Prevention (DLP): ")

doc.save(DOCX_PATH)
print("Word document created successfully at:", DOCX_PATH)

# ==============================================================================
# 3. PDF GENERATION VIA REPORTLAB
# ==============================================================================
print("Generating PDF document (.pdf) via ReportLab...")

# Register fonts
pdfmetrics.registerFont(TTFont('Latha', 'C:/Windows/Fonts/latha.ttf'))
pdfmetrics.registerFont(TTFont('LathaBold', 'C:/Windows/Fonts/lathab.ttf'))

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#64748B"))
        
        # Header (pages 2+)
        if self._pageNumber > 1:
            self.drawString(40, 11 * inch - 36, "UnifAI Enterprise Architecture & Feature Master Guide")
            self.drawRightString(8.5 * inch - 40, 11 * inch - 36, "Confidential — System Documentation")
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(40, 11 * inch - 40, 8.5 * inch - 40, 11 * inch - 40)
            
        # Footer
        self.drawString(40, 32, "UnifAI Unified AI Gateway & Governance Platform")
        page_str = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(8.5 * inch - 40, 32, page_str)
        self.setStrokeColor(colors.HexColor("#CBD5E1"))
        self.setLineWidth(0.5)
        self.line(40, 42, 8.5 * inch - 40, 42)
        
        self.restoreState()

pdf_doc = SimpleDocTemplate(
    PDF_PATH,
    pagesize=letter,
    leftMargin=40,
    rightMargin=40,
    topMargin=48,
    bottomMargin=48
)

styles = getSampleStyleSheet()

style_title = ParagraphStyle(
    'PdfTitle',
    fontName='LathaBold',
    fontSize=20,
    leading=24,
    textColor=colors.HexColor('#1E3A8A'),
    spaceAfter=4
)

style_subtitle = ParagraphStyle(
    'PdfSubtitle',
    fontName='Latha',
    fontSize=10,
    leading=14,
    textColor=colors.HexColor('#475569'),
    spaceAfter=12
)

style_h1 = ParagraphStyle(
    'PdfH1',
    fontName='LathaBold',
    fontSize=13,
    leading=16,
    textColor=colors.HexColor('#1E3A8A'),
    spaceBefore=14,
    spaceAfter=6,
    keepWithNext=True
)

style_h2 = ParagraphStyle(
    'PdfH2',
    fontName='LathaBold',
    fontSize=10.5,
    leading=13,
    textColor=colors.HexColor('#2563EB'),
    spaceBefore=10,
    spaceAfter=4,
    keepWithNext=True
)

style_body = ParagraphStyle(
    'PdfBody',
    fontName='Latha',
    fontSize=8.5,
    leading=11.5,
    textColor=colors.HexColor('#1E293B'),
    spaceAfter=4
)

style_bullet = ParagraphStyle(
    'PdfBullet',
    fontName='Latha',
    fontSize=8.5,
    leading=11.5,
    textColor=colors.HexColor('#1E293B'),
    leftIndent=12,
    spaceAfter=3
)

style_th = ParagraphStyle(
    'PdfTH',
    fontName='LathaBold',
    fontSize=7.5,
    leading=9.5,
    textColor=colors.white
)

style_td = ParagraphStyle(
    'PdfTD',
    fontName='Latha',
    fontSize=7.5,
    leading=9.5,
    textColor=colors.HexColor('#0F172A')
)

style_td_code = ParagraphStyle(
    'PdfTDCode',
    fontName='Helvetica-Bold',
    fontSize=7.5,
    leading=9.5,
    textColor=colors.HexColor('#0F172A')
)

story = []

# Title
story.append(Paragraph("UnifAI Enterprise Architecture & Feature Master Guide", style_title))
story.append(Paragraph("விரிவான கணினி கட்டமைப்பு, ஹெடர்ஸ் மற்றும் அம்சங்களின் முழு கையேடு (Tamil & English Technical Manual)", style_subtitle))
story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#0284C7"), spaceAfter=10))

# Executive Summary
story.append(Paragraph("1. Executive Summary & Core Value Proposition", style_h1))
story.append(Paragraph(
    "UnifAI is a high-performance Unified AI Gateway, Router, Governance, and Observability Control Plane. "
    "It centralizes AI consumption across OpenAI, Anthropic Claude, Google Gemini, AWS Bedrock, Mistral, and local models. "
    "UnifAI provides <b>Cost Optimization (up to 80%)</b>, <b>Zero-Downtime Reliability (99.99%)</b>, <b>Data Loss Prevention (DLP)</b>, "
    "and <b>Agentic Tool Extensibility via MCP Gateway</b>.",
    style_body
))
story.append(Paragraph("• <b>Cost Reduction:</b> Automatic Complexity Routing and Semantic Caching deliver up to 80% spend reduction.", style_bullet))
story.append(Paragraph("• <b>High Availability:</b> Circuit Breaker automatically fails over to backup providers during cloud outages.", style_bullet))
story.append(Paragraph("• <b>Governance:</b> Virtual Keys (sk-uf-*), RBAC permissions, and team budget enforcement.", style_bullet))
story.append(Paragraph("• <b>Security:</b> Pre-LLM and Post-LLM Guardrails for PII redaction and prompt injection defense.", style_bullet))
story.append(Paragraph("• <b>Observability:</b> Real-time logs and streaming telemetry connectors to Datadog, Kafka, and BigQuery.", style_bullet))

# Lifecycle
story.append(Paragraph("2. End-to-End Request Lifecycle (7-Stage Pipeline)", style_h1))
story.append(Paragraph(
    "1. <b>Transport & Governance:</b> FastHTTP router parses headers; validates Virtual Key and checks budget quota.<br/>"
    "2. <b>Pre-LLM Guardrails:</b> Evaluates Google CEL rules and Presidio DLP to block sensitive data and prompt injections.<br/>"
    "3. <b>Semantic Caching:</b> Checks vector store for existing answers. If cosine similarity matches, returns cached answer.<br/>"
    "4. <b>Intelligent Routing:</b> Complexity Router classifies prompt (Simple/Reasoning); Circuit Breaker checks provider health.<br/>"
    "5. <b>Upstream Execution:</b> Dispatches prompt to target LLM or executes MCP tools with OAuth credentials.<br/>"
    "6. <b>Post-LLM Guardrails:</b> Scans generated output for toxic content or accidental credential leakage.<br/>"
    "7. <b>Telemetry & Headers:</b> Appends record to logs, streams to connectors, and returns response with x-unifai-* headers.",
    style_body
))

# Headers Guide
story.append(Paragraph("3. Exhaustive HTTP Headers Reference Guide", style_h1))
story.append(Paragraph("HTTP headers control routing, caching, security, and logging per request without modifying the body payload:", style_body))

def make_pdf_table(data_tuples, col_widths, headers):
    table_data = [[Paragraph(h, style_th) for h in headers]]
    for row in data_tuples:
        col0 = Paragraph(row[0], style_td_code)
        col1 = Paragraph(row[1], style_td)
        col2 = Paragraph(row[2], style_td)
        table_data.append([col0, col1, col2])
    t = Table(table_data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    return t

story.append(Paragraph("3.1 Authentication & Identity Headers", style_h2))
story.append(make_pdf_table(headers_auth_data, [110, 80, 340], ["Header Name", "Sample Format", "Description & Technical Purpose"]))
story.append(Spacer(1, 6))

story.append(Paragraph("3.2 Routing & Model Selection Headers", style_h2))
story.append(make_pdf_table(headers_route_data, [125, 85, 320], ["Header Name", "Sample Format", "Description & Technical Purpose"]))
story.append(Spacer(1, 6))

story.append(Paragraph("3.3 Semantic Caching Headers", style_h2))
story.append(make_pdf_table(headers_cache_data, [120, 80, 330], ["Header Name", "Sample Format", "Description & Technical Purpose"]))
story.append(Spacer(1, 6))

story.append(Paragraph("3.4 MCP & Tool Execution Headers", style_h2))
story.append(make_pdf_table(headers_mcp_data, [125, 85, 320], ["Header Name", "Sample Format", "Description & Technical Purpose"]))
story.append(Spacer(1, 6))

story.append(Paragraph("3.5 Observability, Logging & Response Headers", style_h2))
story.append(make_pdf_table(headers_obs_data, [125, 85, 320], ["Header Name", "Sample Format", "Description & Technical Purpose"]))
story.append(Spacer(1, 8))

# 7 Pillars
story.append(Paragraph("4. In-Depth Feature Catalog (7 Core Pillars - 38 Features)", style_h1))

def make_pillar_table(items, col_widths, headers):
    table_data = [[Paragraph(h, style_th) for h in headers]]
    for name, desc in items:
        table_data.append([Paragraph(name, style_td_code), Paragraph(desc, style_td)])
    t = Table(table_data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('LEFTPADDING', (0, 0), (-1, -1), 4),
        ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    return t

for p_title, p_items in pillars:
    story.append(Paragraph(p_title, style_h2))
    story.append(make_pillar_table(p_items, [130, 400], ["Feature Name", "Core Capabilities & Practical Utility"]))
    story.append(Spacer(1, 6))

# Technology Stack
story.append(Paragraph("5. Technology Stack & Programming Languages", style_h1))
story.append(Paragraph("Engineered for ultra-low latency, memory safety, and massive concurrency:", style_body))

tech_table_data = [[Paragraph(h, style_th) for h in ["Technology / Language", "System Component", "Architectural Role & Benefit"]]]
for tech, comp, role in tech_data:
    tech_table_data.append([Paragraph(tech, style_td_code), Paragraph(comp, style_td), Paragraph(role, style_td)])

t_tech = Table(tech_table_data, colWidths=[110, 110, 310])
t_tech.setStyle(TableStyle([
    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ('TOPPADDING', (0, 0), (-1, -1), 3),
    ('LEFTPADDING', (0, 0), (-1, -1), 4),
    ('RIGHTPADDING', (0, 0), (-1, -1), 4),
    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
]))
story.append(t_tech)
story.append(Spacer(1, 8))

# Production Scenarios
story.append(Paragraph("6. Production Architectural Scenarios & Playbooks", style_h1))
story.append(Paragraph("• <b>Scenario A (Multi-Tenant B2B Attribution):</b> SaaS platform appends <code>x-uf-customer-id: cust_42</code> on requests. UnifAI attributes token spend per customer, exported to BigQuery for invoicing.", style_body))
story.append(Paragraph("• <b>Scenario B (Zero-Downtime Failover):</b> When OpenAI returns 5xx errors, Circuit Breaker trips and reroutes requests to Claude-3.5-Sonnet in <100ms. Clients receive <code>x-unifai-fallback-index: 1</code>.", style_body))
story.append(Paragraph("• <b>Scenario C (Complexity & Cache Savings):</b> Semantic Caching resolves 25% of queries at $0 cost; Complexity Router sends 70% of remaining queries to micro-models, cutting total spend by 78%.", style_body))
story.append(Paragraph("• <b>Scenario D (Enterprise DLP):</b> Browser AI local proxy scans employee browser queries for credit card numbers or API secrets, immediately blocking leaks and recording security audit trails.", style_body))

pdf_doc.build(story, canvasmaker=NumberedCanvas)
print("PDF document created successfully at:", PDF_PATH)
print("ALL DOCUMENTS GENERATED SUCCESSFULLY IN:", DOC_DIR)
