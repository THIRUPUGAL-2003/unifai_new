import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether, PageBreak, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

# Reconfigure console output to UTF-8
if sys.stdout and hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

WORKSPACE_DIR = r"c:\Users\sakth\OneDrive\ドキュメント\Desktop\unify - Copy"
DOC_DIR = os.path.join(WORKSPACE_DIR, "document")
os.makedirs(DOC_DIR, exist_ok=True)

MD_PATH = os.path.join(DOC_DIR, "UnifAI_Governance_and_Plugins_Master_Guide.md")
DOCX_PATH = os.path.join(DOC_DIR, "UnifAI_Governance_and_Plugins_Master_Guide.docx")
PDF_PATH = os.path.join(DOC_DIR, "UnifAI_Governance_and_Plugins_Master_Guide.pdf")

# ==============================================================================
# DATA DEFINITIONS FOR PLUGINS & 9 GOVERNANCE FEATURES (THANGLISH + ENGLISH)
# ==============================================================================

GOV_FEATURES = [
    {
        "id": "plugins",
        "name_en": "Plugins",
        "name_tanglish": "Plugins (Gateway Extensions & Lifecycle Hooks)",
        "route": "/workspace/plugins",
        "analogy": "Chrome browser extensions allathu WordPress plugins mathiri.",
        "explanation": "UnifAI AI Gateway-kulla ungalukku thevaiyana custom logic-ah (e.g. custom token counter, sentiment checking, internal billing webhook, proprietary data tagging) add panni gateway functionality-ah expand pandra extensible plugin engine. Gateway core code-ah maathaama plug-and-play-ah puthu features load pannikalam.",
        "business_value": "Company-oda proprietary security, legal compliance, and custom billing rules-ah AI request flow-kulla zero friction-oda seamlessly integrate pannalam.",
        "tech_arch": "WebAssembly (Wasm) and Go-based lifecycle interceptor pipeline. Hooks into 4 critical execution phases: `pre-request` (header validation, prompt enrichment), `post-request` (response manipulation), `response-streaming` (SSE chunk interception), and `error-handling`. Managed via dynamic plugin registry and ordered via execution sequence priority.",
        "endpoints": [
            "GET /api/v1/plugins",
            "POST /api/v1/plugins",
            "PATCH /api/v1/plugins/{plugin_id}",
            "DELETE /api/v1/plugins/{plugin_id}",
            "PUT /api/v1/plugins/sequence"
        ],
        "ui_elements": {
            "top_bar": [
                "Header Title with Puzzle Icon: Visual indicator of plugin management.",
                "Add New Plugin Button: Opens the plugin registration sheet.",
                "Plugin Sequence Button (ListOrdered): Opens visual re-ordering dialog to configure execution pipeline precedence."
            ],
            "tabs_and_views": [
                "Plugins Sidebar: List of custom plugins with active selection highlight, icon, and plugin name.",
                "PluginsView: Main editor panel showing Plugin Name, Description, Version, Schema definition, and Active Toggle Switch.",
                "Execution Hook Selectors: Toggles for Pre-Request, Post-Request, Response-Streaming, and Error-Handling."
            ],
            "bottom_elements": [
                "AddNewPluginSheet: Slide-over form to register plugin name, code/URL, target hooks, and initial configuration parameters.",
                "PluginSequenceSheet: Drag-and-drop or order selector determining which plugin executes first, second, or last in the proxy chain.",
                "Delete Plugin Confirmation Dialog: Safeguard preventing accidental deletion of active production plugins."
            ]
        },
        "connections": {
            "receives_from": "Intercepts raw incoming client requests in FastHTTP Proxy before Routing Rules.",
            "triggers_and_affects": "Enriches request headers, modifies prompt payloads, triggers external webhooks, and logs execution traces to LLM Logs."
        },
        "use_case": "Adding an internal compliance plugin that intercepts all prompts, runs a custom legal disclaimer check, and injects billing tags before forwarding to OpenAI."
    },
    {
        "id": "virtual_keys",
        "name_en": "Virtual Keys",
        "name_tanglish": "Virtual Keys (Proxy API Keys & Granular Quotas)",
        "route": "/workspace/governance/virtual-keys",
        "analogy": "Bank master account-ku multiple sub-debit cards kudukura mathiri (ovvoru card-kum thani thani limit & PIN).",
        "explanation": "Unga unmaiyana master OpenAI / Anthropic API keys-ah developers kitta direct-ah kudukka thevaiye illa! Athukku bathila 'Virtual Key' create panni kudupom. Intha virtual key-la monthly $100 budget, specific models mattum use panra permission, and rate limit set pannalam. Oru key leak aanaalum single click-la revoke allathu rotate pannikalam, real keys safe-ah irukkum.",
        "business_value": "Zero credential leakage risk; 100% granular cost attribution by team/app; instant security revocation without breaking other applications.",
        "tech_arch": "High-throughput cryptographic API key gateway middleware. Keys are generated with cryptographically secure tokens (`uf-key-...`) and stored as one-way SHA-256 hashes in PostgreSQL/Redis. FastHTTP auth middleware validates tokens in < 1ms, loads associated Access Profiles, checks budget balances, and enforces RPM/TPM token rate limits before proxying.",
        "endpoints": [
            "GET /api/v1/governance/virtual-keys",
            "POST /api/v1/governance/virtual-keys",
            "PATCH /api/v1/governance/virtual-keys/{key_id}",
            "DELETE /api/v1/governance/virtual-keys/{key_id}",
            "POST /api/v1/governance/virtual-keys/rotate",
            "POST /api/v1/governance/virtual-keys/bulk-rotate"
        ],
        "ui_elements": {
            "top_bar": [
                "Create Virtual Key Button: Opens key generation sheet.",
                "Bulk Rotate Keys Button: Triggers multi-key credential rotation dialog.",
                "Search & Filter Bar: Debounced search by key name, filter by Team, Customer, or Status (Active / Suspended)."
            ],
            "tabs_and_views": [
                "Virtual Keys Table: Columns for Key Name, Masked Secret Value (with copy button and eye toggle reveal), Team / Customer association, Budget Progress Bar (BudgetDisplay), Rate Limit badge (RateLimitDisplay), Status Toggle Switch (Active/Suspended), Actions Menu.",
                "Pinned Actions Column: Quick actions for Edit, Rotate Key, View Logs deep-link, and Delete."
            ],
            "bottom_elements": [
                "VirtualKeySheet (Slide-Over): Name, Description, Team assignment, Customer link, Access Profile attachment, Budget Cap ($), Rate Limits (Requests/min, Tokens/min), Model Allowlist/Denylist, IP Whitelist, Expiration Date.",
                "Key Secret Reveal Modal: One-time high-security display modal showing the full unmasked key upon creation with copy-to-clipboard button."
            ]
        },
        "connections": {
            "receives_from": "Attached to Teams, Customers, and Access Profiles configured in Governance.",
            "triggers_and_affects": "Authorizes incoming requests in FastHTTP Proxy; tracks spend into LLM Logs, Dashboard, and Budgets & Limits."
        },
        "use_case": "Issuing a virtual key with a $200/month cap and Gemini-only access to an external vendor development agency."
    },
    {
        "id": "users",
        "name_en": "Users",
        "name_tanglish": "Users (Platform User Directory & Identity Management)",
        "route": "/workspace/governance/users",
        "analogy": "Company employee directory & staff ID register mathiri.",
        "explanation": "UnifAI platform-la account irukkura ellam developers, team leads, finance managers, and admins list. Yaar yaar enna email, enna role (Admin, Member, Viewer), entha team-la irukaanga nu paathu manage pandra central user directory.",
        "business_value": "Identity governance, frictionless employee onboarding, and automated access termination when personnel depart.",
        "tech_arch": "User identity and authentication manager. Integrated with enterprise SSO (SAML 2.0 / OIDC / OAuth 2.0) and automated SCIM provisioning. Sessions are cryptographically signed using JWTs. Enforces multi-factor authentication (MFA) and maintains foreign key relations to Teams, Virtual Keys, and RBAC roles.",
        "endpoints": [
            "GET /api/v1/governance/users",
            "POST /api/v1/governance/users/invite",
            "PATCH /api/v1/governance/users/{user_id}",
            "DELETE /api/v1/governance/users/{user_id}",
            "POST /api/v1/governance/users/{user_id}/reset-mfa"
        ],
        "ui_elements": {
            "top_bar": [
                "Invite User Button: Opens email invitation dialog with role assignment.",
                "Search Bar: Search users by name, email, or role.",
                "Filter Dropdown: Filter by Role (Admin, Member, Viewer) or Status (Active, Invited, Suspended)."
            ],
            "tabs_and_views": [
                "Users Table: Columns for User Avatar & Name, Email Address, Assigned RBAC Role badge, Team Memberships, Last Active Timestamp, Account Status badge (Active/Pending), Actions Menu.",
                "User Activity Card: Summary of user's personal virtual keys, aggregate token consumption, and audit trail link."
            ],
            "bottom_elements": [
                "InviteUserDialog: Form for User Email, Full Name, Initial Role selection, and Assigned Teams.",
                "EditUserSheet: Slide-over form to update assigned roles, toggle admin privileges, change team assignments, or suspend user access.",
                "Deactivate User Confirmation Modal: Revokes all active user tokens and reassigns created virtual keys."
            ]
        },
        "connections": {
            "receives_from": "Synced automatically from SCIM / SSO or manually invited by Admins.",
            "triggers_and_affects": "Binds to Teams, Virtual Keys, and RBAC roles; logs all platform actions into Audit Logs."
        },
        "use_case": "Onboarding 50 new data science engineers simultaneously and assigning them to the 'AI Research Team' with Member-level permissions."
    },
    {
        "id": "teams",
        "name_en": "Teams",
        "name_tanglish": "Teams (Departmental Squads & Shared Resource Pools)",
        "route": "/workspace/governance/teams",
        "analogy": "Company departments allathu project squads mathiri (e.g. Mobile App Squad, Core Search Team, AI Labs).",
        "explanation": "Developers-ah project allathu department vaariya group panni team-ah pirikalam. Oru team-ku common monthly budget, shared virtual keys, and collaborative model access kudukalam. Team-la irukkura yaar token use panninaalum antha team budget-la irunthu கழிவு aagum.",
        "business_value": "Departmental accountability; eliminates individual credential confusion; enables team-level AI spend tracking for CFO budget reviews.",
        "tech_arch": "Multi-tenant team resource isolation layer. Aggregates usage across all member users and attached virtual keys. Enforces team-level budget limits in PostgreSQL, controls shared access profiles, and maps seamlessly to enterprise SCIM groups.",
        "endpoints": [
            "GET /api/v1/governance/teams",
            "POST /api/v1/governance/teams",
            "PATCH /api/v1/governance/teams/{team_id}",
            "DELETE /api/v1/governance/teams/{team_id}",
            "POST /api/v1/governance/teams/{team_id}/members"
        ],
        "ui_elements": {
            "top_bar": [
                "Create Team Button: Opens team creation dialog.",
                "Search Input: Search teams by name or business unit.",
                "Business Unit Filter: Filter teams by parent division."
            ],
            "tabs_and_views": [
                "Teams Table: Columns for Team Name, Description, Members Count badge, Associated Virtual Keys count, Monthly Budget Progress Bar, Parent Business Unit, Actions Menu.",
                "Team Details Sheet: Comprehensive breakdown showing Team Leads, Active Members, Assigned Virtual Keys, and 30-day spend trajectory graph."
            ],
            "bottom_elements": [
                "CreateTeamDialog: Form specifying Team Name, Description, Parent Business Unit selection, Monthly Budget Allocation ($), and Initial Team Lead.",
                "ManageMembersSheet: Modal to add or remove members and adjust user team roles (Team Lead vs Member)."
            ]
        },
        "connections": {
            "receives_from": "Belongs to a parent Business Unit and groups multiple Users.",
            "triggers_and_affects": "Owns shared Virtual Keys; feeds departmental spend rollups into Dashboard and Connectors (Datadog/Kafka)."
        },
        "use_case": "Allocating a shared $3,000 monthly budget to the 'E-Commerce Recommendations Team' with 12 developers using 4 shared virtual keys."
    },
    {
        "id": "business_units",
        "name_en": "Business Units",
        "name_tanglish": "Business Units (Enterprise Divisions & P&L Centers)",
        "route": "/workspace/governance/business-units",
        "analogy": "Periya corporate conglomerate-oda divisions mathiri (e.g. Retail Division, Banking Division, Healthcare Division).",
        "explanation": "Teams-ku mela irukura periya enterprise division. CFO and executive leadership company-oda multiple business units-kulla AI budget分配 pannavum, P&L (Profit & Loss) tracking seiyavum use aagum. Oru division kulla 10 teams irukkalam.",
        "business_value": "High-level corporate financial governance; automated chargebacks across major business divisions; strategic AI budget planning.",
        "tech_arch": "Top-level organizational container in the multi-tenant governance tree (`BusinessUnit -> Teams -> Users -> Virtual Keys`). Aggregates financial consumption metrics via database rollups and provides master cost-allocation tags for corporate accounting systems (SAP, NetSuite).",
        "endpoints": [
            "GET /api/v1/governance/business-units",
            "POST /api/v1/governance/business-units",
            "PATCH /api/v1/governance/business-units/{bu_id}",
            "DELETE /api/v1/governance/business-units/{bu_id}"
        ],
        "ui_elements": {
            "top_bar": [
                "Create Business Unit Button: Opens division registration sheet.",
                "Search Input: Search divisions by name or corporate code."
            ],
            "tabs_and_views": [
                "Business Units Table: Columns for Unit Name, Corporate Cost Code, Description, Total Sub-Teams count, Aggregate Spend ($), Head of Business Unit, Actions Menu.",
                "Hierarchical Rollup View: Expandable accordion showing all teams, members, and keys nested under this business unit."
            ],
            "bottom_elements": [
                "BusinessUnitSheet: Form specifying Business Unit Name, Cost Center Code, Executive Owner email, Annual/Quarterly AI Budget Cap ($).",
                "Delete Confirmation Modal: Reassigns or archives child teams before division deletion."
            ]
        },
        "connections": {
            "receives_from": "Configured by executive administrators and finance directors.",
            "triggers_and_affects": "Contains multiple Teams; feeds top-level cost analytics into Executive Dashboards and Billing Connectors."
        },
        "use_case": "Tracking that the 'Global Wealth Management' division consumed $45,000 in AI tokens this quarter across its 8 engineering teams."
    },
    {
        "id": "customers",
        "name_en": "Customers",
        "name_tanglish": "Customers (B2B Client Tenants & External Accounts)",
        "route": "/workspace/governance/customers",
        "analogy": "SaaS platform-la irukura external client accounts mathiri (B2B Multi-tenancy).",
        "explanation": "Neenga unga software-ah vera external clients-ku AI features vachu SaaS product-ah vikkiringa na, ovvoru client-aiyum 'Customer'-ah add panni avanga usage-ku bill podalam. Oru customer innoru customer-oda data-vai paarka mudiyathu.",
        "business_value": "Monetize AI applications; seamless B2B client isolation; automated billing invoices based on client consumption.",
        "tech_arch": "External tenant isolation boundary. Associates virtual keys with external customer identifiers (`customer_id`). Enforces hard tenant isolation policies in PostgreSQL using Row-Level Security (RLS) and powers multi-tenant usage metering pipelines.",
        "endpoints": [
            "GET /api/v1/governance/customers",
            "POST /api/v1/governance/customers",
            "PATCH /api/v1/governance/customers/{customer_id}",
            "DELETE /api/v1/governance/customers/{customer_id}"
        ],
        "ui_elements": {
            "top_bar": [
                "Create Customer Button: Opens customer onboarding modal.",
                "Search & Filter: Search by customer name, external company ID, or account tier."
            ],
            "tabs_and_views": [
                "Customers Table: Columns for Customer Name, Account Status badge (Active/Trial/Suspended), Assigned Teams, Number of Virtual Keys, Cumulative Spend ($), Account Tier, Actions Menu.",
                "Usage Metering Panel: Detailed graph of customer's token consumption broken down by day and model."
            ],
            "bottom_elements": [
                "CustomerSheet: Form specifying Customer Name, External Account ID, Contact Email, Assigned Pricing Tier, and Hard Spend Quota.",
                "Suspend Customer Modal: Temporarily halts all virtual keys assigned to this customer during billing delinquency."
            ]
        },
        "connections": {
            "receives_from": "Created by sales or synced from CRM (Salesforce, Stripe).",
            "triggers_and_affects": "Scopes Virtual Keys and filters LLM Logs for customer-specific audit and billing exports."
        },
        "use_case": "Providing 100 enterprise B2B customers their own dedicated virtual keys with strictly metered billing for your AI-powered legal document review SaaS."
    },
    {
        "id": "scim",
        "name_en": "User Provisioning (SCIM)",
        "name_tanglish": "User Provisioning / SCIM (Automated Enterprise SSO & Lifecycle Sync)",
        "route": "/workspace/scim",
        "analogy": "Automatic company ID card issuing & revoking machine mathiri (HR portal-la employee join aana automatic-ah account create aagum).",
        "explanation": "Company HRMS allathu Identity Provider (Okta, Microsoft Entra ID / Azure AD, Keycloak)-la pudhu aal sertha automatic-ah UnifAI-layum login create aagum. Employee velaiya vittu ponavudane Okta-la offboard pannina, automatic-ah UnifAI access cut aagidum! Manual-ah user create panna thevaiye illa.",
        "business_value": "Zero orphan accounts; 100% SOC-2 & ISO-27001 compliance for employee lifecycle security; automated enterprise IT operations.",
        "tech_arch": "Standard SCIM v2.0 (RFC 7643 & RFC 7644) server implementation. Exposes `/scim/v2/Users`, `/scim/v2/Groups`, and `/scim/v2/ServiceProviderConfig`. Authenticated via cryptographically generated Bearer tokens. Parses incoming SCIM JSON payloads to automatically create, update, deactivate users, and synchronize directory group memberships to UnifAI Teams.",
        "endpoints": [
            "GET /scim/v2/Users",
            "POST /scim/v2/Users",
            "GET /scim/v2/Users/{id}",
            "PUT /scim/v2/Users/{id}",
            "PATCH /scim/v2/Users/{id}",
            "DELETE /scim/v2/Users/{id}",
            "GET /scim/v2/Groups",
            "GET /api/v1/scim/config",
            "PUT /api/v1/scim/config"
        ],
        "ui_elements": {
            "top_bar": [
                "Header Title with UserRoundCog Icon: Visual indicator of SCIM provisioning.",
                "SCIM Master Enable Switch: Instant activation/deactivation toggle."
            ],
            "tabs_and_views": [
                "SCIM Endpoints Card: Copyable endpoint URLs for IdP configuration:",
                "• Base URL: `https://<domain>/scim/v2` (with Copy button)",
                "• Users Endpoint: `https://<domain>/scim/v2/Users` (with Copy button)",
                "• ServiceProviderConfig: `https://<domain>/scim/v2/ServiceProviderConfig`",
                "IdP Provider Selector: Toggle between Okta, Microsoft Entra ID (Azure AD), Keycloak, and Custom SCIM v2 Provider."
            ],
            "bottom_elements": [
                "Bearer Token Management: Secret token generator with copy button and rotation warning.",
                "Group Mapping Card: Rules mapping external directory groups (e.g. `okta-ai-engineers`) directly into UnifAI Teams and RBAC roles.",
                "Save SCIM Config Button: Commits configuration with toast confirmation."
            ]
        },
        "connections": {
            "receives_from": "External Identity Providers (Okta, Microsoft Entra ID, PingIdentity).",
            "triggers_and_affects": "Automatically provisions and deprovisions Users and synchronizes Team memberships."
        },
        "use_case": "Automatically creating UnifAI accounts for 2,000 corporate employees via Okta SCIM sync and immediately revoking access when an employee leaves the company."
    },
    {
        "id": "rbac",
        "name_en": "Roles & Permissions (RBAC)",
        "name_tanglish": "Roles & Permissions (RBAC Security & Granular Access Matrix)",
        "route": "/workspace/governance/rbac",
        "analogy": "Office security access badge permissions mathiri (Receptionist lobby-ku mattum entry, Server room-ku Admin mattum entry).",
        "explanation": "Yaaru platform-la enna seiyalaam nu kattu paduthum access control matrix. Super Admin-ku full control irukkum; Developer-ku API key create panna mattum permission; Finance Manager-ku logs and billing charts mattum paarka permission (read-only); Normal user-ku settings maatha permission irukaathu.",
        "business_value": "Least-privilege security principle. Accidental deletions, unauthorized key tampering, and unauthorized configuration changes-ah 100% prevent pannum.",
        "tech_arch": "Declarative Role-Based Access Control (RBAC) engine in Go and React. Enforces fine-grained resource-action pairs (`Resource: ModelProvider, VirtualKeys, RoutingRules, Settings, AuditLogs, Plugins, Governance` × `Action: View, Create, Update, Delete`). Evaluates permissions on every FastHTTP management endpoint and hides UI components conditionally via `useRbac()` hook.",
        "endpoints": [
            "GET /api/v1/governance/rbac/roles",
            "POST /api/v1/governance/rbac/roles",
            "PATCH /api/v1/governance/rbac/roles/{role_id}",
            "DELETE /api/v1/governance/rbac/roles/{role_id}",
            "GET /api/v1/governance/rbac/permissions"
        ],
        "ui_elements": {
            "top_bar": [
                "Create Custom Role Button: Opens role builder dialog.",
                "Role Filter: Filter by System Roles (Admin, Member, Read-Only) vs Custom Defined Roles."
            ],
            "tabs_and_views": [
                "Roles Table: Columns for Role Name, Description, Assigned Users count badge, Role Type (System / Custom), Actions Menu.",
                "RBAC Permission Matrix Grid: Comprehensive grid displaying Resources on rows (Virtual Keys, Model Providers, Routing Rules, Budgets, Audit Logs) and Actions on columns (View, Create, Update, Delete) with interactive checkboxes."
            ],
            "bottom_elements": [
                "RoleBuilderSheet: Slide-over form to name custom role, define description, and toggle specific permission checkboxes.",
                "AssignUsersModal: Dialog to bulk-assign the role to specific users or teams."
            ]
        },
        "connections": {
            "receives_from": "Assigned to Users and SCIM directory groups.",
            "triggers_and_affects": "Enforces authorization gates across all UnifAI UI navigation, buttons, and backend REST APIs."
        },
        "use_case": "Creating a 'Finance Auditor' custom role that can View Dashboard, View LLM Logs, and View Pricing Overrides, but cannot Create, Edit, or Delete any keys or models."
    },
    {
        "id": "access_profiles",
        "name_en": "Access Profiles",
        "name_tanglish": "Access Profiles (Reusable Policy Bundles & Guardrail Packs)",
        "route": "/workspace/governance/access-profiles",
        "analogy": "Pre-packaged security bundle allathu employee safety passport mathiri.",
        "explanation": "Oru standard policy template create panni vachikalam. Example: 'Intern Security Profile'—ithula GPT-4o-mini mattum allow panni, $50 budget limit, sensitive PII data redaction guardrails, and office IP restriction potruppom. Pudhu virtual key create pannum pothu intha profile-ah select panna ella security rules-um automatic-ah apply aagidum!",
        "business_value": "Massive administrative time savings; eliminates configuration mistakes; guarantees that no virtual key goes live without corporate security guardrails.",
        "tech_arch": "Reusable policy entity bundling multiple governance constraints: Allowed Model Patterns, Allowed Providers, Budget Maximums & Reset Cadences, Request/Token Rate Limits (RPM/TPM), Attached MCP Tool Clients, and DLP Guardrail policies. Attached to Virtual Keys via relational join, allowing instant fleet-wide policy updates when a profile is edited.",
        "endpoints": [
            "GET /api/v1/governance/access-profiles",
            "POST /api/v1/governance/access-profiles",
            "PATCH /api/v1/governance/access-profiles/{profile_id}",
            "DELETE /api/v1/governance/access-profiles/{profile_id}",
            "POST /api/v1/governance/access-profiles/{profile_id}/clone"
        ],
        "ui_elements": {
            "top_bar": [
                "Create Access Profile Button: Opens profile builder dialog.",
                "Search & Tag Filter: Filter profiles by tag (e.g. `production`, `intern`, `customer-facing`)."
            ],
            "tabs_and_views": [
                "Access Profiles Table: Columns for Profile Name, Tags badges, Allowed Models summary, Attached Virtual Keys count, Budget & Rate Limit summary, Actions Menu.",
                "Quick Clone Button: Instantly duplicates an existing profile for fast customization."
            ],
            "bottom_elements": [
                "AccessProfileDialog: Comprehensive form configuring:",
                "• Profile Name, Description, and Categorization Tags.",
                "• Allowed Providers & Model Multiselect (or regex wildcard `gpt-4o-*`).",
                "• Budget Caps: Max Dollar Limit ($) & Reset Duration (Daily, Monthly).",
                "• Rate Limits: Max Requests per Minute (RPM) and Max Tokens per Minute (TPM).",
                "• MCP Client Permissions: Select which autonomous tools this profile can execute.",
                "• Attached Virtual Keys selector."
            ]
        },
        "connections": {
            "receives_from": "Created by Security & Compliance administrators.",
            "triggers_and_affects": "Governs Virtual Keys behavior, restricts Model Providers, and enforces MCP Gateway tool permissions."
        },
        "use_case": "Attaching a single 'Customer Facing Chatbot Profile' to 20 virtual keys, guaranteeing strict PII redaction and 50 RPM rate limits across all 20 keys simultaneously."
    },
    {
        "id": "audit_logs",
        "name_en": "Audit Logs",
        "name_tanglish": "Audit Logs (Tamper-Proof Compliance & Activity Trail)",
        "route": "/workspace/audit-logs",
        "analogy": "Bank locker CCTV camera & register note mathiri.",
        "explanation": "Platform-la yaaru yaaru entha nerathula enna maathinaanga nu maatha mudiyatha (tamper-proof) pathivu. Yaarachum API key delete pannalaam, budget limit maathinaalo, pudhu routing rule pottaalo exact timestamp, user name, and client IP address-oda capture aagum. 'Who did what and when' nu eppovum proof irukkum.",
        "business_value": "100% SOC-2, HIPAA, and GDPR audit readiness. Insider threat detection and rapid forensic debugging during system anomalies.",
        "tech_arch": "Immutable, append-only security event audit trail. Intercepts all administrative REST mutations in middleware. Captures Actor User ID, Action Type (CREATE, UPDATE, DELETE), Target Resource, Resource ID, Timestamp, Origin Client IP address, User-Agent, and detailed before/after JSON diffs. Stored in partitioned, tamper-evident PostgreSQL tables with optional S3 Glacier write-once-read-many (WORM) archival.",
        "endpoints": [
            "GET /api/v1/audit-logs",
            "GET /api/v1/audit-logs/{log_id}",
            "GET /api/v1/audit-logs/export"
        ],
        "ui_elements": {
            "top_bar": [
                "Date Range Picker: Quick filters for 24h, 7d, 30d, 90d, and custom calendar range.",
                "Actor Search Input: Search by user email or user ID.",
                "Action Type Filter: Filter by CREATE, UPDATE, DELETE, ROTATE, LOGIN.",
                "Resource Filter: Filter by Virtual Keys, Routing Rules, Model Providers, Users, Teams, etc.",
                "Export Button: Download audit logs in CSV or JSON format for external auditors."
            ],
            "tabs_and_views": [
                "Audit Logs Table: Columns for Timestamp (UTC/Local), Actor Name & Email, Action Badge (Create [green], Update [blue], Delete [red]), Resource Type, Target Resource Name/ID, Client IP Address, Details Button.",
                "Pagination Controls: Fast offset/cursor pagination handling hundreds of thousands of audit events."
            ],
            "bottom_elements": [
                "Audit Detail Drawer / Modal: Displays the complete before-and-after JSON delta diff highlighting exact fields that were modified during the administrative action."
            ]
        },
        "connections": {
            "receives_from": "Automatically intercepts administrative operations across all Governance, Models, and Observability features.",
            "triggers_and_affects": "Feeds security incident alerts into Alert Channels and provides evidence for external compliance auditors."
        },
        "use_case": "Forensic audit demonstrating to external SOC-2 auditors exactly who rotated the production Bedrock credentials on August 14th at 03:22 UTC."
    }
]

GOV_CONN_DATA = [
    ("Plugins", "FastHTTP Proxy", "Intercepts pre/post request payloads to run custom enterprise compliance logic."),
    ("Virtual Keys", "Access Profiles", "Inherits model restrictions, budget limits, and MCP tool permissions from assigned profiles."),
    ("Users", "Teams", "Organized into collaborative squads that share aggregate budgets and virtual keys."),
    ("Teams", "Business Units", "Rolls up departmental spend and resource counts into corporate division P&L centers."),
    ("Customers", "Virtual Keys", "Isolates external B2B client applications with dedicated credentials and usage tracking."),
    ("SCIM", "Users & Teams", "Automatically provisions users and syncs directory groups from Okta and Microsoft Entra ID."),
    ("RBAC", "All UI & Endpoints", "Enforces fine-grained View/Create/Update/Delete permissions on every button and API."),
    ("Access Profiles", "Virtual Keys", "Enforces model allowlists, token quotas, and rate-limits across fleets of keys."),
    ("Audit Logs", "All Features", "Immutably records every configuration change, key rotation, and administrative event.")
]

GOV_MATRIX_DATA = [
    ("Plugins", "Custom rules and company billing logic-ah gateway-kulla plug-and-play-ah add panna mudiyuma?", "Wasm/Go lifecycle interceptor pipeline hooking into pre-request, post-request, and streaming chunks."),
    ("Virtual Keys", "Real OpenAI master keys-ah kudukama safe-ana dummy keys with budget limits kudukalama?", "Cryptographic SHA-256 hashed API tokens with sub-millisecond FastHTTP auth and rate limiters."),
    ("Users", "Company-la irukkura ellam developers and staff directory-ah ore idathula manage pannalama?", "Central identity directory integrated with enterprise SSO (OIDC/SAML) and JWT session tokens."),
    ("Teams", "Developers-ah group panni project teams-kaga shared budget and keys allocate pannalama?", "Multi-tenant team resource isolation layer aggregating spend and managing shared virtual keys."),
    ("Business Units", "CFO and leadership periya divisions (Retail, Banking) vaariya AI spend track panna mudiyuma?", "Top-level corporate cost center container rolling up child teams for SAP/NetSuite accounting."),
    ("Customers", "Vera external B2B clients-ku AI features vithu avangalukku thaniya bill poda mudiyuma?", "External tenant isolation boundary enforcing RLS data separation and customer-metered billing."),
    ("SCIM", "Okta allathu Entra ID-la employee join aana automatic-ah account create aagi, leave aana cut aaguma?", "RFC 7643/7644 SCIM v2.0 server handling automated user lifecycle and group synchronization."),
    ("RBAC", "Admin-ku mattum full rights, mathavangalukku view-only permission pottu security maintain pannalama?", "Declarative resource-action permission matrix evaluating gates across all UI views and REST APIs."),
    ("Access Profiles", "Reusable standard security template create panni 20 keys-ku single click-la apply pannalama?", "Bundled governance policy entity containing model allowlists, budget caps, rate limits, and MCP permissions."),
    ("Audit Logs", "Yaaru entha key-ah maathinaanga, eppo delete panninaanga nu maatha mudiyatha proof irukka?", "Append-only immutable audit trail capturing actor, action, timestamp, IP, and before/after JSON diffs.")
]

# ==============================================================================
# 1. GENERATE MARKDOWN DOCUMENT (THANGLISH + ENGLISH)
# ==============================================================================
def generate_gov_markdown():
    print("Writing Governance & Plugins Markdown document (Thanglish + English)...")
    lines = []
    lines.append("# UnifAI Governance & Plugins Master Deep-Dive Guide")
    lines.append("## UnifAI Governance, Security & Plugins Amaippin Muzhumaiyana Manual")
    lines.append("")
    lines.append("**Pillars Covered:** Plugins & Governance (Virtual Keys, Users, Teams, Business Units, Customers, SCIM, RBAC, Access Profiles, Audit Logs)  ")
    lines.append("**Target Audience:** Security Architects, CTOs, CFOs, Compliance Officers, Developers & System Admins  ")
    lines.append("**Language:** Bilingual (Thanglish - Tamil in English letters + Clean English)  ")
    lines.append("**Generated At:** 2026-09-05  ")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Table of Contents")
    lines.append("1. [Governance & Plugins Architecture Map](#1-governance--plugins-architecture-map)")
    lines.append("2. [Detailed Feature Dissection (10 Core Features)](#2-detailed-feature-dissection)")
    for f in GOV_FEATURES:
        lines.append(f"   - [{f['name_en']} ({f['name_tanglish']}) ({f['route']})](#{f['id']})")
    lines.append("3. [Cross-Feature Interconnections & Data Flow](#3-cross-feature-interconnections--data-flow)")
    lines.append("4. [Tech vs Non-Tech Comparative Matrix](#4-tech-vs-non-tech-comparative-matrix)")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 1. Governance & Plugins Architecture Map")
    lines.append("### Enterprise Identity, Security Boundaries & Request Lifecycle Flow")
    lines.append("")
    lines.append("```")
    lines.append("             [ ENTERPRISE IDENTITY (Okta / Entra ID / IdP) ]")
    lines.append("                                   │")
    lines.append("                    (SCIM v2 Sync) ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │                USER PROVISIONING (SCIM)                   │")
    lines.append("       │  • Auto-creates Users, Maps directory groups to Teams     │")
    lines.append("       └───────────────────────────┬───────────────────────────────┘")
    lines.append("                                   │")
    lines.append("                                   ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │             ENTERPRISE TENANT HIERARCHY                   │")
    lines.append("       │                                                           │")
    lines.append("       │   [ BUSINESS UNITS ] (Divisions / P&L Centers)            │")
    lines.append("       │          │                                                │")
    lines.append("       │          ▼                                                │")
    lines.append("       │      [ TEAMS ] (Departments / Squads) ◄──► [ CUSTOMERS ]   │")
    lines.append("       │          │                                                │")
    lines.append("       │          ▼                                                │")
    lines.append("       │      [ USERS ] (Developers / Admins / Viewers)             │")
    lines.append("       └───────────────────────────┬───────────────────────────────┘")
    lines.append("                                   │")
    lines.append("                                   ▼ (Governed by RBAC Matrix)")
    lines.append("       ┌───────────────────────────┐   ┌───────────────────────────┐")
    lines.append("       │      ACCESS PROFILES      │──►│       VIRTUAL KEYS        │")
    lines.append("       │  • Model Allowlists       │   │  • Cryptographic Tokens   │")
    lines.append("       │  • Budget & Rate Limits   │   │  • Granular Cost Caps     │")
    lines.append("       │  • Guardrails & MCP Tools │   │  • Instant Revocation     │")
    lines.append("       └───────────────────────────┘   └─────────────┬─────────────┘")
    lines.append("                                                     │ (Bearer uf-key-...)")
    lines.append("                                                     ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │               FastHTTP PROXY GATEWAY                      │")
    lines.append("       │  • Token Auth, Budget check, IP verification              │")
    lines.append("       └───────────────────────────┬───────────────────────────────┘")
    lines.append("                                   │")
    lines.append("                                   ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │                      PLUGINS ENGINE                       │")
    lines.append("       │  • Pre-Request Hooks: Custom Auth, Headers, Validation    │")
    lines.append("       │  • Post-Request & Streaming Hooks: Billing, Scrubbing     │")
    lines.append("       └───────────────────────────┬───────────────────────────────┘")
    lines.append("                                   │")
    lines.append("                                   ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │                        AUDIT LOGS                         │")
    lines.append("       │  • Immutable Record: Actor, Action, Target, Diff, IP      │")
    lines.append("       └───────────────────────────────────────────────────────────┘")
    lines.append("```")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 2. Detailed Feature Dissection (10 Core Features)")
    lines.append("### Pathu Features-oda Aazhamana Vivaram (Thanglish + English)")
    lines.append("")

    for f in GOV_FEATURES:
        lines.append(f"<a name='{f['id']}'></a>")
        lines.append(f"## {f['name_en']} — {f['name_tanglish']}")
        lines.append(f"**UI Route:** `{f['route']}`\n")
        lines.append("### 👤 Non-Tech Perspective (Makkal Puriyura Mathiri Eliya Vilakkam)")
        lines.append(f"- **Uruvagam (Analogy):** {f['analogy']}")
        lines.append(f"- **Vilakkam (Explanation):** {f['explanation']}")
        lines.append(f"- **Business Value (Vaniga Payan):** {f['business_value']}\n")
        lines.append("### 💻 Tech Perspective (Engineers-kaga Technical Architecture)")
        lines.append(f"- **Backend Architecture:** {f['tech_arch']}")
        lines.append("- **Backend Endpoints:**")
        for ep in f['endpoints']:
            lines.append(f"  * `{ep}`")
        lines.append("")
        lines.append("### 🖥️ Screen Layout & Interactive Elements (Thirai Koorugal & Buttons)")
        lines.append("**1. Melpura Koorugal (Top Bar Controls):**")
        for itm in f['ui_elements']['top_bar']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("**2. Mathiya Koorugal & Tables (Tabs & Views):**")
        for itm in f['ui_elements']['tabs_and_views']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("**3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):**")
        for itm in f['ui_elements']['bottom_elements']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("### 🔗 Connections & Structure Map (Inaippugal)")
        lines.append(f"- **Data-vai Perumidam (Receives From):** {f['connections']['receives_from']}")
        lines.append(f"- **Iyakkum Koorugal (Triggers & Affects):** {f['connections']['triggers_and_affects']}")
        lines.append("")
        lines.append(f"### 💡 Production Use Case: {f['use_case']}\n")
        lines.append("---\n")

    lines.append("# 3. Cross-Feature Interconnections & Data Flow")
    lines.append("### Governance & Plugins Features Kulla Irukura Direct Connections\n")
    lines.append("| Source Feature | Connected To | Data Flow & Trigger Action |")
    lines.append("| :--- | :--- | :--- |")
    for src, dst, flow in GOV_CONN_DATA:
        lines.append(f"| **{src}** | **{dst}** | {flow} |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 4. Tech vs Non-Tech Comparative Matrix")
    lines.append("### Thozhilnutpam vs Vanigam Parvai Oppeedo\n")
    lines.append("| Feature | Non-Tech View (Manager / CFO Parvai) | Tech View (DevOps / Architect Parvai) |")
    lines.append("| :--- | :--- | :--- |")
    for ft, nt, tv in GOV_MATRIX_DATA:
        lines.append(f"| **{ft}** | \"{nt}\" | \"{tv}\" |")
    lines.append("")

    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Governance & Plugins Markdown written to: {MD_PATH}")

# ==============================================================================
# 2. GENERATE DOCX DOCUMENT (THANGLISH + ENGLISH)
# ==============================================================================
def generate_gov_docx():
    print("Writing Governance & Plugins Word Document (.docx)...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        hp = section.header.paragraphs[0]
        hp.text = "UnifAI Governance & Plugins Master Deep-Dive Technical Manual"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if hp.runs:
            hp.runs[0].font.name = "Segoe UI"
            hp.runs[0].font.size = Pt(8.5)
            hp.runs[0].font.color.rgb = RGBColor(100, 116, 139)
            
        fp = section.footer.paragraphs[0]
        fp.text = "Confidential & Proprietary — UnifAI Governance & Plugins Control Plane (Thanglish + English)"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if fp.runs:
            fp.runs[0].font.name = "Segoe UI"
            fp.runs[0].font.size = Pt(8.5)
            fp.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    def style_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        return p

    def style_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def style_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.italic = italic
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def format_table(table, col_widths, col_names):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(col_names):
            hdr_cells[i].text = name
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
            tcPr.append(shd)
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = "Segoe UI"
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for r_idx, row in enumerate(table.rows[1:]):
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                tcPr.append(shd)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                    
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    rt = p_title.add_run("UnifAI Governance & Plugins Master Deep-Dive Guide")
    rt.font.name = "Segoe UI"
    rt.font.size = Pt(22)
    rt.bold = True
    rt.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    rsub = p_sub.add_run("Plugins, Virtual Keys, Users, Teams, Business Units, Customers, SCIM, RBAC, Access Profiles & Audit Logs (Thanglish + English)")
    rsub.font.name = "Segoe UI"
    rsub.font.size = Pt(11.5)
    rsub.font.color.rgb = RGBColor(71, 85, 105)

    style_h1("1. Governance & Plugins Architecture Map")
    add_p("Enterprise identity (Okta/Entra) SCIM v2 vazhiya sync aagi, Business Units, Teams, Users, Virtual Keys, Access Profiles, Plugins, matrum Audit Logs vazhiya operate aagura security control plane:")

    style_h1("2. Detailed Feature Dissection (10 Core Features)")
    for f in GOV_FEATURES:
        style_h2(f"{f['name_en']} — {f['name_tanglish']}")
        add_p(f['route'], bold_prefix="Route: ", italic=True)

        style_h3("Non-Tech Perspective (Eliya Vilakkam)")
        add_p(f['analogy'], bold_prefix="Uruvagam (Analogy): ")
        add_p(f['explanation'], bold_prefix="Vilakkam: ")
        add_p(f['business_value'], bold_prefix="Business Value: ")

        style_h3("Tech Perspective (Technical Architecture)")
        add_p(f['tech_arch'], bold_prefix="Backend Architecture: ")
        add_p("Backend Endpoints:")
        for ep in f['endpoints']:
            add_bullet(ep)

        style_h3("Screen Layout & Interactive Elements (Thirai Koorugal)")
        add_p("1. Melpura Koorugal (Top Bar Controls):", bold_prefix="")
        for itm in f['ui_elements']['top_bar']:
            add_bullet(itm)
        add_p("2. Mathiya Koorugal & Tables (Tabs & Views):", bold_prefix="")
        for itm in f['ui_elements']['tabs_and_views']:
            add_bullet(itm)
        add_p("3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):", bold_prefix="")
        for itm in f['ui_elements']['bottom_elements']:
            add_bullet(itm)

        style_h3("Module Connections (Inaippugal)")
        add_p(f['connections']['receives_from'], bold_prefix="Data-vai Perumidam: ")
        add_p(f['connections']['triggers_and_affects'], bold_prefix="Iyakkum Koorugal: ")

        style_h3("Production Use Case (Nadamurai Udharanam)")
        add_p(f['use_case'])

    style_h1("3. Cross-Feature Interconnections")
    t_conn = doc.add_table(rows=len(GOV_CONN_DATA)+1, cols=3)
    format_table(t_conn, [1.8, 1.8, 3.4], ["Source Feature", "Connected To", "Data Flow & Trigger Action"])
    for idx, (src, dst, flow) in enumerate(GOV_CONN_DATA):
        t_conn.rows[idx+1].cells[0].text = src
        t_conn.rows[idx+1].cells[1].text = dst
        t_conn.rows[idx+1].cells[2].text = flow

    style_h1("4. Tech vs Non-Tech Comparative Matrix")
    t_mat = doc.add_table(rows=len(GOV_MATRIX_DATA)+1, cols=3)
    format_table(t_mat, [1.5, 2.7, 2.8], ["Feature", "Non-Tech View (Manager / CFO)", "Tech View (DevOps / Architect)"])
    for idx, (ft, nt, tv) in enumerate(GOV_MATRIX_DATA):
        t_mat.rows[idx+1].cells[0].text = ft
        t_mat.rows[idx+1].cells[1].text = nt
        t_mat.rows[idx+1].cells[2].text = tv

    doc.save(DOCX_PATH)
    print(f"Governance & Plugins Word document written to: {DOCX_PATH}")

# ==============================================================================
# 3. GENERATE PDF DOCUMENT VIA REPORTLAB (THANGLISH + ENGLISH)
# ==============================================================================
def generate_gov_pdf():
    print("Writing Governance & Plugins PDF Document (.pdf)...")

    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748B"))
            
            if self._pageNumber > 1:
                self.drawString(40, 11 * inch - 36, "UnifAI Governance & Plugins Master Deep-Dive Technical Manual (Thanglish + English)")
                self.drawRightString(8.5 * inch - 40, 11 * inch - 36, "Confidential — Engineering Guide")
                self.setStrokeColor(colors.HexColor("#CBD5E1"))
                self.setLineWidth(0.5)
                self.line(40, 11 * inch - 40, 8.5 * inch - 40, 11 * inch - 40)
                
            self.drawString(40, 32, "UnifAI Unified Governance, Identity, Security & Plugins Control Plane")
            page_str = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(8.5 * inch - 40, 32, page_str)
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(40, 42, 8.5 * inch - 40, 42)
            
            self.restoreState()

    pdf_doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=48,
        bottomMargin=48
    )

    styles = getSampleStyleSheet()

    p_title_style = ParagraphStyle(
        'PdfTitle',
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=3
    )

    p_subtitle_style = ParagraphStyle(
        'PdfSubTitle',
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#475569'),
        spaceAfter=10
    )

    h1_style = ParagraphStyle(
        'PdfH1',
        fontName='Helvetica-Bold',
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=12,
        spaceAfter=5,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'PdfH2',
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'PdfH3',
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=5,
        spaceAfter=2,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'PdfBody',
        fontName='Helvetica',
        fontSize=8,
        leading=11,
        textColor=colors.HexColor('#334155'),
        spaceAfter=3
    )

    bullet_style = ParagraphStyle(
        'PdfBullet',
        fontName='Helvetica',
        fontSize=7.8,
        leading=10.5,
        textColor=colors.HexColor('#334155'),
        leftIndent=12,
        spaceAfter=2
    )

    th_style = ParagraphStyle(
        'PdfTH',
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.white
    )

    td_style = ParagraphStyle(
        'PdfTD',
        fontName='Helvetica',
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor('#0F172A')
    )

    td_code_style = ParagraphStyle(
        'PdfTDCode',
        fontName='Helvetica-Bold',
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor('#1E3A8A')
    )

    story = []

    # Title & Subtitle
    story.append(Paragraph("UnifAI Governance & Plugins Master Deep-Dive Technical Manual", p_title_style))
    story.append(Paragraph("Plugins, Virtual Keys, Users, Teams, Business Units, Customers, SCIM, RBAC, Access Profiles & Audit Logs (Thanglish + English)", p_subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#2563EB'), spaceBefore=2, spaceAfter=8))

    # Architecture Overview
    story.append(Paragraph("1. Governance & Plugins Architecture Map", h1_style))
    story.append(Paragraph(
        "UnifAI Governance and Plugins system forms the multi-tenant enterprise boundary, identity lifecycle, and runtime policy enforcement plane. "
        "Enterprise identities from Okta, Microsoft Entra ID, or Keycloak are synchronized via <b>SCIM v2.0</b> into a multi-tier hierarchy: "
        "<b>Business Units</b> (corporate cost centers) &gt; <b>Teams</b> (departmental squads) &gt; <b>Users</b>. "
        "Access is governed by fine-grained <b>Roles & Permissions (RBAC)</b>. Client applications connect using <b>Virtual Keys</b> "
        "configured with reusable <b>Access Profiles</b> (model allowlists, budget caps, rate limits, guardrails). "
        "Incoming requests trigger the <b>Plugins Engine</b> for custom logic, while every administrative action is immutably recorded in <b>Audit Logs</b>.",
        body_style
    ))
    story.append(Spacer(1, 4))

    # Features Loop
    story.append(Paragraph("2. Detailed Feature Dissection (10 Core Features)", h1_style))
    for f in GOV_FEATURES:
        feat_elements = []
        feat_elements.append(Paragraph(f"<b>{f['name_en']} — {f['name_tanglish']}</b>", h2_style))
        feat_elements.append(Paragraph(f"<b>Route:</b> <font color='#2563EB'>{f['route']}</font>", body_style))

        feat_elements.append(Paragraph("<b>👤 Non-Tech Perspective (Eliya Vilakkam):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Uruvagam (Analogy):</b> {f['analogy']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Vilakkam:</b> {f['explanation']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Business Value:</b> {f['business_value']}", body_style))

        feat_elements.append(Paragraph("<b>💻 Tech Perspective (Technical Architecture):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Backend Architecture:</b> {f['tech_arch']}", body_style))
        ep_text = ", ".join([f"<code>{e}</code>" for e in f['endpoints']])
        feat_elements.append(Paragraph(f"• <b>Endpoints:</b> {ep_text}", body_style))

        feat_elements.append(Paragraph("<b>🖥️ Screen Layout & Interactive Elements (Thirai Koorugal):</b>", h3_style))
        feat_elements.append(Paragraph("<b>1. Melpura Koorugal (Top Bar Controls):</b>", body_style))
        for itm in f['ui_elements']['top_bar']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))
        feat_elements.append(Paragraph("<b>2. Mathiya Koorugal & Tables (Tabs & Views):</b>", body_style))
        for itm in f['ui_elements']['tabs_and_views']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))
        feat_elements.append(Paragraph("<b>3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):</b>", body_style))
        for itm in f['ui_elements']['bottom_elements']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))

        feat_elements.append(Paragraph("<b>🔗 Connections & Structure Map (Inaippugal):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Data-vai Perumidam:</b> {f['connections']['receives_from']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Iyakkum Koorugal:</b> {f['connections']['triggers_and_affects']}", body_style))

        feat_elements.append(Paragraph(f"<b>💡 Production Use Case:</b> {f['use_case']}", body_style))
        feat_elements.append(Spacer(1, 6))
        story.append(KeepTogether(feat_elements))

    # Cross Connections
    story.append(Paragraph("3. Cross-Feature Interconnections & Data Flow", h1_style))
    conn_table_data = [[Paragraph(h, th_style) for h in ["Source Feature", "Connected To", "Data Flow & Trigger Action"]]]
    for src, dst, flow in GOV_CONN_DATA:
        conn_table_data.append([Paragraph(src, td_code_style), Paragraph(dst, td_code_style), Paragraph(flow, td_style)])
    tbl_conn = Table(conn_table_data, colWidths=[105, 105, 320])
    tbl_conn.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_conn)
    story.append(Spacer(1, 8))

    # Comparative Matrix
    story.append(Paragraph("4. Tech vs Non-Tech Comparative Matrix", h1_style))
    mat_table_data = [[Paragraph(h, th_style) for h in ["Feature", "Non-Tech View (Manager / CFO Parvai)", "Tech View (DevOps / Architect Parvai)"]]]
    for ft, nt, tv in GOV_MATRIX_DATA:
        mat_table_data.append([Paragraph(ft, td_code_style), Paragraph(nt, td_style), Paragraph(tv, td_style)])
    tbl_mat = Table(mat_table_data, colWidths=[95, 215, 220])
    tbl_mat.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_mat)

    pdf_doc.build(story, canvasmaker=NumberedCanvas)
    print(f"Governance & Plugins PDF document written to: {PDF_PATH}")

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================
if __name__ == "__main__":
    print("Building dedicated Governance & Plugins Master Documentation (Thanglish + English)...")
    generate_gov_markdown()
    generate_gov_docx()
    generate_gov_pdf()
    print("Governance & Plugins Master Documentation successfully built in:", DOC_DIR)
