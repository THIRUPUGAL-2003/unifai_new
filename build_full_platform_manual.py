import os
import sys

# Ensure UTF-8 output encoding for Windows console
if sys.stdout and hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether, PageBreak, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

WORKSPACE_DIR = r"c:\Users\sakth\OneDrive\ドキュメント\Desktop\unify - Copy"
DOC_DIR = os.path.join(WORKSPACE_DIR, "document")
os.makedirs(DOC_DIR, exist_ok=True)

MD_PATH = os.path.join(DOC_DIR, "UnifAI_Full_Platform_Master_Manual.md")
DOCX_PATH = os.path.join(DOC_DIR, "UnifAI_Full_Platform_Master_Manual.docx")
PDF_PATH = os.path.join(DOC_DIR, "UnifAI_Full_Platform_Master_Manual.pdf")

# ==============================================================================
# EXHAUSTIVE DATA DEFINITION: ALL 40 FEATURES ACROSS 9 PILLARS (THANGLISH + ENGLISH)
# ==============================================================================

PILLARS = [
    {
        "pillar_id": "observability",
        "pillar_name_en": "Observability & Telemetry",
        "pillar_name_tanglish": "Observability (Kangaanippu & Seyalthiran Paguppaayvu)",
        "description": "Company-la nadakkura ellam AI transactions, cost, latency, token spend, matrum error logs-ah real-time-la track panni external monitoring tools (Datadog, Kafka) kooda connect pandra control plane.",
        "features": [
            {
                "id": "obs_dashboard",
                "name_en": "Dashboard",
                "name_tanglish": "Dashboard (Main Metrics & Analytics Cockpit)",
                "route": "/workspace/dashboard",
                "analogy": "Vimanathoda Cockpit allathu luxury car speedometer dashboard mathiri.",
                "explanation": "Car ottum pothu speed, petrol level, engine condition ore meter-la theriyura mathiri, company-oda total AI spend, ethana lakh requests vandhathu, entha AI model fast-ah irukku, entha team athigama spend pandranga nu ore glance-la real-time graphs vazhiya pakkalam.",
                "business_value": "Immediate budget transparency, cost spike prevention, and executive performance reporting.",
                "tech_arch": "WebSocket connection (`useWebSocket`) synchronized with TanStack Query. FastHTTP lock-free atomic counters (sync/atomic) aggregate time-series buckets into PostgreSQL. Computes p50, p90, p95, p99 latencies and Time-To-First-Token (TTFT).",
                "endpoints": ["GET /api/v1/dashboard/stats", "GET /api/v1/dashboard/histogram", "WS /api/v1/ws/dashboard"],
                "ui_elements": {
                    "top_bar": ["Time Range Selector (1h, 24h, 7d, 30d, Custom)", "Timezone Toggle (UTC / Local)", "Export Popover (CSV/JSON/PNG)", "Filter Sidebar Toggle"],
                    "tabs_and_views": ["Overview Tab (Total Spend, Requests, Tokens, p95 Latency)", "Provider Usage Tab", "Model Rankings Tab", "Dimension Rankings (x-uf-dim-*)", "MCP Tool Metrics"],
                    "bottom_elements": ["Line vs Bar Chart Toggle", "Model Performance Leaderboard", "Live WebSocket Status Indicator"]
                },
                "connections": {"receives_from": "Aggregates raw telemetry from LLM Logs and FastHTTP Proxy.", "triggers_and_affects": "Alerts finance managers and provides data for Budget enforcement."},
                "use_case": "CFO weekly review analyzing whether the company is staying within the $20,000 monthly AI budget."
            },
            {
                "id": "obs_llm_logs",
                "name_en": "LLM Logs",
                "name_tanglish": "LLM Logs (AI Request & Response Transaction Ledger)",
                "route": "/workspace/logs",
                "analogy": "Bank passbook statement allathu CCTV footage log book mathiri.",
                "explanation": "Customers and employees AI kitta enna kelvi kettanga, AI enna bathil sonnathu, ethana tokens aachu, evlo dollar bill aachu nu second-by-second capture seiyura complete transaction ledger.",
                "business_value": "100% auditability, prompt debugging, compliance evidence, and customer dispute resolution.",
                "tech_arch": "High-speed batch-insert pipeline into PostgreSQL `llm_logs` table. Stores prompt text, model response, TTFT, token usage, status codes, and virtual key tags with sub-millisecond overhead.",
                "endpoints": ["GET /api/v1/logs", "GET /api/v1/logs/{id}", "POST /api/v1/logs/export"],
                "ui_elements": {
                    "top_bar": ["Search Input (Prompt/Response query)", "Column Visibility Picker", "Export Button (CSV/JSON)", "Auto-Refresh Dropdown"],
                    "tabs_and_views": ["Logs Table: Status Code, Model, Latency ms, Total Tokens, Spend $, Virtual Key, User ID, Timestamp"],
                    "bottom_elements": ["Log Detail Sheet: Formatted Message Tree, Raw Headers, Guardrail Verdicts, Raw JSON Copy Button"]
                },
                "connections": {"receives_from": "Intercepts all calls passing through FastHTTP Proxy Gateway.", "triggers_and_affects": "Feeds live data into Dashboard, Connectors, and Budgets & Limits."},
                "use_case": "Customer support team debugging why an AI agent gave an incorrect shipping estimate to a VIP customer."
            },
            {
                "id": "obs_mcp_logs",
                "name_en": "MCP Logs",
                "name_tanglish": "MCP Logs (Autonomous Agent Tool Call Audit)",
                "route": "/workspace/mcp-logs",
                "analogy": "Factory supervisor duty register mathiri.",
                "explanation": "AI Agent thaanaga mudiveduthu company database-ah paathucha, email anupucha, file-ah delete pannucha nu AI seitha ovvoru external action-aiyum monitor seiyura tool execution audit book.",
                "business_value": "Autonomous agent accountability, security verification, and third-party API call monitoring.",
                "tech_arch": "Model Context Protocol (MCP) JSON-RPC 2.0 frame capture engine. Records tool names, arguments, execution duration, and stdout/stderr outputs across stdio and SSE transports.",
                "endpoints": ["GET /api/v1/mcp/logs", "GET /api/v1/mcp/logs/{id}"],
                "ui_elements": {
                    "top_bar": ["Filter by MCP Server", "Tool Name Search", "Execution Status Filter (Success/Fail/Timeout)"],
                    "tabs_and_views": ["MCP Logs Table: Server Name, Tool Executed, Duration ms, Payload Size KB, Invocation Status"],
                    "bottom_elements": ["Tool Execution Drawer: Input JSON Arguments, Tool Response JSON, Error Stack Trace"]
                },
                "connections": {"receives_from": "Intercepts tool calls executed in MCP Gateway.", "triggers_and_affects": "Feeds tool reliability statistics into MCP Gateway and Dashboard."},
                "use_case": "Verifying that an automated sales agent only searched the CRM database and did not modify any customer records."
            },
            {
                "id": "obs_browser_ai",
                "name_en": "Browser AI",
                "name_tanglish": "Browser AI (Employee Web ChatGPT DLP Proxy)",
                "route": "/workspace/browser-ai",
                "analogy": "Office entrance security guard & bag scanner mathiri.",
                "explanation": "Employees company laptop-la web ChatGPT allathu Claude use pannum pothu, company passwords, client details, allathu confidential code-ah copy-paste panni leak seiyatha mathiri thadukkura security wall.",
                "business_value": "Prevents corporate data leaks to public web AI services without having to ban AI tools for employees.",
                "tech_arch": "Client-side mitmproxy daemon intercepting web traffic to `chatgpt.com` and `claude.ai`. Inspects clipboard paste events and file uploads using local regex and Presidio DLP algorithms.",
                "endpoints": ["GET /api/v1/browser-ai/sessions", "GET /api/v1/browser-ai/violations", "POST /api/v1/browser-ai/sync"],
                "ui_elements": {
                    "top_bar": ["Proxy Status Badge (Active/Inactive)", "Policy Switcher (Block / Redact / Audit)", "Extension Heartbeat"],
                    "tabs_and_views": ["Intercepted Sessions Tab", "DLP Violations Tab", "Attachment Scans Tab"],
                    "bottom_elements": ["Violation Incident Viewer: Diff showing original confidential text vs redacted asterisk text"]
                },
                "connections": {"receives_from": "Employee browser extensions and local proxy daemons.", "triggers_and_affects": "Blocks confidential pastes and logs security violations to Audit Logs."},
                "use_case": "Employee accidentally copying a customer's credit card number into ChatGPT gets warned and the data is masked automatically."
            },
            {
                "id": "obs_connectors",
                "name_en": "Connectors",
                "name_tanglish": "Connectors (External Observability & Telemetry Pipelines)",
                "route": "/workspace/observability",
                "analogy": "Water pipeline network mathiri (Thanniya thevaiyana tank-ku anuppura mathiri).",
                "explanation": "UnifAI-la create aagura AI logs and metrics-ah company-oda main monitoring tools (Datadog, Splunk, Dynatrace, Apache Kafka, Google BigQuery)-ku automatic-ah push seiyura export pipelines.",
                "business_value": "Zero data silos; seamless integration with enterprise security and billing infrastructure.",
                "tech_arch": "Async streaming telemetry push engine. Batches events into configurable queue workers with retry backoff and Dead Letter Queue (DLQ). Pushes via OpenTelemetry (OTel), HTTP webhooks, or Kafka producers.",
                "endpoints": ["GET /api/v1/observability/connectors", "POST /api/v1/observability/connectors", "DELETE /api/v1/observability/connectors/{id}"],
                "ui_elements": {
                    "top_bar": ["Add Connector Button", "Connector Type Filter (Datadog, Kafka, BigQuery, OTel)", "Health Indicator"],
                    "tabs_and_views": ["Connectors List Cards with endpoint URL, flush interval ms, and delivery status badge"],
                    "bottom_elements": ["AddConnectorDialog: Endpoint URL, Bearer Token, Batch Size (100-5000), Send Test Ping Button"]
                },
                "connections": {"receives_from": "Subscribes to FastHTTP Proxy event stream and LLM Logs.", "triggers_and_affects": "Streams enterprise telemetry to corporate BigQuery and Kafka topics."},
                "use_case": "Streaming every AI transaction into corporate BigQuery for monthly automated billing invoices."
            },
            {
                "id": "obs_logs_settings",
                "name_en": "Logs Settings",
                "name_tanglish": "Logs Settings (Log Governance & Retention Policies)",
                "route": "/workspace/config/logging",
                "analogy": "Office record room archival & paper shredding rules mathiri.",
                "explanation": "Pazhaya AI logs-ah ethana naalaikku apram automatic-ah delete pannanum (e.g. 30 days retention), yaarachum PII data anupuna epdi mask pannanum, cold storage S3-ku epdi move pannanum nu set pandra control room.",
                "business_value": "100% GDPR, HIPAA, and SOC-2 data compliance; saves huge database storage costs.",
                "tech_arch": "Log lifecycle engine. Schedules automated PostgreSQL partition drops, triggers `VACUUM` maintenance, manages multi-part uploads to AWS S3 Glacier, and controls traffic sampling rates (1% to 100%).",
                "endpoints": ["GET /api/v1/config/logging", "PUT /api/v1/config/logging", "POST /api/v1/config/logging/purge"],
                "ui_elements": {
                    "top_bar": ["Master Logging Switch", "Data Retention Input (Days)", "Traffic Sampling Slider (0-100%)"],
                    "tabs_and_views": ["Privacy Controls: Mask PII Switch, Store Prompts Toggle, S3 Cold Storage Backup Bucket input"],
                    "bottom_elements": ["Purge Expired Logs Now Button", "Save Logging Settings Button"]
                },
                "connections": {"receives_from": "Administrative configuration inputs.", "triggers_and_affects": "Enforces storage cleanup in PostgreSQL and automated masking in LLM Logs."},
                "use_case": "Setting a strict 30-day retention rule so user chat history is permanently wiped to comply with European privacy laws."
            }
        ]
    },
    {
        "pillar_id": "models",
        "pillar_name_en": "Models & Traffic Control",
        "pillar_name_tanglish": "Models (AI Model Catalog, Routing & Resilience)",
        "route_group": "/workspace/models",
        "description": "Company-oda complete AI models inventory, vendor API keys, smart routing rules, cost triage, and zero-downtime failover systems.",
        "features": [
            {
                "id": "mod_catalog",
                "name_en": "Model Catalog",
                "name_tanglish": "Model Catalog (AI Models Inventory & Menu)",
                "route": "/workspace/model-catalog",
                "analogy": "Amazon product catalog allathu hotel menu card mathiri.",
                "explanation": "Company-la configure panni irukkura ellam AI models (GPT-4o, Claude 3.5 Sonnet, Gemini 1.5, Llama 3) list, athoda context length, 24-hour traffic, total cost ellame ore idathula pakkalam.",
                "business_value": "Eliminates duplicate model usage; standardizes AI models across teams; provides full pricing visibility.",
                "tech_arch": "FastHTTP model registry endpoint (`/api/v1/models/catalog`). TanStack Query client caching with real-time token metrics aggregation. Ingests 24h usage telemetry from PostgreSQL.",
                "endpoints": ["GET /api/v1/models/catalog", "GET /api/v1/models/attributes", "PATCH /api/v1/models/{id}/attributes"],
                "ui_elements": {
                    "top_bar": ["4 Summary Cards (Total Providers, Total Models, Requests 24h, Cost 24h)", "Provider Filter Dropdown", "Overview vs Attributes Tabs"],
                    "tabs_and_views": ["Overview Table: Provider, Models Used, Traffic 24h, Cost 24h", "Attributes Table: Model ID, Display Name, Pricing, Context Window, Output Tokens"],
                    "bottom_elements": ["AttributeSheet: Edit display name, description, context length, modality tags"]
                },
                "connections": {"receives_from": "Model Providers and LLM Logs.", "triggers_and_affects": "Supplies model options to Routing Rules and Complexity Router."},
                "use_case": "Auditing and deprecating older, expensive models like GPT-4 32k in favor of cheaper GPT-4o-mini."
            },
            {
                "id": "mod_providers",
                "name_en": "Model Providers",
                "name_tanglish": "Model Providers (AI Vendor Credentials & Gateway Settings)",
                "route": "/workspace/providers",
                "analogy": "Telecom SIM card & network switchboard mathiri (Airtel, Jio, Vodafone connections).",
                "explanation": "OpenAI, Anthropic, AWS Bedrock, Google Vertex AI, Azure OpenAI, Groq, Ollama aagiya companies-oda connect aaga thevaiyana API keys, secret credentials, network speed, timeout settings configure pandra control panel.",
                "business_value": "Zero vendor lock-in; secure central key vault; support for private on-premise AI models.",
                "tech_arch": "Multi-vendor credential management and FastHTTP connection pool. Configures custom base URLs, proxy tunnels, exponential backoff retries, and raw payload flags. Instant connection ping testing.",
                "endpoints": ["GET /api/v1/providers", "POST /api/v1/providers", "PATCH /api/v1/providers/{name}", "DELETE /api/v1/providers/{name}"],
                "ui_elements": {
                    "top_bar": ["Provider Sidebar with status badges (Active/Warning/Error)", "Add Provider Dropdown (+ Add Custom Provider)"],
                    "tabs_and_views": ["Provider Header: Icon, Name, Edit Provider Config button, Delete button", "ModelProviderKeysTableView: Key Name, Masked Key Secret, Quota Limit, Ping Test button"],
                    "bottom_elements": ["ProviderConfigSheet: Concurrency limits, buffer size, network timeouts, retry delays, proxy settings", "AddCustomProviderSheet: Custom base URL and Keyless mode toggle"]
                },
                "connections": {"receives_from": "Admin inputs.", "triggers_and_affects": "Powers outbound AI traffic in FastHTTP Proxy and feeds Routing Rules."},
                "use_case": "Safely rotating production OpenAI keys and adding a local private Ollama instance for sensitive data."
            },
            {
                "id": "mod_limits",
                "name_en": "Budgets & Limits",
                "name_tanglish": "Budgets & Limits (Cost Caps & Token Rate Limits)",
                "route": "/workspace/model-limits",
                "analogy": "Credit card daily & monthly spend limit mathiri.",
                "explanation": "AI use panna panna bill laatchakanakula pogama irukka, ovvoru team-kum, user-kum, athavathu ovvoru AI model-kum monthly budget ($500) allathu daily token limit set pandra edham. Limit thanduna automatic-ah request stop aagidum.",
                "business_value": "Eliminates bill shock; guarantees 100% budget compliance across departments.",
                "tech_arch": "High-throughput token bucket rate limiter and budget enforcer middleware. Atomic Redis/PostgreSQL counters (`INCRBY`). Pre-flight evaluation rejects requests with HTTP 429 or routes to a cheaper fallback.",
                "endpoints": ["GET /api/v1/governance/model-limits", "POST /api/v1/governance/model-limits", "DELETE /api/v1/governance/model-limits/{id}"],
                "ui_elements": {
                    "top_bar": ["Create Model Limit Button", "Search Bar", "Provider Filter", "Scope Filter (Global, Virtual Key, User, Team)"],
                    "tabs_and_views": ["Model Limits Table: Scope Badge, Target Entity, Provider, Model, Budget Progress Bar $, Token Progress Bar, Reset Duration"],
                    "bottom_elements": ["ModelLimitSheet: Scope Selector, Target pickers, Max Budget $, Max Tokens, Reset Duration, Threshold Alert Rules (80% warning, 100% hard block)"]
                },
                "connections": {"receives_from": "Rates from Pricing Overrides and live spend from LLM Logs.", "triggers_and_affects": "Blocks or reroutes calls in FastHTTP Proxy when quota is exhausted."},
                "use_case": "Giving intern developers a strict $50/month cap while granting the production app an elastic $5,000 budget."
            },
            {
                "id": "mod_routing",
                "name_en": "Routing Rules",
                "name_tanglish": "Routing Rules (Intelligent Traffic Steering & Fallbacks)",
                "route": "/workspace/routing-rules",
                "analogy": "Railway track switcher mathiri (Train-ah thevaikku etha track-la thiruppura mathiri).",
                "explanation": "Oru customer request varum pothu, athu entha AI model-ku poganum nu rules poduvom. Example: Premium user-ku GPT-4o, Free user-ku Claude 3.5 Haiku nu condition base panni smart-ah route pandrathu.",
                "business_value": "Custom user experience; cost optimization based on customer value; zero downtime via fallback chains.",
                "tech_arch": "CEL (Common Expression Language) evaluation engine written in Go. Evaluates conditions against request headers (`x-uf-*`), prompt metadata, user roles, and token counts to dispatch to prioritized target pools.",
                "endpoints": ["GET /api/v1/routing-rules", "POST /api/v1/routing-rules", "PATCH /api/v1/routing-rules/{id}", "DELETE /api/v1/routing-rules/{id}"],
                "ui_elements": {
                    "top_bar": ["Create Routing Rule Button", "Search Input", "Priority Sort"],
                    "tabs_and_views": ["Routing Rules Table: Rule Name, Priority Badge, CEL Condition, Target Model/Provider, Enabled Switch, Actions", "Visual Rule Tree View"],
                    "bottom_elements": ["RoutingRuleSheet: Rule Name, Priority slider (1-100), Visual/Raw CEL Condition Builder, Target Provider/Model, Fallback Priority list"]
                },
                "connections": {"receives_from": "FastHTTP Proxy requests and Complexity Router tiers.", "triggers_and_affects": "Directs traffic to upstream Model Providers and triggers Circuit Breaker on failure."},
                "use_case": "Routing all European requests to EU-hosted Azure OpenAI instances for strict GDPR data residency."
            },
            {
                "id": "mod_complexity",
                "name_en": "Complexity Router",
                "name_tanglish": "Complexity Router (AI Query Triage & Cost Optimizer)",
                "route": "/workspace/complexity-router",
                "analogy": "Hospital Triage System mathiri (Fever-ku junior doctor, surgery-ku specialist mathiri).",
                "explanation": "Simple-ana kelvi ('Hi', 'What is Python?') ketta costly GPT-4o thevailla, athukku fast & cheap-ana GPT-4o-mini pothum. Aana periya complex coding allathu math problem ketta mattum o1 allathu Claude Opus-ku automatic-ah anupum. Ithanaala 70% AI bill micham aagum!",
                "business_value": "Saves up to 70% in monthly AI costs with zero compromise on quality.",
                "tech_arch": "Multi-factor prompt complexity analysis engine. Computes a continuous complexity score (0.0 to 1.0) using token length heuristics, vocabulary entropy, code regexes (`def`, `class`, `function`), and keyword weighting. Maps into 4 tiers: `SIMPLE`, `MEDIUM`, `COMPLEX`, `REASONING`.",
                "endpoints": ["GET /api/v1/governance/complexity-analyzer", "PUT /api/v1/governance/complexity-analyzer", "POST /api/v1/governance/complexity-analyzer/reset"],
                "ui_elements": {
                    "top_bar": ["Save Configuration Button", "Reset to Defaults Button (RotateCcw)", "Docs Link"],
                    "tabs_and_views": ["Progressive okLCH Palette Bar (Simple, Medium, Complex, Reasoning)", "Tier Boundary Inputs (Simple->Med, Med->Comp, Comp->Reason)", "Keyword Lists (TagInput for Code, Reasoning, Simple)"],
                    "bottom_elements": ["Tier Target Model Mappings (Assign default model per tier)", "Interactive Prompt Test Sandbox"]
                },
                "connections": {"receives_from": "Intercepts incoming requests in FastHTTP Proxy before Routing Rules.", "triggers_and_affects": "Selects appropriate model tier and logs classification in LLM Logs."},
                "use_case": "Customer support bot where 85% routine queries go to lightweight models and 15% complex technical queries go to reasoning models."
            },
            {
                "id": "mod_circuit_breaker",
                "name_en": "Circuit Breaker",
                "name_tanglish": "Circuit Breaker (Outage Protection & Auto-Failover)",
                "route": "/workspace/circuit-breaker",
                "analogy": "Veetla irukkura Electric MCB / Fuse mathiri.",
                "explanation": "OpenAI allathu Anthropic server crash aanaallathu rate limit aagi error vantha, user-ku error screen kaatama, automatic-ah fraction of second-la backup model-ku (e.g. AWS Bedrock allathu Azure)-ku switch panni 24x7 non-stop-ah vela seiya vaikkum.",
                "business_value": "99.99% Enterprise High Availability; zero customer-facing downtime during AI vendor outages.",
                "tech_arch": "Distributed Circuit Breaker finite state machine (Closed -> Open -> Half-Open). Monitored via real-time polling every 8 seconds (`pollingInterval: 8000ms`). Detects response signals (HTTP 429/503, `x-ratelimit-remaining: 0`). Instantly redirects calls to fallback provider with automated cooldown probing.",
                "endpoints": ["GET /api/v1/circuit-breaker/policies", "POST /api/v1/circuit-breaker/policies", "PUT /api/v1/circuit-breaker/policies/{name}", "GET /api/v1/circuit-breaker/state", "POST /api/v1/circuit-breaker/policies/{name}/reset"],
                "ui_elements": {
                    "top_bar": ["Header with Shield Icon", "Create Policy Button"],
                    "tabs_and_views": ["Circuit Breaker Table: Policy Name, Primary Provider/Model, Fallback Provider/Model, Trigger Condition, Cooldown, State Badge (Closed/Open/Half-Open), Actions"],
                    "bottom_elements": ["Circuit Policy Dialog: Policy Name, Primary/Fallback comboboxes, Response Header Signal matchers, Cooldown duration", "Manual Circuit Reset Button (RotateCcw)"]
                },
                "connections": {"receives_from": "Monitors HTTP response headers and errors from FastHTTP Proxy.", "triggers_and_affects": "Reroutes failed requests and notifies Alert Channels."},
                "use_case": "During peak hours, when OpenAI returns HTTP 429 Too Many Requests, traffic automatically pivots to AWS Bedrock with zero disruption."
            },
            {
                "id": "mod_pricing_overrides",
                "name_en": "Pricing Overrides",
                "name_tanglish": "Pricing Overrides (Custom Enterprise Rates & Discounts)",
                "route": "/workspace/custom-pricing/overrides",
                "analogy": "Special corporate discount allathu wholesale negotiated price contract mathiri.",
                "explanation": "OpenAI website-la 1 million token-ku $5 nu irunthaalum, unga company-ku special enterprise discount allathu committed spend agreement iruntha, antha custom rate-ah inga enter pannikalam. Athuku thagundha mathiri unga internal dashboard-la exact real spend calculate aagum.",
                "business_value": "100% accurate financial chargeback and departmental billing based on actual negotiated contracts.",
                "tech_arch": "Multi-tier hierarchical pricing precedence engine: (1) Virtual Key Override -> (2) User/Team Override -> (3) Provider Global Override -> (4) Default Provider Catalog. Calculates prompt tokens, cached input tokens, output tokens, reasoning tokens, and fixed invocation fees.",
                "endpoints": ["GET /api/v1/pricing/overrides", "POST /api/v1/pricing/overrides", "PUT /api/v1/pricing/overrides/{id}", "DELETE /api/v1/pricing/overrides/{id}"],
                "ui_elements": {
                    "top_bar": ["Add Pricing Override Button", "Scope Filter (Global, Virtual Key, Team)", "Search Input"],
                    "tabs_and_views": ["Pricing Overrides Table: Scope Badge, Target Entity, Provider, Model, Input Cost $/1M, Cached Input Cost, Output Cost $/1M, Request Fee, Actions"],
                    "bottom_elements": ["PricingOverrideSheet: Scope, Provider, Model, and PricingFieldSelector inputs for granular token pricing"]
                },
                "connections": {"receives_from": "Finance administrators.", "triggers_and_affects": "Feeds real-time spend calculations into Dashboard, LLM Logs, and Budgets & Limits."},
                "use_case": "Applying an AWS 20% committed spend discount to Bedrock Claude models so internal reporting matches cloud invoices."
            },
            {
                "id": "mod_settings",
                "name_en": "Model Settings",
                "name_tanglish": "Model Settings (Master Catalog Sync & Recursion Guard)",
                "route": "/workspace/custom-pricing",
                "analogy": "Global system clock & master registry synchronizer mathiri.",
                "explanation": "Market-la puthu puthu AI models varum pothu, athoda official pricing automatic-ah download aagi update aaga thevaiyana master website link, sync frequency (24 hours), mathum routing rules-la infinite loop vanthu server hang aagidama thadukkura max depth settings configure pandra edham.",
                "business_value": "Zero-maintenance operations; automated price drop ingestion without code deployments.",
                "tech_arch": "Core configuration manager in PostgreSQL (`framework_config`, `client_config`). Orchestrates asynchronous cron workers to fetch remote pricing datasheets (CSV/JSON). Enforces `routing_chain_max_depth` to prevent circular fallback loops.",
                "endpoints": ["GET /api/v1/config/core", "PUT /api/v1/config/core", "POST /api/v1/config/pricing/sync"],
                "ui_elements": {
                    "top_bar": ["Header Title", "Dirty State Tracker highlighting unsaved changes"],
                    "tabs_and_views": ["ModelSettingsView: Pricing Datasheet URL, Pricing Sync Interval (Hours), Model Parameters URL, Routing Chain Max Depth (slider/input)"],
                    "bottom_elements": ["Force Pricing Sync Now Button (with loading spinner)", "Save Configuration Button"]
                },
                "connections": {"receives_from": "Admin inputs and remote pricing feeds.", "triggers_and_affects": "Updates foundational pricing for Model Catalog and prevents routing recursion loops."},
                "use_case": "Instantly ingesting OpenAI's latest price reduction across 30 enterprise models with a single click of 'Force Pricing Sync Now'."
            }
        ]
    },
    {
        "pillar_id": "mcp_gateway",
        "pillar_name_en": "MCP Gateway & Tool Execution",
        "pillar_name_tanglish": "MCP Gateway (Model Context Protocol & Agent Tools)",
        "route_group": "/workspace/mcp",
        "description": "Autonomous AI Agents company-oda internal databases, APIs, codebases, and external SaaS tools kooda interact panna thevaiyana standardized gateway.",
        "features": [
            {
                "id": "mcp_catalog",
                "name_en": "MCP Catalog",
                "name_tanglish": "MCP Catalog (Tool Registry & App Store)",
                "route": "/workspace/mcp-registry",
                "analogy": "App Store allathu Google Play Store mathiri.",
                "explanation": "AI Agents use panna koodiya ellam certified tools (GitHub MCP, PostgreSQL MCP, Slack MCP, Google Drive MCP, Puppeteer Web Scraper) list irukkura marketplace. Single click-la puthu tool install pannikalam.",
                "business_value": "Fast-tracks AI agent deployment; verified secure tool marketplace; zero boilerplate integration.",
                "tech_arch": "Curated MCP registry client. Fetches certified tool schemas, version tags, environment variable templates, and execution commands (stdio binary or Docker container) from central repositories.",
                "endpoints": ["GET /api/v1/mcp/registry", "POST /api/v1/mcp/registry/install"],
                "ui_elements": {
                    "top_bar": ["Search Catalog Input", "Category Filter (Databases, Version Control, Communication, DevTools)", "Installed vs Available Filter"],
                    "tabs_and_views": ["Tool Cards Grid: Icon, Tool Name, Description, Verified Badge, Version, 'Install' button"],
                    "bottom_elements": ["InstallToolDialog: Environment variable inputs (API keys, connection strings) and transport mode selector"]
                },
                "connections": {"receives_from": "External MCP package repositories.", "triggers_and_affects": "Deploys tools into Installed Servers."},
                "use_case": "Installing the PostgreSQL MCP tool to let internal analytics AI query corporate sales data safely."
            },
            {
                "id": "mcp_servers",
                "name_en": "Installed Servers",
                "name_tanglish": "Installed Servers (Active MCP Instances & Health)",
                "route": "/workspace/mcp-gateway",
                "analogy": "Computer-la install aana software apps dashboard mathiri.",
                "explanation": "Company-kulla ippo active-ah run aagikittu irukura MCP tool servers list. Entha server online-la irukku, ethu crashed, ethana tools antha server tharuthu nu paathu restart allathu configure pandra idham.",
                "business_value": "Tool availability monitoring; instant recovery from crashed tool daemons; lifecycle control.",
                "tech_arch": "MCP Server lifecycle daemon supervisor. Manages child OS processes (`stdio`) or HTTP Server-Sent Events (`SSE`) endpoints. Performs automatic process health checks, restarts crashed servers, and registers exported tool schemas.",
                "endpoints": ["GET /api/v1/mcp/servers", "POST /api/v1/mcp/servers/{id}/restart", "DELETE /api/v1/mcp/servers/{id}"],
                "ui_elements": {
                    "top_bar": ["Add Server Button", "Search Servers", "Health Filter (Running, Degraded, Stopped)"],
                    "tabs_and_views": ["Installed Servers Table: Server Name, Transport (stdio/SSE), Status Badge, Tool Count, Memory Usage, Actions (Restart, Edit, Delete)"],
                    "bottom_elements": ["ServerConfigSheet: Command arguments, working directory, environment variables, timeout limits"]
                },
                "connections": {"receives_from": "MCP Catalog installations.", "triggers_and_affects": "Exposes tools to Tool Groups and logs calls to MCP Logs."},
                "use_case": "Restarting the GitHub MCP server daemon after generating a new personal access token."
            },
            {
                "id": "mcp_sessions",
                "name_en": "MCP Sessions",
                "name_tanglish": "Sessions (Live Stateful Agent Connections)",
                "route": "/workspace/mcp-sessions",
                "analogy": "Active phone calls allathu live chat sessions list mathiri.",
                "explanation": "Ippo entha entha AI agents live-ah company tools kooda connect panni vela senjukittu irukku nu kaatara active connection monitor. Thevaillatha session-ah force-ah terminate pannalam.",
                "business_value": "Prevents runaway autonomous agents; immediate containment of malfunctioning automation scripts.",
                "tech_arch": "Stateful session tracker for SSE and WebSocket MCP connections. Tracks session ID, client IP, connected agent identity, active context windows, and tool execution lock state.",
                "endpoints": ["GET /api/v1/mcp/sessions", "DELETE /api/v1/mcp/sessions/{id}"],
                "ui_elements": {
                    "top_bar": ["Active Sessions Count Badge", "Kill All Sessions Emergency Button"],
                    "tabs_and_views": ["Sessions Table: Session ID, Connected Agent/User, Duration, Tools Invoked Count, Memory State, Actions (Terminate)"],
                    "bottom_elements": ["Session Trace Modal: Live stream of JSON-RPC tool frames exchanged in this session"]
                },
                "connections": {"receives_from": "Live incoming agent connections.", "triggers_and_affects": "Can terminate active sessions to protect downstream infrastructure."},
                "use_case": "Killing an agent session that entered an infinite loop repeatedly calling the Slack notification tool."
            },
            {
                "id": "mcp_tool_groups",
                "name_en": "Tool Groups",
                "name_tanglish": "Tool Groups (Permission Bundles & Tool Access Control)",
                "route": "/workspace/mcp-tool-groups",
                "analogy": "Security clearance badge groups mathiri (Junior dev-ku read-only tools, Senior dev-ku write tools).",
                "explanation": "Multiple tools-ah onna sethu bundle pannalam. Example: 'Read-Only Tools' bundle (view database, read files) allathu 'Production Tools' bundle (write database, send email). Virtual key-ku etha mathiri intha groups-ah assign pannikalam.",
                "business_value": "Least-privilege security for autonomous AI; prevents AI from executing dangerous destructive commands.",
                "tech_arch": "Access control abstraction layer over MCP tool schemas. Maps Virtual Keys to permitted Tool Groups. Gateway filters out unpermitted tool definitions from the LLM prompt's system message dynamically.",
                "endpoints": ["GET /api/v1/mcp/tool-groups", "POST /api/v1/mcp/tool-groups", "PATCH /api/v1/mcp/tool-groups/{id}"],
                "ui_elements": {
                    "top_bar": ["Create Tool Group Button", "Search Groups"],
                    "tabs_and_views": ["Tool Groups Cards: Group Name, Description, Included Tools count, Assigned Virtual Keys count, Actions"],
                    "bottom_elements": ["ToolGroupSheet: Group Name, Description, Interactive Tool Selector Checkboxes across all installed servers"]
                },
                "connections": {"receives_from": "Installed Servers tools.", "triggers_and_affects": "Binds to Virtual Keys and Access Profiles."},
                "use_case": "Restricting an intern's chatbot to only search internal documentation while preventing it from modifying the production database."
            },
            {
                "id": "mcp_auth_config",
                "name_en": "Auth Config",
                "name_tanglish": "Auth Config (MCP Security & Client Tokens)",
                "route": "/workspace/mcp-auth-config",
                "analogy": "Security passport & digital signature verification mathiri.",
                "explanation": "External agents and client software MCP Gateway-kulla connect aaga thevaiyana OAuth 2.0, Bearer token, and mutual TLS authentication settings configure pandra edham.",
                "business_value": "Prevents unauthorized agents from connecting to internal enterprise tools.",
                "tech_arch": "MCP Gateway authentication middleware. Validates incoming connection credentials against OAuth 2.0 introspection endpoints, API keys, or mTLS client certificates before allowing tool execution.",
                "endpoints": ["GET /api/v1/mcp/auth-config", "PUT /api/v1/mcp/auth-config"],
                "ui_elements": {
                    "top_bar": ["Master Auth Enforcement Toggle", "Save Auth Config Button"],
                    "tabs_and_views": ["Auth Settings Form: Allowed Auth Types (Bearer, mTLS, OAuth2), Token Expiration Duration, Whitelisted Client IPs"],
                    "bottom_elements": ["Client Token Generator: Create dedicated secure connection tokens for trusted agent processes"]
                },
                "connections": {"receives_from": "Security administrators.", "triggers_and_affects": "Authorizes or rejects incoming connections to MCP Gateway."},
                "use_case": "Requiring mutual TLS certificates for all agent connections originating from external cloud environments."
            },
            {
                "id": "mcp_settings",
                "name_en": "Gateway Settings",
                "name_tanglish": "Gateway Settings (MCP Engine Timeouts & Network Ports)",
                "route": "/workspace/mcp-settings",
                "analogy": "Factory master electricity & network switch settings mathiri.",
                "explanation": "MCP Gateway-oda core engine settings: Oru tool call maximum ethana seconds run aagalam (timeout), maximum payload size evlo irukkalam, background worker threads ethana nu configure pandra control panel.",
                "business_value": "Prevents slow tools from hanging server resources; guarantees gateway stability.",
                "tech_arch": "Core runtime configuration for the MCP Gateway subsystem. Manages maximum concurrent sessions, global execution timeouts, SSE heartbeat intervals, and payload size limits.",
                "endpoints": ["GET /api/v1/mcp/settings", "PUT /api/v1/mcp/settings"],
                "ui_elements": {
                    "top_bar": ["Header Title", "Save Settings Button"],
                    "tabs_and_views": ["Settings Form: Global Tool Timeout ms, Max Payload Size MB, Max Concurrent Sessions, SSE Heartbeat Seconds"],
                    "bottom_elements": ["Reset to Recommended Defaults Button"]
                },
                "connections": {"receives_from": "System Administrators.", "triggers_and_affects": "Governs the execution environment of all Installed Servers and Sessions."},
                "use_case": "Setting a strict 15-second timeout on all tool calls so a slow database query doesn't freeze the AI user experience."
            }
        ]
    },
    {
        "pillar_id": "prompt_management",
        "pillar_name_en": "Prompt Management & Playground",
        "pillar_name_tanglish": "Prompt Management (Prompt Repo, Skills & Chat Playground)",
        "route_group": "/workspace/prompts",
        "description": "Enterprise prompt templates version control, reusable agent skills, and multi-model testing playground.",
        "features": [
            {
                "id": "prompt_repo",
                "name_en": "Prompt Repo",
                "name_tanglish": "Prompt Repo (Version Controlled Prompt Templates)",
                "route": "/workspace/prompt-repo",
                "analogy": "Software code-kaga irukkura GitHub repository mathiri, prompts-kaga irukkura Git repository.",
                "explanation": "Developers application-la use pandra prompt templates-ah hardcode pannama, inga central-ah store panni version control (v1, v2, v3) pannalam. Code redeploy pannama prompt-ah change panni A/B test pannalam.",
                "business_value": "Decouples prompt engineering from backend deployments; enables non-technical domain experts to optimize prompts safely.",
                "tech_arch": "Prompt template versioning engine with variable interpolation (`{{user_name}}`, `{{account_id}}`). Supports commit messages, semantic diffing, rollback to previous versions, and production release tagging.",
                "endpoints": ["GET /api/v1/prompts", "POST /api/v1/prompts", "GET /api/v1/prompts/{id}/versions", "POST /api/v1/prompts/{id}/deploy"],
                "ui_elements": {
                    "top_bar": ["Create Prompt Button", "Search Prompts", "Tag Filter (Sales, Support, Code, Legal)"],
                    "tabs_and_views": ["Prompts Table: Prompt Name, Latest Version, Production Tag, Variables count, Last Updated, Actions", "Visual Diff Viewer: Side-by-side comparison of v1 vs v2"],
                    "bottom_elements": ["PromptEditorSheet: Template textarea, Variable detector, Model parameter settings (temperature, top_p)"]
                },
                "connections": {"receives_from": "Prompt engineers and developers.", "triggers_and_affects": "Supplies production templates to client apps via FastHTTP Proxy and Guardrail Rules."},
                "use_case": "Updating the company's customer support prompt to include holiday return policies instantly without touching backend server code."
            },
            {
                "id": "skills_repo",
                "name_en": "Skills Repo",
                "name_tanglish": "Skills Repo (Reusable Agent Competencies & Knowledge Packs)",
                "route": "/workspace/skills-repo",
                "analogy": "Specialized employee training certificates & skills handbook mathiri.",
                "explanation": "AI Agent seiya koodiya complex multi-step workflows-ah 'Skill'-ah define panni store pannalam (e.g. 'Generate SQL Query', 'Summarize Financial PDF', 'Draft Legal Contract'). Entha agent venumnaalum intha skill-ah borrow pannikalam.",
                "business_value": "Massive reuse of engineered AI workflows across multiple company applications.",
                "tech_arch": "Modular skill definition registry. Packages system instructions, few-shot demonstration examples, required MCP tool bindings, and validation schemas into an executable skill bundle.",
                "endpoints": ["GET /api/v1/skills", "POST /api/v1/skills", "PATCH /api/v1/skills/{id}"],
                "ui_elements": {
                    "top_bar": ["Create Skill Button", "Search Skills", "Category Filter"],
                    "tabs_and_views": ["Skills Grid: Skill Name, Version, Attached Tools count, Few-Shot Examples count, Actions Menu"],
                    "bottom_elements": ["SkillBuilderSheet: System instructions, Required tool bindings, Demonstration input/output pairs"]
                },
                "connections": {"receives_from": "AI Engineers.", "triggers_and_affects": "Enriches agent context during chat and automated agent workflows."},
                "use_case": "Creating a standard 'Summarize Earnings Call' skill that automatically invokes financial tools and formats outputs consistently."
            },
            {
                "id": "chat_playground",
                "name_en": "Chat / Playground",
                "name_tanglish": "Playground (Multi-Model Testing & Comparison Sandbox)",
                "route": "/workspace/chat",
                "analogy": "Test drive track mathiri (Car vaangurathukku munnadi test drive panni pakkura mathiri).",
                "explanation": "Different AI models-ah (GPT-4o vs Claude 3.5 Sonnet vs Gemini 1.5) side-by-side compare panni, speed, answer quality, and cost-ah live-ah test panni pakkura interactive testing playground.",
                "business_value": "Fast experimentation; empirical model evaluation before committing to production deployments.",
                "tech_arch": "Interactive multi-model streaming sandbox. Dispatches concurrent SSE requests to multiple providers simultaneously. Displays real-time Time-To-First-Token (TTFT), tokens per second (TPS), and dollar cost per response.",
                "endpoints": ["POST /api/v1/chat/completions", "GET /api/v1/models"],
                "ui_elements": {
                    "top_bar": ["Model Selector Dropdowns (Multi-Model Split Screen)", "Temperature & Max Tokens Sliders", "System Prompt Input"],
                    "tabs_and_views": ["Split-Screen Chat View: Compare Model A vs Model B side-by-side in real time", "Live Latency & Cost Counter under each bubble"],
                    "bottom_elements": ["Export Chat to Prompt Repo Button", "Parameter Configuration Sheet"]
                },
                "connections": {"receives_from": "Direct user prompts and templates from Prompt Repo.", "triggers_and_affects": "Executes calls across Model Providers and displays telemetry."},
                "use_case": "Testing whether Claude 3.5 Sonnet writes better SQL queries than GPT-4o on proprietary company database schemas."
            }
        ]
    },
    {
        "pillar_id": "plugins",
        "pillar_name_en": "Plugins & Extensibility",
        "pillar_name_tanglish": "Plugins (Gateway Extensions & Custom Logic)",
        "route_group": "/workspace/plugins",
        "description": "Gateway functionality-ah expand panna custom logic, billing webhooks, and proprietary headers inject pandra extensible engine.",
        "features": [
            {
                "id": "plg_plugins",
                "name_en": "Plugins",
                "name_tanglish": "Plugins (Gateway Extensions & Lifecycle Hooks)",
                "route": "/workspace/plugins",
                "analogy": "Chrome extensions allathu WordPress plugins mathiri.",
                "explanation": "Gateway core code-ah maathaama, ungalukku thevaiyana custom rules (e.g. proprietary token counter, custom headers, billing webhooks) add panna thevaiyana extensible plugin store.",
                "business_value": "Zero friction integration of custom enterprise logic and legal policies directly into the proxy request pipeline.",
                "tech_arch": "WebAssembly (Wasm) and Go-based lifecycle interceptor pipeline. Hooks into `pre-request`, `post-request`, `response-streaming`, and `error-handling` execution phases with ordered priority.",
                "endpoints": ["GET /api/v1/plugins", "POST /api/v1/plugins", "PUT /api/v1/plugins/sequence", "DELETE /api/v1/plugins/{id}"],
                "ui_elements": {
                    "top_bar": ["Add New Plugin Button", "Plugin Sequence Button (ListOrdered)"],
                    "tabs_and_views": ["Plugins Sidebar with status badges", "PluginsView: Name, Description, Version, Schema, Active Toggle"],
                    "bottom_elements": ["AddNewPluginSheet: Name, Code/URL, Hook selectors", "PluginSequenceSheet: Visual drag-and-drop execution ordering"]
                },
                "connections": {"receives_from": "FastHTTP Proxy raw request.", "triggers_and_affects": "Enriches headers, mutates payloads, and logs traces to LLM Logs."},
                "use_case": "Injecting an enterprise compliance plugin that validates legal disclaimers on every financial advice prompt."
            }
        ]
    },
    {
        "pillar_id": "governance",
        "pillar_name_en": "Governance & Identity",
        "pillar_name_tanglish": "Governance (Identity, RBAC, Virtual Keys & Security)",
        "route_group": "/workspace/governance",
        "description": "Multi-tenant enterprise hierarchy, synthetic Virtual Keys, automated SCIM provisioning, RBAC permissions, and immutable audit trails.",
        "features": [
            {
                "id": "gov_virtual_keys",
                "name_en": "Virtual Keys",
                "name_tanglish": "Virtual Keys (Proxy API Keys & Granular Quotas)",
                "route": "/workspace/governance/virtual-keys",
                "analogy": "Bank master account-ku sub-debit cards mathiri (ovvoru card-kum thani limit).",
                "explanation": "Master OpenAI keys-ah developers kitta direct-ah tharama, dummy Virtual Key create panni tharuvom. Monthly $100 limit, specific models permission potruppom. Key leak aanaalum single click-la revoke pannikalam, real keys safe-ah irukkum.",
                "business_value": "Zero credential leakage risk; 100% cost attribution by application; instant security revocation.",
                "tech_arch": "High-throughput cryptographic API key gateway middleware. Keys are generated with cryptographically secure tokens (`uf-key-...`) and stored as SHA-256 hashes. FastHTTP middleware validates tokens in < 1ms, checks budget balances, and enforces RPM/TPM limits.",
                "endpoints": ["GET /api/v1/governance/virtual-keys", "POST /api/v1/governance/virtual-keys", "POST /api/v1/governance/virtual-keys/rotate"],
                "ui_elements": {
                    "top_bar": ["Create Virtual Key Button", "Bulk Rotate Keys Button", "Search & Filter Bar"],
                    "tabs_and_views": ["Virtual Keys Table: Name, Masked Secret (copy/eye toggle), Team, Customer, BudgetDisplay progress bar, RateLimitDisplay badge, Status switch, Actions"],
                    "bottom_elements": ["VirtualKeySheet: Team, Access Profile, Budget Cap $, Rate Limits, Model Allowlist, Expiration Date", "One-Time Secret Reveal Modal"]
                },
                "connections": {"receives_from": "Teams, Customers, and Access Profiles.", "triggers_and_affects": "Authorizes incoming requests in FastHTTP Proxy and tracks spend into LLM Logs."},
                "use_case": "Issuing a virtual key with a $150/month cap and Gemini-only access to an external vendor development squad."
            },
            {
                "id": "gov_users",
                "name_en": "Users",
                "name_tanglish": "Users (Platform User Directory & Identity Management)",
                "route": "/workspace/governance/users",
                "analogy": "Company employee directory & staff ID register mathiri.",
                "explanation": "UnifAI platform-la account irukkura ellam developers, team leads, finance managers, and admins list. Yaar yaar enna email, enna role, entha team-la irukaanga nu manage pandra central user directory.",
                "business_value": "Centralized identity governance; seamless onboarding and automated access termination when staff leave.",
                "tech_arch": "User identity and authentication manager integrated with enterprise SSO (SAML 2.0 / OIDC) and automated SCIM provisioning. Sessions are signed via JWTs with mandatory MFA support.",
                "endpoints": ["GET /api/v1/governance/users", "POST /api/v1/governance/users/invite", "PATCH /api/v1/governance/users/{id}"],
                "ui_elements": {
                    "top_bar": ["Invite User Button", "Search Users", "Filter by Role / Status"],
                    "tabs_and_views": ["Users Table: Avatar, Name, Email, Assigned RBAC Role badge, Team Memberships, Last Active, Status badge, Actions"],
                    "bottom_elements": ["InviteUserDialog: Email, Full Name, Initial Role, Assigned Teams", "EditUserSheet: Role adjust, Suspend user toggle"]
                },
                "connections": {"receives_from": "SCIM / SSO or Admin invites.", "triggers_and_affects": "Binds to Teams, Virtual Keys, and RBAC roles; actions logged to Audit Logs."},
                "use_case": "Inviting 25 software engineers to the platform with Member-level permissions assigned to the Search Squad."
            },
            {
                "id": "gov_teams",
                "name_en": "Teams",
                "name_tanglish": "Teams (Departmental Squads & Shared Resource Pools)",
                "route": "/workspace/governance/teams",
                "analogy": "Company project squads allathu departments mathiri (e.g. Mobile Squad, Search Team).",
                "explanation": "Developers-ah group panni team-ah pirikalam. Oru team-ku common monthly budget, shared virtual keys, and collaborative model access kudukalam.",
                "business_value": "Departmental accountability; team-level AI spend tracking for CFO budget reviews.",
                "tech_arch": "Multi-tenant team resource isolation layer. Aggregates usage across all member users and attached virtual keys. Enforces team-level budget caps in PostgreSQL.",
                "endpoints": ["GET /api/v1/governance/teams", "POST /api/v1/governance/teams", "PATCH /api/v1/governance/teams/{id}"],
                "ui_elements": {
                    "top_bar": ["Create Team Button", "Search Teams", "Business Unit Filter"],
                    "tabs_and_views": ["Teams Table: Team Name, Description, Members count badge, Virtual Keys count, Monthly Budget Progress Bar, Parent Business Unit"],
                    "bottom_elements": ["CreateTeamDialog: Name, Description, Parent Business Unit, Monthly Budget $, Team Lead", "ManageMembersSheet"]
                },
                "connections": {"receives_from": "Belongs to a Business Unit; groups Users.", "triggers_and_affects": "Owns shared Virtual Keys; feeds spend into Dashboard."},
                "use_case": "Allocating a shared $4,000 monthly budget to the 'Mobile App Team' with 15 engineers."
            },
            {
                "id": "gov_business_units",
                "name_en": "Business Units",
                "name_tanglish": "Business Units (Enterprise Divisions & P&L Centers)",
                "route": "/workspace/governance/business-units",
                "analogy": "Corporate conglomerate divisions mathiri (Retail Division, Banking Division).",
                "explanation": "Teams-ku mela irukura periya enterprise division. CFO and executive leadership company-oda multiple business units-kulla AI budget分配 pannavum, P&L tracking seiyavum use aagum.",
                "business_value": "High-level corporate financial governance; automated chargebacks across major business divisions.",
                "tech_arch": "Top-level container in the multi-tenant governance tree (`BusinessUnit -> Teams -> Users -> Virtual Keys`). Aggregates financial consumption for corporate ERP systems (SAP, NetSuite).",
                "endpoints": ["GET /api/v1/governance/business-units", "POST /api/v1/governance/business-units"],
                "ui_elements": {
                    "top_bar": ["Create Business Unit Button", "Search Divisions"],
                    "tabs_and_views": ["Business Units Table: Unit Name, Cost Code, Total Sub-Teams count, Aggregate Spend $, Head of Division"],
                    "bottom_elements": ["BusinessUnitSheet: Unit Name, Cost Center Code, Owner Email, Quarterly Budget Cap $"]
                },
                "connections": {"receives_from": "Executive management.", "triggers_and_affects": "Contains multiple Teams; feeds top-level cost analytics into Executive Dashboards."},
                "use_case": "Tracking that the 'Healthcare Services' division consumed $60,000 in AI tokens across its 10 engineering teams."
            },
            {
                "id": "gov_customers",
                "name_en": "Customers",
                "name_tanglish": "Customers (B2B Client Tenants & External Accounts)",
                "route": "/workspace/governance/customers",
                "analogy": "SaaS B2B client accounts mathiri.",
                "explanation": "Neenga unga software-ah vera external clients-ku AI features vachu SaaS product-ah vikkiringa na, ovvoru client-aiyum 'Customer'-ah add panni avanga usage-ku bill podalam. Oru customer innoru customer-oda data-vai paarka mudiyathu.",
                "business_value": "Monetize AI applications; seamless B2B client isolation; automated billing invoices based on client consumption.",
                "tech_arch": "External tenant isolation boundary. Associates virtual keys with external customer identifiers (`customer_id`). Enforces hard tenant isolation policies in PostgreSQL using Row-Level Security (RLS).",
                "endpoints": ["GET /api/v1/governance/customers", "POST /api/v1/governance/customers"],
                "ui_elements": {
                    "top_bar": ["Create Customer Button", "Search & Filter Bar"],
                    "tabs_and_views": ["Customers Table: Customer Name, Status badge, Assigned Teams, Virtual Keys count, Cumulative Spend $, Tier, Actions"],
                    "bottom_elements": ["CustomerSheet: Customer Name, External Account ID, Contact Email, Assigned Pricing Tier, Hard Spend Quota"]
                },
                "connections": {"receives_from": "Sales/CRM sync.", "triggers_and_affects": "Scopes Virtual Keys and filters LLM Logs for customer-specific billing."},
                "use_case": "Giving 50 B2B enterprise clients their own dedicated virtual keys with strictly metered billing for your AI legal review software."
            },
            {
                "id": "gov_scim",
                "name_en": "User Provisioning (SCIM)",
                "name_tanglish": "User Provisioning / SCIM (Automated Enterprise SSO Lifecycle)",
                "route": "/workspace/scim",
                "analogy": "Automatic company ID card issuing & revoking machine mathiri.",
                "explanation": "Okta, Microsoft Entra ID / Azure AD, Keycloak-la pudhu aal sertha automatic-ah UnifAI-layum login create aagum. Employee velaiya vittu ponavudane Okta-la offboard pannina, automatic-ah UnifAI access cut aagidum!",
                "business_value": "Zero orphan accounts; 100% SOC-2 compliance for employee lifecycle security; automated enterprise IT operations.",
                "tech_arch": "Standard SCIM v2.0 (RFC 7643 & RFC 7644) server. Handles `/scim/v2/Users` and `/scim/v2/Groups` with Bearer Token auth. Automatically creates, updates, deactivates users, and maps directory groups to Teams.",
                "endpoints": ["GET /scim/v2/Users", "POST /scim/v2/Users", "GET /scim/v2/Groups", "GET /api/v1/scim/config"],
                "ui_elements": {
                    "top_bar": ["SCIM Master Enable Switch", "IdP Provider Selector (Okta, Entra ID, Keycloak)"],
                    "tabs_and_views": ["SCIM Endpoints Card: Base URL, Users endpoint, ServiceProviderConfig (with copy buttons)", "Group Mapping Card"],
                    "bottom_elements": ["Bearer Token Generator with copy button and rotation warning", "Save SCIM Config Button"]
                },
                "connections": {"receives_from": "External Identity Providers (Okta, Entra ID).", "triggers_and_affects": "Provisions and deprovisions Users and synchronizes Teams."},
                "use_case": "Auto-provisioning UnifAI accounts for 1,500 corporate employees via Okta SCIM sync."
            },
            {
                "id": "gov_rbac",
                "name_en": "Roles & Permissions (RBAC)",
                "name_tanglish": "Roles & Permissions (RBAC Security & Access Matrix)",
                "route": "/workspace/governance/rbac",
                "analogy": "Office security access badge permissions mathiri.",
                "explanation": "Yaaru platform-la enna seiyalaam nu kattu paduthum access control matrix. Super Admin-ku full control; Developer-ku API key create panna mattum permission; Finance-ku logs and billing charts mattum paarka permission (read-only).",
                "business_value": "Least-privilege security principle. Prevents accidental deletions and unauthorized key tampering.",
                "tech_arch": "Declarative Role-Based Access Control (RBAC) engine in Go and React. Enforces fine-grained resource-action pairs (`Resource: ModelProvider, VirtualKeys, RoutingRules, Settings, AuditLogs, Plugins` × `Action: View, Create, Update, Delete`).",
                "endpoints": ["GET /api/v1/governance/rbac/roles", "POST /api/v1/governance/rbac/roles", "GET /api/v1/governance/rbac/permissions"],
                "ui_elements": {
                    "top_bar": ["Create Custom Role Button", "Role Filter (System vs Custom)"],
                    "tabs_and_views": ["Roles Table: Role Name, Description, Assigned Users count, Type badge, Actions", "RBAC Permission Matrix Grid: Resources on rows, Actions on columns with checkboxes"],
                    "bottom_elements": ["RoleBuilderSheet: Role Name, Description, Permission checkboxes", "AssignUsersModal"]
                },
                "connections": {"receives_from": "Assigned to Users and Teams.", "triggers_and_affects": "Controls authorization on every API and UI view."},
                "use_case": "Creating a 'Finance Auditor' role that can View Dashboard and View LLM Logs, but cannot create or delete any keys or models."
            },
            {
                "id": "gov_access_profiles",
                "name_en": "Access Profiles",
                "name_tanglish": "Access Profiles (Reusable Policy Bundles & Guardrail Packs)",
                "route": "/workspace/governance/access-profiles",
                "analogy": "Pre-packaged security passport allathu policy bundle mathiri.",
                "explanation": "Oru standard policy template create panni vachikalam (e.g. 'Intern Profile': GPT-4o-mini mattum allow panni, $50 budget limit, PII redaction ON). Pudhu virtual key create pannum pothu intha profile-ah select panna ella security rules-um automatic-ah apply aagidum!",
                "business_value": "Administrative time savings; eliminates configuration mistakes; guarantees that no virtual key goes live without security guardrails.",
                "tech_arch": "Reusable policy entity bundling Allowed Models, Allowed Providers, Budget Caps, Rate Limits (RPM/TPM), MCP Tool Groups, and Guardrails. Attached to Virtual Keys via relational join.",
                "endpoints": ["GET /api/v1/governance/access-profiles", "POST /api/v1/governance/access-profiles", "POST /api/v1/governance/access-profiles/{id}/clone"],
                "ui_elements": {
                    "top_bar": ["Create Access Profile Button", "Search & Tag Filter"],
                    "tabs_and_views": ["Access Profiles Table: Profile Name, Tags badges, Allowed Models, Attached Keys count, Budget/Rate limit summary, Actions"],
                    "bottom_elements": ["AccessProfileDialog: Allowed Models multiselect, Budget Caps $, RPM/TPM Rate limits, MCP Tool Client pickers", "Quick Clone Button"]
                },
                "connections": {"receives_from": "Security administrators.", "triggers_and_affects": "Governs Virtual Keys behavior and restricts Model Providers and MCP tools."},
                "use_case": "Attaching a single 'Customer Chatbot Profile' to 30 virtual keys, guaranteeing PII redaction and 60 RPM limits across all 30 keys simultaneously."
            },
            {
                "id": "gov_audit_logs",
                "name_en": "Audit Logs",
                "name_tanglish": "Audit Logs (Tamper-Proof Compliance & Activity Trail)",
                "route": "/workspace/audit-logs",
                "analogy": "Bank locker CCTV camera & register note mathiri.",
                "explanation": "Platform-la yaaru yaaru entha nerathula enna maathinaanga nu maatha mudiyatha (tamper-proof) pathivu. Yaarachum API key delete pannalaam, budget limit maathinaalo, pudhu routing rule pottaalo exact timestamp, user name, and client IP address-oda capture aagum.",
                "business_value": "100% SOC-2, HIPAA, and GDPR audit readiness; insider threat detection; rapid forensic debugging.",
                "tech_arch": "Immutable, append-only security event audit trail. Captures Actor User ID, Action Type (CREATE, UPDATE, DELETE), Target Resource, Timestamp, Client IP address, User-Agent, and before/after JSON diffs in PostgreSQL.",
                "endpoints": ["GET /api/v1/audit-logs", "GET /api/v1/audit-logs/{id}", "GET /api/v1/audit-logs/export"],
                "ui_elements": {
                    "top_bar": ["Date Range Picker", "Actor Search", "Action Type Filter (CREATE, UPDATE, DELETE)", "Resource Filter", "Export Button (CSV/JSON)"],
                    "tabs_and_views": ["Audit Logs Table: Timestamp, Actor Name & Email, Action Badge, Resource Type, Target ID, Client IP, Details Button"],
                    "bottom_elements": ["Audit Detail Drawer: Displays complete before-and-after JSON delta diff highlighting modified fields"]
                },
                "connections": {"receives_from": "Intercepts administrative mutations across all Governance, Models, and Observability features.", "triggers_and_affects": "Feeds security alerts and provides compliance proof for auditors."},
                "use_case": "Demonstrating to external SOC-2 auditors exactly who rotated the production AWS Bedrock credentials on July 10th."
            }
        ]
    },
    {
        "pillar_id": "guardrails",
        "pillar_name_en": "Security Guardrails",
        "pillar_name_tanglish": "Guardrails (Content Safety & Attack Prevention)",
        "route_group": "/workspace/guardrails",
        "description": "Prompt injection prevention, PII redaction, toxicity filtering, and multi-engine security scanners.",
        "features": [
            {
                "id": "grd_rules",
                "name_en": "Guardrail Rules",
                "name_tanglish": "Guardrail Rules (AI Safety Policies & Screening)",
                "route": "/workspace/guardrails/configuration",
                "analogy": "Airport baggage scanner & security checkpoint mathiri.",
                "explanation": "Customer allathu employee AI kitta bad words, company secrets, credit card numbers, allathu jailbreak / prompt injection attacks seiyatha mathiri thadukkura security rules. Rules meera patta udane request-ah block pannum allathu PII data-vai asterisk (*) pottu mask pannum.",
                "business_value": "Zero data leakage (PII/Secrets); regulatory compliance (HIPAA/GDPR); brand reputation protection.",
                "tech_arch": "Pre-request and post-completion content security policy engine in Go. Evaluates prompt scope (`all`, specific Prompt Repo IDs, or custom CEL expressions). Rejects violations with HTTP 400/403 with verdict codes (`PII_DETECTED`, `PROMPT_INJECTION`).",
                "endpoints": ["GET /api/v1/guardrails/config", "PUT /api/v1/guardrails/config", "POST /api/v1/guardrails/test"],
                "ui_elements": {
                    "top_bar": ["Header with ShieldAlert Icon", "Add Rule Button", "Active Rules Count Badge"],
                    "tabs_and_views": ["Guardrail Rules Table: Rule Name, Active Toggle, Prompt Scope, Linked Providers, Evaluation Action, Actions"],
                    "bottom_elements": ["GuardrailRuleDialog: Rule Name, Linked Providers multiselect, Prompt Scope Selector (All, Selected Prompts, Custom CEL), CEL Editor"]
                },
                "connections": {"receives_from": "FastHTTP Proxy and Prompt Repository.", "triggers_and_affects": "Blocks attacks, redacts sensitive text, triggers Alert Channels, and logs incidents to LLM Logs."},
                "use_case": "Preventing chatbot users from extracting system prompts or submitting SQL injection payloads into AI text fields."
            },
            {
                "id": "grd_providers",
                "name_en": "Guardrail Providers",
                "name_tanglish": "Guardrail Providers (Security Engines & Regex Scanners)",
                "route": "/workspace/guardrails/providers",
                "analogy": "Security inspection scanner machinery brands mathiri.",
                "explanation": "Security rules run aagurathukku thevaiyana scanning engines (Presidio PII Engine, Llama-Guard Toxicity Engine, Lakera Prompt Injection Scanner, AWS Bedrock Guardrails, allathu Custom Regex Pattern matching engines) configure pandra edham.",
                "business_value": "Customizable data defense; multi-layered scanning capability; adaptability to enterprise-specific regex patterns.",
                "tech_arch": "Multi-engine security scanner registry. Supports local RE2 regex matching, lightweight embeddings classifiers, and external gRPC/HTTP safety microservices (Microsoft Presidio, Meta Llama-Guard). Concurrent Go worker execution (< 15ms).",
                "endpoints": ["GET /api/v1/guardrails/config", "PUT /api/v1/guardrails/config"],
                "ui_elements": {
                    "top_bar": ["Add Provider Button", "Navigation Tabs (Rules vs Providers)"],
                    "tabs_and_views": ["Guardrail Providers Table: Provider ID, Policy Name, Engine Type (Regex/Presidio/Bedrock), Pattern Count badge, Linked Rules count badge, Actions"],
                    "bottom_elements": ["GuardrailProviderDialog: Provider ID, Policy Name, Engine Type, Multi-line Regex Patterns textarea, Pattern Syntax Validator"]
                },
                "connections": {"receives_from": "DevSecOps engineers.", "triggers_and_affects": "Supplies scanning algorithms and regex patterns to Guardrail Rules."},
                "use_case": "Configuring custom regex patterns to detect and mask proprietary internal project code names before prompts reach public models."
            }
        ]
    },
    {
        "pillar_id": "infrastructure",
        "pillar_name_en": "High Availability & Infrastructure",
        "pillar_name_tanglish": "Infrastructure (Cluster Mesh, Alerts & OAuth)",
        "route_group": "/workspace/infrastructure",
        "description": "Multi-region cluster synchronization, Slack/PagerDuty alert channels, and third-party OAuth access grants.",
        "features": [
            {
                "id": "inf_cluster",
                "name_en": "Cluster Config",
                "name_tanglish": "Cluster Config (Distributed Mesh & High-Availability Sync)",
                "route": "/workspace/cluster",
                "analogy": "Multi-aircraft flight formation allathu multi-branch bank network mathiri.",
                "explanation": "Multiple UnifAI servers (USA, Europe, India regions) run aagum pothu, antha ellam servers-um onnukonnu pesikittu, rate limits, virtual key balances, and circuit breaker states instant-ah sync aaguratha uruthi seiyura master cluster control panel.",
                "business_value": "Zero single point of failure; multi-region active-active deployment; global consistency.",
                "tech_arch": "Distributed Mesh / Gossip protocol (HashiCorp Memberlist / Serf). Nodes communicate via TCP/UDP peer gossip over port 7946. Synchronizes in-memory rate-limit buckets and circuit breaker states (< 50ms) across regions.",
                "endpoints": ["GET /api/v1/cluster/config", "PUT /api/v1/cluster/config", "GET /api/v1/cluster/nodes"],
                "ui_elements": {
                    "top_bar": ["Header with Network Icon", "Cluster Master Enable Switch"],
                    "tabs_and_views": ["Cluster Topology Card: Cluster Type (Mesh/Gossip), Region ID, Peer Seed List textarea (host:port, e.g. 10.0.0.12:7946)", "Node Info Card: Node ID, Bind Address, Role, Active Peers count"],
                    "bottom_elements": ["Peer Syntax Validator with inline warnings", "Save Configuration Button"]
                },
                "connections": {"receives_from": "DevOps / SRE teams.", "triggers_and_affects": "Synchronizes Virtual Key rate limits and Circuit Breaker states across all gateway pods."},
                "use_case": "Deploying 10 UnifAI gateway pods in Kubernetes, ensuring a user's 60 RPM rate limit is enforced globally rather than allowing 600 requests."
            },
            {
                "id": "inf_alerts",
                "name_en": "Alert Channels",
                "name_tanglish": "Alert Channels (Incident Notifications & Webhooks)",
                "route": "/workspace/alert-channels",
                "analogy": "Emergency fire alarm & automated SMS siren mathiri.",
                "explanation": "AI provider crash aanaalo, budget 90% reach aanaalo, allathu prompt injection attack nadanthalo, on-duty engineers-ku instant-ah Slack, PagerDuty, Microsoft Teams, allathu Email alert anuppura emergency notification hub.",
                "business_value": "Rapid incident response; minimizes downtime; prevents surprise cloud billing spikes.",
                "tech_arch": "Multi-channel incident notification dispatcher. Formats structured alerts into Slack webhook blocks, PagerDuty Events v2 JSON, or SMTP email alerts. Implements notification de-duplication and rate dampening.",
                "endpoints": ["GET /api/v1/alert-channels", "POST /api/v1/alert-channels", "POST /api/v1/alert-channels/{id}/test"],
                "ui_elements": {
                    "top_bar": ["Add Alert Channel Button", "Channel Type Filter (Slack, PagerDuty, Email, Webhook)"],
                    "tabs_and_views": ["Alert Channels Table: Channel Name, Type badge, Webhook Endpoint, Subscribed Events badges, Health Status, Actions"],
                    "bottom_elements": ["AddChannelDialog: Channel Name, Type selector, Webhook URL, Event Subscriptions checkboxes, Send Test Notification button"]
                },
                "connections": {"receives_from": "Triggered by Circuit Breaker trips, Budgets & Limits warnings, and Guardrail violations.", "triggers_and_affects": "Dispatches messages to corporate Slack and PagerDuty."},
                "use_case": "Triggering a high-urgency PagerDuty incident when OpenAI fails and Circuit Breaker activates the backup Azure route."
            },
            {
                "id": "inf_oauth",
                "name_en": "OAuth Grants",
                "name_tanglish": "OAuth Grants (Third-Party App Authorizations)",
                "route": "/workspace/oauth-grants",
                "analogy": "Mobile phone-la 'Allow app to access camera/contacts' permission screen mathiri.",
                "explanation": "External software and AI developer tools (Cursor IDE, Continue.dev, internal apps) UnifAI gateway-kulla login panni access kettu varum pothu, antha third-party app-ku permission grant pandra OAuth 2.0 authorization center.",
                "business_value": "Standardized, secure enterprise developer tool integration without sharing permanent secret keys.",
                "tech_arch": "OAuth 2.0 / OIDC Authorization Server. Manages OAuth client IDs, client secrets, redirect URIs, authorization code exchanges, and refresh token rotation. Issues scoped access tokens.",
                "endpoints": ["GET /api/v1/oauth/grants", "POST /api/v1/oauth/clients", "DELETE /api/v1/oauth/grants/{id}"],
                "ui_elements": {
                    "top_bar": ["Register OAuth Client Button", "Search Grants"],
                    "tabs_and_views": ["Active Grants Table: Client App Name, User, Scopes Granted badges, Issued At, Expires At, Revoke Access Button"],
                    "bottom_elements": ["RegisterClientDialog: App Name, Redirect URIs, Allowed Grant Types, Requested Scopes"]
                },
                "connections": {"receives_from": "External developer tools (Cursor, LangChain apps).", "triggers_and_affects": "Issues scoped Virtual Key equivalent tokens for API access."},
                "use_case": "Authorizing developer Cursor IDE instances to route requests through UnifAI via OAuth 2.0 login."
            }
        ]
    },
    {
        "pillar_id": "adaptive_routing",
        "pillar_name_en": "Adaptive Routing & Load Balancing",
        "pillar_name_tanglish": "Adaptive Routing (Dynamic Latency & Score-Based Balancing)",
        "route_group": "/workspace/adaptive-routing",
        "description": "Multi-Armed Bandit real-time latency scoring, automated key weighting, and unhealthy route pruning.",
        "features": [
            {
                "id": "adp_dashboard",
                "name_en": "Adaptive Routing Dashboard",
                "name_tanglish": "Adaptive Routing Dashboard (Live Health Scoring & Traffic Steering)",
                "route": "/workspace/adaptive-routing",
                "analogy": "Google Maps Live Traffic Navigation mathiri.",
                "explanation": "Oru AI model (OpenAI GPT-4o) ippo slow aagiduchuna allathu error adikkithu na, system live-ah detect panni, automatically athe tharam konda innoru fast-ana provider-ku (Azure OpenAI allathu Anthropic) traffic-ah thiruppi vittu user-ku semma fast response tharum.",
                "business_value": "Guarantees lowest possible latency; automated self-healing routing during vendor slowdowns.",
                "tech_arch": "Multi-Armed Bandit (MAB) reinforcement learning load balancer. Evaluates provider directions and individual API key routes based on live p50/p95 latency and error rates. Polled every 8 seconds (`pollingInterval: 8000ms`). Dynamically assigns routing weights.",
                "endpoints": ["GET /api/v1/load-balancer/routes", "PUT /api/v1/load-balancer/config", "GET /api/v1/load-balancer/metrics"],
                "ui_elements": {
                    "top_bar": ["Header with Gauge Icon", "Settings Button deep-link"],
                    "tabs_and_views": ["3 Summary Metric Cards: Load Balancer Switch, Live Scoring Strategy, Active Dynamic Routes", "Live Adaptive Routes Table: Provider/Direction, Model Name, Health Score (0-100 color badge), p50/p95 Latency ms, Error Rate %, Traffic Weight %, Status"],
                    "bottom_elements": ["8-second background auto-refresh poller reflecting live shifting traffic weights"]
                },
                "connections": {"receives_from": "Latency and error telemetry from FastHTTP Proxy and Model Providers.", "triggers_and_affects": "Dynamically overrides Routing Rules traffic weights."},
                "use_case": "When OpenAI latency spikes from 800ms to 4,500ms, adaptive routing automatically shifts 85% of traffic to Azure OpenAI maintaining 650ms p95 latency."
            },
            {
                "id": "adp_settings",
                "name_en": "Adaptive Routing Settings",
                "name_tanglish": "Adaptive Routing Settings (Load Balancing Policy & Pruning Rules)",
                "route": "/workspace/adaptive-routing/settings",
                "analogy": "Car cruise control & autopilot driving preference settings mathiri.",
                "explanation": "Adaptive load balancer eppadi nadanthukanum nu rules poduvom: Slow-ana vendor-ah automatic-ah ignore pannanuma? Oru vendor fail aana automatic-ah adutha vendor-ku reroute pannanuma? Multi-key balancing on pannanuma nu control pandra settings page.",
                "business_value": "Predictable routing behavior; fine-grained engineering control over automated failovers.",
                "tech_arch": "Load balancer configuration and policy persistence engine in PostgreSQL (`LoadBalancerConfig`). Controls direction selection algorithms, key-level round-robin weighting, failed direction rerouting, and fallback pruning.",
                "endpoints": ["GET /api/v1/load-balancer/config", "PUT /api/v1/load-balancer/config"],
                "ui_elements": {
                    "top_bar": ["Header Title", "Back Button to Dashboard", "Save Settings Button"],
                    "tabs_and_views": ["Settings Card with 5 Policy Toggle Rows: (1) Enable adaptive load balancing, (2) Direction selection, (3) Route selection, (4) Reroute failed directions, (5) Prune failed fallbacks"],
                    "bottom_elements": ["Save Configuration Button with instant toast confirmation"]
                },
                "connections": {"receives_from": "DevOps / Infrastructure leads.", "triggers_and_affects": "Governs the algorithmic behavior of Adaptive Routing Dashboard."},
                "use_case": "Enabling 'Reroute failed directions' ensuring critical production apps never suffer timeouts from degraded upstream vendors."
            }
        ]
    }
]

# Total feature count check
TOTAL_FEATURES_COUNT = sum(len(p['features']) for p in PILLARS)

# Flattened list of all features
ALL_FEATURES = []
for p in PILLARS:
    for f in p['features']:
        f['pillar_ref'] = p['pillar_name_en']
        ALL_FEATURES.append(f)

# Master Connections Matrix
MASTER_CONNECTIONS = [
    ("Dashboard", "LLM Logs", "Observability", "Visualizes real-time request volume, token usage, latencies, and costs."),
    ("LLM Logs", "Connectors", "Observability", "Streams structured transaction events to Datadog, Kafka, and BigQuery."),
    ("MCP Logs", "MCP Gateway", "MCP", "Records JSON-RPC arguments, tool execution traces, and stdout/stderr outputs."),
    ("Browser AI", "Audit Logs", "Security", "Intercepts employee web AI paste events and logs confidential DLP violations."),
    ("Model Catalog", "Routing Rules", "Models", "Provides verified model identifiers and context parameters for routing targets."),
    ("Model Providers", "FastHTTP Proxy", "Core", "Supplies decrypted API keys, custom base URLs, and connection pools."),
    ("Budgets & Limits", "Virtual Keys", "Governance", "Enforces strict dollar quotas and token rate limits on API keys."),
    ("Complexity Router", "Routing Rules", "Models", "Scores query complexity (0-1) and assigns model tier (Simple to Reasoning)."),
    ("Circuit Breaker", "Alert Channels", "Infrastructure", "Triggers high-priority PagerDuty alerts when a primary provider fails."),
    ("Pricing Overrides", "Dashboard & Billing", "Finance", "Calculates real dollar spend using negotiated corporate discount rates."),
    ("MCP Gateway", "Tool Groups", "MCP", "Organizes certified tools into permission groups for Virtual Key assignment."),
    ("Prompt Repo", "Guardrail Rules", "Security", "Links prompt templates to content screening and jailbreak detection rules."),
    ("Plugins", "FastHTTP Proxy", "Extensibility", "Executes Wasm/Go pre-request, post-request, and streaming lifecycle hooks."),
    ("Virtual Keys", "Access Profiles", "Governance", "Inherits model restrictions, budget limits, and MCP permissions."),
    ("SCIM", "Users & Teams", "Identity", "Auto-provisions users and syncs directory groups from Okta and Entra ID."),
    ("RBAC", "All UI & APIs", "Security", "Enforces declarative View/Create/Update/Delete permissions on every endpoint."),
    ("Audit Logs", "All Features", "Compliance", "Immutably records every configuration change, key rotation, and admin action."),
    ("Guardrail Rules", "Guardrail Providers", "Security", "Executes RE2 regex, Presidio, or Bedrock scanner engines on prompt text."),
    ("Cluster Config", "Virtual Keys & Circuit", "Infrastructure", "Synchronizes rate-limit buckets and circuit trip states globally (:7946)."),
    ("Adaptive Routing", "Model Providers", "Performance", "Continuously scores provider latencies (ms) and shifts traffic dynamically.")
]

# Master Comparative Matrix
MASTER_MATRIX = [
    ("Dashboard", "AI spend evlo aaguthu, budget control-la irukka nu ore screen-la pakkalam.", "WebSocket telemetry & atomic counters aggregating p95 latencies and TTFT into Postgres."),
    ("LLM Logs", "Yaaru enna kelvi kettanga, AI enna bathil sonnathu nu passbook mathiri check pannalam.", "Sub-millisecond PostgreSQL batch-insert ledger capturing raw payloads, tokens, and status codes."),
    ("MCP Logs", "AI Agent company database-la enna vela senjithu nu inspect pannalam.", "JSON-RPC 2.0 tool execution tracing across stdio and SSE agent transport protocols."),
    ("Browser AI", "Employees company secrets-ah web ChatGPT-la copy-paste pannama thadukkalam.", "Client-side mitmproxy daemon intercepting web AI pastes with Presidio DLP regexes."),
    ("Connectors", "AI spend and logs automatic-ah namma corporate BigQuery/Datadog-ku poiduma?", "Async streaming queue worker pushing events via OpenTelemetry, Kafka, and HTTP webhooks."),
    ("Model Catalog", "Entha entha AI models namakku irukku? Ethu active-ah irukku nu menu card mathiri pakkalam.", "FastHTTP model registry with capability filtering (vision, tool calling) and 24h usage metrics."),
    ("Model Providers", "OpenAI, Anthropic, Bedrock API keys safe-ah connect aagi run aagutha?", "Multi-vendor credential vault, FastHTTP client connection pools, and exponential retry backoffs."),
    ("Budgets & Limits", "Excessive bill varama ovvoru team-kum monthly budget cap potu stop panna mudiyuma?", "Atomic Redis/Postgres token bucket rate limiters evaluating multi-tiered limits with HTTP 429."),
    ("Complexity Router", "Simple kelvikku cheap model, tough kelvikku matum costly model use panni 70% bill micham aagutha?", "Heuristic multi-factor scoring engine categorizing queries into Simple, Med, Complex, Reasoning."),
    ("Circuit Breaker", "OpenAI crash aanaal kooda customer-ku error theriyaama backup model-ku switch aaguma?", "Distributed finite state machine (Closed/Open/Half-Open) monitoring response header signals."),
    ("MCP Catalog", "AI agent-ku thevaiyana tools (GitHub, SQL, Slack) app store mathiri install pannalama?", "Curated MCP registry client fetching validated schemas and execution configurations."),
    ("Prompt Repo", "Prompt templates-ah code mathiri GitHub style-la version control panni release pannalama?", "Prompt template engine with variable interpolation (`{{var}}`), semantic diffs, and release tags."),
    ("Plugins", "Custom company rules and billing webhooks gateway-kulla plug-and-play-ah poda mudiyuma?", "Wasm/Go lifecycle interceptor pipeline hooking into pre-request, post-request, and streaming chunks."),
    ("Virtual Keys", "Real OpenAI master keys-ah tharama safe-ana dummy keys with budget limits tharalama?", "Cryptographic SHA-256 hashed API tokens with sub-millisecond FastHTTP auth and rate limiters."),
    ("SCIM", "Okta-la employee join aana automatic-ah account create aagi, leave aana cut aaguma?", "RFC 7643/7644 SCIM v2.0 server handling automated user lifecycle and group synchronization."),
    ("RBAC", "Admin-ku mattum full rights, mathavangalukku view-only permission pottu lock pannalama?", "Declarative resource-action permission matrix evaluating gates across all UI views and REST APIs."),
    ("Access Profiles", "Reusable standard security template create panni 20 keys-ku single click-la apply pannalama?", "Bundled governance policy entity containing model allowlists, budget caps, rate limits, and MCP tools."),
    ("Audit Logs", "Yaaru entha key-ah maathinaanga, eppo delete panninaanga nu maatha mudiyatha proof irukka?", "Append-only immutable audit trail capturing actor, action, timestamp, IP, and before/after JSON diffs."),
    ("Guardrail Rules", "Prompt injection, bad words, and company secrets AI kitta pogama thadukkuma?", "CEL conditional content security engine scanning prompts/responses and returning 400/403 on violation."),
    ("Cluster Config", "Multiple servers run aanaalum budget and rate limits correct-ah sync aaguma?", "Distributed Mesh/Gossip consensus protocol (port 7946) synchronizing state across multi-region nodes."),
    ("Adaptive Routing", "Oru AI vendor slow aana automatic-ah fastest backup vendor-ku reroute aaguma?", "Real-time Multi-Armed Bandit load balancer calculating live health scores (0-100) from p95 latency.")
]

# ==============================================================================
# 1. GENERATE FULL PLATFORM MARKDOWN DOCUMENT (THANGLISH + ENGLISH)
# ==============================================================================
def generate_full_markdown():
    print(f"Writing Complete Platform Master Markdown document ({TOTAL_FEATURES_COUNT} Features)...")
    lines = []
    lines.append("# UnifAI Complete Platform Master Deep-Dive Technical Manual")
    lines.append("## UnifAI Muzhumaiyana Architecture, Features, Layout & Operations Guide")
    lines.append("")
    lines.append(f"**Document Status:** Complete & Exhaustive ({TOTAL_FEATURES_COUNT} Core Features Across 9 Pillars)  ")
    lines.append("**Target Audience:** Developers, System Architects, CTOs, CFOs, DevSecOps, SREs & Product Managers  ")
    lines.append("**Language:** Bilingual (Thanglish - Tamil in English letters + Clean English)  ")
    lines.append("**Generated At:** 2026-09-05  ")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Table of Contents")
    lines.append("1. [UnifAI Master System Architecture Map](#1-unifai-master-system-architecture-map)")
    lines.append("2. [Comprehensive Pillars & Features Inventory (All 40 Features)](#2-comprehensive-pillars--features-inventory)")
    for p in PILLARS:
        lines.append(f"   - **{p['pillar_name_en']} ({p['pillar_name_tanglish']})**")
        for f in p['features']:
            lines.append(f"     * [{f['name_en']} — {f['name_tanglish']} ({f['route']})](#{f['id']})")
    lines.append("3. [Cross-Feature Interconnections & Data Flow Matrix](#3-cross-feature-interconnections--data-flow-matrix)")
    lines.append("4. [Master Tech vs Non-Tech Comparative Matrix](#4-master-tech-vs-non-tech-comparative-matrix)")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 1. UnifAI Master System Architecture Map")
    lines.append("### Muzhumaiyana 9 Pillars End-to-End Traffic Steering, Security & Observability Flow")
    lines.append("")
    lines.append("```")
    lines.append("             [ ENTERPRISE IDENTITY (Okta / Microsoft Entra ID / IdP) ]")
    lines.append("                                        │")
    lines.append("                         (SCIM v2 Sync) ▼")
    lines.append("       ┌────────────────────────────────────────────────────────────────┐")
    lines.append("       │     GOVERNANCE: USERS, TEAMS, BUSINESS UNITS & CUSTOMERS       │")
    lines.append("       │  • Multi-tenant organizational hierarchy & RBAC security       │")
    lines.append("       │  • Access Profiles: Reusable model allowlists, budgets, tools  │")
    lines.append("       └────────────────────────────────┬───────────────────────────────┘")
    lines.append("                                        │ (Generates uf-key-...)")
    lines.append("                                        ▼")
    lines.append("       ┌────────────────────────────────────────────────────────────────┐")
    lines.append("       │                       VIRTUAL KEYS                             │")
    lines.append("       │  • Cryptographic SHA-256 API tokens with strict dollar quotas  │")
    lines.append("       └────────────────────────────────┬───────────────────────────────┘")
    lines.append("                                        │ (Client Bearer Request)")
    lines.append("                                        ▼")
    lines.append("       ┌────────────────────────────────────────────────────────────────┐")
    lines.append("       │              FastHTTP HIGH-PERFORMANCE PROXY GATEWAY           │")
    lines.append("       │  • Sub-millisecond token auth, Rate limit bucket verification  │")
    lines.append("       └───────┬────────────────────────────────────────────────┬───────┘")
    lines.append("               │                                                │")
    lines.append("   (Cluster)   ▼                                                ▼ (Extensions)")
    lines.append("       ┌────────────────────────┐                      ┌────────────────────────┐")
    lines.append("       │     CLUSTER CONFIG     │                      │     PLUGINS ENGINE     │")
    lines.append("       │  • Mesh Gossip (:7946) │                      │  • Wasm/Go Pre-Request │")
    lines.append("       │  • Global Bucket Sync  │                      │  • Post-Req / Stream   │")
    lines.append("       └────────────────────────┘                      └───────────┬────────────┘")
    lines.append("                                                                   │")
    lines.append("                                                                   ▼")
    lines.append("       ┌────────────────────────────────────────────────────────────────┐")
    lines.append("       │                  SECURITY GUARDRAILS ENGINE                    │")
    lines.append("       │  • Guardrail Rules (CEL expressions, Prompt Repo bindings)     │")
    lines.append("       │  • Guardrail Providers (Presidio, RE2 Regex, Llama-Guard)      │")
    lines.append("       └───────┬────────────────────────────────────────────────┬───────┘")
    lines.append("        (Fail) │                                                │ (Pass)")
    lines.append("               ▼                                                ▼")
    lines.append("       ┌────────────────────────┐                      ┌────────────────────────┐")
    lines.append("       │ HTTP 400/403 Security  │                      │   COMPLEXITY ROUTER    │")
    lines.append("       │ Violation (PII / Attack│                      │  • Score: 0.0 to 1.0   │")
    lines.append("       └────────────────────────┘                      │  • Simple/Med/Comp/Reas│")
    lines.append("                                                       └───────────┬────────────┘")
    lines.append("                                                                   │")
    lines.append("                                                                   ▼")
    lines.append("       ┌────────────────────────────────────────────────────────────────┐")
    lines.append("       │             ROUTING RULES & ADAPTIVE LOAD BALANCER             │")
    lines.append("       │  • CEL expression evaluation, Priority target pools            │")
    lines.append("       │  • Adaptive Routing: Real-time Multi-Armed Bandit health scores│")
    lines.append("       │  • Circuit Breaker: Outage protection & zero-downtime failover │")
    lines.append("       └───────────────────────────────┬────────────────────────────────┘")
    lines.append("                                       │")
    lines.append("                                       ▼")
    lines.append("       ┌────────────────────────────────────────────────────────────────┐")
    lines.append("       │              MODEL PROVIDERS & MCP TOOL GATEWAY                │")
    lines.append("       │  • OpenAI / Anthropic / AWS Bedrock / Google Vertex / Ollama   │")
    lines.append("       │  • MCP Gateway: Autonomous Agent Tool Execution (stdio / SSE)  │")
    lines.append("       └───────────────────────────────┬────────────────────────────────┘")
    lines.append("                                       │")
    lines.append("                                       ▼")
    lines.append("       ┌────────────────────────────────────────────────────────────────┐")
    lines.append("       │             UNIFIED OBSERVABILITY & AUDIT RECORDING            │")
    lines.append("       │  • LLM Logs & MCP Logs: Full prompts, outputs, tokens, TTFT    │")
    lines.append("       │  • Dashboard: Realtime latency percentiles (p95) & total spend │")
    lines.append("       │  • Connectors: Streaming push to Datadog, Kafka, and BigQuery  │")
    lines.append("       │  • Audit Logs: Immutable before/after JSON diffs of admin ops  │")
    lines.append("       └────────────────────────────────────────────────────────────────┘")
    lines.append("```")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 2. Comprehensive Pillars & Features Inventory")
    lines.append("### All 40 Features Deep-Dive (Thanglish + English)")
    lines.append("")

    for p in PILLARS:
        lines.append(f"## Pillar: {p['pillar_name_en']} — {p['pillar_name_tanglish']}")
        lines.append(f"*{p['description']}*\n")

        for f in p['features']:
            lines.append(f"<a name='{f['id']}'></a>")
            lines.append(f"### {f['name_en']} — {f['name_tanglish']}")
            lines.append(f"**UI Route:** `{f['route']}`\n")
            lines.append("#### 👤 Non-Tech Perspective (Makkal Puriyura Eliya Vilakkam)")
            lines.append(f"- **Uruvagam (Analogy):** {f['analogy']}")
            lines.append(f"- **Vilakkam (Explanation):** {f['explanation']}")
            lines.append(f"- **Business Value (Vaniga Payan):** {f['business_value']}\n")
            lines.append("#### 💻 Tech Perspective (Technical Architecture)")
            lines.append(f"- **Backend Architecture:** {f['tech_arch']}")
            lines.append("- **Backend Endpoints:**")
            for ep in f['endpoints']:
                lines.append(f"  * `{ep}`")
            lines.append("")
            lines.append("#### 🖥️ Screen Layout & Interactive Elements (Thirai Koorugal)")
            lines.append("**1. Melpura Koorugal (Top Bar Controls):**")
            for itm in f['ui_elements']['top_bar']:
                lines.append(f"- {itm}")
            lines.append("")
            lines.append("**2. Mathiya Koorugal & Tables (Tabs & Views):**")
            for itm in f['ui_elements']['tabs_and_views']:
                lines.append(f"- {itm}")
            lines.append("")
            lines.append("**3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):**")
            for itm in f['ui_elements']['bottom_elements']:
                lines.append(f"- {itm}")
            lines.append("")
            lines.append("#### 🔗 Connections & Structure Map (Inaippugal)")
            lines.append(f"- **Data-vai Perumidam (Receives From):** {f['connections']['receives_from']}")
            lines.append(f"- **Iyakkum Koorugal (Triggers & Affects):** {f['connections']['triggers_and_affects']}")
            lines.append("")
            lines.append(f"#### 💡 Production Use Case: {f['use_case']}\n")
            lines.append("---\n")

    lines.append("# 3. Cross-Feature Interconnections & Data Flow Matrix")
    lines.append("### UnifAI Platform Features Kulla Irukura Master Connections\n")
    lines.append("| Source Feature | Connected To | Category | Data Flow & Trigger Action |")
    lines.append("| :--- | :--- | :--- | :--- |")
    for src, dst, cat, flow in MASTER_CONNECTIONS:
        lines.append(f"| **{src}** | **{dst}** | `{cat}` | {flow} |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 4. Master Tech vs Non-Tech Comparative Matrix")
    lines.append("### Thozhilnutpam vs Vanigam Parvai Master Oppeedo\n")
    lines.append("| Feature | Non-Tech View (Manager / CFO Parvai) | Tech View (DevOps / Architect Parvai) |")
    lines.append("| :--- | :--- | :--- |")
    for ft, nt, tv in MASTER_MATRIX:
        lines.append(f"| **{ft}** | \"{nt}\" | \"{tv}\" |")
    lines.append("")

    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Master Markdown written to: {MD_PATH}")

# ==============================================================================
# 2. GENERATE FULL PLATFORM DOCX DOCUMENT (THANGLISH + ENGLISH)
# ==============================================================================
def generate_full_docx():
    print(f"Writing Complete Platform Master Word Document (.docx) ({TOTAL_FEATURES_COUNT} Features)...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        hp = section.header.paragraphs[0]
        hp.text = "UnifAI Complete Platform Master Technical Manual (Thanglish + English)"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if hp.runs:
            hp.runs[0].font.name = "Segoe UI"
            hp.runs[0].font.size = Pt(8.5)
            hp.runs[0].font.color.rgb = RGBColor(100, 116, 139)
            
        fp = section.footer.paragraphs[0]
        fp.text = "Confidential & Proprietary — UnifAI Enterprise Architecture & Feature Specification"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if fp.runs:
            fp.runs[0].font.name = "Segoe UI"
            fp.runs[0].font.size = Pt(8.5)
            fp.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    def style_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        return p

    def style_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def style_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10.5)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(9.5)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(9.5)
        r.italic = italic
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(9.5)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def format_table(table, col_widths, col_names):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(col_names):
            hdr_cells[i].text = name
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
            tcPr.append(shd)
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = "Segoe UI"
                run.font.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for r_idx, row in enumerate(table.rows[1:]):
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                tcPr.append(shd)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                    
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    rt = p_title.add_run("UnifAI Complete Platform Master Manual")
    rt.font.name = "Segoe UI"
    rt.font.size = Pt(22)
    rt.bold = True
    rt.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    rsub = p_sub.add_run(f"Exhaustive Technical & Operational Manual Covering All {TOTAL_FEATURES_COUNT} Features across 9 Pillars (Thanglish + English)")
    rsub.font.name = "Segoe UI"
    rsub.font.size = Pt(11.5)
    rsub.font.color.rgb = RGBColor(71, 85, 105)

    style_h1("1. Master System Architecture Overview")
    add_p("UnifAI acts as the centralized high-throughput AI gateway, policy enforcement plane, and multi-cloud router. Every incoming request passes through authentication, rate-limiting, guardrails, complexity scoring, and dynamic load balancing before reaching upstream AI models or MCP agent tool execution servers.")

    style_h1(f"2. Comprehensive Feature Breakdown ({TOTAL_FEATURES_COUNT} Features)")
    for p in PILLARS:
        style_h1(f"Pillar: {p['pillar_name_en']} ({p['pillar_name_tanglish']})")
        add_p(p['description'], italic=True)

        for f in p['features']:
            style_h2(f"{f['name_en']} — {f['name_tanglish']}")
            add_p(f['route'], bold_prefix="Route: ", italic=True)

            style_h3("Non-Tech Perspective (Eliya Vilakkam)")
            add_p(f['analogy'], bold_prefix="Uruvagam (Analogy): ")
            add_p(f['explanation'], bold_prefix="Vilakkam: ")
            add_p(f['business_value'], bold_prefix="Business Value: ")

            style_h3("Tech Perspective (Technical Architecture)")
            add_p(f['tech_arch'], bold_prefix="Backend Architecture: ")
            add_p("Backend Endpoints:")
            for ep in f['endpoints']:
                add_bullet(ep)

            style_h3("Screen Layout & Interactive Elements (Thirai Koorugal)")
            add_p("1. Melpura Koorugal (Top Bar Controls):", bold_prefix="")
            for itm in f['ui_elements']['top_bar']:
                add_bullet(itm)
            add_p("2. Mathiya Koorugal & Tables (Tabs & Views):", bold_prefix="")
            for itm in f['ui_elements']['tabs_and_views']:
                add_bullet(itm)
            add_p("3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):", bold_prefix="")
            for itm in f['ui_elements']['bottom_elements']:
                add_bullet(itm)

            style_h3("Module Connections (Inaippugal)")
            add_p(f['connections']['receives_from'], bold_prefix="Data-vai Perumidam: ")
            add_p(f['connections']['triggers_and_affects'], bold_prefix="Iyakkum Koorugal: ")

            style_h3("Production Use Case (Nadamurai Udharanam)")
            add_p(f['use_case'])

    style_h1("3. Master Cross-Feature Interconnections")
    t_conn = doc.add_table(rows=len(MASTER_CONNECTIONS)+1, cols=4)
    format_table(t_conn, [1.5, 1.5, 1.2, 2.8], ["Source Feature", "Connected To", "Category", "Data Flow & Trigger Action"])
    for idx, (src, dst, cat, flow) in enumerate(MASTER_CONNECTIONS):
        t_conn.rows[idx+1].cells[0].text = src
        t_conn.rows[idx+1].cells[1].text = dst
        t_conn.rows[idx+1].cells[2].text = cat
        t_conn.rows[idx+1].cells[3].text = flow

    style_h1("4. Master Tech vs Non-Tech Comparative Matrix")
    t_mat = doc.add_table(rows=len(MASTER_MATRIX)+1, cols=3)
    format_table(t_mat, [1.5, 2.7, 2.8], ["Feature", "Non-Tech View (Manager / CFO Parvai)", "Tech View (DevOps / Architect Parvai)"])
    for idx, (ft, nt, tv) in enumerate(MASTER_MATRIX):
        t_mat.rows[idx+1].cells[0].text = ft
        t_mat.rows[idx+1].cells[1].text = nt
        t_mat.rows[idx+1].cells[2].text = tv

    doc.save(DOCX_PATH)
    print(f"Master Word document written to: {DOCX_PATH}")

# ==============================================================================
# 3. GENERATE FULL PLATFORM PDF DOCUMENT VIA REPORTLAB (THANGLISH + ENGLISH)
# ==============================================================================
def generate_full_pdf():
    print(f"Writing Complete Platform Master PDF Document (.pdf) ({TOTAL_FEATURES_COUNT} Features)...")

    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748B"))
            
            if self._pageNumber > 1:
                self.drawString(40, 11 * inch - 36, "UnifAI Complete Platform Master Manual (Thanglish + English)")
                self.drawRightString(8.5 * inch - 40, 11 * inch - 36, "Confidential — Engineering & Operations Guide")
                self.setStrokeColor(colors.HexColor("#CBD5E1"))
                self.setLineWidth(0.5)
                self.line(40, 11 * inch - 40, 8.5 * inch - 40, 11 * inch - 40)
                
            self.drawString(40, 32, "UnifAI Unified AI Control Plane, Governance, Security & Performance")
            page_str = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(8.5 * inch - 40, 32, page_str)
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(40, 42, 8.5 * inch - 40, 42)
            
            self.restoreState()

    pdf_doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=48,
        bottomMargin=48
    )

    styles = getSampleStyleSheet()

    p_title_style = ParagraphStyle(
        'PdfTitle',
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=3
    )

    p_subtitle_style = ParagraphStyle(
        'PdfSubTitle',
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#475569'),
        spaceAfter=10
    )

    h1_style = ParagraphStyle(
        'PdfH1',
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=11,
        spaceAfter=4,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'PdfH2',
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=7,
        spaceAfter=3,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'PdfH3',
        fontName='Helvetica-Bold',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=4,
        spaceAfter=2,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'PdfBody',
        fontName='Helvetica',
        fontSize=7.8,
        leading=10.5,
        textColor=colors.HexColor('#334155'),
        spaceAfter=2.5
    )

    bullet_style = ParagraphStyle(
        'PdfBullet',
        fontName='Helvetica',
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor('#334155'),
        leftIndent=10,
        spaceAfter=1.5
    )

    th_style = ParagraphStyle(
        'PdfTH',
        fontName='Helvetica-Bold',
        fontSize=7.5,
        leading=9.5,
        textColor=colors.white
    )

    td_style = ParagraphStyle(
        'PdfTD',
        fontName='Helvetica',
        fontSize=7,
        leading=9.5,
        textColor=colors.HexColor('#0F172A')
    )

    td_code_style = ParagraphStyle(
        'PdfTDCode',
        fontName='Helvetica-Bold',
        fontSize=7,
        leading=9.5,
        textColor=colors.HexColor('#1E3A8A')
    )

    story = []

    # Title & Subtitle
    story.append(Paragraph("UnifAI Complete Platform Master Technical Manual", p_title_style))
    story.append(Paragraph(f"Exhaustive Technical & Operational Specification Covering All {TOTAL_FEATURES_COUNT} Features Across 9 Pillars (Thanglish + English)", p_subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#2563EB'), spaceBefore=2, spaceAfter=8))

    # Architecture Overview
    story.append(Paragraph("1. Master System Architecture Overview", h1_style))
    story.append(Paragraph(
        f"This master document represents the single source of truth for the entire UnifAI platform, covering all {TOTAL_FEATURES_COUNT} features "
        "across 9 core pillars: Observability, Models & Traffic Control, MCP Gateway & Tool Execution, Prompt Management & Playground, "
        "Plugins & Extensibility, Governance & Identity, Security Guardrails, High Availability Infrastructure, and Adaptive Routing. "
        "Every feature is detailed from both an intuitive non-technical perspective (real-world analogies, plain Thanglish explanation, business value) "
        "and an in-depth technical engineering perspective (FastHTTP architecture, endpoints, UI layouts, and connection structure maps).",
        body_style
    ))
    story.append(Spacer(1, 4))

    # Features Loop
    story.append(Paragraph(f"2. Comprehensive Feature Breakdown ({TOTAL_FEATURES_COUNT} Features Across 9 Pillars)", h1_style))
    for p in PILLARS:
        story.append(Paragraph(f"<b>Pillar: {p['pillar_name_en']} ({p['pillar_name_tanglish']})</b>", h1_style))
        story.append(Paragraph(f"<i>{p['description']}</i>", body_style))
        story.append(Spacer(1, 3))

        for f in p['features']:
            feat_elements = []
            feat_elements.append(Paragraph(f"<b>{f['name_en']} — {f['name_tanglish']}</b>", h2_style))
            feat_elements.append(Paragraph(f"<b>Route:</b> <font color='#2563EB'>{f['route']}</font>", body_style))

            feat_elements.append(Paragraph("<b>👤 Non-Tech Perspective (Eliya Vilakkam):</b>", h3_style))
            feat_elements.append(Paragraph(f"• <b>Uruvagam (Analogy):</b> {f['analogy']}", body_style))
            feat_elements.append(Paragraph(f"• <b>Vilakkam:</b> {f['explanation']}", body_style))
            feat_elements.append(Paragraph(f"• <b>Business Value:</b> {f['business_value']}", body_style))

            feat_elements.append(Paragraph("<b>💻 Tech Perspective (Technical Architecture):</b>", h3_style))
            feat_elements.append(Paragraph(f"• <b>Backend Architecture:</b> {f['tech_arch']}", body_style))
            ep_text = ", ".join([f"<code>{e}</code>" for e in f['endpoints']])
            feat_elements.append(Paragraph(f"• <b>Endpoints:</b> {ep_text}", body_style))

            feat_elements.append(Paragraph("<b>🖥️ Screen Layout & Interactive Elements (Thirai Koorugal):</b>", h3_style))
            feat_elements.append(Paragraph("<b>1. Melpura Koorugal (Top Bar Controls):</b>", body_style))
            for itm in f['ui_elements']['top_bar']:
                feat_elements.append(Paragraph(f"• {itm}", bullet_style))
            feat_elements.append(Paragraph("<b>2. Mathiya Koorugal & Tables (Tabs & Views):</b>", body_style))
            for itm in f['ui_elements']['tabs_and_views']:
                feat_elements.append(Paragraph(f"• {itm}", bullet_style))
            feat_elements.append(Paragraph("<b>3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):</b>", body_style))
            for itm in f['ui_elements']['bottom_elements']:
                feat_elements.append(Paragraph(f"• {itm}", bullet_style))

            feat_elements.append(Paragraph("<b>🔗 Connections & Structure Map (Inaippugal):</b>", h3_style))
            feat_elements.append(Paragraph(f"• <b>Data-vai Perumidam:</b> {f['connections']['receives_from']}", body_style))
            feat_elements.append(Paragraph(f"• <b>Iyakkum Koorugal:</b> {f['connections']['triggers_and_affects']}", body_style))

            feat_elements.append(Paragraph(f"<b>💡 Production Use Case:</b> {f['use_case']}", body_style))
            feat_elements.append(Spacer(1, 5))
            story.append(KeepTogether(feat_elements))

    # Cross Connections
    story.append(Paragraph("3. Master Cross-Feature Interconnections Matrix", h1_style))
    conn_table_data = [[Paragraph(h, th_style) for h in ["Source Feature", "Connected To", "Category", "Data Flow & Trigger Action"]]]
    for src, dst, cat, flow in MASTER_CONNECTIONS:
        conn_table_data.append([Paragraph(src, td_code_style), Paragraph(dst, td_code_style), Paragraph(cat, td_style), Paragraph(flow, td_style)])
    tbl_conn = Table(conn_table_data, colWidths=[95, 95, 75, 265])
    tbl_conn.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('LEFTPADDING', (0, 0), (-1, -1), 2.5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2.5),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_conn)
    story.append(Spacer(1, 8))

    # Comparative Matrix
    story.append(Paragraph("4. Master Tech vs Non-Tech Comparative Matrix", h1_style))
    mat_table_data = [[Paragraph(h, th_style) for h in ["Feature", "Non-Tech View (Manager / CFO Parvai)", "Tech View (DevOps / Architect Parvai)"]]]
    for ft, nt, tv in MASTER_MATRIX:
        mat_table_data.append([Paragraph(ft, td_code_style), Paragraph(nt, td_style), Paragraph(tv, td_style)])
    tbl_mat = Table(mat_table_data, colWidths=[90, 220, 220])
    tbl_mat.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('LEFTPADDING', (0, 0), (-1, -1), 2.5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 2.5),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_mat)

    pdf_doc.build(story, canvasmaker=NumberedCanvas)
    print(f"Master PDF document written to: {PDF_PATH}")

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================
if __name__ == "__main__":
    print(f"Building Complete Platform Master Documentation ({TOTAL_FEATURES_COUNT} Features in Thanglish + English)...")
    generate_full_markdown()
    generate_full_docx()
    generate_full_pdf()
    print("Complete Platform Master Documentation successfully built in:", DOC_DIR)
