import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether, PageBreak, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

# Reconfigure console output to UTF-8
if sys.stdout and hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

WORKSPACE_DIR = r"c:\Users\sakth\OneDrive\ドキュメント\Desktop\unify - Copy"
DOC_DIR = os.path.join(WORKSPACE_DIR, "document")
os.makedirs(DOC_DIR, exist_ok=True)

MD_PATH = os.path.join(DOC_DIR, "UnifAI_Guardrails_Cluster_AdaptiveRouting_Guide.md")
DOCX_PATH = os.path.join(DOC_DIR, "UnifAI_Guardrails_Cluster_AdaptiveRouting_Guide.docx")
PDF_PATH = os.path.join(DOC_DIR, "UnifAI_Guardrails_Cluster_AdaptiveRouting_Guide.pdf")

# ==============================================================================
# DATA DEFINITIONS FOR GUARDRAILS, CLUSTER & ADAPTIVE ROUTING (THANGLISH + ENGLISH)
# ==============================================================================

GUARDRAIL_FEATURES = [
    {
        "id": "guardrails_rules",
        "name_en": "Guardrail Rules",
        "name_tanglish": "Guardrail Rules (AI Safety Policies & Prompt Screening)",
        "route": "/workspace/guardrails/configuration",
        "analogy": "Airport baggage scanner & customs security checkpoint mathiri.",
        "explanation": "Customer allathu employee AI kitta thappana kelvi kettu, bad words, company secrets, credit card numbers, allathu jailbreak / prompt injection attacks seiyatha mathiri thadukkura security rules. Incoming prompts matrum outgoing responses-ah scan panni, rules meera patta udane request-ah block pannum allathu PII data-vai asterisk (*) pottu mask pannum.",
        "business_value": "Zero data leakage (PII/Secrets), regulatory compliance (HIPAA/GDPR), and brand reputation protection from toxic AI responses.",
        "tech_arch": "Pre-request and post-completion content security policy evaluation engine in Go. Evaluates prompt scope (`all`, specific prompt IDs from Prompt Repository, or custom Common Expression Language (CEL) expressions). Dispatches content to linked Guardrail Providers in parallel. Rejects violations with HTTP 400 Bad Request / 403 Forbidden containing structured guardrail verdict codes (`PII_DETECTED`, `PROMPT_INJECTION`, `TOXICITY`).",
        "endpoints": [
            "GET /api/v1/guardrails/config",
            "PUT /api/v1/guardrails/config",
            "POST /api/v1/guardrails/test"
        ],
        "ui_elements": {
            "top_bar": [
                "Header with ShieldAlert Icon: Visual indicator of content safety controls.",
                "Add Rule Button: Opens the rule configuration modal dialog.",
                "Rules Count Badge: Displays the number of active safety rules."
            ],
            "tabs_and_views": [
                "Guardrail Rules Table: Columns for Rule Name, Active Toggle Switch, Prompt Scope (All, Selected Prompts, Custom CEL), Linked Providers summary, Evaluation Action, Actions Menu (Edit, Delete).",
                "Instant Toggle: Enable or disable individual rules without service downtime."
            ],
            "bottom_elements": [
                "GuardrailRuleDialog (Modal): Form configuring Rule Name, Description, Active Switch, Linked Guardrail Providers multiselect, Prompt Scope Selector (All Prompts, Pick Prompts, Custom CEL Expression), and CEL Editor textarea.",
                "Delete Rule Confirmation Modal: Prevents accidental removal of production safety guardrails."
            ]
        },
        "connections": {
            "receives_from": "Evaluates incoming requests in FastHTTP Proxy and checks against templates in Prompt Repository.",
            "triggers_and_affects": "Blocks violating requests, redacts sensitive text, triggers Alert Channels, and logs security incidents into LLM Logs and Audit Logs."
        },
        "use_case": "Preventing customer service chatbot users from extracting system prompts or submitting SQL injection payloads into AI text fields."
    },
    {
        "id": "guardrails_providers",
        "name_en": "Guardrail Providers",
        "name_tanglish": "Guardrail Providers (Security Engines & Regex Scanners)",
        "route": "/workspace/guardrails/providers",
        "analogy": "Security inspection machinery & scanner vendor brands mathiri (X-ray machine, metal detector brands).",
        "explanation": "Security rules run aagurathukku thevaiyana scanning engines (e.g. Presidio PII Engine, Llama-Guard Toxicity Engine, Lakera Prompt Injection Scanner, AWS Bedrock Guardrails, allathu Custom Regex Pattern matching engines) configure pandra edham. Namakku thevaiyana regex patterns (credit cards, Aadhaar, SSN, passwords) inga add pannikalam.",
        "business_value": "Customizable data defense; multi-layered scanning capability; adaptability to enterprise-specific regex patterns and third-party security vendors.",
        "tech_arch": "Multi-engine security scanner registry. Supports local regex compiled matching (Google RE2 engine), lightweight embeddings classifiers, and external gRPC/HTTP safety microservices (Microsoft Presidio, Meta Llama-Guard). Scanners execute concurrently using Go worker routines with strict execution timeout deadlines (< 15ms) to prevent latency spikes.",
        "endpoints": [
            "GET /api/v1/guardrails/config",
            "PUT /api/v1/guardrails/config",
            "POST /api/v1/guardrails/providers/{id}/test"
        ],
        "ui_elements": {
            "top_bar": [
                "Add Provider Button: Opens provider engine creation dialog.",
                "Navigation Tabs: Seamless switching between Rules Configuration and Providers Configuration."
            ],
            "tabs_and_views": [
                "Guardrail Providers Table: Columns for Provider ID, Policy Name, Engine Type (Regex / Presidio / Bedrock), Pattern Count badge, Linked Rules count badge, Actions (Edit, Delete).",
                "Dependency Safety Check: Prevents deletion of any provider currently linked to an active rule."
            ],
            "bottom_elements": [
                "GuardrailProviderDialog: Form configuring Provider ID, Policy Name, Engine Type selector, and multi-line Regex Patterns textarea (one regex pattern per line, e.g. `\\b\\d{4}-\\d{4}-\\d{4}-\\d{4}\\b` for credit cards).",
                "Pattern Syntax Validator: Real-time RE2 regex syntax checker with instant error highlighting."
            ]
        },
        "connections": {
            "receives_from": "Configured by DevSecOps engineers with enterprise regex patterns and API credentials.",
            "triggers_and_affects": "Supplies scanning algorithms and regex patterns to Guardrail Rules."
        },
        "use_case": "Configuring custom regex patterns to detect and mask proprietary internal project code names (e.g. 'Project-Titanium', 'Apollo-Core') before prompts reach public OpenAI models."
    },
    {
        "id": "cluster_config",
        "name_en": "Cluster Config",
        "name_tanglish": "Cluster Config (Distributed Mesh & High-Availability Sync)",
        "route": "/workspace/cluster",
        "analogy": "Multi-aircraft flight formation allathu multi-branch bank network mathiri (ellam branches-um real-time sync-la irukkum).",
        "explanation": "Oru periya company-la lakshakanakana requests handle panna multiple UnifAI servers (USA, Europe, India regions) run aagum pothu, antha ellam servers-um onnukonnu pesikittu, rate limits, virtual key balances, and circuit breaker states instant-ah sync aaguratha uruthi seiyura master cluster control panel.",
        "business_value": "Zero single point of failure; multi-region active-active deployment; global consistency across worldwide AI infrastructure.",
        "tech_arch": "Distributed Mesh / Gossip protocol (HashiCorp Memberlist / Serf consensus). Cluster nodes communicate via TCP/UDP peer gossip over a dedicated clustering port (default: 7946). Synchronizes in-memory token rate-limit buckets, distributed circuit breaker states, and virtual key revocations with eventual consistency (< 50ms) across regions. Automatic split-brain avoidance and dead-node ejection.",
        "endpoints": [
            "GET /api/v1/cluster/config",
            "PUT /api/v1/cluster/config",
            "GET /api/v1/cluster/nodes",
            "POST /api/v1/cluster/join"
        ],
        "ui_elements": {
            "top_bar": [
                "Header with Network Icon: Visual representation of distributed clustering.",
                "Cluster Mode Master Switch: Instant toggle between Standalone and Distributed Cluster mode."
            ],
            "tabs_and_views": [
                "Cluster Topology Configuration Card:",
                "• Cluster Type Selector: Mesh vs Gossip/Raft.",
                "• Region Identifier Input: e.g. `us-east-1`, `eu-west-1`, `ap-south-1`.",
                "• Peer Seed List Textarea: Multi-line list of peer node addresses (`host:port`, e.g. `10.0.0.12:7946`).",
                "Node Information Card: Node ID, Local Bind Address, Node Role (Leader, Follower, Peer), Active Healthy Peers count."
            ],
            "bottom_elements": [
                "Peer Syntax Validator: Real-time check verifying valid IP:Port formats with inline warning alerts.",
                "Save Configuration Button: Commits cluster topology with reload notification."
            ]
        },
        "connections": {
            "receives_from": "Configured by Cloud Infrastructure & Kubernetes DevOps teams.",
            "triggers_and_affects": "Synchronizes Virtual Key rate limits, Circuit Breaker trip states, and Adaptive Routing scores across all gateway instances."
        },
        "use_case": "Deploying 10 load-balanced UnifAI gateway pods in Kubernetes, ensuring that a user's 60 RPM rate limit is globally enforced across all 10 pods rather than allowing 600 requests."
    },
    {
        "id": "adaptive_routing_dashboard",
        "name_en": "Adaptive Routing Dashboard",
        "name_tanglish": "Adaptive Routing Dashboard (Live Health Scoring & Traffic Steering)",
        "route": "/workspace/adaptive-routing",
        "analogy": "Google Maps Live Traffic Navigation mathiri (Traffic irukkura road-ah thavirthu fast-ana road-la auto-reroute pandra mathiri).",
        "explanation": "Oru AI model (e.g. OpenAI GPT-4o) ippo slow aagiduchuna allathu error adikkithu na, system athai live-ah detect panni, automatically athe tharam konda innoru fast-ana provider-ku (e.g. Azure OpenAI allathu Anthropic Claude) traffic-ah thiruppi vittu user-ku semma fast response tharum. Ovvoru vendor-oda live health score-aiyum inga live-ah pakkalam.",
        "business_value": "Fastest AI response times for users; zero human intervention during vendor slowdowns; automated performance optimization.",
        "tech_arch": "Multi-Armed Bandit (MAB) reinforcement learning load balancer. Continuously evaluates provider directions and individual API key routes based on live telemetry (p50/p95 latency, error rates, throttle headers, HTTP status codes). Polled via UI every 8 seconds (`pollingInterval: 8000ms`). Dynamically adjusts traffic weighting coefficients in real time to steer traffic to the highest-scoring upstream paths.",
        "endpoints": [
            "GET /api/v1/load-balancer/routes",
            "PUT /api/v1/load-balancer/config",
            "GET /api/v1/load-balancer/metrics"
        ],
        "ui_elements": {
            "top_bar": [
                "Header with Gauge Icon: Adaptive Routing and Performance Steering.",
                "Settings Button: Deep-link to Adaptive Routing Settings (`/workspace/adaptive-routing/settings`)."
            ],
            "tabs_and_views": [
                "Summary Metric Cards (3 Cards):",
                "• Load Balancer Switch: Master toggle displaying Enabled/Disabled status with instant switch.",
                "• Scoring Strategy: Displays active algorithm (`latency_p95`, `balanced`, `least_errors`, `cost_aware`).",
                "• Active Dynamic Routes Count: Total monitored provider directions.",
                "Live Adaptive Routes Table: Columns for Provider / Direction, Model Name, Health Score (0-100 color badge), p50 Latency (ms), p95 Latency (ms), Error Rate (%), Traffic Weight (%), Active Status badge."
            ],
            "bottom_elements": [
                "Auto-Refresh Poller: Live 8-second background polling reflecting shifting traffic weights as upstream latencies change."
            ]
        },
        "connections": {
            "receives_from": "Continuously ingests latency and error metrics from FastHTTP Proxy and Model Providers.",
            "triggers_and_affects": "Dynamically overrides Routing Rules traffic weights and cooperates with Circuit Breaker."
        },
        "use_case": "During OpenAI latency spikes (jumping from 800ms to 4,500ms), adaptive routing automatically shifts 85% of traffic to Azure OpenAI instance maintaining a 650ms p95 latency."
    },
    {
        "id": "adaptive_routing_settings",
        "name_en": "Adaptive Routing Settings",
        "name_tanglish": "Adaptive Routing Settings (Load Balancing Policy & Pruning Rules)",
        "route": "/workspace/adaptive-routing/settings",
        "analogy": "Car cruise control & autopilot driving preference settings mathiri.",
        "explanation": "Adaptive load balancer eppadi nadanthukanum nu rules poduvom: Slow-ana vendor-ah automatic-ah ignore pannanuma? Oru vendor fail aana automatic-ah adutha vendor-ku reroute pannanuma? Multi-key balancing on pannanuma nu control pandra settings page.",
        "business_value": "Predictable routing behavior; fine-grained engineering control over automated failovers; custom tuning for sensitive production workloads.",
        "tech_arch": "Load balancer configuration and policy persistence engine. Manages `LoadBalancerConfig` stored in PostgreSQL. Controls direction selection algorithms, key-level round-robin weighting, failed direction rerouting policies, and automatic pruning of degraded fallback nodes from dynamic routing tables.",
        "endpoints": [
            "GET /api/v1/load-balancer/config",
            "PUT /api/v1/load-balancer/config"
        ],
        "ui_elements": {
            "top_bar": [
                "Header Title: Adaptive Routing Settings.",
                "Back Button: Return navigation to Adaptive Routing Dashboard (`/workspace/adaptive-routing`).",
                "Save Settings Button: Commits configuration updates."
            ],
            "tabs_and_views": [
                "Settings Card with 5 Policy Toggle Rows:",
                "1. Enable adaptive load balancing: Master switch using live route scores when selecting provider and key.",
                "2. Direction selection: Automatically pick the healthiest provider for a requested model.",
                "3. Route selection: Weight individual API keys inside the chosen provider to balance quota.",
                "4. Reroute failed directions: If a pinned provider is unhealthy, automatically reroute the request to a healthy one.",
                "5. Prune failed fallbacks: Automatically drop unhealthy directions from a request's configured fallbacks list."
            ],
            "bottom_elements": [
                "Save Configuration Button (Save): Persists configuration with instant toast confirmation and state synchronization."
            ]
        },
        "connections": {
            "receives_from": "Configured by Lead Architects and DevOps engineers.",
            "triggers_and_affects": "Governs the live algorithmic behavior of Adaptive Routing Dashboard and upstream request dispatching."
        },
        "use_case": "Enabling 'Reroute failed directions' and 'Prune failed fallbacks' ensuring critical banking apps never experience timeouts due to stale upstream connections."
    }
]

GUARDRAIL_CONN_DATA = [
    ("Guardrail Rules", "Guardrail Providers", "Executes regex, Presidio, or Bedrock scanner engines linked to the rule."),
    ("Guardrail Rules", "FastHTTP Proxy", "Intercepts and scans incoming prompts and outgoing completions before forwarding."),
    ("Cluster Config", "Virtual Keys", "Synchronizes in-memory token rate-limit buckets across all cluster nodes."),
    ("Cluster Config", "Circuit Breaker", "Propagates tripped circuit breaker states globally within 50ms."),
    ("Adaptive Routing", "Model Providers", "Measures live latency and error rates across all vendor API keys."),
    ("Adaptive Routing", "Routing Rules", "Dynamically adjusts weights and bypasses slow routes based on live scores."),
    ("Adaptive Settings", "Adaptive Routing", "Defines the pruning, rerouting, and key weighting policies for the load balancer.")
]

GUARDRAIL_MATRIX_DATA = [
    ("Guardrail Rules", "Prompt injection, bad words, and company secrets AI kitta pogama thadukkuma?", "CEL conditional content security engine scanning prompts/responses and returning 400/403 on violation."),
    ("Guardrail Providers", "Custom regex patterns (credit cards, Aadhaar, internal codes) add panna mudiyuma?", "Multi-engine scanner registry supporting RE2 regexes, Presidio, Llama-Guard, and Bedrock Guardrails."),
    ("Cluster Config", "Multiple servers run aanaalum budget and rate limits correct-ah sync aaguma?", "Distributed Mesh/Gossip consensus protocol (port 7946) synchronizing state across multi-region nodes."),
    ("Adaptive Dashboard", "Oru AI vendor slow aana automatic-ah fastest backup vendor-ku reroute aaguma?", "Real-time Multi-Armed Bandit load balancer calculating live health scores (0-100) from p95 latency."),
    ("Adaptive Settings", "Slow vendor-ah thavirkka, fallbacks prune panna policies configure pannalama?", "Persisted LoadBalancerConfig controlling direction selection, route selection, and failed route pruning.")
]

# ==============================================================================
# 1. GENERATE MARKDOWN DOCUMENT (THANGLISH + ENGLISH)
# ==============================================================================
def generate_guardrail_markdown():
    print("Writing Guardrails, Cluster & Adaptive Routing Markdown document (Thanglish + English)...")
    lines = []
    lines.append("# UnifAI Guardrails, Cluster & Adaptive Routing Master Guide")
    lines.append("## UnifAI Security Guardrails, Cluster Sync & Adaptive Routing Manual")
    lines.append("")
    lines.append("**Pillars Covered:** Guardrails (Rules & Providers), Cluster Config, and Adaptive Routing (Dashboard & Settings)  ")
    lines.append("**Target Audience:** Security Engineers, System Architects, SREs, CTOs, and Infrastructure Admins  ")
    lines.append("**Language:** Bilingual (Thanglish - Tamil in English letters + Clean English)  ")
    lines.append("**Generated At:** 2026-09-05  ")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Table of Contents")
    lines.append("1. [System Architecture & Traffic Flow Map](#1-system-architecture--traffic-flow-map)")
    lines.append("2. [Detailed Feature Dissection (5 Core Features)](#2-detailed-feature-dissection)")
    for f in GUARDRAIL_FEATURES:
        lines.append(f"   - [{f['name_en']} ({f['name_tanglish']}) ({f['route']})](#{f['id']})")
    lines.append("3. [Cross-Feature Interconnections & Data Flow](#3-cross-feature-interconnections--data-flow)")
    lines.append("4. [Tech vs Non-Tech Comparative Matrix](#4-tech-vs-non-tech-comparative-matrix)")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 1. System Architecture & Traffic Flow Map")
    lines.append("### Safety Screening, Cluster Consensus & Dynamic Load Balancing Flow")
    lines.append("")
    lines.append("```")
    lines.append("                     [ INCOMING CLIENT AI REQUEST ]")
    lines.append("                                   │")
    lines.append("                                   ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │                CLUSTER CONFIG SYNCHRONIZATION             │")
    lines.append("       │  • Distributed Mesh / Gossip protocol (TCP/UDP :7946)     │")
    lines.append("       │  • Synchronizes Token Buckets & Rate Limits across nodes  │")
    lines.append("       └───────────────────────────┬───────────────────────────────┘")
    lines.append("                                   │")
    lines.append("                                   ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │                  GUARDRAILS POLICY ENGINE                 │")
    lines.append("       │  • Evaluates Guardrail Rules (CEL, Prompt Repo scopes)    │")
    lines.append("       │  • Executes Guardrail Providers (Regex, Presidio, Lakera) │")
    lines.append("       └─────────────┬───────────────────────────────┬─────────────┘")
    lines.append("        (Violates)   │                               │ (Clean / Pass)")
    lines.append("                     ▼                               ▼")
    lines.append("       ┌───────────────────────────┐   ┌───────────────────────────┐")
    lines.append("       │ HTTP 400 / 403 Security   │   │   ADAPTIVE LOAD BALANCER  │")
    lines.append("       │ Violation (PII / Attack)  │   │  • Multi-Armed Bandit     │")
    lines.append("       └───────────────────────────┘   │  • Real-time Health Scores│")
    lines.append("                                       │  • Latency & Error Weight │")
    lines.append("                                       └─────────────┬─────────────┘")
    lines.append("                                                     │ (Fastest Direction)")
    lines.append("                                                     ▼")
    lines.append("                                       ┌───────────────────────────┐")
    lines.append("                                       │  UPSTREAM MODEL PROVIDER  │")
    lines.append("                                       │  • OpenAI / Bedrock / GCP │")
    lines.append("                                       └─────────────┬─────────────┘")
    lines.append("                                                     │")
    lines.append("                                                     ▼ (Response Stream)")
    lines.append("                                       ┌───────────────────────────┐")
    lines.append("                                       │   OUTBOUND GUARDRAILS     │")
    lines.append("                                       │  • Hallucination / Secrets│")
    lines.append("                                       └─────────────┬─────────────┘")
    lines.append("                                                     │")
    lines.append("                                                     ▼")
    lines.append("                                            [ CLIENT RESPONSE ]")
    lines.append("```")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 2. Detailed Feature Dissection (5 Core Features)")
    lines.append("### Ainthu Features-oda Aazhamana Vivaram (Thanglish + English)")
    lines.append("")

    for f in GUARDRAIL_FEATURES:
        lines.append(f"<a name='{f['id']}'></a>")
        lines.append(f"## {f['name_en']} — {f['name_tanglish']}")
        lines.append(f"**UI Route:** `{f['route']}`\n")
        lines.append("### 👤 Non-Tech Perspective (Makkal Puriyura Mathiri Eliya Vilakkam)")
        lines.append(f"- **Uruvagam (Analogy):** {f['analogy']}")
        lines.append(f"- **Vilakkam (Explanation):** {f['explanation']}")
        lines.append(f"- **Business Value (Vaniga Payan):** {f['business_value']}\n")
        lines.append("### 💻 Tech Perspective (Engineers-kaga Technical Architecture)")
        lines.append(f"- **Backend Architecture:** {f['tech_arch']}")
        lines.append("- **Backend Endpoints:**")
        for ep in f['endpoints']:
            lines.append(f"  * `{ep}`")
        lines.append("")
        lines.append("### 🖥️ Screen Layout & Interactive Elements (Thirai Koorugal & Buttons)")
        lines.append("**1. Melpura Koorugal (Top Bar Controls):**")
        for itm in f['ui_elements']['top_bar']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("**2. Mathiya Koorugal & Tables (Tabs & Views):**")
        for itm in f['ui_elements']['tabs_and_views']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("**3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):**")
        for itm in f['ui_elements']['bottom_elements']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("### 🔗 Connections & Structure Map (Inaippugal)")
        lines.append(f"- **Data-vai Perumidam (Receives From):** {f['connections']['receives_from']}")
        lines.append(f"- **Iyakkum Koorugal (Triggers & Affects):** {f['connections']['triggers_and_affects']}")
        lines.append("")
        lines.append(f"### 💡 Production Use Case: {f['use_case']}\n")
        lines.append("---\n")

    lines.append("# 3. Cross-Feature Interconnections & Data Flow")
    lines.append("### Guardrails, Cluster & Adaptive Routing Kulla Irukura Direct Connections\n")
    lines.append("| Source Feature | Connected To | Data Flow & Trigger Action |")
    lines.append("| :--- | :--- | :--- |")
    for src, dst, flow in GUARDRAIL_CONN_DATA:
        lines.append(f"| **{src}** | **{dst}** | {flow} |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 4. Tech vs Non-Tech Comparative Matrix")
    lines.append("### Thozhilnutpam vs Vanigam Parvai Oppeedo\n")
    lines.append("| Feature | Non-Tech View (Manager / CFO Parvai) | Tech View (DevOps / Architect Parvai) |")
    lines.append("| :--- | :--- | :--- |")
    for ft, nt, tv in GUARDRAIL_MATRIX_DATA:
        lines.append(f"| **{ft}** | \"{nt}\" | \"{tv}\" |")
    lines.append("")

    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Guardrails & Adaptive Routing Markdown written to: {MD_PATH}")

# ==============================================================================
# 2. GENERATE DOCX DOCUMENT (THANGLISH + ENGLISH)
# ==============================================================================
def generate_guardrail_docx():
    print("Writing Guardrails, Cluster & Adaptive Routing Word Document (.docx)...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        hp = section.header.paragraphs[0]
        hp.text = "UnifAI Guardrails, Cluster & Adaptive Routing Deep-Dive Manual"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if hp.runs:
            hp.runs[0].font.name = "Segoe UI"
            hp.runs[0].font.size = Pt(8.5)
            hp.runs[0].font.color.rgb = RGBColor(100, 116, 139)
            
        fp = section.footer.paragraphs[0]
        fp.text = "Confidential & Proprietary — UnifAI Security & Performance Control Plane (Thanglish + English)"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if fp.runs:
            fp.runs[0].font.name = "Segoe UI"
            fp.runs[0].font.size = Pt(8.5)
            fp.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    def style_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        return p

    def style_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def style_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.italic = italic
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def format_table(table, col_widths, col_names):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(col_names):
            hdr_cells[i].text = name
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
            tcPr.append(shd)
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = "Segoe UI"
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for r_idx, row in enumerate(table.rows[1:]):
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                tcPr.append(shd)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                    
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    rt = p_title.add_run("UnifAI Guardrails, Cluster & Adaptive Routing Guide")
    rt.font.name = "Segoe UI"
    rt.font.size = Pt(22)
    rt.bold = True
    rt.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    rsub = p_sub.add_run("Guardrail Rules, Providers, Cluster Config, Adaptive Routing Dashboard & Settings (Thanglish + English)")
    rsub.font.name = "Segoe UI"
    rsub.font.size = Pt(11.5)
    rsub.font.color.rgb = RGBColor(71, 85, 105)

    style_h1("1. System Architecture & Traffic Flow Map")
    add_p("Guardrail safety filters, Cluster state synchronization (port 7946), matrum Adaptive Routing Multi-Armed Bandit steering nadakkura dynamic workflow:")

    style_h1("2. Detailed Feature Dissection (5 Core Features)")
    for f in GUARDRAIL_FEATURES:
        style_h2(f"{f['name_en']} — {f['name_tanglish']}")
        add_p(f['route'], bold_prefix="Route: ", italic=True)

        style_h3("Non-Tech Perspective (Eliya Vilakkam)")
        add_p(f['analogy'], bold_prefix="Uruvagam (Analogy): ")
        add_p(f['explanation'], bold_prefix="Vilakkam: ")
        add_p(f['business_value'], bold_prefix="Business Value: ")

        style_h3("Tech Perspective (Technical Architecture)")
        add_p(f['tech_arch'], bold_prefix="Backend Architecture: ")
        add_p("Backend Endpoints:")
        for ep in f['endpoints']:
            add_bullet(ep)

        style_h3("Screen Layout & Interactive Elements (Thirai Koorugal)")
        add_p("1. Melpura Koorugal (Top Bar Controls):", bold_prefix="")
        for itm in f['ui_elements']['top_bar']:
            add_bullet(itm)
        add_p("2. Mathiya Koorugal & Tables (Tabs & Views):", bold_prefix="")
        for itm in f['ui_elements']['tabs_and_views']:
            add_bullet(itm)
        add_p("3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):", bold_prefix="")
        for itm in f['ui_elements']['bottom_elements']:
            add_bullet(itm)

        style_h3("Module Connections (Inaippugal)")
        add_p(f['connections']['receives_from'], bold_prefix="Data-vai Perumidam: ")
        add_p(f['connections']['triggers_and_affects'], bold_prefix="Iyakkum Koorugal: ")

        style_h3("Production Use Case (Nadamurai Udharanam)")
        add_p(f['use_case'])

    style_h1("3. Cross-Feature Interconnections")
    t_conn = doc.add_table(rows=len(GUARDRAIL_CONN_DATA)+1, cols=3)
    format_table(t_conn, [1.8, 1.8, 3.4], ["Source Feature", "Connected To", "Data Flow & Trigger Action"])
    for idx, (src, dst, flow) in enumerate(GUARDRAIL_CONN_DATA):
        t_conn.rows[idx+1].cells[0].text = src
        t_conn.rows[idx+1].cells[1].text = dst
        t_conn.rows[idx+1].cells[2].text = flow

    style_h1("4. Tech vs Non-Tech Comparative Matrix")
    t_mat = doc.add_table(rows=len(GUARDRAIL_MATRIX_DATA)+1, cols=3)
    format_table(t_mat, [1.5, 2.7, 2.8], ["Feature", "Non-Tech View (Manager / CFO)", "Tech View (DevOps / Architect)"])
    for idx, (ft, nt, tv) in enumerate(GUARDRAIL_MATRIX_DATA):
        t_mat.rows[idx+1].cells[0].text = ft
        t_mat.rows[idx+1].cells[1].text = nt
        t_mat.rows[idx+1].cells[2].text = tv

    doc.save(DOCX_PATH)
    print(f"Guardrails & Adaptive Routing Word document written to: {DOCX_PATH}")

# ==============================================================================
# 3. GENERATE PDF DOCUMENT VIA REPORTLAB (THANGLISH + ENGLISH)
# ==============================================================================
def generate_guardrail_pdf():
    print("Writing Guardrails, Cluster & Adaptive Routing PDF Document (.pdf)...")

    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748B"))
            
            if self._pageNumber > 1:
                self.drawString(40, 11 * inch - 36, "UnifAI Guardrails, Cluster & Adaptive Routing Deep-Dive Manual (Thanglish + English)")
                self.drawRightString(8.5 * inch - 40, 11 * inch - 36, "Confidential — Engineering Guide")
                self.setStrokeColor(colors.HexColor("#CBD5E1"))
                self.setLineWidth(0.5)
                self.line(40, 11 * inch - 40, 8.5 * inch - 40, 11 * inch - 40)
                
            self.drawString(40, 32, "UnifAI Unified Security Guardrails & Performance Control Plane")
            page_str = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(8.5 * inch - 40, 32, page_str)
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(40, 42, 8.5 * inch - 40, 42)
            
            self.restoreState()

    pdf_doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=48,
        bottomMargin=48
    )

    styles = getSampleStyleSheet()

    p_title_style = ParagraphStyle(
        'PdfTitle',
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=3
    )

    p_subtitle_style = ParagraphStyle(
        'PdfSubTitle',
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#475569'),
        spaceAfter=10
    )

    h1_style = ParagraphStyle(
        'PdfH1',
        fontName='Helvetica-Bold',
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=12,
        spaceAfter=5,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'PdfH2',
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'PdfH3',
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=5,
        spaceAfter=2,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'PdfBody',
        fontName='Helvetica',
        fontSize=8,
        leading=11,
        textColor=colors.HexColor('#334155'),
        spaceAfter=3
    )

    bullet_style = ParagraphStyle(
        'PdfBullet',
        fontName='Helvetica',
        fontSize=7.8,
        leading=10.5,
        textColor=colors.HexColor('#334155'),
        leftIndent=12,
        spaceAfter=2
    )

    th_style = ParagraphStyle(
        'PdfTH',
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.white
    )

    td_style = ParagraphStyle(
        'PdfTD',
        fontName='Helvetica',
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor('#0F172A')
    )

    td_code_style = ParagraphStyle(
        'PdfTDCode',
        fontName='Helvetica-Bold',
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor('#1E3A8A')
    )

    story = []

    # Title & Subtitle
    story.append(Paragraph("UnifAI Guardrails, Cluster & Adaptive Routing Guide", p_title_style))
    story.append(Paragraph("Security Rules, Providers, Mesh Cluster Sync & Dynamic Adaptive Routing (Thanglish + English)", p_subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#2563EB'), spaceBefore=2, spaceAfter=8))

    # Architecture Overview
    story.append(Paragraph("1. System Architecture & Traffic Flow Map", h1_style))
    story.append(Paragraph(
        "This manual covers UnifAI's runtime security defense and dynamic performance steering. "
        "<b>Cluster Config</b> synchronizes rate-limit token buckets and circuit breaker trip signals across distributed nodes via port 7946 gossip. "
        "Every client request is screened by <b>Guardrail Rules</b> and executed against <b>Guardrail Providers</b> (Presidio, RE2 regexes, Bedrock) "
        "to prevent PII leakage and prompt injection attacks. Clean requests are dynamically steered by the <b>Adaptive Routing Dashboard</b> "
        "using a reinforcement learning Multi-Armed Bandit model to the fastest, lowest-latency upstream model, governed by <b>Adaptive Routing Settings</b>.",
        body_style
    ))
    story.append(Spacer(1, 4))

    # Features Loop
    story.append(Paragraph("2. Detailed Feature Dissection (5 Core Features)", h1_style))
    for f in GUARDRAIL_FEATURES:
        feat_elements = []
        feat_elements.append(Paragraph(f"<b>{f['name_en']} — {f['name_tanglish']}</b>", h2_style))
        feat_elements.append(Paragraph(f"<b>Route:</b> <font color='#2563EB'>{f['route']}</font>", body_style))

        feat_elements.append(Paragraph("<b>👤 Non-Tech Perspective (Eliya Vilakkam):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Uruvagam (Analogy):</b> {f['analogy']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Vilakkam:</b> {f['explanation']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Business Value:</b> {f['business_value']}", body_style))

        feat_elements.append(Paragraph("<b>💻 Tech Perspective (Technical Architecture):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Backend Architecture:</b> {f['tech_arch']}", body_style))
        ep_text = ", ".join([f"<code>{e}</code>" for e in f['endpoints']])
        feat_elements.append(Paragraph(f"• <b>Endpoints:</b> {ep_text}", body_style))

        feat_elements.append(Paragraph("<b>🖥️ Screen Layout & Interactive Elements (Thirai Koorugal):</b>", h3_style))
        feat_elements.append(Paragraph("<b>1. Melpura Koorugal (Top Bar Controls):</b>", body_style))
        for itm in f['ui_elements']['top_bar']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))
        feat_elements.append(Paragraph("<b>2. Mathiya Koorugal & Tables (Tabs & Views):</b>", body_style))
        for itm in f['ui_elements']['tabs_and_views']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))
        feat_elements.append(Paragraph("<b>3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):</b>", body_style))
        for itm in f['ui_elements']['bottom_elements']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))

        feat_elements.append(Paragraph("<b>🔗 Connections & Structure Map (Inaippugal):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Data-vai Perumidam:</b> {f['connections']['receives_from']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Iyakkum Koorugal:</b> {f['connections']['triggers_and_affects']}", body_style))

        feat_elements.append(Paragraph(f"<b>💡 Production Use Case:</b> {f['use_case']}", body_style))
        feat_elements.append(Spacer(1, 6))
        story.append(KeepTogether(feat_elements))

    # Cross Connections
    story.append(Paragraph("3. Cross-Feature Interconnections & Data Flow", h1_style))
    conn_table_data = [[Paragraph(h, th_style) for h in ["Source Feature", "Connected To", "Data Flow & Trigger Action"]]]
    for src, dst, flow in GUARDRAIL_CONN_DATA:
        conn_table_data.append([Paragraph(src, td_code_style), Paragraph(dst, td_code_style), Paragraph(flow, td_style)])
    tbl_conn = Table(conn_table_data, colWidths=[105, 105, 320])
    tbl_conn.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_conn)
    story.append(Spacer(1, 8))

    # Comparative Matrix
    story.append(Paragraph("4. Tech vs Non-Tech Comparative Matrix", h1_style))
    mat_table_data = [[Paragraph(h, th_style) for h in ["Feature", "Non-Tech View (Manager / CFO Parvai)", "Tech View (DevOps / Architect Parvai)"]]]
    for ft, nt, tv in GUARDRAIL_MATRIX_DATA:
        mat_table_data.append([Paragraph(ft, td_code_style), Paragraph(nt, td_style), Paragraph(tv, td_style)])
    tbl_mat = Table(mat_table_data, colWidths=[95, 215, 220])
    tbl_mat.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_mat)

    pdf_doc.build(story, canvasmaker=NumberedCanvas)
    print(f"Guardrails & Adaptive Routing PDF document written to: {PDF_PATH}")

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================
if __name__ == "__main__":
    print("Building dedicated Guardrails, Cluster & Adaptive Routing Documentation (Thanglish + English)...")
    generate_guardrail_markdown()
    generate_guardrail_docx()
    generate_guardrail_pdf()
    print("Guardrails & Adaptive Routing Documentation successfully built in:", DOC_DIR)
