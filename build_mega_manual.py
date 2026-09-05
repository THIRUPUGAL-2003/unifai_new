# -*- coding: utf-8 -*-
"""
UnifAI Enterprise Deep-Dive Architecture & Master Feature Guide Generator
Produces an exhaustive, multi-page technical manual covering every single
feature, architecture layer, HTTP header, tech stack, and production scenario.
"""

import os
import sys

# Ensure UTF-8 output encoding for Windows console
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

DOC_DIR = os.path.join(os.getcwd(), "document")
os.makedirs(DOC_DIR, exist_ok=True)

MD_PATH = os.path.join(DOC_DIR, "UnifAI_Architecture_and_Features_Guide.md")
DOCX_PATH = os.path.join(DOC_DIR, "UnifAI_Architecture_and_Features_Guide.docx")
PDF_PATH = os.path.join(DOC_DIR, "UnifAI_Architecture_and_Features_Guide.pdf")

# Register fonts for Tamil + English in ReportLab
pdfmetrics.registerFont(TTFont('Latha', 'C:/Windows/Fonts/latha.ttf'))
pdfmetrics.registerFont(TTFont('LathaBold', 'C:/Windows/Fonts/lathab.ttf'))

# ==============================================================================
# DATA DEFINITIONS: EXHAUSTIVE SPECIFICATION OF ALL FEATURES & HEADERS
# ==============================================================================

HEADERS_DATA = [
    {
        "category": "Authentication & Identity Headers (அடையாளம் & பாதுகாப்பு)",
        "headers": [
            ("x-uf-vk", "sk-uf-prod_corp_a91b2c3d4e", "String (Virtual Key)",
             "UnifAI Virtual Key representing synthetic API credentials. Binds the caller to an authorized Team, User, Budget Quota, Access Profile, and MCP Tool Groups.",
             "Virtual Key மூலமாக client அங்கீகரிக்கப்படுகிறது. Upstream provider keys மறைக்கப்பட்டு, பட்ஜெட் மற்றும் அனுமதிகள் இதனுடன் இணைக்கப்படும்."),
            ("Authorization", "Bearer sk-uf-prod_corp_a91b2c3d4e", "Standard Header",
             "Standard HTTP Authorization header alternative for OpenAI SDK compatibility. Automatically parsed as a Virtual Key if prefixed with 'sk-uf-'.",
             "OpenAI SDK பயன்படுத்தும்போது 'x-uf-vk'-க்கு மாற்றாக standard Bearer token வடிவத்தில் Virtual Key-ஐ அனுப்பலாம்."),
            ("x-uf-api-key", "key_adm_89f0a21bc9e4", "String (Admin Key)",
             "Platform Control Plane Administrative API Key. Authorizes access to administrative CRUD endpoints (creating keys, editing routing rules, cluster sync).",
             "UnifAI நிர்வாக API-களை அணுக பயன்படும் Platform Admin Key. புதிய ரூல்ஸ், மாடல்கள் மற்றும் கீகளை உருவாக்க உதவும்."),
            ("x-uf-customer-id", "cust_enterprise_7710", "String (UUID/Slug)",
             "Tenant / Customer scoping key for B2B multi-tenant SaaS applications. Routes usage, token consumption, and dollar spend to a specific end-customer account under a shared team virtual key.",
             "B2B SaaS செயலிகளில், எந்த இறுதி வாடிக்கையாளர் (End Customer) இந்த request-ஐ அனுப்புகிறார் என பிரித்து Cost & Analytics track செய்ய உதவும்."),
            ("x-uf-customer-name", "Acme Corporation Global", "String (Human Name)",
             "Display name for the customer tenant, populated in audit records, billing statements, and telemetry dashboards alongside x-uf-customer-id.",
             "வாடிக்கையாளரின் பெயர். Audit logs மற்றும் பில்லிங் அறிக்கைகளில் எளிதாக அடையாளம் காண பயன்படுகிறது."),
            ("x-uf-direct-key", "true / false", "Boolean Flag",
             "Direct Key Pass-Through. When set to 'true', UnifAI bypasses its internal key pool and uses the raw provider key supplied in standard headers (e.g. x-api-key or x-goog-api-key).",
             "UnifAI-ல் சேமிக்கப்பட்ட key-ஐ பயன்படுத்தாமல், Client தன்னிடம் உள்ள raw OpenAI/Anthropic key-ஐ நேரடியாக பாஸ் செய்ய உதவும்."),
            ("X-UnifAI-Temp-Token", "jwt_or_uuid_token", "Short-lived Session Token",
             "Ephemeral temporary session token used during interactive OAuth consent flows, MCP tool authorization callbacks, and restricted UI views.",
             "MCP டூல்ஸ் மற்றும் OAuth login callbacks-ன் போது தற்காலிக அங்கீகாரத்திற்காக பயன்படுத்தப்படும் short-lived டோக்கன்.")
        ]
    },
    {
        "category": "Intelligent Routing & Model Steering Headers (ரூட்டிங் & மாடல் தேர்வு)",
        "headers": [
            ("x-uf-provider", "openai | anthropic | bedrock | vertex | groq | ollama", "Provider Enum",
             "Explicit Provider Pinning. Overrides automated routing rules and forces UnifAI to dispatch the request directly to the specified upstream model provider.",
             "தானியங்கி ரூட்டிங் விதிகளை தவிர்த்து, குறிப்பிட்ட ஒரு AI Provider-க்கு (Ex: Anthropic அல்லது Bedrock) நேரடியாக கோரிக்கையை அனுப்ப உதவும்."),
            ("x-uf-model", "claude-3-7-sonnet | gpt-4o | deepseek-r1", "Model ID",
             "Explicit Model Selection. Overrides the body 'model' parameter or acts as the primary model target when using native SDK conversions.",
             "கோரிக்கை இயக்கப்பட வேண்டிய மாடலின் பெயரை நேரடியாக நிர்ணயிக்க பயன்படுகிறது."),
            ("x-uf-api-key-id", "key_azure_eastus_02", "Credential Pool ID",
             "Target Credential Pinning. When a single provider has multiple keys (e.g., 5 Azure OpenAI endpoints in different regions), pins execution to a specific key ID.",
             "ஒரே provider-ல் பல API keys (multi-region accounts) இருக்கும்போது, குறிப்பிட்ட ஒரு key-ஐ மட்டும் தேர்வு செய்ய உதவும்."),
            ("x-uf-circuit-breaker-bypass", "true", "Diagnostic Flag",
             "Administrative Circuit Breaker Bypass. Forces the gateway to attempt dispatch to an upstream provider even if its circuit breaker is currently in the TRIPPED (OPEN) state.",
             "Circuit Breaker open-ஆக (tripped) இருந்தாலும், சோதனைகளுக்காக அந்த provider-க்கு கட்டாயமாக request அனுப்ப உதவும்.")
        ]
    },
    {
        "category": "Semantic Caching & Performance Headers (கேச்சிங் & லேடன்சி குறைப்பு)",
        "headers": [
            ("x-uf-cache-key", "tenant_alpha:customer_support:v2", "Namespace String",
             "Cache Partition Namespace. Segregates cached response vectors into isolated buckets to guarantee complete data isolation between tenants, departments, or application versions.",
             "கேச் தரவை தனித்தனி பாகங்களாக (partitions) பிரிக்க உதவுகிறது. வெவ்வேறு வாடிக்கையாளர்களின் டேட்டா ஒன்றோடொன்று கலக்காமல் பாதுகாக்கும்."),
            ("x-uf-cache-ttl", "3600 (seconds) | 86400", "Integer (Seconds)",
             "Per-Request Cache Time-To-Live. Overrides the global cache TTL to specify exactly how many seconds this response should remain valid in the cache.",
             "இந்த குறிப்பிட்ட பதில் எத்தனை வினாடிகள் கேச்சில் சேமிக்கப்பட்டிருக்க வேண்டும் என்பதை நிர்ணயிக்கிறது."),
            ("x-uf-cache-threshold", "0.92 (range: 0.0 to 1.0)", "Float Score",
             "Cosine Similarity Threshold Override. Sets the minimum vector similarity required to consider a stored prompt a 'Cache Hit'. Higher = more accurate; Lower = more hits.",
             "Semantic cache-ல் இரண்டு கேள்விகள் எத்தனை சதவீதம் ஒத்துப்போக வேண்டும் (Cosine similarity) என்பதை மாற்ற உதவும்."),
            ("x-uf-cache-type", "direct | semantic", "Cache Mode Enum",
             "'direct' enforces exact SHA-256 prompt string matching. 'semantic' generates an embedding vector and performs approximate nearest neighbor (ANN) search.",
             "'direct' என்பது 100% வார்த்தை பொருத்தம்; 'semantic' என்பது அர்த்த ரீதியான பொருத்தம் (Embedding Vector similarity)."),
            ("x-uf-cache-no-store", "true / false", "Boolean Flag",
             "No-Store Directive. When set to 'true', UnifAI will serve a cached response if a hit exists, but will NOT store the new LLM completion in the cache upon a miss.",
             "ஏற்கனவே கேச்சில் இருந்தால் பதிலை எடு, ஆனால் புதிய பதிலை கேச்சில் சேமிக்காதே என்ற உத்தரவு.")
        ]
    },
    {
        "category": "MCP (Model Context Protocol) & Tool Headers (டூல்ஸ் பயன்பாடு)",
        "headers": [
            ("x-uf-mcp-include-clients", "* | github,postgres,slack", "Comma-separated List",
             "MCP Server Whitelist. Injects only the specified MCP servers into the LLM request. Setting '*' exposes all servers permitted by the caller's Access Profile.",
             "எந்தெந்த MCP servers-ன் டூல்ஸ்களை மாடலுக்கு வழங்க வேண்டும் என கட்டுப்படுத்தும் whitelist (Ex: github, postgres)."),
            ("x-uf-mcp-include-tools", "query_database,create_pull_request", "Comma-separated List",
             "MCP Tool Filter. Restricts the exposed tools to a specific subset of function names within the permitted MCP servers.",
             "சர்வரில் உள்ள டூல்ஸ்களில் குறிப்பிட்ட சில functions-ஐ மட்டும் மாடலுக்கு வெளிப்படுத்த உதவும்."),
            ("x-uf-mcp-session-id", "sess_oauth_usr_8821a", "Session Identifier",
             "Per-User Tool Session Key. Binds the tool execution to a human user's personal OAuth session (e.g. user's personal GitHub token rather than a system token).",
             "பயனரின் தனிப்பட்ட OAuth credentials மூலம் MCP tools-ஐ இயக்க பயன்படும் Session ID."),
            ("x-uf-eh-*", "x-uf-eh-workspace-id: ws_991", "Prefixed Passthrough",
             "Extra Tool Headers. Any header matching the 'x-uf-eh-*' prefix is forwarded directly to upstream MCP servers if included in their allowed_extra_headers configuration.",
             "MCP சர்வருக்கு தேவைப்படும் தனிப்பயன் headers-ஐ பாதுகாப்பாக forward செய்ய உதவும் prefix.")
        ]
    },
    {
        "category": "Observability, Tracing & Debugging Headers (கண்காணிப்பு & பிழைத்திருத்தம்)",
        "headers": [
            ("x-uf-session-id", "sess_chat_conv_12345", "String (Session ID)",
             "Conversation Session Grouping. Groups multiple sequential LLM calls into a single unified trace in OpenTelemetry and LLM Logs.",
             "ஒரு பயனரின் தொடர்ச்சியான சாட் உரையாடல்களை ஒரே Trace-ஆக இணைக்க பயன்படும் Session ID."),
            ("x-uf-dim-<key>", "x-uf-dim-environment: production", "Arbitrary Key-Value",
             "Custom Log Dimension. Injects custom dimensional metadata into internal log records and exported connector events (e.g., department, tier, feature).",
             "Logs-ல் custom filters சேர்க்க பயன்படும் metadata (Ex: x-uf-dim-app: mobile-app)."),
            ("x-uf-lh-<header>", "x-uf-lh-correlation-id: req_091", "Header Capture",
             "Captured Request Header. Explicitly instructs UnifAI's logging engine to record the value of this header into the LLM log entry.",
             "குறிப்பிட்ட ஒரு request header-ஐ logs-ல் பதிவு செய்ய கட்டளையிடும் header."),
            ("x-uf-disable-content-logging", "true / false", "Privacy Flag",
             "Content Privacy Mode. Completely drops prompt and completion text from log storage, recording only token counts, duration, latency, cost, and metadata.",
             "ரகசியத்தன்மையை பாதுகாக்க, prompt மற்றும் பதிலை logs-ல் சேமிக்காமல், வெறும் செலவு மற்றும் வேகத்தை மட்டும் பதிவு செய்ய உதவும்."),
            ("x-uf-store-raw-request-response", "true / false", "Debug Capture",
             "Raw Wire Byte Storage. Instructs the gateway to persist the exact raw JSON byte stream exchanged with the upstream provider for low-level debugging.",
             "Upstream provider-க்கு அனுப்பிய அசல் raw bytes-ஐ logs-ல் சேமிக்க உதவும் debug toggle."),
            ("x-uf-send-back-raw-request", "true / false", "Echo Header",
             "Debug Echo. Injects the converted raw request payload sent to the upstream provider into a debugging HTTP response header.",
             "Gateway மாற்றி அனுப்பிய raw payload-ஐ response header-ல் திரும்பப் பெற உதவும்."),
            ("x-uf-send-back-raw-response", "true / false", "Echo Header",
             "Debug Echo. Injects the unparsed upstream response payload into a debugging HTTP response header.",
             "Upstream provider தந்த அசல் பதிலை response header-ல் திரும்பப் பெற உதவும்."),
            ("traceparent", "00-4bf92f3577b34da6a3ce929d0e0e4736-00f067aa0ba902b7-01", "W3C Trace Context",
             "Distributed Tracing Header. Propagates distributed OpenTelemetry trace context from calling microservices through UnifAI to downstream connectors.",
             "Microservices வழியே OpenTelemetry Distributed Tracing-ஐ தொடர்ந்து கண்காணிக்க உதவும் W3C standard header."),
            ("x-uf-log-repo-id", "repo_maxim_eval_01", "External Repo ID",
             "Evaluation Routing. Directs the request's telemetry log directly to a designated evaluation repository in platforms like Maxim AI.",
             "Maxim AI போன்ற வெளி evaluation தளங்களில் குறிப்பிட்ட repository-க்கு logs-ஐ அனுப்ப உதவும்.")
        ]
    },
    {
        "category": "Gateway Response Headers (UnifAI திருப்பி அனுப்பும் Headers)",
        "headers": [
            ("x-unifai-provider", "anthropic | openai | bedrock", "Provider String",
             "Indicates which model provider actually executed and served the request.",
             "எந்த AI Provider கோரிக்கையை நிறைவேற்றியது என்பதை வாடிக்கையாளருக்கு தெரிவிக்கும்."),
            ("x-unifai-original-model", "gpt-4o", "Model Identifier",
             "The model originally requested by the client application prior to routing.",
             "வாடிக்கையாளர் கேட்ட அசல் மாடலின் பெயர்."),
            ("x-unifai-resolved-model", "claude-3-5-sonnet", "Model Identifier",
             "The actual model architecture that handled the prompt following routing rule or complexity tier resolution.",
             "Routing விதிகள் மற்றும் complexity ஆய்வுக்குப் பிறகு உண்மையில் இயங்கிய மாடல்."),
            ("x-unifai-fallback-index", "0 (primary) | 1, 2... (fallback)", "Integer Index",
             "Indicates whether a fallback route was taken. '0' denotes primary provider success; non-zero indicates which fallback triggered.",
             "Primary provider தோல்வியடைந்து Fallback provider இயங்கியதா என்பதை காட்டும் (0 = No fallback)."),
            ("x-unifai-request-type", "chat | completion | embedding | realtime", "Type String",
             "Identifies the execution modality: chat completions, text completions, embeddings, or realtime voice.",
             "கோரிக்கையின் வகை: chat, embeddings, realtime voice, etc.")
        ]
    }
]

# ==============================================================================
# PILLARS DATA: DEEP DIVE INTO EVERY SINGLE FEATURE
# ==============================================================================

PILLARS_DATA = [
    {
        "pillar_id": "pillar_1_observability",
        "pillar_title": "Pillar 1: Observability (முழுமையான கண்காணிப்பு & செயல்திறன்)",
        "pillar_desc": "The Observability pillar provides sub-millisecond telemetry, cost accounting, request logging, and external data connectors across all AI workloads.",
        "features": [
            {
                "name": "Dashboard",
                "route": "/workspace/dashboard",
                "summary": "Real-time executive analytics console displaying financial, throughput, and performance metrics across models and virtual keys.",
                "architecture": (
                    "The Dashboard is powered by a high-frequency WebSocket connection (useWebSocket) paired with TanStack Query. "
                    "In the backend, metrics are aggregated in-memory using atomic counters and persisted to time-series Postgres tables. "
                    "It calculates p50, p90, p95, and p99 latency percentiles, tracks dollar spend using live Pricing Overrides, "
                    "and measures cache hit ratios across all active tenant nodes."
                ),
                "components": [
                    "Time Range Filter: 1h, 24h, 7d, 30d, and custom calendar windowing.",
                    "Spend Gauges: Total aggregate cost ($), daily burn rate, and projected monthly invoice.",
                    "Latency Percentiles: p50, p95, p99 curves showing TTFT (Time To First Token) and total latency.",
                    "Breakdown Views: Bar and donut charts slicing traffic by Provider, Model, and Virtual Key.",
                    "Cache Efficiency Meter: Real-time calculation of cost saved via Semantic Cache hits."
                ],
                "interconnections": "Receives aggregated telemetry from LLM Logs, Pricing Overrides, and Semantic Caching. Influences Budgets & Limits alert triggers.",
                "use_case": "Engineering leaders and FinOps teams monitor daily token burn, identify slow endpoints (p99 > 3000ms), and detect abnormal traffic surges before budget overruns occur."
            },
            {
                "name": "LLM Logs",
                "route": "/workspace/logs",
                "summary": "Full-fidelity request and response transaction ledger capturing complete prompts, completions, tokens, latency, and guardrail verdicts.",
                "architecture": (
                    "Requests passing through FastHTTP are buffered asynchronously into an in-memory queue (logstore) to guarantee zero impact on proxy latency. "
                    "A dedicated background worker pool writes records in micro-batches to PostgreSQL with pgx v5. "
                    "Each log entry captures request_id, session_id, virtual_key_id, provider, model requested vs resolved, prompt_tokens, "
                    "completion_tokens, exact dollar cost, latency_ms, ttft_ms, HTTP status, and full JSON payloads."
                ),
                "components": [
                    "Full-Text & Regex Search: Search across user prompts and model completions.",
                    "Multidimensional Filtering: Filter by Virtual Key, Customer ID, Provider, Model, Status Code, and Date Range.",
                    "Content Privacy Control: Supports x-uf-disable-content-logging to omit sensitive text while preserving metrics.",
                    "Raw Byte Inspector: Inspects exact wire JSON exchanged with upstream providers.",
                    "Guardrail Annotation: Visual flags indicating whether input or output triggered PII redaction or safety rules."
                ],
                "interconnections": "Fed by every LLM invocation; powers the Dashboard metrics; streams out to Connectors (Datadog, Kafka); validates Guardrails accuracy.",
                "use_case": "Developers troubleshoot failed user requests, investigate hallucinated model answers, verify TTFT performance, and export compliance audit logs."
            },
            {
                "name": "MCP Logs",
                "route": "/workspace/mcp-logs",
                "summary": "Granular execution audit log tracking Model Context Protocol tool invocations, arguments, return payloads, and runtime errors.",
                "architecture": (
                    "Located in 'core/mcp/exec.go', the MCP interceptor instruments every tool execution. "
                    "It records the calling agent's identity, the target MCP server, function name, parsed input JSON arguments, "
                    "execution duration in milliseconds, and the exact tool output or error stack trace. "
                    "Supports both stdio and SSE transport protocols."
                ),
                "components": [
                    "Tool Call Inspector: View structured input parameters and returned output schemas.",
                    "Execution Time Tracking: Measure tool execution latency to identify slow database queries or API scrapers.",
                    "Error Diagnostics: Full exception stack traces for crashed tool processes or timeout breaches.",
                    "Session Binding: Correlates tool calls with specific User Auth Sessions and Virtual Keys."
                ],
                "interconnections": "Tied directly to the MCP Gateway, Tool Groups, and Auth Sessions. Streams tool performance into Dashboard.",
                "use_case": "AI Agent engineers diagnose why an autonomous agent entered an infinite tool loop, verify SQL queries executed by the Database MCP tool, and track 3rd-party API failures."
            },
            {
                "name": "Browser AI",
                "route": "/workspace/browser-ai",
                "summary": "Enterprise Browser Extension and local DLP proxy that intercepts employee AI queries on public web portals (ChatGPT, Claude) to enforce corporate data safety.",
                "architecture": (
                    "Consists of two synchronized layers: (1) A local proxy daemon (browser_ai_proxy.py) with browser extensions deployed via MDM, "
                    "and (2) a centralized control console in the UnifAI UI. "
                    "It intercepts HTTPS requests to web AI portals, runs Presidio DLP and local Ollama classification models, "
                    "and blocks the transmission of proprietary source code, credentials, or customer PII."
                ),
                "components": [
                    "Target Websites Registry: Manages parent and sub-domain policies (e.g. chatgpt.com, claude.ai, perplexity.ai).",
                    "DLP Rule Engine: Regex-based and LLM-assisted policy enforcement on browser paste events.",
                    "Attachment Text Extractor: Inspects and extracts text from uploaded PDF, Word, and text attachments.",
                    "Tamper-Proof Agent Management: Heartbeat tracking, bulk agent deletion, and uninstall key protection.",
                    "Violation Logs: Dedicated log tracking employee DLP violations with source IP, user email, and matched policy."
                ],
                "interconnections": "Shares Guardrail Providers (Presidio, Regex) and routes violation telemetry into Audit Logs and Connectors.",
                "use_case": "Prevents corporate data leakage by stopping engineers from pasting confidential customer databases or private cryptographic keys into public web ChatGPT."
            },
            {
                "name": "Connectors",
                "route": "/workspace/observability",
                "summary": "Real-time streaming export pipelines forwarding UnifAI metrics, logs, and trace spans to enterprise observability platforms.",
                "architecture": (
                    "Implemented in 'framework/connectors/runtime.go'. It operates as an asynchronous worker ring buffer that decouples export operations "
                    "from the gateway request path. It converts internal telemetry events into platform-native formats and delivers them with automatic retries, "
                    "exponential backoff, and circuit-breaking."
                ),
                "components": [
                    "Datadog Connector: Exports APM trace spans and metrics (unifai.requests.count, unifai.tokens.total, unifai.latency).",
                    "New Relic Connector: Streams custom events and transaction metrics for enterprise APM.",
                    "Apache Kafka Connector: Publishes structured JSON event streams to dedicated enterprise Kafka topics.",
                    "Google BigQuery Connector: Direct streaming inserts into data warehouse tables for custom SQL billing analytics.",
                    "Google Cloud PubSub Connector: Pushes event notifications to downstream serverless functions.",
                    "OpenTelemetry (OTel) Exporter: OTLP gRPC/HTTP exporter compatible with Prometheus, Jaeger, and Grafana."
                ],
                "interconnections": "Reads continuously from LLM Logs, MCP Logs, and Audit Logs; forwards data to external corporate SIEMs.",
                "use_case": "Enterprises stream all AI consumption logs to their centralized Datadog dashboard and BigQuery data lake for corporate billing audits and security threat analysis."
            },
            {
                "name": "Logs Settings",
                "route": "/workspace/config/logging",
                "summary": "Centralized policy management console for log retention, sampling ratios, data privacy redaction, and raw byte persistence.",
                "architecture": (
                    "Located in 'ui/app/workspace/config/views/loggingView.tsx'. Controls internal storage garbage collection routines, "
                    "sampling filters in the FastHTTP middleware, and sensitive payload redaction algorithms."
                ),
                "components": [
                    "Retention Policy: Configurable auto-purge schedules (e.g. 7 days, 30 days, 90 days, 365 days).",
                    "Traffic Sampling: Configures percentage-based sampling (1% to 100%) to reduce storage costs in massive volume environments.",
                    "Content Redaction Toggle: Global enforcement of zero-prompt storage for extreme regulatory environments.",
                    "Raw Wire Byte Storage: Toggle for storing full wire-level JSON payloads for diagnostic deep-dives.",
                    "Storage Offloading: Integrates with AWS S3 and Google Cloud Storage for cold log archival."
                ],
                "interconnections": "Directly governs the storage behavior of LLM Logs and MCP Logs.",
                "use_case": "Healthcare and financial institutions enforce 30-day auto-deletion policies and disable prompt storage to satisfy HIPAA and PCI-DSS compliance."
            }
        ]
    },
    {
        "pillar_id": "pillar_2_models",
        "pillar_title": "Pillar 2: Models & Intelligent Routing (மாடல் மேலாண்மை & ரூட்டிங்)",
        "pillar_desc": "The Models pillar orchestrates model cataloging, dynamic multi-provider traffic routing, financial budget quotas, and automated failover.",
        "features": [
            {
                "name": "Model Catalog",
                "route": "/workspace/model-catalog",
                "summary": "Unified inventory of all connected models across cloud and self-hosted providers, displaying context limits, capabilities, and pricing.",
                "architecture": (
                    "Maintained in 'framework/modelcatalog'. Aggregates model metadata from all configured providers. "
                    "Maintains token window boundaries, max generation tokens, supported modalities (text, vision, audio), "
                    "feature capabilities (function calling, JSON mode, reasoning effort), and real-time provider reachability."
                ),
                "components": [
                    "Model Directory: Comprehensive searchable catalog of models from OpenAI, Anthropic, Bedrock, Vertex, Ollama, etc.",
                    "Capability Matrix: Flags for Tool Calling, Vision, Structured Outputs, and Streaming support.",
                    "Token Limits: Context window size (e.g. 128k, 200k, 1M) and maximum output token ceilings.",
                    "Model Aliasing: Maps logical names (e.g. 'production-fast') to specific underlying architectures."
                ],
                "interconnections": "Supplies target definitions to Routing Rules, Complexity Router, and Access Profiles. Displays pricing in Dashboard.",
                "use_case": "Developers query the catalog to discover which models support vision inputs and tool-calling capabilities across enterprise accounts."
            },
            {
                "name": "Model Providers",
                "route": "/workspace/providers",
                "summary": "Credential and endpoint management console for upstream AI vendors, supporting multi-key rotation and custom private endpoints.",
                "architecture": (
                    "Located in 'core/providers'. Contains dedicated adapters for 25+ AI providers (OpenAI, Azure, Anthropic, Bedrock, Vertex, Groq, Cerebras, Cohere, DeepSeek, Mistral, Ollama, vLLM, Replicate). "
                    "Implements connection pooling, custom base URLs, and multi-key rotation pools with round-robin or priority weighting."
                ),
                "components": [
                    "Provider Adapters: Specialized wire protocol implementations for each major AI provider.",
                    "Multi-Key Rotation: Pool multiple API keys to distribute rate limits and avoid provider throttling.",
                    "Private Endpoints: Configure custom base URLs for internal vLLM, Ollama, or VPC-peered deployments.",
                    "Connection Tuning: Per-provider HTTP keep-alive, dial timeouts, and connection pool sizes."
                ],
                "interconnections": "Powers the Model Catalog; feeds health data to Circuit Breaker; executes requests dispatched by Routing Rules.",
                "use_case": "Platform teams rotate OpenAI API keys without downtime and configure dedicated AWS Bedrock IAM role credentials."
            },
            {
                "name": "Budgets & Limits",
                "route": "/workspace/model-limits",
                "summary": "Multi-tier rate limiting and financial budget enforcement engine preventing runaway spending and provider throttling.",
                "architecture": (
                    "Implemented in 'framework/configstore' and 'plugins/governance/tracker.go'. "
                    "Maintains atomic sliding-window rate limit counters (RPM, TPM, RPD) in Redis / local memory. "
                    "Continuously evaluates current token burn against daily and monthly dollar budgets ($ USD). "
                    "Triggers immediate rejection with HTTP 429 when limits are breached."
                ),
                "components": [
                    "Throughput Limits: Requests Per Minute (RPM), Tokens Per Minute (TPM), Requests Per Day (RPD).",
                    "Dollar Spending Caps: Hard and soft monthly/daily financial budget limits in USD.",
                    "Multilevel Scoping: Enforce limits globally, per Virtual Key, per Team, per User, or per Model.",
                    "Breach Actions: Reject with 429 Too Many Requests, downgrade to a cheaper model, or trigger webhook alerts."
                ],
                "interconnections": "Evaluated during Stage 1 of the request lifecycle; updates budget balances from LLM Logs token consumption.",
                "use_case": "Ensures an experimental hackathon project cannot exceed $500/month or flood the company's shared OpenAI tier with 10,000 RPM."
            },
            {
                "name": "Routing Rules",
                "route": "/workspace/routing-rules",
                "summary": "Condition-based dynamic routing engine powered by Google CEL expressions for intelligent traffic routing and A/B testing.",
                "architecture": (
                    "Located in 'plugins/governance/routing.go'. Pre-compiles Google Common Expression Language (CEL) rules into memory. "
                    "Upon request arrival, evaluates rules against request context (model name, prompt content, user department, headers). "
                    "Selects target provider, performs traffic splitting, or triggers fallback chains."
                ),
                "components": [
                    "CEL Expression Builder: Define rules like `request.model == 'fast' && prompt.length < 500`.",
                    "Traffic Splitting: Distribute traffic (e.g. 80% to OpenAI, 20% to Anthropic) for zero-risk model evaluation.",
                    "Header-Based Routing: Route requests based on custom headers like `x-uf-dim-env: staging`.",
                    "Fallback Target Sequences: Ordered array of backup models if primary destination is unavailable."
                ],
                "interconnections": "Works in tandem with Complexity Router; hands off resolved target to Circuit Breaker.",
                "use_case": "Automatically routes customer queries mentioning 'code' to DeepSeek Coder, while routing French queries to Mistral Large."
            },
            {
                "name": "Complexity Router",
                "route": "/workspace/complexity-router",
                "summary": "Intelligent prompt complexity classifier that buckets queries into 4 tiers and routes them to cost-appropriate models, cutting costs by 60–80%.",
                "architecture": (
                    "Located in 'plugins/governance/complexity'. Analyzes incoming user messages across 4 distinct tiers: "
                    "SIMPLE, MEDIUM, COMPLEX, and REASONING. Uses keyword dictionary matching, token length boundaries, "
                    "and lightweight heuristics to identify task difficulty and route to corresponding model architectures."
                ),
                "components": [
                    "Tier Palette: Simple (P1), Medium (P2), Complex (P3), Reasoning (P4).",
                    "Keyword Lists: Reasoning keywords ('prove', 'step by step', 'derivation'), coding terms, and simple greetings.",
                    "Boundary Controls: Token length thresholds separating simple questions from heavy multi-turn reasoning.",
                    "Tier Target Mapping: Simple -> gpt-4o-mini; Medium -> claude-3-5-haiku; Complex -> gpt-4o; Reasoning -> o1."
                ],
                "interconnections": "Feeds classified tier context directly into Routing Rules; updates cost metrics in Dashboard.",
                "use_case": "Slashes enterprise AI bills by routing routine 'What are your store hours?' questions to a $0.15/1M token model instead of a $15/1M token reasoning model."
            },
            {
                "name": "Circuit Breaker",
                "route": "/workspace/circuit-breaker",
                "summary": "Automated failure detection and failover engine that guarantees 99.99% application uptime by rerouting around downed AI providers.",
                "architecture": (
                    "Implemented in 'framework/circuitbreaker/runtime.go'. Employs a 3-state finite state machine (CLOSED, OPEN, HALF-OPEN). "
                    "Tracks consecutive 5xx errors, timeouts, and latency spikes across a sliding sample window. "
                    "When error thresholds are crossed, the circuit TRIPS (OPEN) and diverts traffic to healthy fallbacks."
                ),
                "components": [
                    "Failure Threshold: Configurable error percentage (e.g. 50% errors over 10 requests).",
                    "Latency Threshold: Trips if p95 response time exceeds configured ceiling (e.g. 8000ms).",
                    "Cool-off Duration: Sleep window (e.g. 30s) before transitioning to HALF-OPEN to test provider recovery.",
                    "Fallback Chains: Automatic cascade through secondary and tertiary backup models.",
                    "Header Emission: Returns `x-unifai-fallback-index: 1` to notify callers that a fallback route was executed."
                ],
                "interconnections": "Guards all outbound requests dispatched by Routing Rules and Model Providers.",
                "use_case": "When an OpenAI outage causes 504 Gateway Timeouts, UnifAI automatically shifts chatbot traffic to Claude-3.5-Sonnet in <100ms with zero dropped user sessions."
            },
            {
                "name": "Pricing Overrides",
                "route": "/workspace/custom-pricing/overrides",
                "summary": "Granular custom token pricing table defining exact $/1M token costs for enterprise negotiated discount rates and custom fine-tuned models.",
                "architecture": (
                    "Maintained in 'framework/configstore'. Evaluated during post-request processing. "
                    "Matches the executed provider and model against custom price rules, calculating prompt cost, "
                    "completion cost, and cached token read cost down to micro-cents."
                ),
                "components": [
                    "Input / Output Token Rates: Configurable cost per 1M prompt and completion tokens.",
                    "Cached Token Discount: Custom pricing for prompt cache read hits.",
                    "Scoped Overrides: Define custom rates for specific enterprise accounts or private clusters."
                ],
                "interconnections": "Calculates exact dollar spend recorded in LLM Logs, Budgets & Limits, and Dashboard.",
                "use_case": "Enterprises with custom 30% volume discounts from Microsoft Azure input their exact discounted token rates for precise financial accounting."
            },
            {
                "name": "Model Settings",
                "route": "/workspace/custom-pricing",
                "summary": "Global routing parameters, request timeout defaults, retry policies, and fallback model definitions.",
                "architecture": (
                    "Defines platform-wide default behaviors applied when specific provider or routing rules are omitted. "
                    "Manages HTTP transport timeouts, max retries, exponential backoff with jitter, and default fallback models."
                ),
                "components": [
                    "Global Timeouts: Sets default request and connect timeouts (e.g. 60s).",
                    "Retry Policy: Configures maximum retry attempts and exponential backoff jitter.",
                    "Default Fallback: Universal emergency fallback model when all routing chains fail."
                ],
                "interconnections": "Underpins the execution engine across all Model Providers.",
                "use_case": "Standardizes a 30-second timeout and 2-retry policy across all AI microservices in the organization."
            }
        ]
    },
    {
        "pillar_id": "pillar_3_mcp",
        "pillar_title": "Pillar 3: MCP Gateway (டூல்ஸ் & ஏஜென்ட்கள் - Model Context Protocol)",
        "pillar_desc": "The Model Context Protocol Gateway standardizes agentic tool integration, sandboxed code execution, and enterprise OAuth credential delegation.",
        "features": [
            {
                "name": "MCP Catalog",
                "route": "/workspace/mcp-registry",
                "summary": "Master registry of all registered MCP servers, available tool definitions, schemas, and transport configurations.",
                "architecture": (
                    "Located in 'core/mcp'. Implements the Model Context Protocol JSON-RPC 2.0 specification over stdio, SSE, and in-process transports. "
                    "Maintains active connections, fetches available tool schemas, resources, and prompt templates, "
                    "and dynamically formats tool specifications into model-native tool calling schemas."
                ),
                "components": [
                    "Server Registry: Catalog of active and registered MCP servers.",
                    "Schema Inspector: View function signatures, descriptions, and required JSON parameters.",
                    "Transport Support: Native support for stdio subprocesses and HTTP Server-Sent Events (SSE).",
                    "Health Monitor: Continuously pings MCP servers and flags unresponsive processes."
                ],
                "interconnections": "Supplies tool definitions to Tool Groups and the LLM execution pipeline; records execution in MCP Logs.",
                "use_case": "Agent developers inspect what database query tools and web scraping functions are currently registered and online."
            },
            {
                "name": "MCP Library",
                "route": "/workspace/mcp-registry/library",
                "summary": "One-click curated installer for popular, production-ready open-source MCP servers.",
                "architecture": (
                    "Pre-configured repository of enterprise MCP servers. Allows administrators to install and spin up "
                    "containerized or subprocess MCP servers with pre-tested configurations in a single click."
                ),
                "components": [
                    "Curated Servers: GitHub, PostgreSQL, Slack, Google Drive, Jira, Local Filesystem, Memory, Brave Search.",
                    "Configuration Templates: Pre-filled environment variables and connection string templates.",
                    "1-Click Deploy: Automatically provisions and connects servers to the MCP Catalog."
                ],
                "interconnections": "Directly registers newly installed tools into the MCP Catalog and Tool Groups.",
                "use_case": "A team instantly adds a PostgreSQL database query tool to their AI Agent in 30 seconds without writing custom integration code."
            },
            {
                "name": "Tool Groups",
                "route": "/workspace/mcp-tool-groups",
                "summary": "Logical grouping and security boundary mechanism bundling related tools together for role-based assignment to Virtual Keys.",
                "architecture": (
                    "Maintained in 'framework/mcptoolgroups'. Groups individual tools into logical operational bundles "
                    "(e.g. 'DevOps-Tools', 'Finance-Tools', 'Support-Tools'). "
                    "Attached to Virtual Keys via Access Profiles to enforce strict least-privilege tool access."
                ),
                "components": [
                    "Tool Bundles: Logical collections of tools from one or more MCP servers.",
                    "Granular Whitelisting: Selectively include or exclude individual functions within an MCP server.",
                    "Access Profile Mapping: Bind tool groups directly to specific Virtual Keys."
                ],
                "interconnections": "Bound to Virtual Keys via Access Profiles; controls which tools are injected into LLM requests.",
                "use_case": "Ensures a Customer Support bot only has access to Zendesk tools, completely blocking it from executing Database Drop Table commands."
            },
            {
                "name": "Auth Sessions",
                "route": "/workspace/mcp-sessions",
                "summary": "Stateful credential manager tracking per-user authenticated sessions for tools requiring individual user permissions.",
                "architecture": (
                    "Located in 'core/mcp/credstore'. Associates stateful credentials (e.g. personal access tokens, session tokens) "
                    "with individual human users via `x-uf-mcp-session-id`. When an agent calls a tool, UnifAI injects "
                    "that specific user's credentials rather than a shared corporate master key."
                ),
                "components": [
                    "Session Keying: Keyed to an end-user SSO identity, virtual key, or client session ID.",
                    "Credential Vault: Encrypted at-rest storage for user session tokens.",
                    "Session Lifecycle: Automatic expiration, idle timeouts, and manual revocation."
                ],
                "interconnections": "Works alongside OAuth Grants to supply credentials during MCP execution; logged in MCP Logs.",
                "use_case": "When an executive asks an AI agent to 'Summarize my unread emails', the agent accesses Gmail using that executive's specific session token."
            },
            {
                "name": "OAuth Grants",
                "route": "/workspace/oauth-grants",
                "summary": "Downstream OAuth 2.0 authorization server managing consent screens, token exchange, and refresh token cycles for 3rd-party tools.",
                "architecture": (
                    "Implemented in 'framework/oauth2'. Acts as a full OAuth 2.0 authorization server supporting PKCE. "
                    "Handles authorization code redirects, user consent interfaces, token storage, and automated "
                    "background refresh token cycles for third-party providers (Google, Microsoft, GitHub, Slack)."
                ),
                "components": [
                    "OAuth 2.0 PKCE Flow: Secure authorization flow for browser and desktop clients.",
                    "Consent Management: Displays fine-grained scope approval screens to end users.",
                    "Token Refresh Engine: Automatically refreshes expired OAuth access tokens in the background."
                ],
                "interconnections": "Supplies live access tokens to Auth Sessions for MCP tool execution.",
                "use_case": "End-users securely authorize an AI agent to read their Google Calendar without ever exposing their Google password."
            },
            {
                "name": "MCP Settings",
                "route": "/workspace/mcp-settings",
                "summary": "Global execution parameters for MCP servers, including execution timeouts, concurrency limits, and restart policies.",
                "architecture": (
                    "Configures runtime constraints for MCP tool execution processes. "
                    "Enforces hard execution timeouts, maximum concurrent tool processes, memory limits, and auto-restart policies."
                ),
                "components": [
                    "Execution Timeout: Maximum duration a tool is permitted to run (e.g. 30s) before cancellation.",
                    "Concurrency Limits: Caps simultaneous tool executions to prevent server CPU exhaustion.",
                    "Process Restart Policy: Auto-restart crashed stdio server subprocesses."
                ],
                "interconnections": "Governs the execution runtime of all servers in the MCP Catalog.",
                "use_case": "Terminates rogue or hanging tool execution scripts after 30 seconds to prevent resource exhaustion."
            },
            {
                "name": "Plugins",
                "route": "/workspace/plugins",
                "summary": "Modular extensibility framework supporting custom Go plugins and sandboxed Starlark scripts in the request pipeline.",
                "architecture": (
                    "Located in 'core/pluginpipeline.go' and 'core/mcp/codemode/starlark'. "
                    "Supports both compiled Go plugins implementing BasePlugin and dynamically interpreted Starlark scripts. "
                    "Provides lifecycle hooks: PreLLMHook, PostLLMHook, PreMCPHook, and PostMCPHook."
                ),
                "components": [
                    "Lifecycle Hooks: Intercept and modify requests before and after LLM and MCP execution.",
                    "Starlark Script Sandbox: Secure, non-Turing complete runtime for dynamic in-memory transformations.",
                    "Custom Encryption: Apply proprietary data encryption algorithms to prompts."
                ],
                "interconnections": "Integrates into the core FastHTTP request-response lifecycle.",
                "use_case": "An enterprise injects a proprietary data anonymization algorithm into the PreLLMHook before prompts leave the company network."
            }
        ]
    },
    {
        "pillar_id": "pillar_4_governance",
        "pillar_title": "Pillar 4: Governance & Enterprise Compliance (அணுகல் கட்டுப்பாடு & பாதுகாப்பு)",
        "pillar_desc": "The Governance pillar establishes multi-tenant identity, role-based access control, automated SCIM provisioning, and regulatory audit logging.",
        "features": [
            {
                "name": "Virtual Keys",
                "route": "/workspace/governance/virtual-keys",
                "summary": "Synthetic proxy API keys (sk-uf-*) issued to applications, encapsulating permissions, quotas, and upstream credential masking.",
                "architecture": (
                    "Located in 'plugins/governance/main.go'. Virtual Keys are high-entropy cryptographic strings starting with 'sk-uf-'. "
                    "They hide upstream provider credentials, binding the request to a specific Team, Customer, Budget Limit, "
                    "Access Profile, and Tool Group. Verified in <0.5ms via local cache."
                ),
                "components": [
                    "Key Obfuscation: Completely hides sensitive upstream OpenAI and Anthropic master keys.",
                    "Configurable Expiration: Set automated key expiry dates (e.g. 90-day rotation).",
                    "IP & CIDR Restrictions: Restrict key usage to specific corporate IP address ranges.",
                    "Instant Revocation: Revoke compromised keys in real-time across the cluster without restarting apps."
                ],
                "interconnections": "Entry point for authentication; enforces Budgets & Limits, Access Profiles, and Audit Logs.",
                "use_case": "A microservice uses 'sk-uf-prod-svc-123' to make AI calls; if compromised, it can be revoked instantly without affecting other services."
            },
            {
                "name": "Users",
                "route": "/workspace/governance/users",
                "summary": "User account directory for developers, administrators, and stakeholders with workspace access.",
                "architecture": (
                    "Manages platform user accounts, authentication credentials (bcrypt hashed), "
                    "SSO SAML/OIDC identity bindings, and Multi-Factor Authentication (MFA) status."
                ),
                "components": [
                    "User Directory: Searchable list of workspace members and their status.",
                    "MFA Enforcement: Mandatory Two-Factor Authentication via TOTP authenticator apps.",
                    "Session Invalidation: Force-terminate active user login sessions across devices."
                ],
                "interconnections": "Members of Teams; assigned Roles & Permissions (RBAC); actions tracked in Audit Logs.",
                "use_case": "Workspace administrators invite developers and enforce MFA for all production environment access."
            },
            {
                "name": "Teams",
                "route": "/workspace/governance/teams",
                "summary": "Departmental organizational units managing shared budgets, team-owned virtual keys, and delegated administration.",
                "architecture": (
                    "Organizes users into functional departments (e.g. Frontend Engineering, Data Science, Customer Support). "
                    "Virtual Keys and Budgets are created at the Team level for unified departmental tracking."
                ),
                "components": [
                    "Team Hierarchy: Departmental groupings with dedicated Team Admins.",
                    "Shared Budget Pool: Shared monthly financial spending quota across team members.",
                    "Team Virtual Keys: Keys accessible only by authorized team engineers."
                ],
                "interconnections": "Groups Users; belongs to Business Units; owns Virtual Keys and Budgets.",
                "use_case": "The Data Science team is allocated a $5,000 monthly budget and 5 dedicated GPU model keys."
            },
            {
                "name": "Business Units",
                "route": "/workspace/governance/business-units",
                "summary": "Top-level enterprise divisions grouping multiple teams for macroscopic financial chargeback and corporate reporting.",
                "architecture": (
                    "Represents top-level corporate subsidiaries or operating divisions (e.g. Retail Banking, Commercial Insurance, Corporate IT). "
                    "Aggregates token consumption across multiple child teams."
                ),
                "components": [
                    "Division Grouping: Hierarchical aggregation of teams.",
                    "Corporate Chargeback: High-level financial reporting for CFO and Finance teams."
                ],
                "interconnections": "Contains multiple Teams; reported in Dashboard and BigQuery connector exports.",
                "use_case": "Corporate Finance views the total quarterly AI expenditure of the entire Retail Division across 15 engineering teams."
            },
            {
                "name": "Customers",
                "route": "/workspace/governance/customers",
                "summary": "Multi-tenant client registry for B2B SaaS applications, tracking token consumption and cost per paying customer.",
                "architecture": (
                    "Allows B2B SaaS platforms to track external tenant consumption. When requests include `x-uf-customer-id`, "
                    "UnifAI isolates and attributes token counts and dollar spend to that specific customer entity."
                ),
                "components": [
                    "Customer Registry: Directory of client organizations using the SaaS product.",
                    "Per-Customer Quotas: Enforce rate limits and cost caps on specific customer accounts.",
                    "Billing Reports: Export itemized monthly AI consumption bills per customer."
                ],
                "interconnections": "Scoped via the `x-uf-customer-id` header; reported in Dashboard and LLM Logs.",
                "use_case": "A legal SaaS platform calculates that Customer X consumed $412 in GPT-4o legal doc summaries this month, auto-invoicing them via Stripe."
            },
            {
                "name": "User Provisioning (SCIM)",
                "route": "/workspace/scim",
                "summary": "SCIM v2.0 endpoint automating employee onboarding and deprovisioning from enterprise Identity Providers (Okta, Azure AD).",
                "architecture": (
                    "Implements RFC 7643 and RFC 7644 SCIM 2.0 endpoints (/scim/v2/Users, /scim/v2/Groups). "
                    "Receives webhook push updates from enterprise IdPs (Okta, Microsoft Entra ID / Azure AD, PingIdentity). "
                    "Automatically provisions accounts, assigns team memberships, and immediately deactivates former employees."
                ),
                "components": [
                    "SCIM 2.0 Protocol: Standardized enterprise user lifecycle integration.",
                    "Automated Deprovisioning: Instant revocation of platform access when an employee leaves the company.",
                    "Group Mapping: Automatically maps corporate Okta groups to UnifAI Teams and Roles."
                ],
                "interconnections": "Directly provisions Users and updates Team memberships; actions recorded in Audit Logs.",
                "use_case": "When an engineer is offboarded in Okta, their UnifAI account and personal virtual keys are deactivated within 1 second."
            },
            {
                "name": "Roles & Permissions (RBAC)",
                "route": "/workspace/governance/rbac",
                "summary": "Fine-grained Role-Based Access Control matrix governing administrative actions across all gateway resources.",
                "architecture": (
                    "Implemented in 'framework/rbac'. Enforces a matrix of Resources (VirtualKeys, RoutingRules, Guardrails, Logs, Providers, Settings) "
                    "and Operations (View, Create, Edit, Delete). Evaluated on every UI route match and control plane API call."
                ),
                "components": [
                    "Predefined Roles: Super Admin, Workspace Admin, Developer, Auditor, Viewer.",
                    "Custom Roles: Create bespoke enterprise roles with granular permission toggles.",
                    "Resource Scoping: Restrict permissions to specific Teams or Business Units."
                ],
                "interconnections": "Governs UI visibility (via useRbac hook) and authorizes Admin API calls.",
                "use_case": "Junior developers can view logs and test prompts, but are blocked from creating new Virtual Keys or modifying Routing Rules."
            },
            {
                "name": "Access Profiles",
                "route": "/workspace/governance/access-profiles",
                "summary": "Reusable security policy presets defining allowed models, providers, and MCP tool groups bound to Virtual Keys.",
                "architecture": (
                    "Located in 'framework/configstore'. Encapsulates model and tool permissions into reusable bundles. "
                    "When attached to a Virtual Key, UnifAI rejects any request attempting to use an unapproved model or unlisted MCP tool."
                ),
                "components": [
                    "Approved Models Whitelist: Exact list of allowed model IDs.",
                    "Approved Providers Whitelist: Approved vendor endpoints.",
                    "Approved Tool Groups: Permitted MCP tool groups."
                ],
                "interconnections": "Bound to Virtual Keys; evaluated during Stage 1 request validation.",
                "use_case": "A 'Staging-Profile' guarantees that staging microservices can only use cheap open-source models (Llama 3.1 8B), preventing accidental GPT-4o spend in QA."
            },
            {
                "name": "Audit Logs",
                "route": "/workspace/audit-logs",
                "summary": "Cryptographically timestamped, immutable audit ledger recording all administrative mutations for SOC 2 and ISO 27001 compliance.",
                "architecture": (
                    "Every state mutation across the platform (creating a key, editing a route, altering budget limits) "
                    "triggers an immutable audit event. Captures Actor ID, action type, target resource, JSON diff of changes (before/after), "
                    "source IP address, user agent, and timestamp."
                ),
                "components": [
                    "Immutable Storage: Tamper-resistant append-only database table.",
                    "JSON Diff Inspector: Visual before-and-after comparison of configuration changes.",
                    "Compliance Export: One-click export to CSV / JSON for regulatory audits."
                ],
                "interconnections": "Records all mutations across Governance, Models, Guardrails, and Settings; streams to Connectors.",
                "use_case": "SOC 2 auditors verify who changed the production routing rule on August 12th and view the exact JSON diff of the modification."
            }
        ]
    },
    {
        "pillar_id": "pillar_5_guardrails",
        "pillar_title": "Pillar 5: Guardrails & Content Security (உள்ளடக்கப் பாதுகாப்பு & கொள்கை)",
        "pillar_desc": "The Guardrails pillar intercepts prompt inputs and model outputs to enforce PII redaction, brand safety, prompt injection defense, and cluster sync.",
        "features": [
            {
                "name": "Rules",
                "route": "/workspace/guardrails/configuration",
                "summary": "Safety policy definitions executed pre-LLM (input validation) and post-LLM (output validation) to mask PII and block prompt injections.",
                "architecture": (
                    "Located in 'plugins/guardrails/main.go'. Pre-compiles Google CEL expressions for lightning-fast execution. "
                    "Evaluates input prompts before they reach external LLMs and output responses before they return to clients. "
                    "Supported actions: Block request (400 Bad Request), Redact/Mask matching tokens (e.g. `[REDACTED_SSN]`), or Flag in logs."
                ),
                "components": [
                    "Pre-LLM Rules: Input scanning for prompt injection, jailbreak keywords, and leaked secrets.",
                    "Post-LLM Rules: Output scanning for hallucinated API keys, toxic output, or leaked PII.",
                    "Custom Actions: Block, Redact/Mask, Flag, or Replace.",
                    "Threshold Tuning: Custom sensitivity scores for classification models."
                ],
                "interconnections": "Executes in Stage 2 (Pre-LLM) and Stage 6 (Post-LLM); annotates records in LLM Logs.",
                "use_case": "Automatically detects and masks credit card numbers and Aadhaar numbers with `[REDACTED_PII]` before sending prompts to OpenAI."
            },
            {
                "name": "Providers",
                "route": "/workspace/guardrails/providers",
                "summary": "Configuration console for specialized safety engines, including Presidio DLP, Llama Guard, AWS Bedrock Guardrails, and regex filters.",
                "architecture": (
                    "Adapters integrating dedicated detection engines: "
                    "1. Microsoft Presidio DLP (entity extraction for PII), "
                    "2. Meta Llama Guard (content moderation classification), "
                    "3. AWS Bedrock Guardrails (managed cloud safety), "
                    "4. Lakera AI (prompt injection & jailbreak defense), and "
                    "5. High-speed compiled regex engines."
                ),
                "components": [
                    "Presidio DLP Integration: Recognizes SSN, Credit Cards, Names, Phone Numbers, Email, Medical IDs.",
                    "Llama Guard Integration: Flags hate speech, violence, self-harm, and sexual content.",
                    "Regex Fast-Path: Zero-latency pattern matching for proprietary internal account numbers."
                ],
                "interconnections": "Supplies detection capabilities to Guardrail Rules and Browser AI.",
                "use_case": "Enables Meta Llama Guard to block adversarial jailbreak attempts ('Ignore all previous instructions...')."
            },
            {
                "name": "Cluster Config",
                "route": "/workspace/cluster",
                "summary": "Multi-node distributed cluster coordination engine synchronizing rules, rate limit counters, and configurations in real-time.",
                "architecture": (
                    "Coordinates distributed UnifAI gateway instances deployed in high-availability clusters. "
                    "Uses pub/sub sync to propagate routing rule updates, newly issued virtual keys, and rate limit counters "
                    "across all nodes in <50ms without restarting services."
                ),
                "components": [
                    "Node Discovery: Real-time health monitoring of all active cluster gateway instances.",
                    "Config Synchronization: Instant zero-downtime propagation of rule mutations.",
                    "Distributed Rate Limiting: Synchronized Redis counters preventing quota circumvention."
                ],
                "interconnections": "Synchronizes Governance, Routing Rules, Guardrails, and Caching across all cluster nodes.",
                "use_case": "A global deployment with 10 gateway instances across US-East and EU-West maintains synchronized rate limits and instant rule updates."
            }
        ]
    },
    {
        "pillar_id": "pillar_6_adaptive",
        "pillar_title": "Pillar 6: Adaptive Routing & Assets (ஸ்மார்ட் தேர்வு & அறிவு களஞ்சியம்)",
        "pillar_desc": "Adaptive Routing leverages reinforcement learning to optimize latency and cost while managing versioned prompt templates and agent skills.",
        "features": [
            {
                "name": "Adaptive Routing Dashboard & Settings",
                "route": "/workspace/adaptive-routing",
                "summary": "Reinforcement learning traffic router that continuously evaluates provider latency and error rates to steer requests to the fastest, cheapest model.",
                "architecture": (
                    "Implements Multi-Armed Bandit (MAB) optimization algorithms (Epsilon-Greedy, Thompson Sampling, Upper Confidence Bound UCB1). "
                    "Continuously evaluates real-time latency, error rates, and token costs across equivalent candidate models. "
                    "Dynamically shifts traffic weights to minimize p95 latency while adhering to cost budgets."
                ),
                "components": [
                    "Algorithm Selection: Toggle between Epsilon-Greedy, Thompson Sampling, and Latency-Weighted routing.",
                    "Latency vs Cost Sliders: Configurable weighting favoring speed vs dollar savings.",
                    "Exploration Ratio: Percentage of traffic reserved for testing slower/newer models.",
                    "Performance Dashboard: Visual charts showing dynamic weight distribution and latency improvements."
                ],
                "interconnections": "Acts as an intelligent layer within Routing Rules; reads performance telemetry from LLM Logs.",
                "use_case": "When Azure OpenAI experiences temporary latency spikes (p95 > 4000ms), the Adaptive Router automatically diverts 90% of traffic to AWS Bedrock."
            },
            {
                "name": "Prompt Repository",
                "route": "/workspace/prompt-repo",
                "summary": "Centralized, version-controlled enterprise prompt repository supporting variable interpolation, test suites, and model bindings.",
                "architecture": (
                    "Enterprise prompt management engine. Stores parameterized prompt templates with Git-style versioning (v1, v2, production tag). "
                    "Supports Mustache/Jinja variable interpolation (`{{customer_query}}`), parameter presets (temperature, top_p), "
                    "and test suite benchmarking."
                ),
                "components": [
                    "Version Control: Full commit history, diffs, and instant rollback capabilities.",
                    "Variable Templating: Define required variables with default values and type validation.",
                    "Model Binding: Associate prompts with tested, approved model architectures.",
                    "Playground Testing: Interactive test console to execute templates against real models."
                ],
                "interconnections": "Referenced by applications via prompt IDs; logged in LLM Logs; bound to specific Model Catalog entries.",
                "use_case": "Prompt engineers update customer service prompt templates in UnifAI; microservices automatically consume the updated prompt without code redeployments."
            },
            {
                "name": "Skills Repository",
                "route": "/workspace/skills-repo",
                "summary": "Standardized catalog of autonomous AI Agent instructions, domain personas, and specialized execution skill packages.",
                "architecture": (
                    "Stores modular, reusable agent skills (system instructions, role definitions, operational guardrails). "
                    "Skills can be dynamically injected into LLM requests or bound to specific MCP Tool Groups to build autonomous agents."
                ),
                "components": [
                    "Skill Packages: Modular instructions (e.g. 'SQL-Expert', 'Python-Debugger', 'Compliance-Auditor').",
                    "Tool Bindings: Associate specialized skills directly with approved MCP tools.",
                    "Version History: Track refinements in agent behavior and system instructions."
                ],
                "interconnections": "Combines with MCP Tool Groups and Prompt Repository to configure autonomous agents.",
                "use_case": "Teams share a battle-tested 'Secure Code Reviewer' skill package across 20 different development squads."
            }
        ]
    },
    {
        "pillar_id": "pillar_7_settings",
        "pillar_title": "Pillar 7: Global Settings & Engine Tuning (கட்டமைப்பு & செயல்திறன் ட்யூனிங்)",
        "pillar_desc": "Platform-wide system configuration, wire compatibility translation, semantic vector caching, and network security.",
        "features": [
            {
                "name": "Client Settings",
                "route": "/workspace/config/client-settings",
                "summary": "Global connection timeouts, keep-alive settings, and header forward allowlists.",
                "architecture": (
                    "Configures network-level HTTP client behavior in FastHTTP. "
                    "Defines header forwarding allowlists and blocklists, controlling which client headers are forwarded to upstream AI providers."
                ),
                "components": [
                    "Header Allowlists: Explicitly permit custom headers (x-uf-eh-*) to reach upstream providers.",
                    "Header Blocklists: Strip sensitive internal tracking headers before dispatch.",
                    "Keep-Alive Tuning: Configure client connection persistence and idle socket timeouts."
                ],
                "interconnections": "Governs the inbound and outbound HTTP Transport layer.",
                "use_case": "Allows microservices to forward custom transaction correlation IDs to upstream LLM providers while stripping internal cookies."
            },
            {
                "name": "Compatibility",
                "route": "/workspace/config/compatibility",
                "summary": "Automated JSON API translation engine bridging OpenAI, Anthropic, AWS Bedrock, and Google Gemini schemas.",
                "architecture": (
                    "Located in 'plugins/compat'. Real-time schema converter. "
                    "Enables client applications written with the OpenAI SDK (`/v1/chat/completions`) to interact seamlessly "
                    "with Anthropic (`/v1/messages`), AWS Bedrock Converse, or Google Vertex AI without modifying a single line of client code."
                ),
                "components": [
                    "Schema Translation: Converts messages, roles (system/user/assistant), tools, and choices across provider formats.",
                    "Parameter Adaptation: Automatically converts or safely drops unsupported parameters (e.g. frequency_penalty).",
                    "Streaming SSE Normalization: Normalizes provider-specific SSE streaming chunks into standard OpenAI-compatible chunks."
                ],
                "interconnections": "Executes inside the HTTP Transport layer before provider dispatch and after response receipt.",
                "use_case": "An application built exclusively for OpenAI can switch seamlessly to Anthropic Claude 3.7 Sonnet by changing only the model name."
            },
            {
                "name": "Caching (Semantic Cache)",
                "route": "/workspace/config/caching",
                "summary": "High-speed semantic vector caching engine using Redis or Qdrant to deliver instant, zero-cost responses for similar queries.",
                "architecture": (
                    "Implemented in 'plugins/semanticcache'. Combines direct SHA-256 hash lookup with vector embedding search. "
                    "Generates vector embeddings of user queries using fast embedding models (OpenAI text-embedding-3-small or local BGE). "
                    "Performs cosine similarity search against Qdrant / Redis vector partitions. "
                    "If similarity >= threshold (default 0.90), returns cached answer in <15ms at $0.00 cost."
                ),
                "components": [
                    "Vector Store Backend: Connect to Redis Vector Engine or Qdrant.",
                    "Embedding Generator: Configurable embedding model source for vectorizing incoming prompts.",
                    "Similarity Threshold: Global cosine similarity threshold (e.g. 0.90).",
                    "Partition Keying: Scopes cache entries via `x-uf-cache-key`.",
                    "Streaming Cache: Caches streaming SSE chunks for seamless instant replay."
                ],
                "interconnections": "Evaluated in Stage 3 of the request pipeline; records cache hit savings in Dashboard and LLM Logs.",
                "use_case": "In high-volume customer support, 25% of queries are repeated questions; Semantic Caching answers them instantly, cutting monthly API spend by thousands of dollars."
            },
            {
                "name": "Security",
                "route": "/workspace/config/security",
                "summary": "Enterprise network defense console managing TLS certificates, IP allowlists, CORS policies, and secret encryption.",
                "architecture": (
                    "Located in 'core/schemas/vault.go' and HTTP middleware. "
                    "Enforces TLS 1.3 encryption in transit, encrypts API keys at rest using AES-256-GCM, "
                    "validates incoming IP addresses against CIDR allowlists, and enforces strict Cross-Origin Resource Sharing (CORS) rules."
                ),
                "components": [
                    "AES-256-GCM Encryption: Master key encryption for all stored provider secrets and virtual keys.",
                    "IP CIDR Whitelisting: Restrict gateway access to designated corporate VPN or VPC IP subnets.",
                    "CORS Policy Manager: Whitelist allowed frontend web origins.",
                    "Direct Key Policy: Enable or disable `x-uf-direct-key` pass-through globally."
                ],
                "interconnections": "Protects all incoming connections and secures all stored secrets in PostgreSQL.",
                "use_case": "Security teams restrict UnifAI Gateway access exclusively to internal corporate VPC subnets and enforce TLS 1.3."
            },
            {
                "name": "API Keys",
                "route": "/workspace/config/api-keys",
                "summary": "Administrative API key manager for automation, CI/CD pipelines, and infrastructure-as-code orchestration.",
                "architecture": (
                    "Issues and manages Control Plane Admin API Keys (`x-uf-api-key`). "
                    "Used by Terraform providers, Kubernetes operators, and CI/CD pipelines to programmatically configure UnifAI."
                ),
                "components": [
                    "Admin Key Generation: Generate cryptographically random administrative keys.",
                    "Scope Assignment: Assign read-only or read-write permissions to admin keys.",
                    "Usage Tracking: Audit log of actions executed by each admin key."
                ],
                "interconnections": "Authorizes programmatic control-plane configuration requests; recorded in Audit Logs.",
                "use_case": "Terraform scripts use an Admin API Key to provision Virtual Keys and Routing Rules during automated environment spin-up."
            },
            {
                "name": "Performance Tuning",
                "route": "/workspace/config/performance-tuning",
                "summary": "Low-level runtime performance tuning console for worker pool concurrency, memory pools, and buffer sizes.",
                "architecture": (
                    "Located in 'core/unifai.go'. Controls internal Go runtime parameters: "
                    "FastHTTP request channel pool size (channelMessagePool), response stream pool (responseStreamPool), "
                    "worker goroutine counts per provider queue, and memory buffer allocations."
                ),
                "components": [
                    "Worker Pool Size: Maximum concurrent worker goroutines per upstream provider queue.",
                    "Zero-Allocation Pool Sizing: Pre-allocated buffer capacity in `sync.Pool` to eliminate GC pauses.",
                    "Queue Overflow Policy: Block-and-wait vs immediate rejection (`dropExcessRequests`) under heavy traffic."
                ],
                "interconnections": "Optimizes the execution performance of the entire core gateway engine.",
                "use_case": "Under peak traffic of 50,000 requests/second, performance engineers tune worker pools to achieve zero GC pause times."
            },
            {
                "name": "Feature Flags",
                "route": "/workspace/config/feature-flags",
                "summary": "Dynamic runtime feature toggles allowing administrators to enable, test, or disable capabilities without redeploying the gateway.",
                "architecture": (
                    "Maintained in 'framework/featureflags'. Real-time in-memory configuration flags synchronized across the cluster. "
                    "Allows instant activation or deactivation of beta features (e.g. experimental routers, new provider adapters) with zero downtime."
                ),
                "components": [
                    "Dynamic Toggles: Enable or disable platform features in real-time.",
                    "Gradual Rollout: Enable beta capabilities for specific teams or virtual keys.",
                    "Instant Kill-Switch: Immediately disable problematic features without redeploying code."
                ],
                "interconnections": "Controls code execution paths across all gateway modules.",
                "use_case": "A platform team safely tests an experimental new reasoning router with internal engineers before enabling it organization-wide."
            },
            {
                "name": "Enterprise Outbound Proxy",
                "route": "/workspace/config/proxy",
                "summary": "Corporate forward and reverse proxy configuration routing outbound AI traffic through enterprise egress gateways.",
                "architecture": (
                    "Located in 'ui/app/workspace/config/views/proxyView.tsx' and core transport. "
                    "Enables UnifAI to route all outbound LLM API requests through corporate forward proxies (HTTP, HTTPS, SOCKS5, TCP). "
                    "Supports basic proxy authentication, corporate bypass allowlists (CIDR / hostnames), and custom CA certificate bundles for SSL inspection."
                ),
                "components": [
                    "Proxy Protocol: HTTP, HTTPS, SOCKS5, and TCP egress.",
                    "Authentication: Basic authentication (Username / Password) for proxy gateways.",
                    "No-Proxy Bypass List: Whitelist of hostnames and IP ranges (localhost, 127.0.0.1, internal VPCs) bypassing the proxy.",
                    "Custom CA Bundle: Enterprise root CA certificate upload for corporate SSL inspection proxies."
                ],
                "interconnections": "Wraps all outbound HTTP and WebSocket connections to Model Providers and external MCP servers.",
                "use_case": "Enterprises in banking and defense require all external internet traffic to traverse corporate BlueCoat or Zscaler egress proxies."
            },
            {
                "name": "Large Payload Streaming Engine",
                "route": "/workspace/config/large-payload",
                "summary": "Zero-memory direct streaming engine designed for massive multimodal outputs, audio files, and 100k+ token responses.",
                "architecture": (
                    "Implemented in 'transports/unifai-http/integrations/utils.go' (tryStreamLargeResponse). "
                    "Detects responses exceeding the memory buffer threshold and bypasses gateway memory allocations entirely, "
                    "streaming chunked transfer encoding directly from the upstream provider to the client."
                ),
                "components": [
                    "Payload Threshold Slider: Configures byte ceiling (e.g. 5MB) triggering large payload bypass.",
                    "Direct Pipe Streaming: Eliminates memory buffering to avoid out-of-memory crashes on massive generations.",
                    "Multimodal Asset Transfer: Optimized for high-resolution vision models, image generation, and audio synthesis."
                ],
                "interconnections": "Integrates directly with FastHTTP response writer pool and Upstream LLM Providers.",
                "use_case": "Prevents gateway memory exhaustion when 1,000 concurrent clients download multi-megabyte audio or high-resolution image generation files."
            },
            {
                "name": "Alert Channels (Enterprise Notifications)",
                "route": "/workspace/alert-channels",
                "summary": "Multi-channel automated alerting engine broadcasting budget thresholds, circuit breaker trips, and error spikes to enterprise ops tools.",
                "architecture": (
                    "Implemented in '@enterprise/components/alert-channels/alertChannelsView'. "
                    "Listens to internal event buses for budget exhaustion events, provider circuit breaker trips, and p95 latency spikes. "
                    "Formats incident payloads and delivers them to Slack Webhooks, PagerDuty, Microsoft Teams, Discord, and Email."
                ),
                "components": [
                    "Channel Connectors: Slack Webhooks, PagerDuty Incident API, Microsoft Teams, Email (SMTP), and Custom Webhooks.",
                    "Trigger Rules: Budget Warning (80% / 100% spend), Circuit Breaker Trip, Error Spike (>5% 5xx), Rate Limit Throttle.",
                    "Notification Throttling: Configurable cooldown windows preventing alert storms during major outages."
                ],
                "interconnections": "Receives alert signals from Budgets & Limits, Circuit Breaker, and Observability.",
                "use_case": "When OpenAI goes down and the Circuit Breaker trips to Claude, an immediate alert is dispatched to the DevOps Slack channel and PagerDuty."
            },
            {
                "name": "MCP Authentication Config & Credential Vault",
                "route": "/workspace/mcp-auth-config",
                "summary": "Centralized credential vault managing per-user headers, API keys, and OAuth client credentials for external MCP servers.",
                "architecture": (
                    "Located in 'core/mcp/credstore'. Stores and resolves authentication credentials for MCP servers. "
                    "Supports static API keys, per-user HTTP headers, and dynamic OAuth 2.0 client secrets with AES-256-GCM encryption."
                ),
                "components": [
                    "Auth Mode Selector: Static API Key, Per-User Headers, or OAuth 2.0 PKCE.",
                    "Header Mapping: Maps client request headers to tool execution headers.",
                    "Encrypted Vault: Hardware-grade encryption for third-party service tokens."
                ],
                "interconnections": "Supplies authorized credentials to MCP Catalog, Tool Groups, and Auth Sessions.",
                "use_case": "Enables an AI agent to authenticate with internal Jira and GitHub enterprise servers using dynamic corporate service accounts."
            },
            {
                "name": "Starlark Sandboxed Code Mode",
                "route": "core/mcp/codemode/starlark",
                "summary": "Embedded Python-like sandboxed execution engine enabling AI agents to run multi-tool workflows and data transformations in-memory.",
                "architecture": (
                    "Implemented in 'core/mcp/codemode/starlark'. Leverages Google Starlark (the deterministic language behind Bazel). "
                    "Agents generate Starlark code to orchestrate multiple tools, perform mathematical calculations, "
                    "and transform data in a secure, memory-isolated sandbox without executing untrusted OS commands."
                ),
                "components": [
                    "Deterministic Sandbox: Non-Turing complete runtime preventing infinite loops and filesystem access.",
                    "Tool Binding Interface: Exposes registered MCP tools as native Starlark functions.",
                    "In-Memory Execution: Zero network latency between multi-step tool calls."
                ],
                "interconnections": "Integrates directly into MCP Gateway and Tool Groups.",
                "use_case": "An autonomous data science agent queries a SQL database, filters 10,000 rows, and calculates summary statistics in-memory in under 50ms."
            },
            {
                "name": "Hardware Secrets Vault & Envelope Encryption",
                "route": "core/schemas/vault.go",
                "summary": "Cryptographic envelope encryption vault securing all stored provider API keys, virtual keys, and credentials at rest.",
                "architecture": (
                    "Located in 'core/schemas/vault.go'. Employs AES-256-GCM authenticated encryption. "
                    "Every secret is encrypted with a unique Data Encryption Key (DEK), which is in turn encrypted with a Master Key (KEK) "
                    "stored in AWS KMS, Google Cloud KMS, HashiCorp Vault, or Infisical."
                ),
                "components": [
                    "AES-256-GCM Encryption: Authenticated cipher preventing tampering and eavesdropping.",
                    "Envelope Encryption: Multi-tier key hierarchy (Master Key / Data Keys).",
                    "KMS Integration: Seamless integration with AWS KMS, Google Cloud KMS, and Infisical."
                ],
                "interconnections": "Protects secrets across Model Providers, Virtual Keys, and MCP Gateway.",
                "use_case": "Guarantees that even if the PostgreSQL database backup is stolen, all OpenAI and Anthropic provider keys remain mathematically unreadable."
            },
            {
                "name": "Key Load Balancer & Key Pool Filtering",
                "route": "framework/loadbalancer",
                "summary": "Advanced multi-key load balancing engine distributing traffic across credential pools with session stickiness.",
                "architecture": (
                    "Implemented in 'framework/loadbalancer' and 'core/keyselectors'. "
                    "Supports Round-Robin, Weighted Distribution, and Priority Failover across multiple provider keys. "
                    "Integrates with Redis KVStore for session stickiness, pinning multi-turn conversations to the same key or GPU node."
                ),
                "components": [
                    "Balancing Strategies: Round-Robin, Weighted Ratios, Priority Tiers, and Least-Connections.",
                    "Session Stickiness: Binds conversational sessions (via x-uf-session-id) to the same upstream deployment.",
                    "Key Pool Filters: Custom hooks (keyPoolFilter) to dynamically disqualify degraded or rate-limited keys."
                ],
                "interconnections": "Operates inside the Models and Model Providers dispatch pipeline.",
                "use_case": "Distributes 100,000 RPM evenly across 10 Azure OpenAI deployments in different geographic regions with automatic failover."
            },
            {
                "name": "Realtime Audio & Voice Gateway",
                "route": "core/schemas/realtime.go",
                "summary": "Bidirectional WebRTC and WebSocket streaming gateway supporting ultra-low latency conversational voice AI.",
                "architecture": (
                    "Located in 'core/schemas/realtime.go'. Bridges client WebRTC and WebSocket voice connections to upstream providers "
                    "(e.g. OpenAI Realtime API). Manages full-duplex audio chunking (PCM16, G.711, Opus), "
                    "Voice Activity Detection (VAD), and automated speech-to-speech protocol normalization."
                ),
                "components": [
                    "Full-Duplex Audio Streaming: Sub-300ms voice conversational latency.",
                    "Protocol Normalization: Bidirectional audio frame translation across voice providers.",
                    "Voice Activity Detection (VAD): Server-side speech interruption and turn-taking detection."
                ],
                "interconnections": "Connects clients to Voice-enabled Model Providers and logs audio tokens in LLM Logs.",
                "use_case": "Powers customer service interactive voice response (IVR) phone bots with natural human-like voice conversational speed."
            }
        ]
    }
]

# ==============================================================================
# TECH STACK DETAILED DATA
# ==============================================================================

TECH_STACK_DATA = [
    {
        "layer": "Core Gateway Engine",
        "tech": "Go (Golang 1.23+)",
        "why": "Ultra-low proxy latency (<1ms), lightweight goroutine concurrency, memory safety, and native zero-allocation memory pools (sync.Pool)."
    },
    {
        "layer": "HTTP & Networking",
        "tech": "FastHTTP & Sonic JSON",
        "why": "FastHTTP avoids per-request heap allocations of standard net/http. ByteDance's Sonic JIT compiler parses JSON at near-native assembly speeds."
    },
    {
        "layer": "Rule Evaluation Engine",
        "tech": "Google CEL (Common Expression Language)",
        "why": "Safe, non-Turing complete expression evaluation. Compiles routing rules and guardrail policies into bytecode, executing in microseconds."
    },
    {
        "layer": "Agent Scripting Sandbox",
        "tech": "Starlark (Google Bazel Language)",
        "why": "Deterministic, thread-safe, sandboxed Python-like language for executing dynamic custom plugins without security risk."
    },
    {
        "layer": "Frontend Web Application",
        "tech": "TypeScript 5.x, React 18, Vite",
        "why": "Strict compile-time type safety across all 38 UI views, instant HMR development, and optimized single-page app bundle delivery."
    },
    {
        "layer": "Routing & UI Components",
        "tech": "TanStack Router, Tailwind CSS, Radix UI",
        "why": "TanStack Router provides preload-on-hover routing for instantaneous navigation. Radix UI (Shadcn) delivers accessible, sleek dark-mode components."
    },
    {
        "layer": "State Management",
        "tech": "Redux Toolkit (RTK Query), WebSockets",
        "why": "Manages client-side caching, optimistic UI updates, and real-time live metrics streaming via WebSockets."
    },
    {
        "layer": "Desktop / Browser DLP Agent",
        "tech": "Python 3.11+, mitmproxy / asyncio",
        "why": "Mature networking and proxy ecosystem for local HTTPS interception, regex matching, and integration with local Ollama models."
    },
    {
        "layer": "Relational Database",
        "tech": "PostgreSQL (pgx v5 driver), SQLite, GORM",
        "why": "ACID transactional storage for Governance, Virtual Keys, Audit Logs, and Budgets. pgx provides high-performance connection pooling."
    },
    {
        "layer": "Caching & Vector Store",
        "tech": "Redis (go-redis), Qdrant, pgvector",
        "why": "Sub-millisecond sliding-window rate limiting in Redis, and high-throughput vector similarity search for Semantic Caching in Qdrant."
    },
    {
        "layer": "Telemetry & Observability",
        "tech": "OpenTelemetry (OTel), Prometheus, Datadog, Kafka",
        "why": "W3C standard distributed tracing, Prometheus metric scrapers, and event-driven enterprise streaming to Kafka and Datadog."
    },
    {
        "layer": "DevOps & Build System",
        "tech": "Docker, Docker Compose, Nix / Flake, Makefile",
        "why": "Hermetic, reproducible development and build environments across multi-platform container architectures."
    }
]

# ==============================================================================
# 1. MARKDOWN GENERATION FUNCTION
# ==============================================================================
def generate_markdown():
    print("Generating comprehensive Markdown document...")
    lines = []
    lines.append("# UnifAI Enterprise Architecture, Headers & Features Master Guide")
    lines.append("## விரிவான கணினி கட்டமைப்பு, ஹெடர்ஸ் மற்றும் 38 அம்சங்களின் முழுமையான கையேடு (Tamil & English Technical Manual)\n")
    lines.append("**Document Version:** 2.0 (Enterprise Edition)  ")
    lines.append("**Classification:** Complete Technical Architecture & System Engineering Manual  ")
    lines.append("**Generated At:** 2026-09-05  ")
    lines.append("**Platform:** UnifAI Unified AI Gateway & Governance Control Plane  \n")
    lines.append("---\n")
    lines.append("## Table of Contents (பொருளடக்கம்)")
    lines.append("1. [Executive Summary & Platform Overview (செயல் சுருக்கம்)](#1-executive-summary--platform-overview)")
    lines.append("2. [End-to-End Request Lifecycle & System Architecture (கோரிக்கை வாழ்க்கைச் சுழற்சி)](#2-end-to-end-request-lifecycle--system-architecture)")
    lines.append("3. [Exhaustive HTTP Headers Reference Guide (ஹெடர்ஸ் முழு விவரக் கையேடு)](#3-exhaustive-http-headers-reference-guide)")
    for cat in HEADERS_DATA:
        lines.append(f"   - {cat['category']}")
    lines.append("4. [In-Depth Feature Catalog (7 Core Pillars - 38 Features)](#4-in-depth-feature-catalog)")
    for p in PILLARS_DATA:
        lines.append(f"   - {p['pillar_title']}")
        for f in p["features"]:
            lines.append(f"     * {f['name']} (`{f['route']}`)")
    lines.append("5. [Cross-Feature Interconnection & Data Flow Matrix (இணைப்பு வரைபடம்)](#5-cross-feature-interconnection--data-flow-matrix)")
    lines.append("6. [Technology Stack & Programming Languages Deep Dive (தொழில்நுட்ப கட்டமைப்பு)](#6-technology-stack--programming-languages-deep-dive)")
    lines.append("7. [Enterprise Production Scenarios & Playbooks (நடைமுறை பயன்பாடுகள்)](#7-enterprise-production-scenarios--playbooks)\n")
    lines.append("---\n")

    # Section 1
    lines.append("# 1. Executive Summary & Platform Overview")
    lines.append("### செயல் சுருக்கம் & தளம் கண்ணோட்டம்\n")
    lines.append(
        "UnifAI is a high-performance, enterprise-grade **Unified AI Gateway, Router, Governance, Guardrails, and Observability Control Plane**. "
        "In modern enterprises, AI consumption is fragmented across multiple vendors (OpenAI, Anthropic Claude, Google Gemini, AWS Bedrock, "
        "Mistral, Groq, and self-hosted Ollama/vLLM models). This fragmentation causes unpredictable billing, compliance risks, security vulnerabilities, "
        "and vendor lock-in.\n"
    )
    lines.append("UnifAI unifies all LLM and agentic interactions behind a single, resilient gateway delivering:\n")
    lines.append("1. **Cost Optimization (செலவு குறைப்பு):** Up to 80% cost reduction via **Complexity Routing** and **Semantic Caching**.")
    lines.append("2. **High Availability (99.99% அப்டைம்):** Instant automated failover via **Circuit Breaker** to backup providers during cloud outages.")
    lines.append("3. **Enterprise Governance (அணுகல் கட்டுப்பாடு):** Virtual Keys (`sk-uf-*`), granular RBAC, multi-tenant customer scoping, and team budgets.")
    lines.append("4. **Content Safety & DLP (உள்ளடக்கப் பாதுகாப்பு):** Input and output inspection (PII masking, prompt injection defense) via Google CEL rules and Presidio DLP.")
    lines.append("5. **Agentic Extensibility (டூல்ஸ் பயன்பாடு):** Model Context Protocol (MCP) gateway managing curated tool libraries and OAuth user sessions.")
    lines.append("6. **Real-time Observability (முழுமையான கண்காணிப்பு):** Sub-millisecond logging, live metrics dashboard, and streaming connectors to Datadog, Kafka, and BigQuery.\n")

    # Section 2
    lines.append("# 2. End-to-End Request Lifecycle & System Architecture")
    lines.append("### கோரிக்கை வாழ்க்கைச் சுழற்சி & கணினி கட்டமைப்பு\n")
    lines.append("```")
    lines.append("[ CLIENT APPLICATION / SDK / BROWSER AI ]")
    lines.append("                   │")
    lines.append("                   ▼  (1) HTTP Request with `x-uf-vk: sk-uf-...` & JSON Body")
    lines.append("┌────────────────────────────────────────────────────────────────────────┐")
    lines.append("│ 1. TRANSPORT & GOVERNANCE LAYER                                        │")
    lines.append("│ • FastHTTP router parses headers and TLS connection                    │")
    lines.append("│ • Virtual Key (`sk-uf-*`) validated against PostgreSQL / memory cache   │")
    lines.append("│ • Verify User, Team, Business Unit, and Customer membership             │")
    lines.append("│ • Enforce Rate Limits (RPM, TPM) & Financial Monthly Budgets           │")
    lines.append("│ • Verify Access Profile (allowed models & allowed MCP tool groups)     │")
    lines.append("└──────────────────────────────────┬─────────────────────────────────────┘")
    lines.append("                                   ▼  (2) Authorized Context")
    lines.append("┌────────────────────────────────────────────────────────────────────────┐")
    lines.append("│ 2. PRE-LLM GUARDRAILS & SAFETY LAYER                                   │")
    lines.append("│ • Evaluate Google CEL (Common Expression Language) Rules               │")
    lines.append("│ • Provider scan: Presidio DLP (PII redaction), Llama Guard, Regex      │")
    lines.append("│ • Prompt Injection & Jailbreak detection                               │")
    lines.append("│ • Short-circuit with 400 Bad Request if policy violated                │")
    lines.append("└──────────────────────────────────┬─────────────────────────────────────┘")
    lines.append("                                   ▼  (3) Safe, Sanitized Prompt")
    lines.append("┌────────────────────────────────────────────────────────────────────────┐")
    lines.append("│ 3. SEMANTIC CACHING LAYER                                              │")
    lines.append("│ • Generate vector embedding of prompt or exact SHA256 key              │")
    lines.append("│ • Search Qdrant / PgVector / Redis cache partition (`x-uf-cache-key`)   │")
    lines.append("│ • If Cosine Similarity >= Threshold (`x-uf-cache-threshold`):          │")
    lines.append("│   ───► [CACHE HIT] Return cached response directly (Latency < 20ms)    │")
    lines.append("└──────────────────────────────────┬─────────────────────────────────────┘")
    lines.append("                                   ▼  (4) Cache Miss -> Needs LLM")
    lines.append("┌────────────────────────────────────────────────────────────────────────┐")
    lines.append("│ 4. INTELLIGENT ROUTING & LOAD BALANCING LAYER                          │")
    lines.append("│ • Complexity Router analyzes prompt (Simple, Medium, Complex, Reason)   │")
    lines.append("│ • Routing Rules match CEL conditions (tier, department, tags)          │")
    lines.append("│ • Adaptive Router checks real-time latency & error rate weights        │")
    lines.append("│ • Circuit Breaker checks provider health; routes to fallback if down   │")
    lines.append("│ • Resolved target model & provider determined                          │")
    lines.append("└──────────────────┬────────────────────────────────┬────────────────────┘")
    lines.append("                   ▼                                ▼")
    lines.append("┌────────────────────────────────────┐ ┌─────────────────────────────────┐")
    lines.append("│ 5A. UPSTREAM LLM PROVIDER          │ │ 5B. MCP GATEWAY EXECUTION       │")
    lines.append("│ • OpenAI, Anthropic, Bedrock, etc. │ │ • Model Context Protocol Tools  │")
    lines.append("│ • API key from rotation pool       │ │ • OAuth token session injected  │")
    lines.append("│ • Streaming SSE / Full response    │ │ • Sandboxed Starlark execution  │")
    lines.append("└──────────────────┬─────────────────┘ └────────────────┬────────────────┘")
    lines.append("                   └────────────────┬───────────────────┘")
    lines.append("                                    ▼  (5) Raw Generated AI Output")
    lines.append("┌────────────────────────────────────────────────────────────────────────┐")
    lines.append("│ 6. POST-LLM GUARDRAILS & OUTPUT FILTERING                              │")
    lines.append("│ • Scan response for hallucinated credentials or toxic outputs          │")
    lines.append("│ • PII masking on generated text before sending to client               │")
    lines.append("└──────────────────────────────────┬─────────────────────────────────────┘")
    lines.append("                                   ▼  (6) Final Validated Response")
    lines.append("┌────────────────────────────────────────────────────────────────────────┐")
    lines.append("│ 7. OBSERVABILITY, METRICS & TELEMETRY                                  │")
    lines.append("│ • Append complete record to LLM Logs & MCP Logs                        │")
    lines.append("│ • Calculate exact token cost via Pricing Overrides                     │")
    lines.append("│ • Update real-time Metrics Dashboard via WebSockets                    │")
    lines.append("│ • Stream telemetry event to Connectors (Datadog, Kafka, BigQuery, OTel)│")
    lines.append("│ • Attach `x-unifai-*` response headers and return HTTP 200 OK to Client│")
    lines.append("└────────────────────────────────────────────────────────────────────────┘")
    lines.append("```\n")

    # Section 3
    lines.append("# 3. Exhaustive HTTP Headers Reference Guide")
    lines.append("### ஹெடர்ஸ் முழு விவரக் கையேடு\n")
    lines.append("Headers allow fine-grained control over routing, caching, security, debugging, and logging per request without altering the JSON body.\n")

    for cat in HEADERS_DATA:
        lines.append(f"## {cat['category']}\n")
        lines.append("| Header Name | Sample Value | Data Type | Description (English) | விளக்கம் (Tamil) |")
        lines.append("| :--- | :--- | :--- | :--- | :--- |")
        for h, sample, dtype, desc_en, desc_ta in cat["headers"]:
            lines.append(f"| `{h}` | `{sample}` | {dtype} | {desc_en} | {desc_ta} |")
        lines.append("\n")

    # Section 4
    lines.append("# 4. In-Depth Feature Catalog (7 Core Pillars - 38 Features)")
    lines.append("### அனைத்து 38 அம்சங்களின் முழுமையான தொழில்நுட்ப உடற்கூறு ஆய்வு (Deep-Dive Technical Dissection)\n")

    for p in PILLARS_DATA:
        lines.append(f"## {p['pillar_title']}\n")
        lines.append(f"*{p['pillar_desc']}*\n")
        
        for f in p["features"]:
            lines.append(f"### {f['name']} (`{f['route']}`)\n")
            lines.append(f"**சுருக்கம் / Overview:** {f['summary']}\n")
            lines.append("#### ⚙️ உட்புற கட்டமைப்பு & இயங்கும் முறை (Internal Architecture & Mechanics)")
            lines.append(f"{f['architecture']}\n")
            lines.append("#### 🧩 முக்கிய கூறுகள் & அமைப்புகள் (Key Components & Capabilities)")
            for comp in f["components"]:
                lines.append(f"- **{comp.split(':', 1)[0]}:** {comp.split(':', 1)[1]}")
            lines.append(f"\n#### 🔗 மற்ற கூறுகளுடன் தொடர்பு (Module Interconnections)")
            lines.append(f"{f['interconnections']}\n")
            lines.append(f"#### 💡 நடைமுறை பயன்பாடு (Enterprise Production Use Case)")
            lines.append(f"{f['use_case']}\n")
            lines.append("---\n")

    # Section 5
    lines.append("# 5. Cross-Feature Interconnection & Data Flow Matrix")
    lines.append("### இணைப்பு வரைபடம் (எப்படி ஒன்றுடன் ஒன்று இணைகிறது?)\n")
    lines.append("| மூல கூறு (Source Module) | இணைக்கப்பட்டுள்ள கூறுகள் (Connected Modules) | தொழில்நுட்ப தொடர்பு & Data Flow (Technical Relationship) |")
    lines.append("| :--- | :--- | :--- |")
    lines.append("| **Virtual Keys** | Users, Teams, Customers, Budgets, Access Profiles | Authenticates request; resolves quota limits, team ownership, and access profile. |")
    lines.append("| **Access Profiles** | Model Catalog, MCP Tool Groups | Restricts the virtual key to specific approved models and MCP tool groups. |")
    lines.append("| **Complexity Router** | Routing Rules, Model Catalog | Classifies prompt tier (Simple/Reasoning) and sets context for routing rules. |")
    lines.append("| **Routing Rules** | Model Providers, Circuit Breaker | Evaluates CEL expressions; passes target provider to Circuit Breaker for health check. |")
    lines.append("| **Circuit Breaker** | Routing Rules, Fallback Providers | If primary provider is TRIPPED, automatically switches to next configured fallback. |")
    lines.append("| **Guardrails Rules** | Guardrail Providers, LLM Pipeline | Intercepts prompt pre-LLM and response post-LLM; runs PII/injection validation. |")
    lines.append("| **Semantic Cache** | Vector Store, Models Pipeline | Intercepts request before LLM; returns cached hit or stores new completion on miss. |")
    lines.append("| **MCP Tool Groups** | MCP Catalog, Virtual Keys | Injects approved tool JSON schemas into LLM prompt; audits calls in MCP Logs. |")
    lines.append("| **Browser AI** | Guardrails, Observability | Intercepts browser web AI queries, runs DLP guardrail rules, logs to Browser AI logs. |")
    lines.append("| **Pricing Overrides** | Dashboard, Budgets & Limits, LLM Logs | Computes exact dollar cost per token; updates budget balances and dashboard graphs. |")
    lines.append("| **Connectors** | LLM Logs, MCP Logs, Audit Logs | Streams structured JSON telemetry to external platforms (Datadog, Kafka, BigQuery). |")
    lines.append("| **SCIM Provisioning** | Users, Teams, RBAC | Automatically provisions enterprise IdP accounts and maps them into UnifAI roles. |\n")

    # Section 6
    lines.append("# 6. Technology Stack & Programming Languages Deep Dive")
    lines.append("### தொழில்நுட்ப கட்டமைப்பு & தேர்வு செய்யப்பட்டதற்கான காரணங்கள்\n")
    lines.append("| அடுக்கு (Layer) | தொழில்நுட்பம் / மொழி (Tech Stack) | பயன்பாடு & கட்டடக்கலை நன்மை (Architectural Benefit & Rationale) |")
    lines.append("| :--- | :--- | :--- |")
    for t in TECH_STACK_DATA:
        lines.append(f"| **{t['layer']}** | `{t['tech']}` | {t['why']} |")
    lines.append("\n")

    # Section 7
    lines.append("# 7. Enterprise Production Scenarios & Playbooks")
    lines.append("### நடைமுறை பயன்பாடுகள் & தயாரிப்பு காட்சிகள் (Playbooks)\n")
    lines.append("### Scenario 1: Multi-Tenant B2B SaaS Cost Attribution")
    lines.append("- **சவால் (Challenge):** ஒரு B2B SaaS நிறுவனம் 500 நிறுவன வாடிக்கையாளர்களுக்கு ஒரே AI அசிஸ்டெண்ட்டை வழங்குகிறது. ஒவ்வொரு வாடிக்கையாளரும் எவ்வளவு AI செலவு செய்கிறார்கள் என்று துல்லியமாக பில் செய்ய வேண்டும்.")
    lines.append("- **தீர்வு (Solution):** Backend சேவை `x-uf-customer-id: cust_42` ஹெடரை இணைத்து UnifAI வழியாக அழைக்கிறது. UnifAI ஒவ்வொரு வாடிக்கையாளரின் டோக்கன் மற்றும் டாலர் செலவை தனித்தனியாக பிரித்து BigQuery-க்கு stream செய்கிறது. மாத இறுதியில் தானியங்கி இன்வாய்ஸ் உருவாக்கப்படுகிறது.\n")

    lines.append("### Scenario 2: Zero-Downtime High Availability Failover")
    lines.append("- **சவால் (Challenge):** OpenAI சர்வர்கள் செயலிழக்கும்போது அல்லது 504 Gateway Timeouts வரும்போது, வாடிக்கையாளர் சேவை சாட்பாட் முடங்கி விடுகிறது.")
    lines.append("- **தீர்வு (Solution):** `gpt-4o` மாடலுக்கு Circuit Breaker அமைக்கப்பட்டு, 50% தோல்வி விகிதம் ஏற்பட்டால் Anthropic `claude-3-5-sonnet` மாடலுக்கு தானாக மாறுகிறது. <100ms-ல் failover முடிந்து, அழைப்பாளருக்கு `x-unifai-fallback-index: 1` ஹெடர் திரும்புகிறது. ஒரு பயனர் அமர்வும் பாதிக்கப்படுவதில்லை.\n")

    lines.append("### Scenario 3: Slashing Costs via Complexity Router & Semantic Cache")
    lines.append("- **சவால் (Challenge):** ஒரு நாளைக்கு 1,000,000 கேள்விகள் வருகின்றன. 60% கேள்விகள் எளிய கேள்விகள் ('கடை எப்போது திறக்கும்?'). ஆனால் அனைத்திற்கும் விலையுயர்ந்த Frontier மாடல்கள் பயன்படுத்தப்படுவதால் மாதாந்திர பில் $45,000 ஆகிறது.")
    lines.append("- **தீர்வு (Solution):** Semantic Caching 25% கேள்விகளுக்கு Redis-ல் இருந்து உடனடி பதிலை $0.00 செலவில் அளிக்கிறது. மீதமுள்ள எளிய கேள்விகளை Complexity Router அடையாளம் கண்டு `gpt-4o-mini` ($0.15/1M) மாடலுக்கும், கடினமான கேள்விகளை மட்டும் `o1` மாடலுக்கும் அனுப்புகிறது. மொத்த செலவு $9,800 ஆக (78% குறைவு) குறைகிறது.\n")

    lines.append("### Scenario 4: Enterprise Data Loss Prevention (DLP) using Browser AI")
    lines.append("- **சவால் (Challenge):** நிறுவன ஊழியர்கள் பொது ChatGPT தளத்தில் நிறுவனத்தின் ரகசிய சோர்ஸ் கோட், வாடிக்கையாளர் ஆதார் எண்கள் மற்றும் API சாவி-களை பேஸ்ட் செய்து விடுகிறார்கள்.")
    lines.append("- **தீர்வு (Solution):** ஊழியர்களின் மடிக்கணினிகளில் UnifAI Browser AI agent நிறுவப்படுகிறது. ஊழியர் ChatGPT தளத்தில் ஏதேனும் ரகசிய தகவலை பேஸ்ட் செய்யும்போது, Browser AI உடனடியாக தடுத்து நிறுத்தி, Audit Logs-ல் எச்சரிக்கையை பதிவு செய்கிறது.\n")

    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Markdown successfully written to: {MD_PATH}")

# ==============================================================================
# 2. DOCX GENERATION FUNCTION
# ==============================================================================
def generate_docx():
    print("Generating comprehensive Word Document (.docx)...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        hp = section.header.paragraphs[0]
        hp.text = "UnifAI Enterprise Deep-Dive Architecture & Master Feature Guide"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if hp.runs:
            hp.runs[0].font.name = "Segoe UI"
            hp.runs[0].font.size = Pt(8.5)
            hp.runs[0].font.color.rgb = RGBColor(100, 116, 139)
            
        fp = section.footer.paragraphs[0]
        fp.text = "Confidential & Proprietary — UnifAI Control Plane Engineering Manual"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if fp.runs:
            fp.runs[0].font.name = "Segoe UI"
            fp.runs[0].font.size = Pt(8.5)
            fp.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    def style_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        return p

    def style_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def style_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.italic = italic
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def format_table(table, col_widths, col_names):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(col_names):
            hdr_cells[i].text = name
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
            tcPr.append(shd)
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = "Segoe UI"
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for r_idx, row in enumerate(table.rows[1:]):
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                tcPr.append(shd)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                    
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    rt = p_title.add_run("UnifAI Enterprise Deep-Dive Architecture & Feature Master Guide")
    rt.font.name = "Segoe UI"
    rt.font.size = Pt(22)
    rt.bold = True
    rt.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    rsub = p_sub.add_run("விரிவான கணினி கட்டமைப்பு, ஹெடர்ஸ் மற்றும் 38 அம்சங்களின் முழுமையான கையேடு (Tamil & English Technical Manual)")
    rsub.font.name = "Segoe UI"
    rsub.font.size = Pt(11)
    rsub.font.color.rgb = RGBColor(71, 85, 105)

    # Section 1
    style_h1("1. Executive Summary & Core Architecture")
    add_p(
        "UnifAI unifies all LLM interactions across cloud and self-hosted models behind a high-concurrency Go-based gateway. "
        "It provides cost optimization up to 80% via Complexity Routing and Semantic Caching, 99.99% uptime via automated Circuit Breakers, "
        "comprehensive governance with Virtual Keys, and pre/post-LLM content guardrails."
    )

    # Section 2
    style_h1("2. End-to-End 7-Stage Request Lifecycle")
    stages = [
        ("Stage 1 - Transport & Governance: ", "FastHTTP router validates Virtual Key (`sk-uf-*`), enforces rate limits (RPM, TPM), and checks budget ceilings."),
        ("Stage 2 - Pre-LLM Guardrails: ", "Evaluates Google CEL rules and Presidio DLP to redact PII and block prompt injections."),
        ("Stage 3 - Semantic Caching: ", "Searches Redis / Qdrant for vector similarity. Cache hit returns instantly (<20ms, $0 cost)."),
        ("Stage 4 - Intelligent Routing: ", "Complexity Router classifies prompt; Routing Rules and Circuit Breakers resolve target model."),
        ("Stage 5 - Upstream Execution: ", "Dispatches request to upstream LLM or executes MCP tools with OAuth user credentials."),
        ("Stage 6 - Post-LLM Guardrails: ", "Inspects generated text for toxic outputs or credential leaks before returning to client."),
        ("Stage 7 - Observability & Telemetry: ", "Appends record to LLM Logs, streams to Datadog/Kafka, and returns response with `x-unifai-*` headers.")
    ]
    for st_title, st_desc in stages:
        add_bullet(st_desc, st_title)

    # Section 3: Headers
    style_h1("3. Exhaustive HTTP Headers Reference Guide")
    for cat in HEADERS_DATA:
        style_h2(cat["category"])
        t = doc.add_table(rows=len(cat["headers"])+1, cols=4)
        col_names = ["Header Name", "Sample Value", "Type", "Description (English & Tamil)"]
        format_table(t, [1.5, 1.3, 1.0, 3.2], col_names)
        for idx, (h, sample, dtype, desc_en, desc_ta) in enumerate(cat["headers"]):
            t.rows[idx+1].cells[0].text = h
            t.rows[idx+1].cells[1].text = sample
            t.rows[idx+1].cells[2].text = dtype
            t.rows[idx+1].cells[3].text = f"{desc_en}\n\n[தமிழ்] {desc_ta}"

    # Section 4: All 38 Features Detailed
    style_h1("4. In-Depth Technical Dissection of All 38 Features (7 Pillars)")
    for p in PILLARS_DATA:
        style_h2(p["pillar_title"])
        add_p(p["pillar_desc"], italic=True)
        for f in p["features"]:
            style_h3(f"{f['name']} ({f['route']})")
            add_p(f["summary"], bold_prefix="Overview: ")
            add_p(f["architecture"], bold_prefix="Internal Architecture & Mechanics: ")
            for comp in f["components"]:
                add_bullet(comp.split(':', 1)[1], comp.split(':', 1)[0] + ":")
            add_p(f["interconnections"], bold_prefix="Module Interconnections: ")
            add_p(f["use_case"], bold_prefix="Production Use Case: ")

    # Section 5: Tech Stack
    style_h1("5. Technology Stack & Programming Languages Deep Dive")
    t_tech = doc.add_table(rows=len(TECH_STACK_DATA)+1, cols=3)
    format_table(t_tech, [1.8, 1.8, 3.4], ["Layer", "Technology / Language", "Architectural Role & Benefit"])
    for idx, t in enumerate(TECH_STACK_DATA):
        t_tech.rows[idx+1].cells[0].text = t["layer"]
        t_tech.rows[idx+1].cells[1].text = t["tech"]
        t_tech.rows[idx+1].cells[2].text = t["why"]

    # Section 6: Scenarios
    style_h1("6. Enterprise Production Scenarios & Playbooks")
    scenarios = [
        ("Scenario 1 - Multi-Tenant B2B SaaS Cost Attribution: ", "SaaS backend attaches `x-uf-customer-id: cust_42` on every call. UnifAI attributes token spend per customer, exported to BigQuery for automated monthly invoicing."),
        ("Scenario 2 - Zero-Downtime High Availability Failover: ", "When OpenAI returns 504 errors, Circuit Breaker detects >50% failure rate and automatically shifts traffic to Claude 3.5 Sonnet in <100ms. Zero user sessions drop."),
        ("Scenario 3 - Slashing Costs via Complexity Router & Semantic Cache: ", "Semantic Caching serves 25% of queries from Redis at $0.00 cost; Complexity Router directs simple questions to micro-models, dropping monthly AI bill from $45,000 to $9,800 (78% savings)."),
        ("Scenario 4 - Enterprise Data Loss Prevention (DLP) using Browser AI: ", "Browser AI local agent monitors web AI portals on employee laptops, intercepting clipboard pastes containing credit cards or passwords, preventing corporate IP leaks.")
    ]
    for sc_title, sc_desc in scenarios:
        add_p(sc_desc, bold_prefix=sc_title)

    doc.save(DOCX_PATH)
    print(f"Word document successfully written to: {DOCX_PATH}")

# ==============================================================================
# 3. PDF GENERATION FUNCTION VIA REPORTLAB
# ==============================================================================
def generate_pdf():
    print("Generating comprehensive Multi-Page PDF Document (.pdf)...")

    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748B"))
            
            # Running Header on page 2 onwards
            if self._pageNumber > 1:
                self.drawString(40, 11 * inch - 36, "UnifAI Enterprise Deep-Dive Architecture & Feature Master Guide")
                self.drawRightString(8.5 * inch - 40, 11 * inch - 36, "Confidential — Engineering Manual")
                self.setStrokeColor(colors.HexColor("#CBD5E1"))
                self.setLineWidth(0.5)
                self.line(40, 11 * inch - 40, 8.5 * inch - 40, 11 * inch - 40)
                
            # Running Footer on all pages
            self.drawString(40, 32, "UnifAI Unified AI Gateway & Control Plane")
            page_str = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(8.5 * inch - 40, 32, page_str)
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(40, 42, 8.5 * inch - 40, 42)
            
            self.restoreState()

    pdf_doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=48,
        bottomMargin=48
    )

    styles = getSampleStyleSheet()

    p_title_style = ParagraphStyle(
        'PdfTitle',
        fontName='LathaBold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=3
    )

    p_subtitle_style = ParagraphStyle(
        'PdfSubTitle',
        fontName='Latha',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#475569'),
        spaceAfter=10
    )

    h1_style = ParagraphStyle(
        'PdfH1',
        fontName='LathaBold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=12,
        spaceAfter=4,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'PdfH2',
        fontName='LathaBold',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'PdfH3',
        fontName='LathaBold',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=6,
        spaceAfter=2,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'PdfBody',
        fontName='Latha',
        fontSize=7.5,
        leading=10.5,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=3
    )

    bullet_style = ParagraphStyle(
        'PdfBullet',
        fontName='Latha',
        fontSize=7.5,
        leading=10.5,
        textColor=colors.HexColor('#1E293B'),
        leftIndent=10,
        spaceAfter=2
    )

    th_style = ParagraphStyle(
        'PdfTH',
        fontName='LathaBold',
        fontSize=7,
        leading=9,
        textColor=colors.white
    )

    td_style = ParagraphStyle(
        'PdfTD',
        fontName='Latha',
        fontSize=6.5,
        leading=8.5,
        textColor=colors.HexColor('#0F172A')
    )

    td_code_style = ParagraphStyle(
        'PdfTDCode',
        fontName='Helvetica-Bold',
        fontSize=6.5,
        leading=8.5,
        textColor=colors.HexColor('#0F172A')
    )

    story = []

    # Title & Metadata
    story.append(Paragraph("UnifAI Enterprise Deep-Dive Architecture & Feature Master Guide", p_title_style))
    story.append(Paragraph("விரிவான கணினி கட்டமைப்பு, ஹெடர்ஸ் மற்றும் 38 அம்சங்களின் முழுமையான கையேடு (Tamil & English Technical Manual)", p_subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#0284C7"), spaceAfter=8))

    # Executive Summary
    story.append(Paragraph("1. Executive Summary & Core Platform Overview", h1_style))
    story.append(Paragraph(
        "UnifAI is a high-performance Unified AI Gateway, Router, Governance, and Observability Control Plane. "
        "It centralizes AI consumption across OpenAI, Anthropic Claude, Google Gemini, AWS Bedrock, Mistral, and local models. "
        "It delivers up to <b>80% cost reduction</b> via Complexity Routing and Semantic Caching, <b>99.99% uptime</b> via Circuit Breakers, "
        "comprehensive governance with Virtual Keys (`sk-uf-*`), and pre/post-LLM content guardrails.",
        body_style
    ))
    story.append(Spacer(1, 4))

    # Lifecycle
    story.append(Paragraph("2. End-to-End 7-Stage Request Lifecycle", h1_style))
    story.append(Paragraph("1. <b>Transport & Governance:</b> FastHTTP validates Virtual Key (`sk-uf-*`), enforces rate limits and budget quotas.<br/>"
                           "2. <b>Pre-LLM Guardrails:</b> Evaluates Google CEL rules and Presidio DLP to redact PII and block prompt injections.<br/>"
                           "3. <b>Semantic Caching:</b> Vector similarity search in Redis/Qdrant. Cache hit returns immediately (<20ms, $0 cost).<br/>"
                           "4. <b>Intelligent Routing:</b> Complexity Router classifies prompt (Simple/Reasoning); Circuit Breakers check provider health.<br/>"
                           "5. <b>Upstream Execution:</b> Dispatches prompt to target LLM or executes MCP tools with OAuth user credentials.<br/>"
                           "6. <b>Post-LLM Guardrails:</b> Scans generated output for toxic content or credential leaks before returning.<br/>"
                           "7. <b>Observability & Telemetry:</b> Appends record to LLM Logs, streams to Datadog/Kafka, and returns response with `x-unifai-*` headers.", body_style))
    story.append(Spacer(1, 6))

    # Headers Section
    story.append(Paragraph("3. Exhaustive HTTP Headers Reference Guide", h1_style))
    
    def make_header_table(headers_list):
        t_data = [[Paragraph(h, th_style) for h in ["Header Name", "Sample Value", "Type", "Description (English & Tamil)"]]]
        for h, sample, dtype, desc_en, desc_ta in headers_list:
            col0 = Paragraph(h, td_code_style)
            col1 = Paragraph(sample, td_style)
            col2 = Paragraph(dtype, td_style)
            col3 = Paragraph(f"<b>{desc_en}</b><br/>[தமிழ்] {desc_ta}", td_style)
            t_data.append([col0, col1, col2, col3])
        tbl = Table(t_data, colWidths=[105, 85, 70, 270])
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
            ('TOPPADDING', (0, 0), (-1, -1), 2.5),
            ('LEFTPADDING', (0, 0), (-1, -1), 3),
            ('RIGHTPADDING', (0, 0), (-1, -1), 3),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
        ]))
        return tbl

    for cat in HEADERS_DATA:
        story.append(Paragraph(cat["category"], h2_style))
        story.append(make_header_table(cat["headers"]))
        story.append(Spacer(1, 5))

    # 38 Features Detailed Breakdown
    story.append(Paragraph("4. In-Depth Technical Dissection of All 38 Features (7 Pillars)", h1_style))
    for p in PILLARS_DATA:
        story.append(Paragraph(p["pillar_title"], h2_style))
        story.append(Paragraph(p["pillar_desc"], body_style))
        story.append(Spacer(1, 3))
        
        for f in p["features"]:
            feat_elements = []
            feat_elements.append(Paragraph(f"<b>{f['name']}</b> ({f['route']})", h3_style))
            feat_elements.append(Paragraph(f"<b>Overview:</b> {f['summary']}", body_style))
            feat_elements.append(Paragraph(f"<b>Internal Architecture:</b> {f['architecture']}", body_style))
            for comp in f["components"]:
                feat_elements.append(Paragraph(f"• <b>{comp.split(':', 1)[0]}:</b> {comp.split(':', 1)[1]}", bullet_style))
            feat_elements.append(Paragraph(f"<b>Interconnections:</b> {f['interconnections']}", body_style))
            feat_elements.append(Paragraph(f"<b>Production Use Case:</b> {f['use_case']}", body_style))
            feat_elements.append(Spacer(1, 4))
            story.append(KeepTogether(feat_elements))

    # Technology Stack
    story.append(Paragraph("5. Technology Stack & Programming Languages Deep Dive", h1_style))
    tech_tdata = [[Paragraph(h, th_style) for h in ["Layer", "Technology / Language", "Architectural Role & Benefit"]]]
    for t in TECH_STACK_DATA:
        tech_tdata.append([Paragraph(t["layer"], td_code_style), Paragraph(t["tech"], td_code_style), Paragraph(t["why"], td_style)])
    tbl_tech = Table(tech_tdata, colWidths=[95, 115, 320])
    tbl_tech.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_tech)
    story.append(Spacer(1, 6))

    # Production Scenarios
    story.append(Paragraph("6. Enterprise Production Scenarios & Playbooks", h1_style))
    scenarios_pdf = [
        ("Scenario 1 - Multi-Tenant B2B SaaS Cost Attribution: ", "SaaS backend attaches `x-uf-customer-id: cust_42` on every call. UnifAI attributes token spend per customer, exported to BigQuery for automated monthly invoicing."),
        ("Scenario 2 - Zero-Downtime High Availability Failover: ", "When OpenAI returns 504 errors, Circuit Breaker detects >50% failure rate and automatically shifts traffic to Claude 3.5 Sonnet in <100ms. Zero user sessions drop."),
        ("Scenario 3 - Slashing Costs via Complexity Router & Semantic Cache: ", "Semantic Caching serves 25% of queries from Redis at $0.00 cost; Complexity Router directs simple questions to micro-models, dropping monthly AI bill from $45,000 to $9,800 (78% savings)."),
        ("Scenario 4 - Enterprise Data Loss Prevention (DLP) using Browser AI: ", "Browser AI local agent monitors web AI portals on employee laptops, intercepting clipboard pastes containing credit cards or passwords, preventing corporate IP leaks.")
    ]
    for sc_title, sc_desc in scenarios_pdf:
        story.append(Paragraph(f"<b>{sc_title}</b>{sc_desc}", body_style))
        story.append(Spacer(1, 2))

    pdf_doc.build(story, canvasmaker=NumberedCanvas)
    print(f"PDF document successfully written to: {PDF_PATH}")

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================
if __name__ == "__main__":
    print("Executing master manual build...")
    generate_markdown()
    generate_docx()
    generate_pdf()
    print("ALL DOCUMENTS GENERATED EXHAUSTIVELY IN:", DOC_DIR)

