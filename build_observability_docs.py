# -*- coding: utf-8 -*-
"""
UnifAI Observability Master Deep-Dive Guide Generator
Generates:
  - document/UnifAI_Observability_Master_Guide.md
  - document/UnifAI_Observability_Master_Guide.docx
  - document/UnifAI_Observability_Master_Guide.pdf
"""

import os
import sys

# Ensure UTF-8 encoding for Windows console
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
if sys.stderr.encoding != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

DOC_DIR = os.path.join(os.getcwd(), "document")
os.makedirs(DOC_DIR, exist_ok=True)

MD_PATH = os.path.join(DOC_DIR, "UnifAI_Observability_Master_Guide.md")
DOCX_PATH = os.path.join(DOC_DIR, "UnifAI_Observability_Master_Guide.docx")
PDF_PATH = os.path.join(DOC_DIR, "UnifAI_Observability_Master_Guide.pdf")

# Register fonts for Tamil + English in ReportLab
pdfmetrics.registerFont(TTFont('Latha', 'C:/Windows/Fonts/latha.ttf'))
pdfmetrics.registerFont(TTFont('LathaBold', 'C:/Windows/Fonts/lathab.ttf'))

# ==============================================================================
# OBSERVABILITY MASTER DATA DEFINITION
# ==============================================================================

OBS_FEATURES = [
    {
        "id": "dashboard",
        "name": "Dashboard (கண்காணிப்பு முகப்புத் திரை)",
        "route": "/workspace/dashboard",
        "non_tech": {
            "analogy": "விமானத்தின் காக்பிட் (Airplane Cockpit) அல்லது சொகுசு காரின் டேஷ்போர்டு போன்றது.",
            "explanation": (
                "ஒரு கார் ஓட்டும்போது ஸ்பீடோமீட்டர் வேகம், பெட்ரோல் அளவு மற்றும் இன்ஜின் நிலைமை காட்டுவது போல, "
                "கம்பெனியில் AI பயன்பாடு எப்படி நடக்கிறது என்பதை தலைவர்கள் ஒரே பார்வையில் பார்க்கும் இடம். "
                "இன்று எத்தனை லட்சம் AI கேள்விகள் கேட்கப்பட்டது, எவ்வளவு டாலர் செலவாகியுள்ளது, எந்த AI மாடல் வேகமாக இயங்குகிறது, "
                "எந்த டீம் அதிக டோக்கன் பயன்படுத்துகிறது என்பதை நிகழ்நேரத்தில் (Realtime Graphs) காட்டும்."
            ),
            "business_value": "செலவை உடனுக்குடன் கட்டுப்படுத்தலாம்; திடீரென பில் எகிறாமல் தடுக்கலாம்; AI அமைப்பின் வேகத்தை கண்காணிக்கலாம்."
        },
        "tech": {
            "architecture": (
                "The Dashboard is powered by a persistent WebSocket connection ('useWebSocket') synchronized with TanStack Query. "
                "Backend telemetry is collected at the FastHTTP transport layer using lock-free atomic counters (sync/atomic). "
                "Metrics are aggregated into time-series buckets (1-min, 5-min, 1-hour intervals) and queried via PostgreSQL. "
                "Calculates percentile response times (p50, p90, p95, p99) and Time-To-First-Token (TTFT). "
                "Dynamically prices token volumes using live rules from Pricing Overrides."
            ),
            "endpoints": ["GET /api/v1/dashboard/stats", "GET /api/v1/dashboard/histogram", "WS /api/v1/ws/dashboard"]
        },
        "ui_elements": {
            "top_bar": [
                "Time Period Selector: Quick toggles for 1h, 24h, 7d, 30d, and custom date range.",
                "Date Range Picker (DateTimePickerWithRange): Calendar and time modal for custom historical analysis.",
                "Timezone Dropdown: Switch between UTC and Local Browser Timezone.",
                "Export Popover Button: Download metrics in CSV, JSON, or PNG chart formats.",
                "Filter Sidebar Toggle Button: Slides out 'LogsFilterSidebar' to filter by Virtual Keys, Providers, and Models."
            ],
            "tabs_and_views": [
                "Overview Tab: Request volume charts, total token consumption, total spend ($), and p50/p95/p99 latency curves.",
                "Provider Usage Tab: Comparative traffic distribution across OpenAI, Anthropic, Bedrock, and Ollama.",
                "Model Rankings Tab: Leaderboard of top models ranked by request count, cost, and latency.",
                "Dimension Rankings Tab: Usage breakdown by custom tags passed in 'x-uf-dim-*' (e.g. environment, department).",
                "MCP Tab: Volume and dollar cost incurred specifically by Model Context Protocol tool executions.",
                "Chart Type Toggle: Button switching between Line Graph and Bar Chart views."
            ],
            "bottom_elements": [
                "Model Filter Select: Dropdown isolating individual model trends.",
                "Cache Hit Ratio Meter: Live gauge showing percentage of requests resolved via Semantic Cache at $0 cost."
            ]
        },
        "connections": {
            "receives_from": "Aggregates real-time data from LLM Logs, Pricing Overrides, and Semantic Caching.",
            "triggers_and_affects": "If error rates spike >5% or spending exceeds 80% of budget, triggers automated alerts to Alert Channels (Slack/PagerDuty)."
        },
        "use_case": "ஒரு ஃபின்டெக் கம்பெனியின் CTO காலை 9 மணிக்கு டேஷ்போர்டைப் பார்த்து, நேற்று இரவு OpenAI-ல் ஏற்பட்ட லேடன்சி ஸ்பைக்கை (p99 > 4000ms) கண்டறிந்து, ரூட்டிங் விதியை Claude-க்கு மாற்ற உத்தரவிடுகிறார்."
    },
    {
        "id": "llm_logs",
        "name": "LLM Logs (AI கோரிக்கை பரிவர்த்தனை பதிவேடு)",
        "route": "/workspace/logs",
        "non_tech": {
            "analogy": "வங்கியின் பாஸ்புக் (Bank Statement) + அலுவலகத்தின் சிசிடிவி கேமரா பதிவு (CCTV Footage).",
            "explanation": (
                "யார், எப்போது, எந்த AI மாடலிடம் என்ன கேள்வி கேட்டார்கள்? மாடல் என்ன பதில் தந்தது? "
                "அதற்கு எத்தனை ரூபாய்/டாலர் செலவானது? பதில் வர எத்தனை விநாடி எடுத்தது? ஏதேனும் ரகசிய ஆதார்/கிரெடிட் கார்டு "
                "தகவல்கள் மறைக்கப்பட்டதா? என்பதை ஒவ்வொரு பரிவர்த்தனையாக அக்குவேறு ஆணிவேறாகப் பதிவு செய்து வைக்கும் ரசீது புத்தகம்."
            ),
            "business_value": "பிழைகளை உடனடியாக சரிசெய்யலாம் (Debugging), சட்ட ரீதியான தணிக்கைக்கு (Audit Compliance) முழு ஆதாரம் கிடைக்கும்."
        },
        "tech": {
            "architecture": (
                "Implemented with an asynchronous, non-blocking ring-buffer queue ('framework/logstore'). "
                "FastHTTP request context captures exact request headers, messages, and wire bytes. "
                "A decoupled worker pool writes micro-batches to PostgreSQL using 'pgx v5' connection pooling. "
                "Every log row records: request_id, session_id, virtual_key_id, provider, requested_model, resolved_model, "
                "prompt_tokens, completion_tokens, total_tokens, cost_usd, latency_ms, ttft_ms, status_code, and guardrail_verdict."
            ),
            "endpoints": ["GET /api/v1/logs", "GET /api/v1/logs/:id", "POST /api/v1/logs/delete"]
        },
        "ui_elements": {
            "top_bar": [
                "Top Metric Cards: Live animated counters for Total Requests, Total Cost ($), Total Tokens, and Errors.",
                "Search Input Bar: Full-text and regex search matching keywords inside prompt or completion text.",
                "WebSocket Live Indicator: Animated green pulsing dot indicating real-time log ingestion.",
                "Column Customizer (useColumnConfig): Dropdown to show/hide specific table columns.",
                "Export CSV Button: Download filtered log entries as CSV.",
                "Delete Logs Button: Role-protected button (RBAC) to delete historical records."
            ],
            "tabs_and_views": [
                "Collapsible Logs Volume Chart: Hourly/Daily bar graph displaying traffic spikes above the table.",
                "Logs Data Table: Columns for Timestamp, Status (Badge), Model (Original -> Resolved), Provider Icon, Virtual Key, Customer ID, Latency, Tokens, Cost ($), Actions."
            ],
            "bottom_elements": [
                "LogDetailSheet (Slide-over): Opens on row click. 4 Tabs: (1) Request (headers & prompt), (2) Response (completion & finish reason), (3) Guardrails (verdict & redacted entities), (4) Trace (spans & TTFT breakdown).",
                "SessionDetailsSheet (Slide-over): Chronological multi-turn conversation replay for requests sharing the same 'x-uf-session-id'.",
                "Row Action Dropdown: 'Copy Request ID', 'Copy cURL Command', 'View Raw Wire Bytes'."
            ]
        },
        "connections": {
            "receives_from": "Intercepts every HTTP and WebSocket request flowing through the FastHTTP Gateway.",
            "triggers_and_affects": "Feeds Dashboard metrics, drains event streams to Connectors (Datadog/Kafka), and validates Guardrail rule accuracy."
        },
        "use_case": "ஒரு மருத்துவர் AI-யிடம் ஒரு நோயாளியின் தரவை உள்ளிடும்போது, அவரது பெயர் மற்றும் மொபைல் எண் '[REDACTED_PII]' என மாற்றப்பட்டு பாதுகாப்பாக அனுப்பப்பட்டதா என செக்யூரிட்டி அதிகாரி Logs-ல் சரிபார்க்கிறார்."
    },
    {
        "id": "mcp_logs",
        "name": "MCP Logs (ஏஜென்ட் டூல்ஸ் இயக்கப் பதிவேடு)",
        "route": "/workspace/mcp-logs",
        "non_tech": {
            "analogy": "ஒரு மெக்கானிக் அல்லது ஊழியர் பயன்படுத்திய கருவிகளின் டைரி (Tools Usage Audit Book).",
            "explanation": (
                "இன்றைய AI வெறும் பதில் மட்டும் சொல்வதில்லை; அது டேட்டாபேஸை அணுகுகிறது, GitHub-ல் கோட் உருவாக்குகிறது, "
                "Slack-ல் மெசேஜ் அனுப்புகிறது. அப்படி AI ஏஜென்ட் இயக்கிய ஒவ்வொரு வெளிப்புற டூலின் (External Tool) "
                "செயல்பாடுகளையும், என்ன இன்புட் கொடுத்து என்ன ரிசல்ட் எடுத்தது என்பதையும் பதிவு செய்யும் இடம்."
            ),
            "business_value": "AI ஏஜென்ட் தவறான டேட்டாபேஸ் கமாண்ட் இயக்கிவிட்டதா அல்லது ஹேக் ஆகிவிட்டதா என்பதை கண்காணிக்கலாம்."
        },
        "tech": {
            "architecture": (
                "Implemented in 'core/mcp/exec.go'. Instruments tool execution lifecycles across stdio subprocesses and HTTP SSE transports. "
                "Captures calling agent identity, target MCP server, tool function name, parsed input JSON arguments, "
                "execution duration in milliseconds, output JSON return schema, and any stderr/exception traces."
            ),
            "endpoints": ["GET /api/v1/mcp-logs", "GET /api/v1/mcp-logs/:id"]
        },
        "ui_elements": {
            "top_bar": [
                "Tool Name Filter Dropdown: Filter by specific tool (e.g. 'execute_sql_query', 'read_file').",
                "Server Label Dropdown: Filter by MCP server (e.g. 'github', 'postgres', 'slack').",
                "Time Window Selector: Filter tool invocations by time.",
                "Export Button: Export tool audit logs."
            ],
            "tabs_and_views": [
                "MCP Logs Table: Columns for Timestamp, Tool Name (Badge), Server Name, Status (Success/Failed), Duration (ms), Virtual Key, User Session ID."
            ],
            "bottom_elements": [
                "ToolExecutionDetailSheet (Slide-over): Detailed inspector with: (1) Arguments Tab (JSON inputs sent to the tool), (2) Results Tab (JSON response returned by tool), (3) Error Stack Trace Tab."
            ]
        },
        "connections": {
            "receives_from": "Triggered by MCP Gateway execution whenever an LLM initiates a Tool Calling loop.",
            "triggers_and_affects": "Tied to Tool Groups and Auth Sessions. If a tool fails repeatedly, updates Circuit Breaker and Alert Channels."
        },
        "use_case": "ஒரு AI கோடிங் ஏஜென்ட் GitHub Repositories-ல் தேவையில்லாத கோப்பை நீக்கிவிட்டதா என்பதை MCP Logs-ல் சென்று 'delete_file' டூலின் இன்புட் ஆர்குமெண்ட்டை எடுத்து ஆராய்கிறார்கள்."
    },
    {
        "id": "browser_ai",
        "name": "Browser AI (பிரவுசர் பாதுகாப்பு & DLP ப்ராக்ஸி)",
        "route": "/workspace/browser-ai",
        "non_tech": {
            "analogy": "அலுவலக வாயிலில் இருக்கும் செக்யூரிட்டி கார்டு (Security Checkpost at the Gate).",
            "explanation": (
                "ஊழியர்கள் தங்கள் அலுவலக லேப்டாப்பில் பொது இணையதளங்களான ChatGPT, Claude, Perplexity போன்றவற்றைப் பயன்படுத்தும்போது, "
                "நிறுவனத்தின் ரகசிய சோர்ஸ் கோட், பாஸ்வேர்டுகள், வாடிக்கையாளர் விபரங்கள் ஆகியவற்றை காப்பி-பேஸ்ட் செய்துவிடாமல் "
                "நடுவில் நின்று தடுத்து நிறுத்தும் டிஜிட்டல் செக்யூரிட்டி கார்டு."
            ),
            "business_value": "நிறுவனத்தின் மிகப்பெரிய பாதுகாப்பு அச்சுறுத்தலான 'Employee Data Leakage'-ஐ 100% தடுத்து நிறுத்துகிறது."
        },
        "tech": {
            "architecture": (
                "A hybrid two-layer architecture: (1) A lightweight Python proxy daemon ('apps/browser-guard/proxy/browser_ai_proxy.py') "
                "with Chrome/Edge extensions deployed via corporate MDM, and (2) a central management UI in UnifAI. "
                "Intercepts HTTPS paste events on designated AI domains, evaluates Presidio DLP regexes, runs local Ollama classification models, "
                "extracts text from uploaded attachments (PDF, DOCX), and enforces block/mask actions."
            ),
            "endpoints": ["GET /api/v1/browser-ai/logs", "POST /api/v1/browser-ai/rules", "POST /api/v1/browser-ai/targets"]
        },
        "ui_elements": {
            "top_bar": [
                "5 Main Tabs: Logs Tab, Rules Tab, Target Websites Tab, Agents Tab, Settings Tab.",
                "Agent Heartbeat Status Bar: Shows active, online, and disconnected employee machines."
            ],
            "tabs_and_views": [
                "Rules Tab: List of DLP policies with toggle switches (Enabled/Disabled).",
                "Target Websites Tab: Directory of monitored AI websites (chatgpt.com, claude.ai) with host roles (Allow/Block).",
                "Agents Tab: Table of employee laptops (Hostname, IP, OS, Agent Version, Last Heartbeat).",
                "Logs Tab: Intercepted queries showing employee ID, target domain, matched rule, and blocked text preview."
            ],
            "bottom_elements": [
                "+ Create Rule Button: Modal with Policy-to-Regex AI generator (converts English descriptions to Regex).",
                "+ Add Target Website Dialog: Add new AI websites to monitor with attachment scanning toggle.",
                "Save Uninstall Key Button: Sets tamper-proof password required to remove the agent from laptops.",
                "Bulk Delete Agents Button: Clean up decommissioned employee machines."
            ]
        },
        "connections": {
            "receives_from": "Receives real-time intercepted traffic from employee laptop daemons via WebSocket/HTTP.",
            "triggers_and_affects": "Shares Guardrail Providers (Presidio); writes violations directly into Audit Logs and Alert Channels."
        },
        "use_case": "ஒரு ஜூனியர் டெவலப்பர் தனது நிறுவனத்தின் AWS Secret Key-ஐ ChatGPT-ல் பேஸ்ட் செய்ய முயலும்போது, Browser AI அதை உடனடியாக தடுத்து நிறுத்தி, 'Corporate Policy Violation' எச்சரிக்கையை திரையில் காட்டுகிறது."
    },
    {
        "id": "connectors",
        "name": "Connectors (வெளிப்புற கண்காணிப்பு இணைப்பு குழாய்கள்)",
        "route": "/workspace/observability",
        "non_tech": {
            "analogy": "தொழிற்சாலையிலிருந்து கழிவுநீர் அல்லது பொருட்களை பெரிய குழாய் மூலம் நகராட்சி டேங்கிற்கு அனுப்புவது போன்ற எக்ஸ்போர்ட் பைப்லைன் (Export Pipeline).",
            "explanation": (
                "UnifAI-க்குள் மட்டுமே தகவல்களை வைத்திருக்காமல், நிறுவனத்தின் கார்ப்பரேட் கண்காணிப்பு தளங்களான Datadog, New Relic, "
                "Google BigQuery, அல்லது Kafka-வுக்கு இந்த AI லாக் தகவல்களை நிகழ்நேரத்தில் தானாகவே பம்ப் செய்து அனுப்பும் இணைப்பு குழாய்கள்."
            ),
            "business_value": "நிறுவனத்தின் அனைத்து மென்பொருள் மற்றும் AI செலவுகளை ஒரே பெரிய எண்டர்பிரைஸ் டேஷ்போர்டில் ஒருங்கிணைக்கலாம்."
        },
        "tech": {
            "architecture": (
                "Implemented in 'framework/connectors/runtime.go'. Employs a decoupled ring-buffer worker pipeline. "
                "Separates telemetry export from the critical path of the gateway to maintain zero-latency proxying. "
                "Formats internal telemetry into vendor-specific schemas and streams them asynchronously with configurable "
                "batch sizes, flush intervals (ms), retry logic, and exponential jitter backoff."
            ),
            "endpoints": ["GET /api/v1/connectors", "POST /api/v1/connectors", "POST /api/v1/connectors/:id/test"]
        },
        "ui_elements": {
            "top_bar": [
                "+ Add Connector Button: Opens creation sheet for external platforms.",
                "Active Connectors Status Badges: Displays health indicators (Green = Streaming, Red = Unreachable)."
            ],
            "tabs_and_views": [
                "Connector Cards Grid: Datadog, New Relic, Apache Kafka, Google BigQuery, Google PubSub, OpenTelemetry (OTel)."
            ],
            "bottom_elements": [
                "ConnectorConfigSheet (Slide-over): Form inputs for API Key, Endpoint URL, Buffer Size Slider, Flush Interval (ms), Header Whitelist.",
                "Test Connection Button: Dispatches a synthetic ping payload to verify credentials and connectivity before saving."
            ]
        },
        "connections": {
            "receives_from": "Drains event records continuously from LLM Logs, MCP Logs, and Audit Logs.",
            "triggers_and_affects": "Pushes live AI telemetry to external corporate data lakes and SIEM systems."
        },
        "use_case": "ஒரு வங்கியின் பில்லிங் சாப்ட்வேர், Google BigQuery-ல் UnifAI Connectors வழியாக வரும் வாடிக்கையாளர் டோக்கன் டேட்டாவை வைத்து தானாகவே மாதாந்திர இன்வாய்ஸ் தயாரிக்கிறது."
    },
    {
        "id": "logs_settings",
        "name": "Logs Settings (பதிவேடு அமைப்புகள் & கொள்கைகள்)",
        "route": "/workspace/config/logging",
        "non_tech": {
            "analogy": "ஆவணங்களை பாதுகாக்கும் காப்பகம் மற்றும் ஆவண அழிப்பு கொள்கை (Document Retention & Shredding Policy).",
            "explanation": (
                "பழைய லாக் விபரங்களை எத்தனை நாட்கள் வைத்திருக்க வேண்டும்? எப்போது தானாக அழிக்க வேண்டும்? "
                "மிகவும் ரகசியமான மருத்துவ/வங்கி உரையாடல்களை லாக் செய்யாமல் மறைக்க வேண்டுமா? "
                "பழைய டேட்டாவை மலிவான கிளவுட் ஸ்டோரேஜிற்கு (AWS S3) மாற்ற வேண்டுமா? என்பதை தீர்மானிக்கும் கொள்கை அறை."
            ),
            "business_value": "சர்வர் ஸ்டோரேஜ் செலவு மிச்சமாகும்; அரசாங்கத்தின் GDPR / HIPAA சட்ட விதிகளுக்கு இணங்கலாம்."
        },
        "tech": {
            "architecture": (
                "Located in 'ui/app/workspace/config/views/loggingView.tsx' and backend garbage collection scheduler. "
                "Controls internal database vacuuming routines, FastHTTP request context logging filters, "
                "and AWS S3 / Google Cloud Storage multi-part blob upload routines for cold archive tiering."
            ),
            "endpoints": ["GET /api/v1/config/logging", "PUT /api/v1/config/logging"]
        },
        "ui_elements": {
            "top_bar": [
                "Page Header with Save Configuration Button and Reset to Defaults Option."
            ],
            "tabs_and_views": [
                "Retention Schedule Slider: Dropdown setting log lifespan (7, 14, 30, 90, 365 days).",
                "Traffic Sampling Rate Slider: 1% to 100% sampling controls for ultra-high traffic environments.",
                "Privacy Switches: (1) Disable Content Logging (omits prompts/responses), (2) Store Raw Bytes, (3) Auto PII Redaction in logs.",
                "External Storage Configuration: AWS S3 / GCS bucket name, region, access key, and secret key inputs."
            ],
            "bottom_elements": [
                "Purge Logs Now Button: Administrative action to instantly delete logs matching selected criteria.",
                "Save Logs Configuration Button: Applies changes instantly across all cluster nodes."
            ]
        },
        "connections": {
            "receives_from": "Admin configurations from the UI.",
            "triggers_and_affects": "Governs the storage lifecycle, privacy masking, and deletion behavior of LLM Logs and MCP Logs."
        },
        "use_case": "ஒரு மருத்துவமனை 'HIPAA Compliance' விதிப்படி, நோயாளிகளின் உரையாடல்களை 30 நாட்களில் ஆட்டோ-டெலீட் செய்ய Retention-ஐ 30 நாட்களாக அமைத்து, Content Logging-ஐ ஆஃப் செய்கிறது."
    }
]

# ==============================================================================
# 1. GENERATE MARKDOWN DOCUMENT
# ==============================================================================
def generate_obs_markdown():
    print("Writing Observability Markdown document...")
    lines = []
    lines.append("# UnifAI Observability Master Deep-Dive Guide")
    lines.append("## UnifAI கண்காணிப்பு அமைப்பின் முழுமையான உடற்கூறு ஆய்வு (Tech & Non-Tech Master Manual)\n")
    lines.append("**Module:** Observability (கண்காணிப்பு & செயல்திறன் பகுப்பாய்வு)  ")
    lines.append("**Target Audience:** Developers, System Architects, CTOs, Product Managers & Operations Teams  ")
    lines.append("**Generated At:** 2026-09-05  ")
    lines.append("**Format:** Bilingual (Tamil & English Technical Guide)  \n")
    lines.append("---\n")

    lines.append("## Table of Contents (பொருளடக்கம்)")
    lines.append("1. [Observability System Architecture & Structure Map (முழுமையான கட்டமைப்பு வரைபடம்)](#1-observability-system-architecture--structure-map)")
    lines.append("2. [Detailed Feature Dissection (6 Core Observability Features)](#2-detailed-feature-dissection)")
    for f in OBS_FEATURES:
        lines.append(f"   - [{f['name']} ({f['route']})](#{f['id']})")
    lines.append("3. [Cross-Feature Interconnections & Data Flow (அம்சங்களுக்கிடையேயான தொடர்பு)](#3-cross-feature-interconnections--data-flow)")
    lines.append("4. [Tech vs Non-Tech Comparative Matrix (தொழில்நுட்ப & வணிக பார்வை ஒப்பீடு)](#4-tech-vs-non-tech-comparative-matrix)\n")
    lines.append("---\n")

    # Section 1
    lines.append("# 1. Observability System Architecture & Structure Map")
    lines.append("### முழுமையான கட்டமைப்பு & தரவு ஓட்டம் (Data Flow Map)\n")
    lines.append("```")
    lines.append("                    [ INCOMING AI REQUESTS & AGENT TOOL CALLS ]")
    lines.append("                                        │")
    lines.append("                                        ▼")
    lines.append("      ┌───────────────────────────────────────────────────────────────────┐")
    lines.append("      │               FastHTTP PROXY & MCP INTERCEPTOR                    │")
    lines.append("      │  • Captures Request Headers, Prompts, User Tokens, Virtual Keys   │")
    lines.append("      │  • Measures Time-To-First-Token (TTFT) and Total Latency          │")
    lines.append("      └─────────────────┬───────────────────────────────┬─────────────────┘")
    lines.append("                        │                               │")
    lines.append("          [LLM Calls]   ▼                 [Tool Calls]  ▼")
    lines.append("      ┌─────────────────────────┐           ┌─────────────────────────┐")
    lines.append("      │       LLM LOGS          │           │        MCP LOGS         │")
    lines.append("      │  • Prompts & Responses  │           │  • Tool Arguments       │")
    lines.append("      │  • Cost ($), Tokens     │           │  • Tool Outputs         │")
    lines.append("      │  • Guardrail Verdicts   │           │  • Execution Time (ms)  │")
    lines.append("      └────────────┬────────────┘           └────────────┬────────────┘")
    lines.append("                   │                                     │")
    lines.append("                   ├──────────────────┬──────────────────┤")
    lines.append("                   ▼                  ▼                  ▼")
    lines.append("      ┌─────────────────────────┐ ┌─────────────────────────┐ ┌─────────────────────────┐")
    lines.append("      │       DASHBOARD         │ │       CONNECTORS        │ │      LOGS SETTINGS      │")
    lines.append("      │  • Realtime Metrics     │ │  • Datadog APM Spans    │ │  • Retention Schedules  │")
    lines.append("      │  • p95/p99 Latencies    │ │  • Apache Kafka Topics  │ │  • Traffic Sampling %   │")
    lines.append("      │  • Cost Aggregations    │ │  • BigQuery Billing     │ │  • PII Content Redactor │")
    lines.append("      │  • Cache Hit Ratios     │ │  • OpenTelemetry (OTel) │ │  • S3 / GCS Archival    │")
    lines.append("      └─────────────────────────┘ └─────────────────────────┘ └─────────────────────────┘")
    lines.append("                   ▲")
    lines.append("                   │  (Employee Web AI Interception)")
    lines.append("      ┌────────────┴────────────┐")
    lines.append("      │       BROWSER AI        │")
    lines.append("      │  • Local DLP Proxy      │")
    lines.append("      │  • Attachment Scanner   │")
    lines.append("      │  • Policy Violations    │")
    lines.append("      └─────────────────────────┘")
    lines.append("```\n")

    # Section 2: 6 Features
    lines.append("# 2. Detailed Feature Dissection (6 Core Observability Features)")
    lines.append("### ஆறு அம்சங்களின் விரிவான உடற்கூறு ஆய்வு\n")

    for f in OBS_FEATURES:
        lines.append(f"<a name='{f['id']}'></a>")
        lines.append(f"## {f['name']}")
        lines.append(f"**UI Route:** `{f['route']}`\n")
        
        lines.append("### 👤 Non-Tech Perspective (சாதாரண மனிதர்களுக்கான எளிய விளக்கம்)")
        lines.append(f"- **உருவகம் (Analogy):** {f['non_tech']['analogy']}")
        lines.append(f"- **விளக்கம் (Explanation):** {f['non_tech']['explanation']}")
        lines.append(f"- **வணிக மதிப்பு (Business Value):** {f['non_tech']['business_value']}\n")

        lines.append("### 💻 Tech Perspective (பொறியாளர்களுக்கான தொழில்நுட்ப விளக்கம்)")
        lines.append(f"- **Backend Architecture:** {f['tech']['architecture']}")
        lines.append("- **Backend Endpoints:**")
        for ep in f['tech']['endpoints']:
            lines.append(f"  * `{ep}`")
        lines.append("\n")

        lines.append("### 🖥️ Screen Layout & Bottom Elements (திரை கூறுகள் & பட்டன்கள்)")
        lines.append("**1. மேல்புற கட்டுப்பாடுகள் (Top Bar Controls):**")
        for item in f['ui_elements']['top_bar']:
            lines.append(f"- {item}")
        lines.append("\n**2. மத்திய திரைக் கூறுகள் & வரைபடங்கள் (Tabs & Views):**")
        for item in f['ui_elements']['tabs_and_views']:
            lines.append(f"- {item}")
        lines.append("\n**3. கீழ்புற கூறுகள், படிவங்கள் & ஸ்லைடு-ஓவர் ஷீட்டுகள் (Bottom Elements & Sheets):**")
        for item in f['ui_elements']['bottom_elements']:
            lines.append(f"- {item}")
        lines.append("\n")

        lines.append("### 🔗 Connection & Structure Map (இணைப்புகள் & செயல்பாடுகள்)")
        lines.append(f"- **தரவை எங்கிருந்து பெறுகிறது (Receives Data From):** {f['connections']['receives_from']}")
        lines.append(f"- **எதனை இயக்குகிறது / பாதிக்கிறது (Triggers & Affects):** {f['connections']['triggers_and_affects']}\n")

        lines.append("### 💡 Real-World Enterprise Use Case (நடைமுறை பயன்பாட்டு உதாரணம்)")
        lines.append(f"{f['use_case']}\n")
        lines.append("---\n")

    # Section 3
    lines.append("# 3. Cross-Feature Interconnections & Data Flow")
    lines.append("### அம்சங்களுக்கிடையேயான நேரடித் தொடர்பு வரைபடம்\n")
    lines.append("| மூல கூறு (Source) | இணைக்கப்பட்டுள்ள கூறு (Connected To) | தரவு பரிமாற்றம் & செயல்பாடு (Data Flow & Action) |")
    lines.append("| :--- | :--- | :--- |")
    lines.append("| **LLM Logs** | **Dashboard** | ஒவ்வொரு லாக் வரியிலிருந்தும் Token count, Latency மற்றும் Cost கணக்கிடப்பட்டு Dashboard வரைபடங்கள் புதுப்பிக்கப்படுகின்றன. |")
    lines.append("| **LLM Logs** | **Connectors** | லாக் பதிவுகள் Asynchronous Ring Buffer வழியாக Datadog, Kafka, மற்றும் BigQuery-க்கு நேரலையாக ஸ்ட்ரீம் செய்யப்படுகின்றன. |")
    lines.append("| **MCP Logs** | **MCP Gateway** | AI ஏஜென்ட் இயக்கிய டூல்களின் JSON இன்புட் மற்றும் அவுட்புட் பரிவர்த்தனைகள் MCP Logs-ல் சேமிக்கப்படுகின்றன. |")
    lines.append("| **Browser AI** | **LLM Logs / Audit Logs** | ஊழியர்கள் ChatGPT தளத்தில் ரகசிய டேட்டாவை பேஸ்ட் செய்யும்போது Browser AI தடுத்து நிறுத்தி, அந்த விதிமீறலை Logs-ல் பதிகிறது. |")
    lines.append("| **Logs Settings** | **LLM & MCP Logs** | பதிவுகள் எத்தனை நாட்கள் சேமிப்பில் இருக்க வேண்டும் (Retention) மற்றும் PII மறைப்பு விதிகளையும் Logs Settings நிர்வகிக்கிறது. |")
    lines.append("| **Dashboard** | **Alert Channels** | செலவு அல்லது எரர் விகிதம் அதிகமாகும்போது Dashboard அமைப்புகள் மூலமாக Slack/PagerDuty-க்கு அலர்ட் செல்கிறது. |\n")

    # Section 4
    lines.append("# 4. Tech vs Non-Tech Comparative Matrix")
    lines.append("### தொழில்நுட்ப & வணிக பார்வை ஒப்பீடு (Comparative Matrix)\n")
    lines.append("| Observability அம்சம் | வணிகப் பார்வை (CFO / Manager Perspective) | தொழில்நுட்பப் பார்வை (DevOps / Architect Perspective) |")
    lines.append("| :--- | :--- | :--- |")
    lines.append("| **Dashboard** | 'இந்த வாரம் AI-க்கு எவ்வளவு செலவாகியுள்ளது? பட்ஜெட் மிச்சமிருக்கிறதா?' | 'p95 லேடன்சி எவ்வளவு? FastHTTP மெட்ரிக்ஸ் மற்றும் WebSocket இணைப்புகள் சீராக உள்ளதா?' |")
    lines.append("| **LLM Logs** | 'வாடிக்கையாளர் என்ன கேள்வி கேட்டார்? மாடல் சரியான பதில் தந்ததா?' | 'HTTP Status Code என்ன? TTFT மில்லி விநாடிகள் மற்றும் PostgreSQL இன்செர்ஷன் வேகம் எவ்வளவு?' |")
    lines.append("| **MCP Logs** | 'AI ஏஜென்ட் நம் கம்பெனி டூல்களை சரியாகப் பயன்படுத்துகிறதா?' | 'stdio/SSE transport வழியாக டூல் இயங்கிய latency மற்றும் JSON Schema validation நிலைகள் என்ன?' |")
    lines.append("| **Browser AI** | 'நிறுவன ரகசியங்களை ஊழியர்கள் ChatGPT-ல் லீக் செய்யாமல் தடுக்க முடிகிறதா?' | 'mitmproxy daemon இயங்குகிறதா? Presidio DLP regexes மற்றும் Ollama மாடல் கிளாசிபிகேஷன் துல்லியமா?' |")
    lines.append("| **Connectors** | 'எங்கள் கார்ப்பரேட் பில்லிங் சிஸ்டத்தில் AI செலவு தானாக ஏறி விடுகிறதா?' | 'Kafka topics-க்கு batch size மற்றும் flush interval (ms) முறையில் டேட்டா இழப்பின்றி செல்கிறதா?' |")
    lines.append("| **Logs Settings** | 'சட்ட விதிகளுக்கு ஏற்ப பழைய உரையாடல்கள் தானாக அழிகிறதா (GDPR Compliance)?' | 'Postgres vacuumingGC அட்டவணை மற்றும் S3 cold storage multi-part upload சரியாக நடக்கிறதா?' |\n")

    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Observability Markdown written to: {MD_PATH}")

OBS_CONN_DATA = [
    ("LLM Logs", "Dashboard", "Computes token counts, latencies, and costs for realtime charts."),
    ("LLM Logs", "Connectors", "Pushes asynchronous event payloads to Datadog, Kafka, and BigQuery."),
    ("MCP Logs", "MCP Gateway", "Records input JSON arguments and tool execution outputs."),
    ("Browser AI", "LLM / Audit Logs", "Intercepts web AI paste events and logs DLP violations."),
    ("Logs Settings", "LLM & MCP Logs", "Enforces data retention schedules and automated PII redaction.")
]

MATRIX_DATA = [
    ("Dashboard", "How much did we spend on AI this week? Is budget on track?", "What are the p95 latency curves and WebSocket health metrics?"),
    ("LLM Logs", "What did the customer ask and was the AI answer accurate?", "What was the HTTP status code, TTFT (ms), and pgx insertion rate?"),
    ("MCP Logs", "Is the autonomous agent using company tools responsibly?", "Did the stdio/SSE transport succeed with valid JSON schema outputs?"),
    ("Browser AI", "Are employees leaking company secrets on web ChatGPT?", "Is the local mitmproxy daemon running with active Presidio DLP filters?"),
    ("Connectors", "Is AI consumption appearing automatically in our billing system?", "Are Kafka event streams and BigQuery streaming inserts zero-drop?"),
    ("Logs Settings", "Are we compliant with GDPR and HIPAA 30-day auto-purge rules?", "Is the Postgres vacuuming routine and S3 archival pipeline healthy?")
]

# ==============================================================================
# 2. GENERATE DOCX DOCUMENT
# ==============================================================================
def generate_obs_docx():
    print("Writing Observability Word Document (.docx)...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        hp = section.header.paragraphs[0]
        hp.text = "UnifAI Observability Master Deep-Dive Technical Manual"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if hp.runs:
            hp.runs[0].font.name = "Segoe UI"
            hp.runs[0].font.size = Pt(8.5)
            hp.runs[0].font.color.rgb = RGBColor(100, 116, 139)
            
        fp = section.footer.paragraphs[0]
        fp.text = "Confidential & Proprietary — UnifAI Observability Control Plane"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if fp.runs:
            fp.runs[0].font.name = "Segoe UI"
            fp.runs[0].font.size = Pt(8.5)
            fp.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    def style_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        return p

    def style_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def style_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.italic = italic
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def format_table(table, col_widths, col_names):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(col_names):
            hdr_cells[i].text = name
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
            tcPr.append(shd)
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = "Segoe UI"
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for r_idx, row in enumerate(table.rows[1:]):
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                tcPr.append(shd)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                    
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    rt = p_title.add_run("UnifAI Observability Master Deep-Dive Guide")
    rt.font.name = "Segoe UI"
    rt.font.size = Pt(22)
    rt.bold = True
    rt.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    rsub = p_sub.add_run("UnifAI கண்காணிப்பு அமைப்பின் முழுமையான உடற்கூறு ஆய்வு (Tech & Non-Tech Master Manual)")
    rsub.font.name = "Segoe UI"
    rsub.font.size = Pt(11)
    rsub.font.color.rgb = RGBColor(71, 85, 105)

    style_h1("1. Observability Overview & Architecture")
    add_p(
        "AI Observability provides complete visibility into AI consumption, financial burn rates, latency percentiles, "
        "and security compliance. UnifAI decouples telemetry logging from the critical proxy path using lock-free atomic "
        "buffers and streaming connectors to ensure sub-millisecond response overhead."
    )

    style_h1("2. Deep Dive: The 6 Observability Features")
    for f in OBS_FEATURES:
        style_h2(f"{f['name']} ({f['route']})")
        
        style_h3("Non-Tech Perspective (எளிய வணிக விளக்கம்)")
        add_p(f['non_tech']['analogy'], bold_prefix="உருவகம் (Analogy): ")
        add_p(f['non_tech']['explanation'], bold_prefix="விளக்கம்: ")
        add_p(f['non_tech']['business_value'], bold_prefix="வணிக மதிப்பு: ")

        style_h3("Tech Perspective (தொழில்நுட்ப கட்டமைப்பு)")
        add_p(f['tech']['architecture'], bold_prefix="Backend Architecture: ")
        add_p(", ".join(f['tech']['endpoints']), bold_prefix="Endpoints: ")

        style_h3("Screen Layout & Bottom Elements (UI கூறுகள் & பட்டன்கள்)")
        add_p("1. மேல்புற கட்டுப்பாடுகள் (Top Bar Controls):", bold_prefix="")
        for itm in f['ui_elements']['top_bar']:
            add_bullet(itm)
        add_p("2. மத்திய திரைக் கூறுகள் & வரைபடங்கள் (Tabs & Views):", bold_prefix="")
        for itm in f['ui_elements']['tabs_and_views']:
            add_bullet(itm)
        add_p("3. கீழ்புற கூறுகள் & ஸ்லைடு-ஓவர் ஷீட்டுகள் (Bottom Elements & Sheets):", bold_prefix="")
        for itm in f['ui_elements']['bottom_elements']:
            add_bullet(itm)

        style_h3("Module Connections (தொடர்புகள்)")
        add_p(f['connections']['receives_from'], bold_prefix="தரவை பெறுமிடம்: ")
        add_p(f['connections']['triggers_and_affects'], bold_prefix="இயக்கும் கூறுகள்: ")

        style_h3("Production Use Case (நடைமுறை உதாரணம்)")
        add_p(f['use_case'])

    style_h1("3. Cross-Feature Interconnections")
    t_conn = doc.add_table(rows=len(OBS_CONN_DATA)+1, cols=3)
    format_table(t_conn, [1.8, 1.8, 3.4], ["Source Feature", "Connected To", "Data Flow & Action"])
    for idx, (src, dst, flow) in enumerate(OBS_CONN_DATA):
        t_conn.rows[idx+1].cells[0].text = src
        t_conn.rows[idx+1].cells[1].text = dst
        t_conn.rows[idx+1].cells[2].text = flow

    style_h1("4. Tech vs Non-Tech Comparative Matrix")
    t_mat = doc.add_table(rows=len(MATRIX_DATA)+1, cols=3)
    format_table(t_mat, [1.5, 2.7, 2.8], ["Feature", "Non-Tech View (Manager / CFO)", "Tech View (DevOps / Architect)"])
    for idx, (ft, nt, tv) in enumerate(MATRIX_DATA):
        t_mat.rows[idx+1].cells[0].text = ft
        t_mat.rows[idx+1].cells[1].text = nt
        t_mat.rows[idx+1].cells[2].text = tv

    doc.save(DOCX_PATH)
    print(f"Observability Word document written to: {DOCX_PATH}")

# ==============================================================================
# 3. GENERATE PDF DOCUMENT VIA REPORTLAB
# ==============================================================================
def generate_obs_pdf():
    print("Writing Observability PDF Document (.pdf)...")

    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748B"))
            
            if self._pageNumber > 1:
                self.drawString(40, 11 * inch - 36, "UnifAI Observability Master Deep-Dive Technical Manual")
                self.drawRightString(8.5 * inch - 40, 11 * inch - 36, "Confidential — Engineering Guide")
                self.setStrokeColor(colors.HexColor("#CBD5E1"))
                self.setLineWidth(0.5)
                self.line(40, 11 * inch - 40, 8.5 * inch - 40, 11 * inch - 40)
                
            self.drawString(40, 32, "UnifAI Unified Observability & Telemetry Control Plane")
            page_str = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(8.5 * inch - 40, 32, page_str)
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(40, 42, 8.5 * inch - 40, 42)
            
            self.restoreState()

    pdf_doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=48,
        bottomMargin=48
    )

    styles = getSampleStyleSheet()

    p_title_style = ParagraphStyle(
        'PdfTitle',
        fontName='LathaBold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=3
    )

    p_subtitle_style = ParagraphStyle(
        'PdfSubTitle',
        fontName='Latha',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#475569'),
        spaceAfter=10
    )

    h1_style = ParagraphStyle(
        'PdfH1',
        fontName='LathaBold',
        fontSize=12,
        leading=15,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=12,
        spaceAfter=4,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'PdfH2',
        fontName='LathaBold',
        fontSize=10,
        leading=13,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'PdfH3',
        fontName='LathaBold',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=6,
        spaceAfter=2,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'PdfBody',
        fontName='Latha',
        fontSize=7.5,
        leading=10.5,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=3
    )

    bullet_style = ParagraphStyle(
        'PdfBullet',
        fontName='Latha',
        fontSize=7.5,
        leading=10.5,
        textColor=colors.HexColor('#1E293B'),
        leftIndent=10,
        spaceAfter=2
    )

    th_style = ParagraphStyle(
        'PdfTH',
        fontName='LathaBold',
        fontSize=7,
        leading=9,
        textColor=colors.white
    )

    td_style = ParagraphStyle(
        'PdfTD',
        fontName='Latha',
        fontSize=6.5,
        leading=8.5,
        textColor=colors.HexColor('#0F172A')
    )

    td_code_style = ParagraphStyle(
        'PdfTDCode',
        fontName='Helvetica-Bold',
        fontSize=6.5,
        leading=8.5,
        textColor=colors.HexColor('#0F172A')
    )

    story = []

    # Title & Metadata
    story.append(Paragraph("UnifAI Observability Master Deep-Dive Guide", p_title_style))
    story.append(Paragraph("UnifAI கண்காணிப்பு அமைப்பின் முழுமையான உடற்கூறு ஆய்வு (Tech & Non-Tech Master Manual)", p_subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#0284C7"), spaceAfter=8))

    # Overview
    story.append(Paragraph("1. Observability Architecture & Ecosystem Overview", h1_style))
    story.append(Paragraph(
        "AI Observability provides complete visibility into AI consumption, financial burn rates, latency percentiles, "
        "and security compliance. UnifAI decouples telemetry logging from the critical proxy path using lock-free atomic "
        "buffers and streaming connectors to ensure sub-millisecond response overhead.",
        body_style
    ))
    story.append(Spacer(1, 4))

    # 6 Features Detailed Breakdown
    story.append(Paragraph("2. Detailed Feature Dissection (6 Core Observability Features)", h1_style))
    for f in OBS_FEATURES:
        feat_elements = []
        feat_elements.append(Paragraph(f"<b>{f['name']}</b> ({f['route']})", h2_style))
        
        feat_elements.append(Paragraph("<b>👤 Non-Tech Perspective (எளிய வணிக விளக்கம்):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>உருவகம் (Analogy):</b> {f['non_tech']['analogy']}", body_style))
        feat_elements.append(Paragraph(f"• <b>விளக்கம் (Explanation):</b> {f['non_tech']['explanation']}", body_style))
        feat_elements.append(Paragraph(f"• <b>வணிக மதிப்பு (Business Value):</b> {f['non_tech']['business_value']}", body_style))

        feat_elements.append(Paragraph("<b>💻 Tech Perspective (தொழில்நுட்ப கட்டமைப்பு):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Backend Architecture:</b> {f['tech']['architecture']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Endpoints:</b> <code>{', '.join(f['tech']['endpoints'])}</code>", body_style))

        feat_elements.append(Paragraph("<b>🖥️ Screen Layout & Bottom Elements (UI கூறுகள் & பட்டன்கள்):</b>", h3_style))
        feat_elements.append(Paragraph("<b>1. மேல்புற கட்டுப்பாடுகள் (Top Bar Controls):</b>", body_style))
        for itm in f['ui_elements']['top_bar']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))
        feat_elements.append(Paragraph("<b>2. மத்திய திரைக் கூறுகள் (Tabs & Views):</b>", body_style))
        for itm in f['ui_elements']['tabs_and_views']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))
        feat_elements.append(Paragraph("<b>3. கீழ்புற கூறுகள் & ஷீட்டுகள் (Bottom Elements & Sheets):</b>", body_style))
        for itm in f['ui_elements']['bottom_elements']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))

        feat_elements.append(Paragraph("<b>🔗 Connections & Structure Map (இணைப்புகள்):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>தரவை பெறுமிடம்:</b> {f['connections']['receives_from']}", body_style))
        feat_elements.append(Paragraph(f"• <b>இயக்கும் கூறுகள்:</b> {f['connections']['triggers_and_affects']}", body_style))

        feat_elements.append(Paragraph(f"<b>💡 Production Use Case:</b> {f['use_case']}", body_style))
        feat_elements.append(Spacer(1, 6))
        story.append(KeepTogether(feat_elements))

    # Cross Connections
    story.append(Paragraph("3. Cross-Feature Interconnections & Data Flow", h1_style))
    conn_table_data = [[Paragraph(h, th_style) for h in ["Source Feature", "Connected To", "Data Flow & Action"]]]
    for src, dst, flow in OBS_CONN_DATA:
        conn_table_data.append([Paragraph(src, td_code_style), Paragraph(dst, td_code_style), Paragraph(flow, td_style)])
    tbl_conn = Table(conn_table_data, colWidths=[105, 105, 320])
    tbl_conn.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_conn)
    story.append(Spacer(1, 8))

    # Comparative Matrix
    story.append(Paragraph("4. Tech vs Non-Tech Comparative Matrix", h1_style))
    mat_table_data = [[Paragraph(h, th_style) for h in ["Feature", "Non-Tech View (Manager / CFO)", "Tech View (DevOps / Architect)"]]]
    for ft, nt, tv in MATRIX_DATA:
        mat_table_data.append([Paragraph(ft, td_code_style), Paragraph(nt, td_style), Paragraph(tv, td_style)])
    tbl_mat = Table(mat_table_data, colWidths=[95, 215, 220])
    tbl_mat.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_mat)

    pdf_doc.build(story, canvasmaker=NumberedCanvas)
    print(f"Observability PDF document written to: {PDF_PATH}")

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================
if __name__ == "__main__":
    print("Building dedicated Observability Master Documentation...")
    generate_obs_markdown()
    generate_obs_docx()
    generate_obs_pdf()
    print("Observability Master Documentation successfully built in:", DOC_DIR)
