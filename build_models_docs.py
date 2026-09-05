import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether, PageBreak, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch

# Reconfigure console output to UTF-8
if sys.stdout and hasattr(sys.stdout, 'reconfigure'):
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass

WORKSPACE_DIR = r"c:\Users\sakth\OneDrive\ドキュメント\Desktop\unify - Copy"
DOC_DIR = os.path.join(WORKSPACE_DIR, "document")
os.makedirs(DOC_DIR, exist_ok=True)

MD_PATH = os.path.join(DOC_DIR, "UnifAI_Models_Master_Guide.md")
DOCX_PATH = os.path.join(DOC_DIR, "UnifAI_Models_Master_Guide.docx")
PDF_PATH = os.path.join(DOC_DIR, "UnifAI_Models_Master_Guide.pdf")

# ==============================================================================
# DATA DEFINITIONS FOR THE 8 MODELS FEATURES (THANGLISH + ENGLISH)
# ==============================================================================

MODELS_FEATURES = [
    {
        "id": "model_catalog",
        "name_en": "Model Catalog",
        "name_tanglish": "Model Catalog (AI Models Menu & Inventory)",
        "route": "/workspace/model-catalog",
        "analogy": "Amazon / Flipkart e-commerce product catalog mathiri allathu restaurant menu card mathiri.",
        "explanation": "Oru hotel-la menu card pathu enna items irukku, evlo price nu therinjukurathu mathiri, company-la irukkura ellam AI models (GPT-4o, Claude 3.5, Gemini, Llama 3) list, athoda context length, 24-hour traffic, total cost ellame ore idathula pakkalam. Entha model active-ah irukku, ethu custom model nu oru single screen-la paathu choose pannikalam.",
        "business_value": "Team leads and developers entha model available-ah irukku nu therinjukalam. Duplicate and unnecessary models use panni cost waste aaguratha thadukkalam. Organization full-ah AI standardisation kondu varalam.",
        "tech_arch": "FastHTTP model registry endpoint (`/api/v1/models/catalog`). TanStack Query client caching with real-time token metrics aggregation. Supports filtering by provider, capability flags (vision, audio, tool calling, reasoning), and custom model attributes. Ingests 24-hour aggregated traffic and cost metrics calculated from PostgreSQL telemetry tables.",
        "endpoints": [
            "GET /api/v1/models/catalog",
            "GET /api/v1/models/attributes",
            "PATCH /api/v1/models/{model_id}/attributes"
        ],
        "ui_elements": {
            "top_bar": [
                "Summary Cards (4 metrics): Total Providers count, Total Models count, Total Requests (24h), Total Cost (24h $).",
                "Provider Filter Dropdown: Quick filter by All Providers or specific providers (OpenAI, Anthropic, Bedrock, etc.).",
                "Tab Switcher: Toggle between 'Overview' tab and 'Attributes' tab."
            ],
            "tabs_and_views": [
                "Overview Tab Table: Columns for Provider Name (with icon and custom badge), Models Used (badges with tooltips), Total Traffic 24h, Total Cost 24h ($).",
                "Attributes Tab Table: Columns for Provider, Model Identifier, Display Name, Description (with hover tooltip for truncated text), Pricing (input/output per 1M tokens), Context Window size (128k, 200k), Max Output Tokens, Actions.",
                "Search & Pagination: Debounced search input (300ms) across model names, Page offset controls (25 items per page)."
            ],
            "bottom_elements": [
                "AttributeSheet (Slide-Over): Triggered by Edit button to update Display Name, Description, Context Window, Output Tokens, Modality Tags, and Custom Parameters.",
                "Custom Provider Indicators: Highlight badges for self-hosted or keyless providers."
            ]
        },
        "connections": {
            "receives_from": "Receives provider lists from Model Providers and aggregated usage/cost data from LLM Logs & Dashboard telemetry.",
            "triggers_and_affects": "Feeds model choices into Routing Rules, Complexity Router, and Circuit Breaker."
        },
        "use_case": "Company-wide AI catalog review to audit which teams are using expensive legacy models (e.g. GPT-4 32k) and migrate them to cost-efficient modern models (e.g. GPT-4o-mini)."
    },
    {
        "id": "model_providers",
        "name_en": "Model Providers",
        "name_tanglish": "Model Providers (AI Vendor Credentials & Gateway Settings)",
        "route": "/workspace/providers",
        "analogy": "Telecom SIM card & network provider switchboard mathiri (Airtel, Jio, Vodafone connections configure pandra switchboard).",
        "explanation": "UnifAI platform vera vera AI vendors-oda (OpenAI, Anthropic, AWS Bedrock, Google Vertex AI, Azure OpenAI, Groq, Ollama) connect aaga thevaiyana API keys, secret credentials, network speed, timeout settings configure pandra control panel. Oru pudhu private AI server (vLLM/Ollama) add pannanum naalum inga simple-ah add pannikalam.",
        "business_value": "Single pane of glass for all AI vendor contracts and credentials. Vendor lock-in thadukkum; enterprise security standards-ku etha mathiri API keys-ah mask panni manage pannalam.",
        "tech_arch": "Multi-vendor credential management and FastHTTP transport connection pool. Manages custom base URLs, proxy tunnels, exponential backoff retries, request/response payload tracing flags (`send_back_raw_request`, `send_back_raw_response`), and keyless private deployments. Validates connections via instant ping test mutations.",
        "endpoints": [
            "GET /api/v1/providers",
            "POST /api/v1/providers",
            "GET /api/v1/providers/{provider_name}",
            "PATCH /api/v1/providers/{provider_name}",
            "DELETE /api/v1/providers/{provider_name}",
            "POST /api/v1/providers/{provider_name}/keys/test"
        ],
        "ui_elements": {
            "top_bar": [
                "Provider Navigation Sidebar: Sorted alphabetical list of configured providers with status badges (Active, Warning, Error).",
                "Add Provider Dropdown: List of known providers (OpenAI, Anthropic, AWS Bedrock, Azure, Google, Mistral, Groq, Ollama) + 'Add Custom Provider' action."
            ],
            "tabs_and_views": [
                "Provider Detail Header: Provider icon, display label, status badge, 'Edit Provider Config' button, and 'Delete Provider' button.",
                "ModelProviderKeysTableView: Table displaying API Key Name, Masked Key Secret, Quota Limit, Key Status, Connection Ping Test button, and Delete Key button.",
                "ProviderGovernanceTable: Governance rules and access policies tied specifically to this vendor."
            ],
            "bottom_elements": [
                "ProviderConfigSheet: Slide-over form configuring Concurrency Limits, Buffer Size, Network Config (Request Timeout ms, Max Retries, Retry Delay), Proxy URL, and Raw Payload Toggles.",
                "AddCustomProviderSheet: Form to configure custom / self-hosted provider names, base URL endpoints (e.g. `http://internal-vllm:8000/v1`), and Keyless Mode toggle."
            ]
        },
        "connections": {
            "receives_from": "Receives API keys and network configurations from administrators.",
            "triggers_and_affects": "Powers all outbound AI traffic through FastHTTP Proxy; supplies connection clients to Routing Rules and Circuit Breaker."
        },
        "use_case": "Zero-downtime rotation of compromised OpenAI production keys and setting up private on-premise Ollama instances for confidential internal documents."
    },
    {
        "id": "model_limits",
        "name_en": "Budgets & Limits",
        "name_tanglish": "Budgets & Limits (Cost Caps & Token Rate Limits)",
        "route": "/workspace/model-limits",
        "analogy": "Credit card limit and bank daily withdrawal limit mathiri.",
        "explanation": "AI use panna panna bill laatchakanakula aagira koodathu nu, ovvoru team-kum, user-kum, allathu ovvoru AI model-kum monthly budget ($500) allathu daily token limit set pandra edham. Limit reach aana udane automated warning email varum, allathu request-ah block panni bill shock varama thadukkum.",
        "business_value": "100% financial safety and budget predictability. Rogue scripts allathu infinite loop bugs naala laksha kanakkula bill varatha prevent pannum.",
        "tech_arch": "High-throughput token bucket rate limiter and budget enforcer middleware. Atomic Redis / PostgreSQL counters (`INCRBY`). Evaluates requests against multi-tiered governance limits (Global, Virtual Key, User, Team). Supports sliding rolling windows and calendar-month alignments. Pre-flight evaluation rejects requests with HTTP 429 or routes to a cheaper fallback model.",
        "endpoints": [
            "GET /api/v1/governance/model-limits",
            "POST /api/v1/governance/model-limits",
            "PUT /api/v1/governance/model-limits/{limit_id}",
            "DELETE /api/v1/governance/model-limits/{limit_id}"
        ],
        "ui_elements": {
            "top_bar": [
                "Create Model Limit Button: Opens the limit configuration sheet.",
                "Search Input: Quick search across model names and entity targets.",
                "Filter Dropdowns: Filter by Provider and by Governance Scope (Global, Virtual Key, User, Team/User Group)."
            ],
            "tabs_and_views": [
                "Model Limits Table: Columns for Scope Badge, Target Entity, Provider, Model, Budget Progress Bar (Current Spend $ vs Limit $), Token Progress Bar (Used vs Max Tokens), Reset Cadence, Actions Menu.",
                "Visual Progress Bars: Color-coded usage indicators (green < 70%, orange 70-90%, red > 90%)."
            ],
            "bottom_elements": [
                "ModelLimitSheet (Modal/Slide-over): Configure Scope, Entity Target, Provider/Model, Max Budget ($), Max Tokens, Reset Duration (Daily, Weekly, Monthly, Rolling Window, Calendar-aligned).",
                "Threshold Alert Rules: Email/Slack webhook triggers on 80% warning and 100% hard limit block."
            ]
        },
        "connections": {
            "receives_from": "Calculates spend using live rates from Pricing Overrides and real-time usage from LLM Logs.",
            "triggers_and_affects": "Enforces rate-limits on Virtual Keys and blocks/reroutes calls in FastHTTP Proxy when budgets are exhausted."
        },
        "use_case": "Giving intern developers a strict $30/month budget cap while giving the production customer-facing app an elastic $5,000 budget."
    },
    {
        "id": "routing_rules",
        "name_en": "Routing Rules",
        "name_tanglish": "Routing Rules (Intelligent Traffic Steering & Fallbacks)",
        "route": "/workspace/routing-rules",
        "analogy": "Railway track switcher mathiri (Train-ah thevaikku thagundha track-la thiruppi vidura mechanism).",
        "explanation": "Oru customer request varum pothu, athu entha AI model-ku poganum nu rules poduvom. Example: Customer 'Enterprise User'-ah iruntha super-fast GPT-4o-ku poganum; 'Free User'-ah iruntha cheap Claude 3.5 Haiku-ku poganum. Oru model slow aana automatic-ah adutha model-ku route pannum.",
        "business_value": "Customized user experience; customer tier-ku thagundha mathiri AI quality allocate pannalam; 100% dynamic load balancing.",
        "tech_arch": "CEL (Common Expression Language) evaluation engine written in Go. Compiles and executes deterministic conditions against request headers (`x-uf-*`), prompt metadata, caller identity, and token lengths. Directs requests to target pools supporting priority weights, fallback arrays, and regional data residency constraints.",
        "endpoints": [
            "GET /api/v1/routing-rules",
            "POST /api/v1/routing-rules",
            "PATCH /api/v1/routing-rules/{rule_id}",
            "DELETE /api/v1/routing-rules/{rule_id}"
        ],
        "ui_elements": {
            "top_bar": [
                "Create Routing Rule Button: Opens rule creation sheet.",
                "Search Input: Filter rules by name and condition.",
                "Priority Sort: Order rules by evaluation priority hierarchy."
            ],
            "tabs_and_views": [
                "Routing Rules Table: Columns for Rule Name, Priority Badge (High/Med/Low), CEL Condition Expression (syntax highlighted), Target Model & Provider, Enabled Toggle Switch (instant activation without restart), Actions.",
                "Visual Rule Tree View: Interactive node graph showing incoming traffic splitting across matching condition branches to final destination targets."
            ],
            "bottom_elements": [
                "RoutingRuleSheet: Rule builder with Name, Priority slider (1-100), Visual Condition Builder / Raw CEL editor (e.g. `request.headers['x-tier'] == 'gold' && prompt.tokens > 500`), Target Provider/Model selection, and Fallback Chain priority list.",
                "RoutingRuleInfoSheet: Detailed execution metrics and audit history for individual rules."
            ]
        },
        "connections": {
            "receives_from": "Evaluates incoming FastHTTP requests and integrates with Complexity Router tiers.",
            "triggers_and_affects": "Selects the upstream provider connection from Model Providers and cooperates with Circuit Breaker during outages."
        },
        "use_case": "Routing all European customer requests (`request.headers['cf-ipcountry'] in ['DE', 'FR', 'UK']`) exclusively to EU-based Azure OpenAI endpoints for GDPR compliance."
    },
    {
        "id": "complexity_router",
        "name_en": "Complexity Router",
        "name_tanglish": "Complexity Router (AI Query Triage & Cost Optimizer)",
        "route": "/workspace/complexity-router",
        "analogy": "Hospital Triage System mathiri (General fever-ku junior doctor, heart surgery-ku senior specialist kitta anuppura mathiri).",
        "explanation": "Simple-ana kelvi ('Hi', 'What is the capital of India?') ketta romba costly-ana GPT-4o thevailla, athukku fast & cheap-ana GPT-4o-mini pothum. Aana periya complex coding allathu math problem ketta mattum o1 allathu Claude Opus-ku automatic-ah anupum. Ithanaala quality kuraama 70% AI bill automatic-ah micham aagum!",
        "business_value": "Massive cost reduction (up to 70% savings) with zero compromise on accuracy. Cheap models handle 80% of volume, expensive models handle 20% complex tasks.",
        "tech_arch": "Multi-factor prompt complexity analysis engine. Calculates a continuous complexity score (0.0 to 1.0) using token length heuristics, vocabulary entropy, code-syntax regex detectors (`def`, `class`, `import`, `function`), and keyword weighting. Maps the score into 4 discrete tiers: `SIMPLE`, `MEDIUM`, `COMPLEX`, `REASONING`.",
        "endpoints": [
            "GET /api/v1/governance/complexity-analyzer",
            "PUT /api/v1/governance/complexity-analyzer",
            "POST /api/v1/governance/complexity-analyzer/reset"
        ],
        "ui_elements": {
            "top_bar": [
                "Save Configuration Button: Commits threshold and keyword changes with active loading spinner.",
                "Reset to Defaults Button (RotateCcw): Restores factory recommended boundary scores.",
                "Docs Link: External link to complexity routing documentation."
            ],
            "tabs_and_views": [
                "Progressive okLCH Palette Visualization: Interactive color-coded spectrum showing the 4 tiers: SIMPLE (faint primary), MEDIUM (medium primary), COMPLEX (strong primary), REASONING (full primary).",
                "Tier Boundary Inputs: Numerical score inputs for `Simple → Medium` (e.g. 0.25), `Medium → Complex` (e.g. 0.60), and `Complex → Reasoning` (e.g. 0.85).",
                "Keyword List Management (TagInput): Tag inputs for Code Keywords (sql, regex, docker), Reasoning Keywords (prove, theorem, analyze), and Simple Keywords (hello, summarize)."
            ],
            "bottom_elements": [
                "Tier Target Model Selectors: Map each tier to a default provider & model (e.g., Simple -> GPT-4o-mini, Medium -> Claude 3.5 Haiku, Complex -> Sonnet, Reasoning -> o1).",
                "Test Prompt Sandbox: Real-time interactive prompt tester calculating live complexity scores."
            ]
        },
        "connections": {
            "receives_from": "Intercepts incoming requests from FastHTTP proxy before Routing Rules evaluation.",
            "triggers_and_affects": "Assigns target model tier, overrides requested models, and logs complexity classification into LLM Logs."
        },
        "use_case": "Customer support automation where 85% routine lookup queries are handled at 1/20th the cost by lightweight models, and 15% technical issues go to deep reasoning models."
    },
    {
        "id": "circuit_breaker",
        "name_en": "Circuit Breaker",
        "name_tanglish": "Circuit Breaker (Outage Protection & Auto-Failover)",
        "route": "/workspace/circuit-breaker",
        "analogy": "Veetla irukkura Electric MCB / Fuse mathiri (High voltage allathu short circuit aana fuse off aagi appliances-ah kaapathum).",
        "explanation": "OpenAI allathu Anthropic server crash aanaallathu rate limit aagi error vantha, user-ku error screen kaatama, automatic-ah fraction of second-la backup model-ku (e.g. AWS Bedrock allathu Azure)-ku traffic-ah switch panni system-ah 24x7 non-stop-ah run panna vaikkum.",
        "business_value": "99.99% Enterprise High Availability. Vendor downtime aanaalum business apps oru second kooda stop aagathu.",
        "tech_arch": "Distributed Circuit Breaker finite state machine (Closed -> Open -> Half-Open). Monitored via real-time polling (`pollingInterval: 8000ms`). Detects provider rate-limit response headers (e.g. `x-ratelimit-remaining: 0`), HTTP 429/503 status codes, and network timeouts. When tripped, instantly redirects calls to secondary fallback provider with automated cooldown probing.",
        "endpoints": [
            "GET /api/v1/circuit-breaker/policies",
            "POST /api/v1/circuit-breaker/policies",
            "PUT /api/v1/circuit-breaker/policies/{name}",
            "DELETE /api/v1/circuit-breaker/policies/{name}",
            "GET /api/v1/circuit-breaker/state",
            "POST /api/v1/circuit-breaker/policies/{name}/reset"
        ],
        "ui_elements": {
            "top_bar": [
                "Header Title with Shield Icon: Visual indicator of active failover protection.",
                "Create Policy Button: Opens the circuit breaker policy creation dialog."
            ],
            "tabs_and_views": [
                "Circuit Breaker Policies Table: Columns for Policy Name, Primary Provider & Model, Fallback Provider & Model, Trigger Condition Signals, Cooldown Duration (e.g. 30s), Circuit State Badge (Closed [green], Open [red], Half-Open [yellow]), Actions Menu.",
                "Live State Indicators: Polled every 8 seconds displaying active tripped circuits and error counts."
            ],
            "bottom_elements": [
                "Circuit Policy Dialog: Form configuring Policy Name, Enabled Switch, Primary Provider/Model comboboxes, Fallback Provider/Model comboboxes, Response Header Signal matchers, and Cooldown Duration.",
                "Manual Circuit Reset Button (RotateCcw): Instantly resets a tripped circuit back to Closed status."
            ]
        },
        "connections": {
            "receives_from": "Monitors HTTP response headers and error codes from FastHTTP Proxy during provider execution.",
            "triggers_and_affects": "Reroutes failed requests to fallback providers and dispatches incident alerts to Alert Channels."
        },
        "use_case": "Preventing application outage when OpenAI experiences global HTTP 503 outage on Cyber Monday by instantly switching all traffic to Azure OpenAI."
    },
    {
        "id": "pricing_overrides",
        "name_en": "Pricing Overrides",
        "name_tanglish": "Pricing Overrides (Custom Enterprise Rates & Discounts)",
        "route": "/workspace/custom-pricing/overrides",
        "analogy": "Special corporate discount allathu wholesale negotiated price contract mathiri.",
        "explanation": "OpenAI website-la 1 million token-ku $5 nu official price irunthaalum, unga company-ku special enterprise discount allathu committed spend agreement iruntha, antha custom rate-ah inga enter pannikalam. Athuku thagundha mathiri unga internal dashboard-la exact real spend calculate aagum.",
        "business_value": "100% accurate financial chargeback and departmental billing. True profit margin calculations based on actual vendor contracts rather than list prices.",
        "tech_arch": "Multi-tier hierarchical pricing resolution engine. Precedence hierarchy: (1) Virtual Key Override -> (2) User / Team Override -> (3) Provider Global Override -> (4) Default Provider Catalog. Calculates input token cost, cached prompt discount, output token cost, reasoning token multiplier, and fixed per-call fee.",
        "endpoints": [
            "GET /api/v1/pricing/overrides",
            "POST /api/v1/pricing/overrides",
            "PUT /api/v1/pricing/overrides/{override_id}",
            "DELETE /api/v1/pricing/overrides/{override_id}"
        ],
        "ui_elements": {
            "top_bar": [
                "Add Pricing Override Button: Opens override creation sheet.",
                "Search Input: Search by model name or provider.",
                "Scope Filter: Filter by Global, Virtual Key, or User Group."
            ],
            "tabs_and_views": [
                "Pricing Overrides Table: Columns for Scope Badge, Target Entity, Provider, Model, Input Cost (per 1M tokens), Cached Input Cost, Output Cost (per 1M tokens), Fixed Request Fee, Actions.",
                "Precision Display: Formats currency values up to 6 decimal places ($0.000000)."
            ],
            "bottom_elements": [
                "PricingOverrideSheet (Slide-Over): Modal to configure Scope, Provider, Model, and granular token pricing via `PricingFieldSelector`.",
                "PricingFieldSelector: Dedicated inputs for Prompt Tokens ($/1M), Cached Input Tokens ($/1M), Output Tokens ($/1M), and Per-Request Invocation Surcharge."
            ]
        },
        "connections": {
            "receives_from": "Configured by finance administrators based on legal contracts.",
            "triggers_and_affects": "Supplies live pricing equations to Dashboard charts, LLM Logs cost calculation, and Budgets & Limits enforcer."
        },
        "use_case": "Applying AWS 25% EDP (Enterprise Discount Program) discount to Bedrock Claude models so billing matches actual monthly AWS invoices."
    },
    {
        "id": "model_settings",
        "name_en": "Model Settings",
        "name_tanglish": "Model Settings (Master Catalog Sync & Recursion Guard)",
        "route": "/workspace/custom-pricing",
        "analogy": "Global system clock & master registry synchronizer mathiri.",
        "explanation": "Market-la puthu puthu AI models varum pothu, athoda official pricing automatic-ah download aagi update aaga thevaiyana master website link, sync frequency (24 hours), mathum routing rules-la infinite loop vanthu server hang aagidama thadukkura max depth settings configure pandra edham.",
        "business_value": "Zero-maintenance operations. New AI models and price reductions are synced automatically without engineering code changes or redeployments.",
        "tech_arch": "Core system configuration manager. Manages `framework_config` and `client_config` tables in PostgreSQL. Orchestrates asynchronous cron workers to fetch remote pricing datasheets (CSV/JSON), parses token schemas, and invalidates in-memory caches. Enforces `routing_chain_max_depth` to prevent circular fallback loops in routing rules.",
        "endpoints": [
            "GET /api/v1/config/core",
            "PUT /api/v1/config/core",
            "POST /api/v1/config/pricing/sync"
        ],
        "ui_elements": {
            "top_bar": [
                "Header Title: Model Settings & Global Synchronization.",
                "Dirty State Tracker: Highlights unsaved changes with save prompt."
            ],
            "tabs_and_views": [
                "ModelSettingsView Form:",
                "• Pricing Datasheet URL Input: Remote CSV/JSON endpoint for automatic pricing ingestion.",
                "• Pricing Sync Interval Input (Hours): Automated background sync cadence (default: 24h).",
                "• Model Parameters URL Input: Global repository link defining context sizes and modality capabilities.",
                "• Routing Chain Max Depth Slider/Input: Maximum recursive hops allowed across fallback chains (default: 5)."
            ],
            "bottom_elements": [
                "Force Pricing Sync Now Button: Triggers immediate background worker ingestion with loading spinner.",
                "Save Configuration Button: Commits configuration updates to core PostgreSQL storage."
            ]
        },
        "connections": {
            "receives_from": "Receives administrative inputs and connects to external UnifAI pricing telemetry feeds.",
            "triggers_and_affects": "Updates foundational pricing data for Model Catalog and enforces maximum recursion depth across Routing Rules."
        },
        "use_case": "Instantly ingesting OpenAI's latest 50% price reduction across 30 enterprise models with a single click of 'Force Pricing Sync Now'."
    }
]

MODELS_CONN_DATA = [
    ("Model Catalog", "Routing Rules", "Provides validated provider & model identifiers for rule destination targets."),
    ("Model Providers", "FastHTTP Proxy", "Supplies decrypted API keys, custom base URLs, and connection pools for HTTP calls."),
    ("Budgets & Limits", "Virtual Keys", "Enforces strict dollar and token quotas on client applications and API keys."),
    ("Routing Rules", "Circuit Breaker", "Cooperates during provider outages to bypass broken routes and execute fallbacks."),
    ("Complexity Router", "Model Providers", "Directs simple queries to lightweight models and complex queries to reasoning models."),
    ("Circuit Breaker", "Alert Channels", "Triggers immediate Slack/PagerDuty alerts when a primary provider circuit trips."),
    ("Pricing Overrides", "Dashboard & Logs", "Injects contracted enterprise prices into real-time spend calculations."),
    ("Model Settings", "Model Catalog", "Automatically syncs latest global model capabilities and prevents routing recursion loops.")
]

MODELS_MATRIX_DATA = [
    ("Model Catalog", "Entha entha AI models namakku irukku? Ethu active-ah irukku?", "Model metadata schema, capability flags (vision, audio), and 24h usage telemetry."),
    ("Model Providers", "AI companies-oda API keys and accounts ellam safe-ah connect aagiyirukka?", "FastHTTP client pools, exponential retry backoffs, proxy tunnels, and raw payload tracing."),
    ("Budgets & Limits", "Entha team-kum excessive bill varama budget control-la irukka?", "Atomic Redis/Postgres token bucket counters, sliding windows, and HTTP 429 rate limiters."),
    ("Routing Rules", "Customer tier-ku thagundha mathiri correct model-ku request pogutha?", "CEL (Common Expression Language) engine evaluating request headers and metadata in Go."),
    ("Complexity Router", "Simple kelvikku cheap model, tough kelvikku matum costly model use panni 70% bill micham aagutha?", "Heuristic multi-factor complexity scoring engine categorizing queries into 4 distinct tiers."),
    ("Circuit Breaker", "OpenAI crash aanaal kooda customer-ku error theriyaama backup model-ku switch aagutha?", "Distributed finite state machine (Closed/Open/Half-Open) with automated header signal polling."),
    ("Pricing Overrides", "Namma negotiate panna special discount price-la dashboard calculate aagutha?", "Multi-tier hierarchical pricing precedence engine applying custom token multipliers."),
    ("Model Settings", "Pudhu AI models and price cuts automatic-ah update aagutha?", "Postgres core config manager with automated background sync workers and recursion guard.")
]

# ==============================================================================
# 1. GENERATE MARKDOWN DOCUMENT (THANGLISH + ENGLISH)
# ==============================================================================
def generate_models_markdown():
    print("Writing Models Markdown document (Thanglish + English)...")
    lines = []
    lines.append("# UnifAI Models Pillar Master Deep-Dive Guide")
    lines.append("## UnifAI Models Amaippin Muzhumaiyana Technical & Non-Technical Manual")
    lines.append("")
    lines.append("**Pillar:** Models (AI Model Catalog, Providers, Limits, Routing, & Resilience)  ")
    lines.append("**Target Audience:** Developers, System Architects, CTOs, CFOs, Product Managers & Operations Teams  ")
    lines.append("**Language:** Bilingual (Thanglish - Tamil in English letters + Clean English)  ")
    lines.append("**Generated At:** 2026-09-05  ")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Table of Contents")
    lines.append("1. [Models System Architecture & Data Flow Map](#1-models-system-architecture--data-flow-map)")
    lines.append("2. [Detailed Feature Dissection (8 Core Models Features)](#2-detailed-feature-dissection)")
    for f in MODELS_FEATURES:
        lines.append(f"   - [{f['name_en']} ({f['name_tanglish']}) ({f['route']})](#{f['id']})")
    lines.append("3. [Cross-Feature Interconnections & Data Flow](#3-cross-feature-interconnections--data-flow)")
    lines.append("4. [Tech vs Non-Tech Comparative Matrix](#4-tech-vs-non-tech-comparative-matrix)")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 1. Models System Architecture & Data Flow Map")
    lines.append("### Muzhumaiyana Models Control Plane & Traffic Steering Flow")
    lines.append("")
    lines.append("```")
    lines.append("                     [ INCOMING CLIENT AI REQUEST ]")
    lines.append("                                   │")
    lines.append("                                   ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │               FastHTTP PROXY GATEWAY                      │")
    lines.append("       │  • Header extraction, API Key validation, Auth checks     │")
    lines.append("       └───────────────────────────┬───────────────────────────────┘")
    lines.append("                                   │")
    lines.append("                                   ▼")
    lines.append("       ┌───────────────────────────────────────────────────────────┐")
    lines.append("       │            BUDGETS & LIMITS ENFORCEMENT                   │")
    lines.append("       │  • Atomic token rate check, Budget threshold verification │")
    lines.append("       └─────────────┬───────────────────────────────┬─────────────┘")
    lines.append("        (Within Quota)│                               │ (Exceeded)")
    lines.append("                     ▼                               ▼")
    lines.append("       ┌───────────────────────────┐   ┌───────────────────────────┐")
    lines.append("       │    COMPLEXITY ROUTER      │   │  HTTP 429 / Cheaper Model │")
    lines.append("       │  • Score: 0.0 - 1.0       │   └───────────────────────────┘")
    lines.append("       │  • Simple/Med/Comp/Reason │")
    lines.append("       └─────────────┬─────────────┘")
    lines.append("                     │")
    lines.append("                     ▼")
    lines.append("       ┌───────────────────────────┐")
    lines.append("       │      ROUTING RULES        │◄─────────────────────────────┐")
    lines.append("       │  • CEL Expression Match   │                              │")
    lines.append("       │  • Prioritized Target Pool│                              │")
    lines.append("       └─────────────┬─────────────┘                              │")
    lines.append("                     │                                            │")
    lines.append("                     ▼                                            │")
    lines.append("       ┌───────────────────────────┐   (Outage / 429 Trip)        │")
    lines.append("       │      CIRCUIT BREAKER      │──────────────────────────────┘")
    lines.append("       │  • Closed / Open / Half   │   (Instant Fallback Reroute)")
    lines.append("       └─────────────┬─────────────┘")
    lines.append("                     │ (Healthy)")
    lines.append("                     ▼")
    lines.append("       ┌───────────────────────────┐")
    lines.append("       │      MODEL PROVIDERS      │◄─────── [ MODEL SETTINGS ]")
    lines.append("       │  • FastHTTP Client Pool   │          • Live Pricing Sync")
    lines.append("       │  • Vault Credentials      │          • Recursion Guard")
    lines.append("       └─────────────┬─────────────┘")
    lines.append("                     │")
    lines.append("                     ▼")
    lines.append("       ┌───────────────────────────┐")
    lines.append("       │   UPSTREAM LLM VENDORS    │")
    lines.append("       │  • OpenAI / Claude / AWS  │")
    lines.append("       └─────────────┬─────────────┘")
    lines.append("                     │")
    lines.append("                     ▼")
    lines.append("       ┌───────────────────────────┐")
    lines.append("       │     PRICING OVERRIDES     │───────► [ MODEL CATALOG ]")
    lines.append("       │  • Contracted Rates ($)   │          • 24h Spend & Usage")
    lines.append("       └───────────────────────────┘")
    lines.append("```")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 2. Detailed Feature Dissection (8 Core Models Features)")
    lines.append("### Ettu Features-oda Aazhamana Vivaram (Thanglish + English)")
    lines.append("")

    for f in MODELS_FEATURES:
        lines.append(f"<a name='{f['id']}'></a>")
        lines.append(f"## {f['name_en']} — {f['name_tanglish']}")
        lines.append(f"**UI Route:** `{f['route']}`\n")
        lines.append("### 👤 Non-Tech Perspective (Makkal Puriyura Mathiri Eliya Vilakkam)")
        lines.append(f"- **Uruvagam (Analogy):** {f['analogy']}")
        lines.append(f"- **Vilakkam (Explanation):** {f['explanation']}")
        lines.append(f"- **Business Value (Vaniga Payan):** {f['business_value']}\n")
        lines.append("### 💻 Tech Perspective (Engineers-kaga Technical Architecture)")
        lines.append(f"- **Backend Architecture:** {f['tech_arch']}")
        lines.append("- **Backend Endpoints:**")
        for ep in f['endpoints']:
            lines.append(f"  * `{ep}`")
        lines.append("")
        lines.append("### 🖥️ Screen Layout & Interactive Elements (Thirai Koorugal & Buttons)")
        lines.append("**1. Melpura Koorugal (Top Bar Controls):**")
        for itm in f['ui_elements']['top_bar']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("**2. Mathiya Koorugal & Tables (Tabs & Views):**")
        for itm in f['ui_elements']['tabs_and_views']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("**3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):**")
        for itm in f['ui_elements']['bottom_elements']:
            lines.append(f"- {itm}")
        lines.append("")
        lines.append("### 🔗 Connections & Structure Map (Inaippugal)")
        lines.append(f"- **Data-vai Perumidam (Receives From):** {f['connections']['receives_from']}")
        lines.append(f"- **Iyakkum Koorugal (Triggers & Affects):** {f['connections']['triggers_and_affects']}")
        lines.append("")
        lines.append(f"### 💡 Production Use Case: {f['use_case']}\n")
        lines.append("---\n")

    lines.append("# 3. Cross-Feature Interconnections & Data Flow")
    lines.append("### Models Features Kulla Irukura Direct Connections\n")
    lines.append("| Source Feature | Connected To | Data Flow & Trigger Action |")
    lines.append("| :--- | :--- | :--- |")
    for src, dst, flow in MODELS_CONN_DATA:
        lines.append(f"| **{src}** | **{dst}** | {flow} |")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("# 4. Tech vs Non-Tech Comparative Matrix")
    lines.append("### Thozhilnutpam vs Vanigam Parvai Oppeedo\n")
    lines.append("| Feature | Non-Tech View (Manager / CFO Parvai) | Tech View (DevOps / Architect Parvai) |")
    lines.append("| :--- | :--- | :--- |")
    for ft, nt, tv in MODELS_MATRIX_DATA:
        lines.append(f"| **{ft}** | \"{nt}\" | \"{tv}\" |")
    lines.append("")

    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Models Markdown written to: {MD_PATH}")

# ==============================================================================
# 2. GENERATE DOCX DOCUMENT (THANGLISH + ENGLISH)
# ==============================================================================
def generate_models_docx():
    print("Writing Models Word Document (.docx)...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
        hp = section.header.paragraphs[0]
        hp.text = "UnifAI Models Pillar Master Deep-Dive Technical Manual"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if hp.runs:
            hp.runs[0].font.name = "Segoe UI"
            hp.runs[0].font.size = Pt(8.5)
            hp.runs[0].font.color.rgb = RGBColor(100, 116, 139)
            
        fp = section.footer.paragraphs[0]
        fp.text = "Confidential & Proprietary — UnifAI Models Control Plane Manual (Thanglish + English)"
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if fp.runs:
            fp.runs[0].font.name = "Segoe UI"
            fp.runs[0].font.size = Pt(8.5)
            fp.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    def style_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(16)
        r.bold = True
        r.font.color.rgb = RGBColor(30, 58, 138)
        return p

    def style_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(13)
        r.bold = True
        r.font.color.rgb = RGBColor(37, 99, 235)
        return p

    def style_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(11)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        return p

    def add_p(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.italic = italic
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def add_bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Segoe UI"
            rb.font.size = Pt(10)
            rb.bold = True
            rb.font.color.rgb = RGBColor(15, 23, 42)
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    def format_table(table, col_widths, col_names):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, name in enumerate(col_names):
            hdr_cells[i].text = name
            tcPr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
            tcPr.append(shd)
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.name = "Segoe UI"
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        for r_idx, row in enumerate(table.rows[1:]):
            bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell in enumerate(row.cells):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                tcPr.append(shd)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = "Segoe UI"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(15, 23, 42)
                    
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    rt = p_title.add_run("UnifAI Models Pillar Master Deep-Dive Guide")
    rt.font.name = "Segoe UI"
    rt.font.size = Pt(22)
    rt.bold = True
    rt.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    rsub = p_sub.add_run("UnifAI Models Amaippin Muzhumaiyana Technical & Non-Technical Manual (Thanglish + English)")
    rsub.font.name = "Segoe UI"
    rsub.font.size = Pt(12)
    rsub.font.color.rgb = RGBColor(71, 85, 105)

    style_h1("1. Models System Architecture & Data Flow Map")
    add_p("Incoming client AI requests FastHTTP proxy vazhiya enter aagi, Budgets & Limits, Complexity Router, Routing Rules, Circuit Breaker, matrum Model Providers vazhiya execute aagura dynamic structure:")

    style_h1("2. Detailed Feature Dissection (8 Core Features)")
    for f in MODELS_FEATURES:
        style_h2(f"{f['name_en']} — {f['name_tanglish']}")
        add_p(f['route'], bold_prefix="Route: ", italic=True)

        style_h3("Non-Tech Perspective (Makkal Puriyura Eliya Vilakkam)")
        add_p(f['analogy'], bold_prefix="Uruvagam (Analogy): ")
        add_p(f['explanation'], bold_prefix="Vilakkam: ")
        add_p(f['business_value'], bold_prefix="Business Value: ")

        style_h3("Tech Perspective (Technical Architecture)")
        add_p(f['tech_arch'], bold_prefix="Backend Architecture: ")
        add_p("Backend Endpoints:")
        for ep in f['endpoints']:
            add_bullet(ep)

        style_h3("Screen Layout & Interactive Elements (Thirai Koorugal)")
        add_p("1. Melpura Koorugal (Top Bar Controls):", bold_prefix="")
        for itm in f['ui_elements']['top_bar']:
            add_bullet(itm)
        add_p("2. Mathiya Koorugal & Tables (Tabs & Views):", bold_prefix="")
        for itm in f['ui_elements']['tabs_and_views']:
            add_bullet(itm)
        add_p("3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):", bold_prefix="")
        for itm in f['ui_elements']['bottom_elements']:
            add_bullet(itm)

        style_h3("Module Connections (Inaippugal)")
        add_p(f['connections']['receives_from'], bold_prefix="Data-vai Perumidam: ")
        add_p(f['connections']['triggers_and_affects'], bold_prefix="Iyakkum Koorugal: ")

        style_h3("Production Use Case (Nadamurai Udharanam)")
        add_p(f['use_case'])

    style_h1("3. Cross-Feature Interconnections")
    t_conn = doc.add_table(rows=len(MODELS_CONN_DATA)+1, cols=3)
    format_table(t_conn, [1.8, 1.8, 3.4], ["Source Feature", "Connected To", "Data Flow & Trigger Action"])
    for idx, (src, dst, flow) in enumerate(MODELS_CONN_DATA):
        t_conn.rows[idx+1].cells[0].text = src
        t_conn.rows[idx+1].cells[1].text = dst
        t_conn.rows[idx+1].cells[2].text = flow

    style_h1("4. Tech vs Non-Tech Comparative Matrix")
    t_mat = doc.add_table(rows=len(MODELS_MATRIX_DATA)+1, cols=3)
    format_table(t_mat, [1.5, 2.7, 2.8], ["Feature", "Non-Tech View (Manager / CFO)", "Tech View (DevOps / Architect)"])
    for idx, (ft, nt, tv) in enumerate(MODELS_MATRIX_DATA):
        t_mat.rows[idx+1].cells[0].text = ft
        t_mat.rows[idx+1].cells[1].text = nt
        t_mat.rows[idx+1].cells[2].text = tv

    doc.save(DOCX_PATH)
    print(f"Models Word document written to: {DOCX_PATH}")

# ==============================================================================
# 3. GENERATE PDF DOCUMENT VIA REPORTLAB (THANGLISH + ENGLISH)
# ==============================================================================
def generate_models_pdf():
    print("Writing Models PDF Document (.pdf)...")

    class NumberedCanvas(canvas.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_page_states = []

        def showPage(self):
            self._saved_page_states.append(dict(self.__dict__))
            self._startPage()

        def save(self):
            num_pages = len(self._saved_page_states)
            for state in self._saved_page_states:
                self.__dict__.update(state)
                self.draw_page_decorations(num_pages)
                super().showPage()
            super().save()

        def draw_page_decorations(self, page_count):
            self.saveState()
            self.setFont("Helvetica", 8)
            self.setFillColor(colors.HexColor("#64748B"))
            
            if self._pageNumber > 1:
                self.drawString(40, 11 * inch - 36, "UnifAI Models Pillar Master Deep-Dive Technical Manual (Thanglish + English)")
                self.drawRightString(8.5 * inch - 40, 11 * inch - 36, "Confidential — Engineering Guide")
                self.setStrokeColor(colors.HexColor("#CBD5E1"))
                self.setLineWidth(0.5)
                self.line(40, 11 * inch - 40, 8.5 * inch - 40, 11 * inch - 40)
                
            self.drawString(40, 32, "UnifAI Unified Models & Routing Control Plane")
            page_str = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(8.5 * inch - 40, 32, page_str)
            self.setStrokeColor(colors.HexColor("#CBD5E1"))
            self.setLineWidth(0.5)
            self.line(40, 42, 8.5 * inch - 40, 42)
            
            self.restoreState()

    pdf_doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=letter,
        leftMargin=40,
        rightMargin=40,
        topMargin=48,
        bottomMargin=48
    )

    styles = getSampleStyleSheet()

    p_title_style = ParagraphStyle(
        'PdfTitle',
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=3
    )

    p_subtitle_style = ParagraphStyle(
        'PdfSubTitle',
        fontName='Helvetica',
        fontSize=9.5,
        leading=13,
        textColor=colors.HexColor('#475569'),
        spaceAfter=10
    )

    h1_style = ParagraphStyle(
        'PdfH1',
        fontName='Helvetica-Bold',
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=12,
        spaceAfter=5,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'PdfH2',
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=14,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=8,
        spaceAfter=3,
        keepWithNext=True
    )

    h3_style = ParagraphStyle(
        'PdfH3',
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#0F172A'),
        spaceBefore=5,
        spaceAfter=2,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'PdfBody',
        fontName='Helvetica',
        fontSize=8,
        leading=11,
        textColor=colors.HexColor('#334155'),
        spaceAfter=3
    )

    bullet_style = ParagraphStyle(
        'PdfBullet',
        fontName='Helvetica',
        fontSize=7.8,
        leading=10.5,
        textColor=colors.HexColor('#334155'),
        leftIndent=12,
        spaceAfter=2
    )

    th_style = ParagraphStyle(
        'PdfTH',
        fontName='Helvetica-Bold',
        fontSize=8,
        leading=10,
        textColor=colors.white
    )

    td_style = ParagraphStyle(
        'PdfTD',
        fontName='Helvetica',
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor('#0F172A')
    )

    td_code_style = ParagraphStyle(
        'PdfTDCode',
        fontName='Helvetica-Bold',
        fontSize=7.5,
        leading=10,
        textColor=colors.HexColor('#1E3A8A')
    )

    story = []

    # Title & Subtitle
    story.append(Paragraph("UnifAI Models Pillar Master Deep-Dive Technical Manual", p_title_style))
    story.append(Paragraph("Models Amaippin Muzhumaiyana Architecture, Features, Layout, & Data Flow Guide (Thanglish + English)", p_subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#2563EB'), spaceBefore=2, spaceAfter=8))

    # Architecture Overview
    story.append(Paragraph("1. Models System Architecture & Data Flow Map", h1_style))
    story.append(Paragraph(
        "UnifAI Models pillar is the central intelligence, cost-optimisation, and traffic routing control plane. "
        "Every incoming FastHTTP request passes through <b>Budgets & Limits</b> (token rate checks), "
        "<b>Complexity Router</b> (prompt triage: Simple, Medium, Complex, Reasoning), "
        "<b>Routing Rules</b> (CEL conditional steering), and <b>Circuit Breaker</b> (zero-downtime failover) "
        "before reaching upstream <b>Model Providers</b> (OpenAI, Anthropic, Bedrock, Vertex AI, Ollama). "
        "Rates are governed by <b>Pricing Overrides</b> and <b>Model Settings</b>, with complete visibility in <b>Model Catalog</b>.",
        body_style
    ))
    story.append(Spacer(1, 4))

    # Features Loop
    story.append(Paragraph("2. Detailed Feature Dissection (8 Core Models Features)", h1_style))
    for f in MODELS_FEATURES:
        feat_elements = []
        feat_elements.append(Paragraph(f"<b>{f['name_en']} — {f['name_tanglish']}</b>", h2_style))
        feat_elements.append(Paragraph(f"<b>Route:</b> <font color='#2563EB'>{f['route']}</font>", body_style))

        feat_elements.append(Paragraph("<b>👤 Non-Tech Perspective (Eliya Vilakkam):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Uruvagam (Analogy):</b> {f['analogy']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Vilakkam:</b> {f['explanation']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Business Value:</b> {f['business_value']}", body_style))

        feat_elements.append(Paragraph("<b>💻 Tech Perspective (Technical Architecture):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Backend Architecture:</b> {f['tech_arch']}", body_style))
        ep_text = ", ".join([f"<code>{e}</code>" for e in f['endpoints']])
        feat_elements.append(Paragraph(f"• <b>Endpoints:</b> {ep_text}", body_style))

        feat_elements.append(Paragraph("<b>🖥️ Screen Layout & Interactive Elements (Thirai Koorugal):</b>", h3_style))
        feat_elements.append(Paragraph("<b>1. Melpura Koorugal (Top Bar Controls):</b>", body_style))
        for itm in f['ui_elements']['top_bar']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))
        feat_elements.append(Paragraph("<b>2. Mathiya Koorugal & Tables (Tabs & Views):</b>", body_style))
        for itm in f['ui_elements']['tabs_and_views']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))
        feat_elements.append(Paragraph("<b>3. Keezhpura Koorugal & Sheets (Bottom Elements & Slide-Overs):</b>", body_style))
        for itm in f['ui_elements']['bottom_elements']:
            feat_elements.append(Paragraph(f"• {itm}", bullet_style))

        feat_elements.append(Paragraph("<b>🔗 Connections & Structure Map (Inaippugal):</b>", h3_style))
        feat_elements.append(Paragraph(f"• <b>Data-vai Perumidam:</b> {f['connections']['receives_from']}", body_style))
        feat_elements.append(Paragraph(f"• <b>Iyakkum Koorugal:</b> {f['connections']['triggers_and_affects']}", body_style))

        feat_elements.append(Paragraph(f"<b>💡 Production Use Case:</b> {f['use_case']}", body_style))
        feat_elements.append(Spacer(1, 6))
        story.append(KeepTogether(feat_elements))

    # Cross Connections
    story.append(Paragraph("3. Cross-Feature Interconnections & Data Flow", h1_style))
    conn_table_data = [[Paragraph(h, th_style) for h in ["Source Feature", "Connected To", "Data Flow & Trigger Action"]]]
    for src, dst, flow in MODELS_CONN_DATA:
        conn_table_data.append([Paragraph(src, td_code_style), Paragraph(dst, td_code_style), Paragraph(flow, td_style)])
    tbl_conn = Table(conn_table_data, colWidths=[105, 105, 320])
    tbl_conn.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_conn)
    story.append(Spacer(1, 8))

    # Comparative Matrix
    story.append(Paragraph("4. Tech vs Non-Tech Comparative Matrix", h1_style))
    mat_table_data = [[Paragraph(h, th_style) for h in ["Feature", "Non-Tech View (Manager / CFO Parvai)", "Tech View (DevOps / Architect Parvai)"]]]
    for ft, nt, tv in MODELS_MATRIX_DATA:
        mat_table_data.append([Paragraph(ft, td_code_style), Paragraph(nt, td_style), Paragraph(tv, td_style)])
    tbl_mat = Table(mat_table_data, colWidths=[95, 215, 220])
    tbl_mat.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A8A')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2.5),
        ('TOPPADDING', (0, 0), (-1, -1), 2.5),
        ('LEFTPADDING', (0, 0), (-1, -1), 3),
        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CBD5E1')),
    ]))
    story.append(tbl_mat)

    pdf_doc.build(story, canvasmaker=NumberedCanvas)
    print(f"Models PDF document written to: {PDF_PATH}")

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================
if __name__ == "__main__":
    print("Building dedicated Models Master Documentation (Thanglish + English)...")
    generate_models_markdown()
    generate_models_docx()
    generate_models_pdf()
    print("Models Master Documentation successfully built in:", DOC_DIR)
